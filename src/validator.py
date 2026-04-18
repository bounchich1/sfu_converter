import logging
from pathlib import Path
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from config import SIBFUConfig


class StyleValidator:
    """Валидатор стилей документа на соответствие стандарту СФУ"""
    
    def __init__(self, config_class=SIBFUConfig):
        """Инициализация валидатора с конфигурацией"""
        self.config = config_class
        self.errors = []
        self.warnings = []
        self.logger = logging.getLogger(__name__)

    def _get_pt_value(self, value):
        """Конвертирует значение в пункты (pt)"""
        if value is None:
            return 0
        if hasattr(value, 'pt'):
            return value.pt
        return float(value) if value else 0

    def _is_header_paragraph(self, para):
        """Определяет, является ли абзац заголовком (по выравниванию)"""
        pf = para.paragraph_format
        if pf.alignment == WD_ALIGN_PARAGRAPH.CENTER:
            return True
        return False

    def validate_font(self, run, para_index):
        """Проверяет шрифт и размер текста"""
        issues = []
        if run.font.name and run.font.name != self.config.FONT_NAME:
            issues.append(f"Абзац {para_index}: Шрифт '{run.font.name}' вместо '{self.config.FONT_NAME}'")
        
        current_size = self._get_pt_value(run.font.size)
        if current_size > 0 and abs(current_size - self.config.FONT_SIZE.pt) > 0.5:
            issues.append(f"Абзац {para_index}: Размер {current_size}pt вместо {self.config.FONT_SIZE.pt}pt")
            
        return issues

    def validate_paragraph_spacing(self, para, para_index):
        """Проверяет интервалы между абзацами (должны быть 0)"""
        issues = []
        pf = para.paragraph_format
        
        if pf.space_before and pf.space_before.pt > 1:
            issues.append(f"Абзац {para_index}: Интервал перед {pf.space_before.pt:.1f}pt (должен быть 0)")
        
        if pf.space_after and pf.space_after.pt > 1:
            issues.append(f"Абзац {para_index}: Интервал после {pf.space_after.pt:.1f}pt (должен быть 0)")

        return issues

    def validate_first_line_indent(self, para, para_index):
        """Проверяет отступ первой строки (только для обычных абзацев)"""
        issues = []
        
        # Заголовки не должны иметь отступ первой строки
        if self._is_header_paragraph(para):
            pf = para.paragraph_format
            current_indent = self._get_pt_value(pf.first_line_indent)
            if current_indent > 5:
                issues.append(f"Абзац {para_index}: Заголовок имеет отступ {current_indent:.1f}pt (должен быть 0)")
            return issues
        
        # Обычные абзацы должны иметь отступ 1.25 см
        pf = para.paragraph_format
        current_indent = self._get_pt_value(pf.first_line_indent)
        expected_indent_pt = 1.25 * 28.3465
        
        if abs(current_indent - expected_indent_pt) > 5:
            issues.append(f"Абзац {para_index}: Отступ {current_indent:.1f}pt вместо {expected_indent_pt:.1f}pt")
        
        return issues

    def validate_line_spacing(self, para, para_index, expected_spacing=1.5):
        """Проверяет межстрочный интервал"""
        issues = []
        pf = para.paragraph_format
        
        if pf.line_spacing:
            spacing = self._get_pt_value(pf.line_spacing)
            if spacing < 10:
                if abs(spacing - expected_spacing) > 0.1:
                    issues.append(f"Абзац {para_index}: Интервал {spacing} вместо {expected_spacing}")
        
        return issues

    def validate_file(self, file_path):
        """Выполняет полную валидацию документа"""
        self.logger.info(f"Начало валидации: {file_path}")
        self.errors = []
        self.warnings = []
        
        doc_path = Path(file_path)
        if not doc_path.exists():
            self.errors.append(f"Файл не найден: {file_path}")
            return False
        
        try:
            doc = Document(str(doc_path))
        except Exception as e:
            self.errors.append(f"Не удалось открыть файл: {e}")
            return False

        para_count = 0
        for para in doc.paragraphs:
            if not para.text.strip():
                para_count += 1
                continue
            
            para_count += 1
            
            if para.runs:
                font_issues = self.validate_font(para.runs[0], para_count)
                spacing_issues = self.validate_paragraph_spacing(para, para_count)
                indent_issues = self.validate_first_line_indent(para, para_count)
                
                self.errors.extend(font_issues)
                self.errors.extend(spacing_issues)
                self.errors.extend(indent_issues)

        if doc.tables:
            self.logger.info(f"Найдено таблиц: {len(doc.tables)}")
            for tbl_idx, table in enumerate(doc.tables):
                for row_idx, row in enumerate(table.rows):
                    for cell_idx, cell in enumerate(row.cells):
                        for para in cell.paragraphs:
                            if para.runs:
                                issues = self.validate_font(para.runs[0], f"Таблица {tbl_idx}")
                                self.errors.extend(issues)

        if self.errors:
            self.logger.warning(f"Найдено ошибок: {len(self.errors)}")
            for err in self.errors[:10]:
                self.logger.warning(err)
            return False
        else:
            self.logger.info(f"Валидация пройдена успешно (проверено {para_count} абзацев)")
            return True

    def get_report(self):
        """Возвращает отчет о валидации"""
        return {
            'errors': len(self.errors),
            'warnings': len(self.warnings),
            'error_list': self.errors,
            'warning_list': self.warnings
        }
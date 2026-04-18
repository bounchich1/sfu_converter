import re
import logging
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from config import SIBFUConfig
from utils_image_insert import insert_image


class TextToDocxConverter:
    """Конвертер TXT в DOCX с применением стилей СФУ"""
    
    def __init__(self, config_class=SIBFUConfig, base_dir=None):
        """Инициализация конвертера с конфигурацией"""
        self.config = config_class
        self.base_dir = Path(base_dir) if base_dir else Path.cwd()
        self.doc = None
        self.logger = logging.getLogger(__name__)

    def _set_run_style(self, run, bold=False):
        """Применяет шрифт Times New Roman к символу"""
        run.font.name = self.config.FONT_NAME
        run.font.size = self.config.FONT_SIZE
        run.font.color.rgb = RGBColor(*self.config.FONT_COLOR_RGB)
        run.bold = bold
        run._element.rPr.rFonts.set(qn('w:eastAsia'), self.config.FONT_NAME)

    def _set_paragraph_format(self, para, style_type='normal'):
        """Применяет форматирование абзаца с правильными интервалами"""
        pf = para.paragraph_format
        cfg = self.config
        
        if style_type == 'normal':
            pf.line_spacing = cfg.LINE_SPACING_NORMAL
            pf.alignment = cfg.ALIGNMENT
            pf.first_line_indent = cfg.FIRST_LINE_INDENT
            pf.space_before = Pt(0)
            pf.space_after = Pt(0)
            
        elif style_type == 'h1':
            para.style = self.doc.styles['Heading 1']
            pf.line_spacing = cfg.H1['line_spacing']
            pf.alignment = cfg.H1['align']
            pf.first_line_indent = cfg.H1['indent']
            pf.space_before = cfg.H1['space_before']
            pf.space_after = cfg.H1['space_after']

        elif style_type == 'h2':
            para.style = self.doc.styles['Heading 2']
            pf.line_spacing = cfg.H2['line_spacing']
            pf.alignment = cfg.H2['align']
            pf.first_line_indent = cfg.H2['indent']
            pf.space_before = cfg.H2['space_before']
            pf.space_after = cfg.H2['space_after']
            
        elif style_type == 'h3':
            para.style = self.doc.styles['Heading 3']
            pf.line_spacing = cfg.H3['line_spacing']
            pf.alignment = cfg.H3['align']
            pf.first_line_indent = cfg.H3['indent']
            pf.space_before = cfg.H3['space_before']
            pf.space_after = cfg.H3['space_after']

        elif style_type == 'caption_img':
            pf.line_spacing = cfg.CAPTION_IMAGE['line_spacing']
            pf.alignment = cfg.CAPTION_IMAGE['align']
            pf.first_line_indent = cfg.CAPTION_IMAGE['indent']
            pf.space_before = cfg.CAPTION_IMAGE['space_before']
            pf.space_after = cfg.CAPTION_IMAGE['space_after']
            
        elif style_type == 'caption_table':
            pf.line_spacing = cfg.CAPTION_TABLE['line_spacing']
            pf.alignment = cfg.CAPTION_TABLE['align']
            pf.first_line_indent = cfg.CAPTION_TABLE['indent']
            pf.space_before = cfg.CAPTION_TABLE['space_before']
            pf.space_after = cfg.CAPTION_TABLE['space_after']

        elif style_type == 'empty_before_header':
            pf.line_spacing = cfg.EMPTY_BEFORE_HEADER['line_spacing']
            pf.space_before = cfg.EMPTY_BEFORE_HEADER['space_before']
            pf.space_after = cfg.EMPTY_BEFORE_HEADER['space_after']
            pf.first_line_indent = Cm(0)

        elif style_type == 'empty_after_header':
            pf.line_spacing = cfg.EMPTY_AFTER_HEADER['line_spacing']
            pf.space_before = cfg.EMPTY_AFTER_HEADER['space_before']
            pf.space_after = cfg.EMPTY_AFTER_HEADER['space_after']
            pf.first_line_indent = Cm(0)
        
        elif style_type == 'empty_after_image':
            pf.line_spacing = cfg.EMPTY_AFTER_IMAGE['line_spacing']
            pf.space_before = cfg.EMPTY_AFTER_IMAGE['space_before']
            pf.space_after = cfg.EMPTY_AFTER_IMAGE['space_after']
            pf.first_line_indent = Cm(0)

        elif style_type == 'empty_before_image':
            pf.line_spacing = cfg.EMPTY_BEFORE_IMAGE['line_spacing']
            pf.space_before = Pt(0)
            pf.space_after = Pt(0)
            pf.first_line_indent = Cm(0)
        
        elif style_type == 'empty_before_table':
            pf.line_spacing = cfg.EMPTY_BEFORE_TABLE['line_spacing']
            pf.space_before = Pt(0)
            pf.space_after = Pt(0)
            pf.first_line_indent = Cm(0)
        
        elif style_type == 'empty_after_table':
            pf.line_spacing = cfg.EMPTY_AFTER_TABLE['line_spacing']
            pf.space_before = Pt(0)
            pf.space_after = Pt(0)
            pf.first_line_indent = Cm(0)

        if not para.runs:
            run = para.add_run()
        for run in para.runs:
            is_bold = style_type in ['h1', 'h2']
            self._set_run_style(run, bold=is_bold)

    def _add_empty_paragraph(self, style_type='empty_after_image'):
        """Добавляет пустой абзац с указанным стилем"""
        p = self.doc.add_paragraph()
        self._set_paragraph_format(p, style_type)

    def _insert_image(self, image_path=None, caption=None):
        """Вставляет изображение из директории images с автоматической фиксацией DPI"""
        if not image_path:
            self._add_empty_paragraph('empty_before_image')
            if caption:
                p = self.doc.add_paragraph(caption)
                self._set_paragraph_format(p, 'caption_img')
            self._add_empty_paragraph('empty_after_image')
            return

        full_path = self.base_dir / 'images' / image_path
        self._add_empty_paragraph('empty_before_image')

        if not full_path.exists():
            self.logger.warning(f"Изображение не найдено: {full_path}")
            p = self.doc.add_paragraph(f"[Изображение не найдено: {image_path}]")
            self._set_paragraph_format(p, 'caption_img')
        else:
            try:      
                success = insert_image(
                    doc=self.doc,
                    image_path=full_path,
                    config=self.config.IMAGE,
                    logger=self.logger
                )
                if success:
                    self.logger.info(f"Изображение вставлено: {image_path}")
                else:
                    self.logger.info(f"Ошибка вставки: {image_path}")
                
            except Exception as e:
                self.logger.error(f"Ошибка вставки изображения: {e}")
                p = self.doc.add_paragraph(f"[Ошибка: {image_path}]")
                self._set_paragraph_format(p, 'caption_img')

        if caption:
            p = self.doc.add_paragraph(caption)
            self._set_paragraph_format(p, 'caption_img')
        
        self._add_empty_paragraph('empty_after_image')

    def _parse_table_line(self, line):
        """Разбирает строку таблицы формата | Ячейка 1 | Ячейка 2 |"""
        line = line.strip()
        if not line.startswith('|') or not line.endswith('|'):
            return None
        cells = [cell.strip() for cell in line.split('|')[1:-1]]
        return cells if cells else None

    def _create_table(self, rows_data, caption=None):
        """Создает таблицу с подписью"""
        if not rows_data:
            return
        
        self._add_empty_paragraph('empty_before_table')
        
        if caption:
            p = self.doc.add_paragraph(caption)
            self._set_paragraph_format(p, 'caption_table')
        
        num_cols = len(rows_data[0])
        table = self.doc.add_table(rows=len(rows_data), cols=num_cols)
        table.style = 'Table Grid'
        table.autofit = False
        
        for row_idx, row_cells in enumerate(rows_data):
            row = table.rows[row_idx]
            for col_idx, text in enumerate(row_cells):
                if col_idx < len(row.cells):
                    cell = row.cells[col_idx]
                    cell.text = text
                    for para in cell.paragraphs:
                        is_header = (row_idx == 0)
                        pf = para.paragraph_format
                        pf.alignment = WD_ALIGN_PARAGRAPH.CENTER if is_header else WD_ALIGN_PARAGRAPH.LEFT
                        pf.space_before = self.config.TABLE_CELL_PADDING
                        pf.space_after = self.config.TABLE_CELL_PADDING
                        pf.first_line_indent = Cm(0)
                        
                        if not para.runs:
                            run = para.add_run()
                        for run in para.runs:
                            run.bold = is_header
                            self._set_run_style(run, bold=is_header)
        
        self._add_empty_paragraph('empty_after_table')

    def _setup_document_margins(self):
        """Настраивает поля страницы согласно стандарту"""
        for section in self.doc.sections:
            section.top_margin = self.config.MARGINS['top']
            section.bottom_margin = self.config.MARGINS['bottom']
            section.left_margin = self.config.MARGINS['left']
            section.right_margin = self.config.MARGINS['right']

    def _load_template(self, template_path):
        """Загружает шаблон документа если он существует"""
        template_file = self.base_dir / 'templates' / template_path
        if template_file.exists():
            self.doc = Document(str(template_file))
            self.logger.info(f"Шаблон загружен: {template_file}")
        else:
            self.doc = Document()
            self._setup_document_margins()
            self.logger.info("Создан новый документ")

    def _initialize_document(self, template=None):
        """Подготавливает документ перед наполнением контентом."""
        if template:
            self._load_template(template)
        else:
            self.doc = Document()
            self._setup_document_margins()

    def _render_lines(self, lines):
        """Преобразует строки исходного файла в содержимое DOCX."""
        i = 0
        while i < len(lines):
            line = lines[i].strip()
            
            if not line:
                i += 1
                continue

            if line.startswith('[H1]'):
                text = line.replace('[H1]', '').strip()
                p = self.doc.add_paragraph(text)
                self._set_paragraph_format(p, 'h1')
                self._add_empty_paragraph('empty_after_header')
                i += 1
                continue

            if line.startswith('[H2]'):
                self._add_empty_paragraph('empty_before_header')
                text = line.replace('[H2]', '').strip()
                p = self.doc.add_paragraph(text)
                self._set_paragraph_format(p, 'h2')
                self._add_empty_paragraph('empty_after_header')
                i += 1
                continue

            if line.startswith('[H3]'):
                text = line.replace('[H3]', '').strip()
                p = self.doc.add_paragraph(text)
                self._set_paragraph_format(p, 'h3')
                self._add_empty_paragraph('empty_after_header')
                i += 1
                continue

            if line.startswith('[IMAGE'):
                match = re.match(r'\[IMAGE(?:=([^\]]+))?\]', line)
                if match:
                    image_path = match.group(1).strip() if match.group(1) else None
                    
                    caption = None
                    if i + 1 < len(lines):
                        next_line = lines[i + 1].strip()
                        if next_line.startswith('Рисунок') or next_line.startswith('Figure'):
                            caption = next_line
                            i += 1
                    
                    self._insert_image(image_path, caption)
                i += 1
                continue

            if line.startswith('[TABLE_START]'):
                i += 1
                table_rows = []
                caption = None
                
                while i < len(lines) and lines[i].strip() != '[TABLE_END]':
                    current_line = lines[i].strip()
                    if current_line.startswith('[TABLE_CAPTION]'):
                        caption = current_line.replace('[TABLE_CAPTION]', '').strip()
                    else:
                        row = self._parse_table_line(current_line)
                        if row:
                            table_rows.append(row)
                    i += 1
                
                self._create_table(table_rows, caption)
                i += 1
                continue

            if line.startswith('[TABLE_CAPTION]'):
                text = line.replace('[TABLE_CAPTION]', '').strip()
                p = self.doc.add_paragraph(text)
                self._set_paragraph_format(p, 'caption_table')
                i += 1
                continue

            p = self.doc.add_paragraph(line)
            self._set_paragraph_format(p, 'normal')
            i += 1

    def convert_file(self, input_file: Path, output_file: Path, template: str | None = None):
        """Конвертирует TXT файл по явным входному и выходному путям."""
        input_file = Path(input_file)
        output_file = Path(output_file)

        self.logger.info(f"Начало конвертации: {input_file}")
        self._initialize_document(template)

        with open(input_file, 'r', encoding='utf-8') as f:
            lines = f.readlines()

        self._render_lines(lines)
        output_file.parent.mkdir(parents=True, exist_ok=True)
        self.doc.save(str(output_file))
        self.logger.info(f"Конвертация завершена: {output_file}")
        return str(output_file)

    def convert(self, input_path, output_path=None, template=None):
        """Совместимый режим конвертации через стандартные examples/ и results/."""
        if output_path is None:
            return None

        input_file = self.base_dir / 'examples' / input_path
        output_file = self.base_dir / 'results' / output_path
        return self.convert_file(input_file, output_file, template)

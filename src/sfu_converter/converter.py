import logging
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from sfu_converter.config import SIBFUConfig
from sfu_converter.domain.ast_nodes import (
    AppendixNode,
    BibliographyEntryNode,
    Document as AstDocument,
    FigureNode,
    FormulaNode,
    HeadingLevel,
    HeadingNode,
    ListNode,
    MetadataNode,
    PageBreakNode,
    ParagraphNode,
    RawBlockNode,
    TableCaptionNode,
    TableNode,
)
from sfu_converter.domain.diagnostics import Severity
from sfu_converter.parser.v1_parser import V1Parser
from sfu_converter.utils_image_insert import insert_image


class TextToDocxConverter:
    """Конвертер TXT в DOCX с применением стилей СФУ"""
    
    def __init__(self, config_class=SIBFUConfig, base_dir=None):
        """Инициализация конвертера с конфигурацией"""
        self.config = config_class
        self.base_dir = Path(base_dir) if base_dir else Path.cwd()
        self.doc = None
        self.logger = logging.getLogger(__name__)

    def _set_run_style(self, run, bold=False):
        """Применяет шрифт Times New Roman к символу"""
        run.font.name = self.config.FONT_NAME
        run.font.size = self.config.FONT_SIZE
        run.font.color.rgb = RGBColor(*self.config.FONT_COLOR_RGB)
        run.bold = bold
        run._element.rPr.rFonts.set(qn('w:eastAsia'), self.config.FONT_NAME)

    def _set_paragraph_format(self, para, style_type='normal'):
        """Применяет форматирование абзаца с правильными интервалами"""
        pf = para.paragraph_format
        cfg = self.config
        
        if style_type == 'normal':
            pf.line_spacing = cfg.LINE_SPACING_NORMAL
            pf.alignment = cfg.ALIGNMENT
            pf.first_line_indent = cfg.FIRST_LINE_INDENT
            pf.space_before = Pt(0)
            pf.space_after = Pt(0)
            
        elif style_type == 'h1':
            para.style = self.doc.styles['Heading 1']
            pf.line_spacing = cfg.H1['line_spacing']
            pf.alignment = cfg.H1['align']
            pf.first_line_indent = cfg.H1['indent']
            pf.space_before = cfg.H1['space_before']
            pf.space_after = cfg.H1['space_after']

        elif style_type == 'h2':
            para.style = self.doc.styles['Heading 2']
            pf.line_spacing = cfg.H2['line_spacing']
            pf.alignment = cfg.H2['align']
            pf.first_line_indent = cfg.H2['indent']
            pf.space_before = cfg.H2['space_before']
            pf.space_after = cfg.H2['space_after']
            
        elif style_type == 'h3':
            para.style = self.doc.styles['Heading 3']
            pf.line_spacing = cfg.H3['line_spacing']
            pf.alignment = cfg.H3['align']
            pf.first_line_indent = cfg.H3['indent']
            pf.space_before = cfg.H3['space_before']
            pf.space_after = cfg.H3['space_after']

        elif style_type == 'caption_img':
            pf.line_spacing = cfg.CAPTION_IMAGE['line_spacing']
            pf.alignment = cfg.CAPTION_IMAGE['align']
            pf.first_line_indent = cfg.CAPTION_IMAGE['indent']
            pf.space_before = cfg.CAPTION_IMAGE['space_before']
            pf.space_after = cfg.CAPTION_IMAGE['space_after']
            
        elif style_type == 'caption_table':
            pf.line_spacing = cfg.CAPTION_TABLE['line_spacing']
            pf.alignment = cfg.CAPTION_TABLE['align']
            pf.first_line_indent = cfg.CAPTION_TABLE['indent']
            pf.space_before = cfg.CAPTION_TABLE['space_before']
            pf.space_after = cfg.CAPTION_TABLE['space_after']

        elif style_type == 'empty_before_header':
            pf.line_spacing = cfg.EMPTY_BEFORE_HEADER['line_spacing']
            pf.space_before = cfg.EMPTY_BEFORE_HEADER['space_before']
            pf.space_after = cfg.EMPTY_BEFORE_HEADER['space_after']
            pf.first_line_indent = Cm(0)

        elif style_type == 'empty_after_header':
            pf.line_spacing = cfg.EMPTY_AFTER_HEADER['line_spacing']
            pf.space_before = cfg.EMPTY_AFTER_HEADER['space_before']
            pf.space_after = cfg.EMPTY_AFTER_HEADER['space_after']
            pf.first_line_indent = Cm(0)
        
        elif style_type == 'empty_after_image':
            pf.line_spacing = cfg.EMPTY_AFTER_IMAGE['line_spacing']
            pf.space_before = cfg.EMPTY_AFTER_IMAGE['space_before']
            pf.space_after = cfg.EMPTY_AFTER_IMAGE['space_after']
            pf.first_line_indent = Cm(0)

        elif style_type == 'empty_before_image':
            pf.line_spacing = cfg.EMPTY_BEFORE_IMAGE['line_spacing']
            pf.space_before = Pt(0)
            pf.space_after = Pt(0)
            pf.first_line_indent = Cm(0)
        
        elif style_type == 'empty_before_table':
            pf.line_spacing = cfg.EMPTY_BEFORE_TABLE['line_spacing']
            pf.space_before = Pt(0)
            pf.space_after = Pt(0)
            pf.first_line_indent = Cm(0)
        
        elif style_type == 'empty_after_table':
            pf.line_spacing = cfg.EMPTY_AFTER_TABLE['line_spacing']
            pf.space_before = Pt(0)
            pf.space_after = Pt(0)
            pf.first_line_indent = Cm(0)

        if not para.runs:
            run = para.add_run()
        for run in para.runs:
            is_bold = style_type in ['h1', 'h2']
            self._set_run_style(run, bold=is_bold)

    def _add_empty_paragraph(self, style_type='empty_after_image'):
        """Добавляет пустой абзац с указанным стилем"""
        p = self.doc.add_paragraph()
        self._set_paragraph_format(p, style_type)

    def _insert_image(self, image_path=None, caption=None):
        """Вставляет изображение из директории images с автоматической фиксацией DPI"""
        if not image_path:
            self._add_empty_paragraph('empty_before_image')
            if caption:
                p = self.doc.add_paragraph(caption)
                self._set_paragraph_format(p, 'caption_img')
            self._add_empty_paragraph('empty_after_image')
            return

        full_path = self.base_dir / 'images' / image_path
        self._add_empty_paragraph('empty_before_image')

        if not full_path.exists():
            self.logger.warning(f"Изображение не найдено: {full_path}")
            p = self.doc.add_paragraph(f"[Изображение не найдено: {image_path}]")
            self._set_paragraph_format(p, 'caption_img')
        else:
            try:      
                success = insert_image(
                    doc=self.doc,
                    image_path=full_path,
                    config=self.config.IMAGE,
                    logger=self.logger
                )
                if success:
                    self.logger.info(f"Изображение вставлено: {image_path}")
                else:
                    self.logger.info(f"Ошибка вставки: {image_path}")
                
            except Exception as e:
                self.logger.error(f"Ошибка вставки изображения: {e}")
                p = self.doc.add_paragraph(f"[Ошибка: {image_path}]")
                self._set_paragraph_format(p, 'caption_img')

        if caption:
            p = self.doc.add_paragraph(caption)
            self._set_paragraph_format(p, 'caption_img')
        
        self._add_empty_paragraph('empty_after_image')

    def _parse_table_line(self, line):
        """Разбирает строку таблицы формата | Ячейка 1 | Ячейка 2 |"""
        line = line.strip()
        if not line.startswith('|') or not line.endswith('|'):
            return None
        cells = [cell.strip() for cell in line.split('|')[1:-1]]
        return cells if cells else None

    def _create_table(self, rows_data, caption=None):
        """Создает таблицу с подписью"""
        if not rows_data:
            return
        
        self._add_empty_paragraph('empty_before_table')
        
        if caption:
            p = self.doc.add_paragraph(caption)
            self._set_paragraph_format(p, 'caption_table')
        
        num_cols = len(rows_data[0])
        table = self.doc.add_table(rows=len(rows_data), cols=num_cols)
        table.style = 'Table Grid'
        table.autofit = False
        
        for row_idx, row_cells in enumerate(rows_data):
            row = table.rows[row_idx]
            for col_idx, text in enumerate(row_cells):
                if col_idx < len(row.cells):
                    cell = row.cells[col_idx]
                    cell.text = text
                    for para in cell.paragraphs:
                        is_header = (row_idx == 0)
                        pf = para.paragraph_format
                        pf.alignment = WD_ALIGN_PARAGRAPH.CENTER if is_header else WD_ALIGN_PARAGRAPH.LEFT
                        pf.space_before = self.config.TABLE_CELL_PADDING
                        pf.space_after = self.config.TABLE_CELL_PADDING
                        pf.first_line_indent = Cm(0)
                        
                        if not para.runs:
                            run = para.add_run()
                        for run in para.runs:
                            run.bold = is_header
                            self._set_run_style(run, bold=is_header)
        
        self._add_empty_paragraph('empty_after_table')

    def _setup_document_margins(self):
        """Настраивает поля страницы согласно стандарту"""
        for section in self.doc.sections:
            section.top_margin = self.config.MARGINS['top']
            section.bottom_margin = self.config.MARGINS['bottom']
            section.left_margin = self.config.MARGINS['left']
            section.right_margin = self.config.MARGINS['right']

    def _load_template(self, template_path):
        """Загружает шаблон документа если он существует"""
        template_file = self.base_dir / 'templates' / template_path
        if template_file.exists():
            self.doc = Document(str(template_file))
            self.logger.info(f"Шаблон загружен: {template_file}")
        else:
            self.doc = Document()
            self._setup_document_margins()
            self.logger.info("Создан новый документ")

    def _initialize_document(self, template=None):
        """Подготавливает документ перед наполнением контентом."""
        if template:
            self._load_template(template)
        else:
            self.doc = Document()
            self._setup_document_margins()

    def _render_lines(self, lines):
        """Преобразует строки исходного файла или AST в содержимое DOCX."""
        if isinstance(lines, AstDocument):
            self._render_from_ast(lines)
            return

        source = self._lines_to_source(lines)
        result = V1Parser().parse(source)
        self._log_parser_diagnostics(result.diagnostics)
        self._render_from_ast(result.document)

    def _lines_to_source(self, lines):
        if isinstance(lines, str):
            return lines

        source_lines = list(lines)
        if any(line.endswith(('\n', '\r')) for line in source_lines):
            return ''.join(source_lines)
        return '\n'.join(source_lines)

    def _log_parser_diagnostics(self, diagnostics):
        for diagnostic in diagnostics:
            line = diagnostic.source.line_start if diagnostic.source else "?"
            message = f"{diagnostic.code} at line {line}: {diagnostic.message}"
            if diagnostic.severity in (Severity.ERROR, Severity.FATAL):
                self.logger.error(message)
            else:
                self.logger.warning(message)

    def _render_from_ast(self, document):
        """Render a parsed domain document into the current document."""
        for block in document.blocks:
            if isinstance(block, HeadingNode):
                self._render_heading(block)
            elif isinstance(block, ParagraphNode):
                self._render_paragraph(block)
            elif isinstance(block, TableNode):
                rows = [[cell.text for cell in row.cells] for row in block.rows]
                self._create_table(rows, block.caption)
            elif isinstance(block, TableCaptionNode):
                p = self.doc.add_paragraph(block.text)
                self._set_paragraph_format(p, 'caption_table')
            elif isinstance(block, FigureNode):
                self._insert_image(block.src, block.caption)
            elif isinstance(block, PageBreakNode):
                self.doc.add_page_break()
            elif isinstance(block, FormulaNode):
                self._render_text_block(block.content)
            elif isinstance(block, ListNode):
                for item in block.items:
                    self._render_text_block(item.text)
            elif isinstance(block, AppendixNode):
                self._render_heading(HeadingNode(level=HeadingLevel.H1, text=block.title))
                self._render_from_ast(AstDocument(blocks=block.blocks))
            elif isinstance(block, BibliographyEntryNode):
                self._render_text_block(f"{block.number}. {block.text}")
            elif isinstance(block, RawBlockNode):
                self._render_text_block(block.text)
            elif isinstance(block, MetadataNode):
                continue

    def _render_heading(self, block):
        if block.level is HeadingLevel.H2:
            self._add_empty_paragraph('empty_before_header')

        style_type = {
            HeadingLevel.H1: 'h1',
            HeadingLevel.H2: 'h2',
            HeadingLevel.H3: 'h3',
        }[block.level]
        p = self.doc.add_paragraph(block.text)
        self._set_paragraph_format(p, style_type)
        self._add_empty_paragraph('empty_after_header')

    def _render_paragraph(self, block):
        self._render_text_block(''.join(run.text for run in block.runs))

    def _render_text_block(self, text):
        p = self.doc.add_paragraph(text)
        self._set_paragraph_format(p, 'normal')

    def convert_file(self, input_file: Path, output_file: Path, template: str | None = None):
        """Конвертирует TXT файл по явным входному и выходному путям."""
        input_file = Path(input_file)
        output_file = Path(output_file)

        self.logger.info(f"Начало конвертации: {input_file}")
        self._initialize_document(template)

        with open(input_file, 'r', encoding='utf-8') as f:
            lines = f.readlines()

        self._render_lines(lines)
        output_file.parent.mkdir(parents=True, exist_ok=True)
        self.doc.save(str(output_file))
        self.logger.info(f"Конвертация завершена: {output_file}")
        return str(output_file)

    def convert(self, input_path, output_path=None, template=None):
        """Совместимый режим конвертации через стандартные examples/ и results/."""
        if output_path is None:
            return None

        input_file = self.base_dir / 'examples' / input_path
        output_file = self.base_dir / 'results' / output_path
        return self.convert_file(input_file, output_file, template)

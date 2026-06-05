from __future__ import annotations

import logging
import re
from io import BytesIO
from pathlib import Path

from docx import Document as DocxDocument
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

from sfu_converter.config import PathConfig, SIBFUConfig
from sfu_converter.domain.constants import (
    APPENDIX_TITLE,
    DASH_SEPARATOR,
    EM_DASH,
    EN_DASH,
)
from sfu_converter.domain.ast_nodes import (
    AppendixNode,
    BibliographyEntryNode,
    CitationNode,
    ContinuationLabel,
    Document,
    DrawingSheetNode,
    FigureNode,
    FrameType,
    FootnoteAnchor,
    FootnoteNode,
    FormulaNode,
    HeadingLevel,
    HeadingNode,
    ListNode,
    ListType,
    MetadataNode,
    PageBreakNode,
    ParagraphNode,
    PosterNode,
    ProjectDesignationNode,
    RawBlockNode,
    ReferenceNode,
    SectionSetupNode,
    SectionOrientation,
    SheetFormat,
    SourceRecordNode,
    SlideDeckNode,
    StructuralSectionNode,
    StructuralSectionType,
    TableCaptionNode,
    TableNote,
    TableNode,
    TableOfContentsNode,
    TitleBlockForm,
    TitlePageNode,
)
from sfu_converter.domain.diagnostics import Diagnostic, DiagnosticCodes, Severity
from sfu_converter.domain.formatting import FormattingProfile, unsupported_rule_diagnostics
from sfu_converter.domain.reference_graph import build_reference_graph
from sfu_converter.infrastructure.abbreviations import abbreviations_for_document, explicit_abbreviations
from sfu_converter.infrastructure.appendix import assign_appendix_letters
from sfu_converter.infrastructure.bibliography import format_record, validate_records
from sfu_converter.infrastructure import docx_styles
from sfu_converter.infrastructure.figure_layout import (
    figure_caption_text,
    figure_reference_diagnostics,
    normalize_caption_dashes,
)
from sfu_converter.infrastructure.formula_layout import (
    explanation_lines,
    split_formula_lines,
)
from sfu_converter.infrastructure.footnotes import add_footnote_reference, patch_docx_bytes, patch_docx_file
from sfu_converter.infrastructure import frames, main_inscription, section_setup
from sfu_converter.infrastructure.docx_measurements import ABBREVIATIONS_COLUMN_WIDTHS, NO_INDENT_CM
from sfu_converter.infrastructure.list_layout import apply_list_item_layout, list_marker
from sfu_converter.infrastructure.numbering import NumberingContext, build_numbering_context
from sfu_converter.infrastructure.page_numbering import (
    Location,
    PageNumberingSection,
    configure as configure_page_numbering,
)
from sfu_converter.infrastructure.project_designation import format_designation
from sfu_converter.infrastructure.toc import TocEntry, TocField, build_toc_field
from sfu_converter.parser.citations import format_citation_node
from sfu_converter.ports.renderer import RendererPort
from sfu_converter.registry import get_profile
from sfu_converter.utils_image_insert import insert_image


_SFU_STYLE_BY_TYPE = {
    "caption_img": docx_styles.FIGURE_CAPTION,
    "figure_explanatory": docx_styles.FIGURE_EXPLANATORY,
    "caption_table": docx_styles.TABLE_CAPTION,
    "table_unit": docx_styles.TABLE_UNIT,
    "formula": docx_styles.FORMULA_BODY,
    "formula_explanation": docx_styles.FORMULA_EXPLANATION,
    "bibliography_entry": docx_styles.BIBLIOGRAPHY_ENTRY,
    "list_item": docx_styles.LIST_ITEM,
    "structural_section": docx_styles.STRUCTURAL_HEADING,
    "toc_heading": docx_styles.TOC_HEADING,
    "appendix_heading": docx_styles.APPENDIX_HEADING,
}

_ITALIC_LETTER_RE = re.compile(r"^[A-Za-zα-ωΑ-Ω]$")


class SectionNumberer:
    """Tracks hierarchical section numbers for H1-H4 headings."""

    def __init__(self):
        self._counters = [0, 0, 0, 0]

    def next_number(self, level: int) -> str:
        if level < 1 or level > len(self._counters):
            raise ValueError(f"Unsupported heading level: {level}")

        index = level - 1
        for parent_index in range(index):
            if self._counters[parent_index] == 0:
                self._counters[parent_index] = 1

        self._counters[index] += 1
        for lower_index in range(index + 1, len(self._counters)):
            self._counters[lower_index] = 0

        return ".".join(str(part) for part in self._counters[:level])

    def reset(self):
        self._counters = [0, 0, 0, 0]


class DocxRenderer(RendererPort):
    """python-docx renderer for the domain AST."""

    def __init__(self, config_class=SIBFUConfig, base_dir=None, logger=None):
        self.config = config_class
        self.base_dir = Path(base_dir) if base_dir else Path.cwd()
        self.doc = None
        self.logger = logger or logging.getLogger(__name__)
        self._style_map = self._build_style_map()
        self._bold_styles = frozenset({"h1", "h2", "structural_section", "toc_heading", "appendix_heading"})
        self._section_numberer = SectionNumberer()
        self._rendered_body_blocks = False

    def render(
        self,
        document: Document,
        profile: FormattingProfile,
        template_path: str | None = None,
        template_mode: str = "append",
    ) -> bytes:
        document, appendix_diagnostics = assign_appendix_letters(document)
        self._initialize_document(template_path, template_mode=template_mode, profile=profile)
        self._appendix_diagnostics = appendix_diagnostics
        self._prepare_document_level_state(document)
        self._render_from_ast(document)
        buffer = BytesIO()
        self.doc.save(buffer)
        return patch_docx_bytes(buffer.getvalue(), self._rendered_footnotes)

    def render_to_file(
        self,
        document: Document,
        profile: FormattingProfile,
        output_path: str,
        template_path: str | None = None,
        template_mode: str = "append",
    ) -> list[Diagnostic]:
        document, appendix_diagnostics = assign_appendix_letters(document)
        diagnostics = unsupported_rule_diagnostics(profile, component="renderer")
        self._initialize_document(template_path, template_mode=template_mode, profile=profile)
        self._appendix_diagnostics = appendix_diagnostics
        self._prepare_document_level_state(document)
        reference_graph = build_reference_graph(document)
        self._reference_diagnostics = reference_graph.diagnostics()
        self._figure_diagnostics = figure_reference_diagnostics(document, reference_graph)
        self._render_from_ast(document)
        destination = Path(output_path)
        destination.parent.mkdir(parents=True, exist_ok=True)
        self.doc.save(str(destination))
        patch_docx_file(destination, self._rendered_footnotes)
        return (
            diagnostics
            + self._appendix_diagnostics
            + self._title_page_diagnostics
            + self._abbreviation_diagnostics
            + self._toc_diagnostics
            + self._reference_diagnostics
            + self._figure_diagnostics
            + self._bibliography_diagnostics
            + self._footnote_diagnostics
        )

    def _set_run_style(self, run, bold=False, italic=False):
        run.font.name = self.config.FONT_NAME
        run.font.size = self.config.FONT_SIZE
        run.font.color.rgb = RGBColor(*self.config.FONT_COLOR_RGB)
        run.bold = bold
        run.italic = italic
        run._element.rPr.rFonts.set(qn("w:eastAsia"), self.config.FONT_NAME)

    def _build_style_map(self):
        """Return style_type -> formatting parameters dict.

        Each entry may carry the optional keys ``word_style``, ``align``,
        ``indent``, ``line_spacing``, ``space_before``, ``space_after``. Only
        present keys are applied to the paragraph; missing keys are left
        untouched. Resolved once per renderer instance and cached.
        """

        cfg = self.config
        zero = Pt(0)
        no_indent = NO_INDENT_CM

        return {
            "normal": {
                "align": cfg.ALIGNMENT,
                "indent": cfg.FIRST_LINE_INDENT,
                "line_spacing": cfg.LINE_SPACING_NORMAL,
                "space_before": zero,
                "space_after": zero,
            },
            "h1": {
                "word_style": "Heading 1",
                "align": cfg.H1["align"],
                "indent": cfg.H1["indent"],
                "line_spacing": cfg.H1["line_spacing"],
                "space_before": cfg.H1["space_before"],
                "space_after": cfg.H1["space_after"],
            },
            "h2": {
                "word_style": "Heading 2",
                "align": cfg.H2["align"],
                "indent": cfg.H2["indent"],
                "line_spacing": cfg.H2["line_spacing"],
                "space_before": cfg.H2["space_before"],
                "space_after": cfg.H2["space_after"],
            },
            "h3": {
                "word_style": "Heading 3",
                "align": cfg.H3["align"],
                "indent": cfg.H3["indent"],
                "line_spacing": cfg.H3["line_spacing"],
                "space_before": cfg.H3["space_before"],
                "space_after": cfg.H3["space_after"],
            },
            "h4": {
                "word_style": "Heading 4",
                "align": cfg.H4["align"],
                "indent": cfg.H4["indent"],
                "line_spacing": cfg.H4["line_spacing"],
                "space_before": cfg.H4["space_before"],
                "space_after": cfg.H4["space_after"],
            },
            "list_item": {
                "align": cfg.LIST_ITEM["align"],
                "indent": cfg.LIST_ITEM["indent"],
                "line_spacing": cfg.LIST_ITEM["line_spacing"],
                "space_before": cfg.LIST_ITEM["space_before"],
                "space_after": cfg.LIST_ITEM["space_after"],
            },
            "structural_section": {
                "align": cfg.STRUCTURAL_SECTION["align"],
                "indent": cfg.STRUCTURAL_SECTION["indent"],
                "line_spacing": cfg.STRUCTURAL_SECTION["line_spacing"],
                "space_before": cfg.STRUCTURAL_SECTION["space_before"],
                "space_after": cfg.STRUCTURAL_SECTION["space_after"],
            },
            "toc_heading": {
                "align": cfg.STRUCTURAL_SECTION["align"],
                "indent": cfg.STRUCTURAL_SECTION["indent"],
                "line_spacing": cfg.STRUCTURAL_SECTION["line_spacing"],
                "space_before": cfg.STRUCTURAL_SECTION["space_before"],
                "space_after": cfg.STRUCTURAL_SECTION["space_after"],
            },
            "appendix_heading": {
                "align": cfg.STRUCTURAL_SECTION["align"],
                "indent": cfg.STRUCTURAL_SECTION["indent"],
                "line_spacing": cfg.STRUCTURAL_SECTION["line_spacing"],
                "space_before": cfg.STRUCTURAL_SECTION["space_before"],
                "space_after": cfg.STRUCTURAL_SECTION["space_after"],
            },
            "caption_img": {
                "align": cfg.CAPTION_IMAGE["align"],
                "indent": cfg.CAPTION_IMAGE["indent"],
                "line_spacing": cfg.CAPTION_IMAGE["line_spacing"],
                "space_before": cfg.CAPTION_IMAGE["space_before"],
                "space_after": cfg.CAPTION_IMAGE["space_after"],
            },
            "figure_explanatory": {
                "align": WD_ALIGN_PARAGRAPH.CENTER,
                "indent": no_indent,
                "line_spacing": cfg.CAPTION_IMAGE["line_spacing"],
                "space_before": zero,
                "space_after": zero,
            },
            "caption_table": {
                "align": cfg.CAPTION_TABLE["align"],
                "indent": cfg.CAPTION_TABLE["indent"],
                "line_spacing": cfg.CAPTION_TABLE["line_spacing"],
                "space_before": cfg.CAPTION_TABLE["space_before"],
                "space_after": cfg.CAPTION_TABLE["space_after"],
            },
            "table_unit": {
                "align": WD_ALIGN_PARAGRAPH.RIGHT,
                "indent": no_indent,
                "line_spacing": 1.0,
                "space_before": zero,
                "space_after": zero,
            },
            "empty_before_header": {
                "indent": no_indent,
                "line_spacing": cfg.EMPTY_BEFORE_HEADER["line_spacing"],
                "space_before": cfg.EMPTY_BEFORE_HEADER["space_before"],
                "space_after": cfg.EMPTY_BEFORE_HEADER["space_after"],
            },
            "empty_after_header": {
                "indent": no_indent,
                "line_spacing": cfg.EMPTY_AFTER_HEADER["line_spacing"],
                "space_before": cfg.EMPTY_AFTER_HEADER["space_before"],
                "space_after": cfg.EMPTY_AFTER_HEADER["space_after"],
            },
            "empty_before_image": {
                "indent": no_indent,
                "line_spacing": cfg.EMPTY_BEFORE_IMAGE["line_spacing"],
                "space_before": zero,
                "space_after": zero,
            },
            "empty_after_image": {
                "indent": no_indent,
                "line_spacing": cfg.EMPTY_AFTER_IMAGE["line_spacing"],
                "space_before": cfg.EMPTY_AFTER_IMAGE["space_before"],
                "space_after": cfg.EMPTY_AFTER_IMAGE["space_after"],
            },
            "empty_before_table": {
                "indent": no_indent,
                "line_spacing": cfg.EMPTY_BEFORE_TABLE["line_spacing"],
                "space_before": zero,
                "space_after": zero,
            },
            "empty_after_table": {
                "indent": no_indent,
                "line_spacing": cfg.EMPTY_AFTER_TABLE["line_spacing"],
                "space_before": zero,
                "space_after": zero,
            },
            "empty_before_formula": {
                "indent": no_indent,
                "line_spacing": cfg.EMPTY_BEFORE_FORMULA["line_spacing"],
                "space_before": zero,
                "space_after": zero,
            },
            "empty_after_formula": {
                "indent": no_indent,
                "line_spacing": cfg.EMPTY_AFTER_FORMULA["line_spacing"],
                "space_before": zero,
                "space_after": zero,
            },
            "formula": {
                "align": cfg.FORMULA["align"],
                "indent": cfg.FORMULA["indent"],
                "line_spacing": cfg.FORMULA["line_spacing"],
                "space_before": cfg.FORMULA["space_before"],
                "space_after": cfg.FORMULA["space_after"],
            },
            "formula_explanation": {
                "align": cfg.FORMULA_EXPLANATION["align"],
                "indent": cfg.FORMULA_EXPLANATION["indent"],
                "line_spacing": cfg.FORMULA_EXPLANATION["line_spacing"],
                "space_before": cfg.FORMULA_EXPLANATION["space_before"],
                "space_after": cfg.FORMULA_EXPLANATION["space_after"],
            },
            "bibliography_entry": {
                "align": cfg.BIBLIOGRAPHY_ENTRY["align"],
                "indent": cfg.BIBLIOGRAPHY_ENTRY["indent"],
                "line_spacing": cfg.BIBLIOGRAPHY_ENTRY["line_spacing"],
                "space_before": cfg.BIBLIOGRAPHY_ENTRY["space_before"],
                "space_after": cfg.BIBLIOGRAPHY_ENTRY["space_after"],
            },
        }

    def _set_paragraph_format(self, para, style_type="normal"):
        style = self._style_map.get(style_type)
        if style is None:
            self.logger.warning("Unknown style_type: %s", style_type)
            return

        word_style = style.get("word_style")
        sfu_style = _SFU_STYLE_BY_TYPE.get(style_type)
        if sfu_style is not None and self.doc is not None and sfu_style in [s.name for s in self.doc.styles]:
            para.style = self.doc.styles[sfu_style]
        elif word_style is not None and self.doc is not None:
            docx_styles.apply_word_heading_style(self.doc, para, word_style)

        pf = para.paragraph_format
        if "align" in style:
            pf.alignment = style["align"]
        if "indent" in style:
            pf.first_line_indent = style["indent"]
        if "line_spacing" in style:
            pf.line_spacing = style["line_spacing"]
        if "space_before" in style:
            pf.space_before = style["space_before"]
        if "space_after" in style:
            pf.space_after = style["space_after"]

        if not para.runs:
            para.add_run()
        bold = style_type in self._bold_styles
        for run in para.runs:
            self._set_run_style(run, bold=bold)

    def _add_empty_paragraph(self, style_type="empty_after_image"):
        p = self.doc.add_paragraph()
        self._set_paragraph_format(p, style_type)

    def _insert_image(self, image_path=None, caption=None, *, explanatory_data=None):
        if not image_path:
            self._add_empty_paragraph("empty_before_image")
            self._add_figure_explanatory_data(explanatory_data)
            if caption:
                p = self.doc.add_paragraph(caption)
                self._set_paragraph_format(p, "caption_img")
            self._add_empty_paragraph("empty_after_image")
            self._rendered_body_blocks = True
            return

        full_path = self._resolve_image_path(image_path)
        self._add_empty_paragraph("empty_before_image")

        if not full_path.exists():
            self.logger.warning(f"Изображение не найдено: {full_path}")
            p = self.doc.add_paragraph(f"[Изображение не найдено: {image_path}]")
            self._set_paragraph_format(p, "caption_img")
            if docx_styles.FIGURE_PLACEHOLDER in [s.name for s in self.doc.styles]:
                p.style = self.doc.styles[docx_styles.FIGURE_PLACEHOLDER]
        else:
            try:
                success = insert_image(
                    doc=self.doc,
                    image_path=full_path,
                    config=self.config.IMAGE,
                    logger=self.logger,
                )
                if success:
                    self.logger.info(f"Изображение вставлено: {image_path}")
                else:
                    self.logger.info(f"Ошибка вставки: {image_path}")
            except Exception as exc:
                self.logger.error(f"Ошибка вставки изображения: {exc}")
                p = self.doc.add_paragraph(f"[Ошибка: {image_path}]")
                self._set_paragraph_format(p, "caption_img")
                if docx_styles.FIGURE_PLACEHOLDER in [s.name for s in self.doc.styles]:
                    p.style = self.doc.styles[docx_styles.FIGURE_PLACEHOLDER]

        self._add_figure_explanatory_data(explanatory_data)

        if caption:
            p = self.doc.add_paragraph(caption)
            self._set_paragraph_format(p, "caption_img")

        self._add_empty_paragraph("empty_after_image")
        self._rendered_body_blocks = True

    def _add_figure_explanatory_data(self, explanatory_data) -> None:
        for line in explanatory_data or ():
            text = str(line).strip()
            if not text:
                continue
            paragraph = self.doc.add_paragraph(text)
            self._set_paragraph_format(paragraph, "figure_explanatory")
            for run in paragraph.runs:
                run.font.size = Pt(12)

    def _parse_table_line(self, line):
        line = line.strip()
        if not line.startswith("|") or not line.endswith("|"):
            return None
        cells = [cell.strip() for cell in line.split("|")[1:-1]]
        return cells if cells else None

    def _create_table(
        self,
        rows_data,
        caption=None,
        *,
        header_row_count: int = 1,
        notes: tuple[TableNote, ...] = (),
        style_name: str = docx_styles.TABLE,
        borderless: bool = False,
    ):
        if not rows_data:
            return

        docx_styles.register_styles(self.doc)
        self._add_empty_paragraph("empty_before_table")

        if caption:
            p = self.doc.add_paragraph(caption)
            self._set_paragraph_format(p, "caption_table")

        table_cfg = self.config.TABLE
        num_cols = len(rows_data[0])
        table = self.doc.add_table(rows=len(rows_data), cols=num_cols)
        table.style = style_name
        table.autofit = False
        self._set_table_borders(table, borderless=borderless)

        for row_idx, row_cells in enumerate(rows_data):
            row = table.rows[row_idx]
            is_header = row_idx < header_row_count
            for col_idx, text in enumerate(row_cells):
                if col_idx < len(row.cells):
                    cell = row.cells[col_idx]
                    cell.text = text
                    self._set_cell_margins(cell, top=120, bottom=120)
                    for para in cell.paragraphs:
                        pf = para.paragraph_format
                        pf.alignment = WD_ALIGN_PARAGRAPH.CENTER if is_header else WD_ALIGN_PARAGRAPH.LEFT
                        pf.space_before = table_cfg["cell_padding"]
                        pf.space_after = table_cfg["cell_padding"]
                        pf.first_line_indent = NO_INDENT_CM
                        pf.line_spacing = table_cfg["line_spacing"]

                        for run in para.runs:
                            self._set_run_style(
                                run,
                                bold=is_header and table_cfg["header_bold"],
                            )
                            run.font.size = table_cfg["header_font_size"] if is_header else table_cfg["font_size"]
                            if not is_header and _ITALIC_LETTER_RE.fullmatch(text.strip()):
                                run.italic = True

        if table_cfg["header_repeat_on_pages"]:
            for row in table.rows[:header_row_count]:
                self._set_repeat_header_row(row)

        if header_row_count > 0 and len(table.rows) >= header_row_count:
            self._set_row_bottom_border(table.rows[header_row_count - 1], "double")

        for note in notes:
            self._append_table_note_row(table, note)

        self._add_empty_paragraph("empty_after_table")
        self._rendered_body_blocks = True

    def _render_table(self, block: TableNode):
        table_number = self._table_number(block)
        is_continuation = block.continuation is not None

        if block.unit_label:
            self._render_table_unit(block.unit_label)

        if is_continuation:
            label = self._format_table_continuation_label(block.continuation, table_number)
            p = self.doc.add_paragraph(label)
            self._set_paragraph_format(p, "caption_table")
            caption = None
        else:
            caption = self._format_table_caption(block.caption, table_number)

        rows = self._table_rows_for_render(block, is_continuation=is_continuation)
        header_row_count = 1 if is_continuation else max(0, min(block.header_row_count, len(rows)))
        self._create_table(
            rows,
            caption,
            header_row_count=header_row_count,
            notes=block.notes,
        )

    def _table_rows_for_render(self, block: TableNode, *, is_continuation: bool) -> list[list[str]]:
        rows = [[cell.text for cell in row.cells] for row in block.rows]
        rows = self._apply_column_units(rows, block.column_units, block.header_row_count)
        if not is_continuation:
            return rows
        if not rows:
            return []
        column_count = len(rows[0])
        body_rows = rows[block.header_row_count :] if block.header_row_count else rows
        return [[str(index) for index in range(1, column_count + 1)], *body_rows]

    def _apply_column_units(
        self,
        rows: list[list[str]],
        column_units: tuple[str | None, ...],
        header_row_count: int,
    ) -> list[list[str]]:
        if not rows or not column_units or header_row_count <= 0:
            return rows
        updated = [list(row) for row in rows]
        first_header = updated[0]
        for index, unit in enumerate(column_units):
            if unit is None or index >= len(first_header):
                continue
            suffix = f", {unit}"
            if not first_header[index].endswith(suffix):
                first_header[index] = f"{first_header[index]}{suffix}"
        return updated

    def _table_number(self, block: TableNode) -> str:
        if block.number and block.number != "auto":
            return str(block.number)
        return str(self._numbering.next_table_number())

    def _format_table_continuation_label(
        self,
        continuation: ContinuationLabel | None,
        table_number: str,
    ) -> str:
        if continuation is ContinuationLabel.FINAL:
            return f"Окончание таблицы {table_number}"
        return f"Продолжение таблицы {table_number}"

    def _render_table_unit(self, unit_label: str) -> None:
        text = unit_label.strip()
        if text and not text.startswith(","):
            text = f", {text}"
        p = self.doc.add_paragraph(text)
        self._set_paragraph_format(p, "table_unit")
        for run in p.runs:
            run.font.size = Pt(12)

    def _append_table_note_row(self, table, note: TableNote) -> None:
        row = table.add_row()
        if len(row.cells) > 1:
            row.cells[0].merge(row.cells[-1])
        cell = row.cells[0]
        self._set_cell_margins(cell, top=120, bottom=120)
        paragraph = cell.paragraphs[0]
        paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        paragraph.paragraph_format.first_line_indent = NO_INDENT_CM
        paragraph.paragraph_format.line_spacing = self.config.TABLE["line_spacing"]
        marker = paragraph.add_run(note.marker)
        self._set_run_style(marker, bold=False)
        marker.font.size = self.config.TABLE["font_size"]
        marker.font.superscript = True
        text = paragraph.add_run(f" {note.text}")
        self._set_run_style(text, bold=False)
        text.font.size = self.config.TABLE["font_size"]

    def _set_table_borders(self, table, *, borderless: bool) -> None:
        tbl_pr = table._tbl.tblPr
        existing = tbl_pr.find(qn("w:tblBorders"))
        if existing is not None:
            tbl_pr.remove(existing)

        borders = OxmlElement("w:tblBorders")
        if borderless:
            values = {
                "top": "nil",
                "left": "nil",
                "bottom": "nil",
                "right": "nil",
                "insideH": "nil",
                "insideV": "nil",
            }
        else:
            values = {
                "top": "nil",
                "left": "single",
                "bottom": "single",
                "right": "single",
                "insideH": "nil",
                "insideV": "single",
            }
        for edge, value in values.items():
            element = OxmlElement(f"w:{edge}")
            element.set(qn("w:val"), value)
            element.set(qn("w:sz"), "4")
            element.set(qn("w:space"), "0")
            element.set(qn("w:color"), "000000")
            borders.append(element)
        tbl_pr.append(borders)

    def _set_row_bottom_border(self, row, value: str) -> None:
        for cell in row.cells:
            tc_pr = cell._tc.get_or_add_tcPr()
            borders = tc_pr.find(qn("w:tcBorders"))
            if borders is None:
                borders = OxmlElement("w:tcBorders")
                tc_pr.append(borders)
            bottom = borders.find(qn("w:bottom"))
            if bottom is None:
                bottom = OxmlElement("w:bottom")
                borders.append(bottom)
            bottom.set(qn("w:val"), value)
            bottom.set(qn("w:sz"), "6" if value == "double" else "4")
            bottom.set(qn("w:color"), "000000")

    def _set_cell_margins(self, cell, *, top: int, bottom: int) -> None:
        tc_pr = cell._tc.get_or_add_tcPr()
        margins = tc_pr.find(qn("w:tcMar"))
        if margins is None:
            margins = OxmlElement("w:tcMar")
            tc_pr.append(margins)
        for edge, value in (("top", top), ("bottom", bottom)):
            element = margins.find(qn(f"w:{edge}"))
            if element is None:
                element = OxmlElement(f"w:{edge}")
                margins.append(element)
            element.set(qn("w:w"), str(value))
            element.set(qn("w:type"), "dxa")

    def _set_repeat_header_row(self, row):
        """Mark a row so Word repeats it on every continuation page."""

        tr = row._tr
        tr_pr = tr.get_or_add_trPr()
        tbl_header = OxmlElement("w:tblHeader")
        tr_pr.append(tbl_header)

    def _format_table_caption(self, caption, table_number):
        if not caption:
            return None
        text = caption.strip()
        if not text:
            return None
        if text.startswith("Таблица"):
            return _normalize_caption_dashes(text)
        return f"Таблица {table_number}{DASH_SEPARATOR}{text}"

    def _format_figure_caption(self, block: FigureNode):
        if not block.caption and not (block.sheet is not None and block.sheet >= 2):
            return None
        if block.caption and block.caption.strip().startswith("Рисунок") and not block.sheet:
            return normalize_caption_dashes(block.caption.strip())
        figure_number = self._numbering.next_figure_number()
        return figure_caption_text(block, figure_number)

    def _setup_document_margins(self):
        for section in self.doc.sections:
            section.top_margin = self.config.MARGINS["top"]
            section.bottom_margin = self.config.MARGINS["bottom"]
            section.left_margin = self.config.MARGINS["left"]
            section.right_margin = self.config.MARGINS["right"]

    def _add_page_numbering(self):
        descriptors = [
            PageNumberingSection(
                start_at=1,
                hide_first_page=bool(self.config.PAGE_NUMBERING["skip_first_page"]),
                suppress_in_section=False,
                location=Location.FOOTER_CENTER,
            )
            for _ in self.doc.sections
        ]
        configure_page_numbering(self.doc, descriptors, self.config)

    def _load_template(self, template_path):
        template_file = self._resolve_template_path(template_path)
        if template_file.exists():
            self.doc = DocxDocument(str(template_file))
            self.logger.info(f"Шаблон загружен: {template_file}")
        else:
            self.doc = DocxDocument()
            self._setup_document_margins()
            self.logger.info("Создан новый документ")

    def _initialize_document(self, template=None, *, template_mode: str = "append", profile: FormattingProfile | None = None):
        if template:
            self._load_template(template)
        else:
            self.doc = DocxDocument()
        self._setup_document_margins()
        docx_styles.register_styles(self.doc)
        self._add_page_numbering()
        self._section_numberer.reset()
        self._rendered_body_blocks = False
        self._template_mode = template_mode
        self._title_page_emitted = False
        self._profile = profile
        self._title_page_diagnostics = []
        self._abbreviation_entries = ()
        self._abbreviation_explicit = False
        self._abbreviation_diagnostics = []
        self._figure_diagnostics = []
        self._appendix_diagnostics = []
        self._reference_diagnostics = []
        self._bibliography_diagnostics = []
        self._footnote_diagnostics = []
        self._toc_diagnostics = []
        self._toc_field: TocField | None = None
        self._footnote_bodies: dict[str, str] = {}
        self._footnote_ids: dict[str, int] = {}
        self._rendered_footnotes: dict[int, str] = {}
        self._formula_symbol_numbers: dict[str, str] = {}
        self._formula_numbers_by_id: dict[str, str] = {}
        self._numbering: NumberingContext = build_numbering_context(profile)

    def _prepare_document_level_state(self, document: Document) -> None:
        self._embed_core_metadata(document.metadata)
        self._abbreviation_entries = abbreviations_for_document(document)
        self._abbreviation_explicit = explicit_abbreviations(document) is not None
        self._abbreviation_diagnostics = []
        if self._profile is not None:
            self._toc_field = build_toc_field(
                document,
                profile=self._profile,
                total_pages=_document_total_pages(document),
            )
            self._toc_diagnostics = list(self._toc_field.diagnostics)
        source_records = tuple(block for block in document.blocks if isinstance(block, SourceRecordNode))
        self._bibliography_diagnostics = validate_records(source_records) if source_records else []
        self._prepare_footnotes(document)

    def _embed_core_metadata(self, metadata) -> None:
        if self.doc is None:
            return
        metadata = dict(metadata or {})
        core = self.doc.core_properties
        if metadata.get("title"):
            core.title = metadata["title"]
        if metadata.get("student"):
            core.author = metadata["student"]
        if metadata.get("subject"):
            core.subject = metadata["subject"]
        keywords = _core_metadata_keywords(metadata)
        if keywords:
            core.keywords = keywords

    def _prepare_footnotes(self, document: Document) -> None:
        self._footnote_diagnostics = []
        self._footnote_bodies = {
            block.marker: block.text
            for block in document.blocks
            if isinstance(block, FootnoteNode)
        }
        anchors: dict[str, list[FootnoteAnchor]] = {}
        for block in document.blocks:
            if not isinstance(block, ParagraphNode):
                continue
            for run in block.runs:
                if isinstance(run, FootnoteAnchor):
                    anchors.setdefault(run.marker, []).append(run)

        for marker, marker_anchors in anchors.items():
            if len(marker_anchors) > 1:
                self._footnote_diagnostics.append(
                    Diagnostic(
                        code=DiagnosticCodes.FOOTNOTE_DUPLICATE,
                        message=f"Footnote anchor '{marker}' appears more than once",
                        severity=Severity.ERROR,
                        source=marker_anchors[1].source,
                        rule_id="common.reference.footnote",
                        data={"marker": marker},
                    )
                )
            if marker not in self._footnote_bodies:
                self._footnote_diagnostics.append(
                    Diagnostic(
                        code=DiagnosticCodes.FOOTNOTE_UNMATCHED_ANCHOR,
                        message=f"Footnote anchor '{marker}' has no matching FN_BODY",
                        severity=Severity.ERROR,
                        source=marker_anchors[0].source,
                        rule_id="common.reference.footnote",
                        data={"marker": marker},
                    )
                )

        for marker in self._footnote_bodies:
            if marker not in anchors:
                self._footnote_diagnostics.append(
                    Diagnostic(
                        code=DiagnosticCodes.FOOTNOTE_UNMATCHED_BODY,
                        message=f"Footnote body '{marker}' has no matching anchor",
                        severity=Severity.ERROR,
                        rule_id="common.reference.footnote",
                        data={"marker": marker},
                    )
                )

        self._footnote_ids = {
            marker: _footnote_id(marker, index)
            for index, marker in enumerate(self._footnote_bodies, start=1)
        }
        self._rendered_footnotes = {}

    def _render_from_ast(
        self,
        document,
        *,
        allow_auto_toc: bool = True,
        default_project_designation: ProjectDesignationNode | None = None,
    ):
        blocks = tuple(document.blocks)
        project_designation = _first_project_designation(blocks) or default_project_designation
        auto_toc_pending = (
            allow_auto_toc
            and self._toc_field is not None
            and self._toc_field.should_insert
            and not self._toc_field.explicit
        )
        for index, block in enumerate(blocks):
            previous_block = blocks[index - 1] if index else None
            next_block = blocks[index + 1] if index + 1 < len(blocks) else None
            if auto_toc_pending and _is_toc_insertion_point(block):
                self._render_table_of_contents(
                    TableOfContentsNode(title=self._toc_field.title, levels=self._toc_field.levels),
                    field=self._toc_field,
                )
                auto_toc_pending = False

            if isinstance(block, StructuralSectionNode):
                self._render_structural_section(block)
            elif isinstance(block, HeadingNode):
                self._render_heading(block)
            elif isinstance(block, ParagraphNode):
                self._render_paragraph(block)
            elif isinstance(block, TableNode):
                self._render_table(block)
            elif isinstance(block, TableCaptionNode):
                p = self.doc.add_paragraph(block.text)
                self._set_paragraph_format(p, "caption_table")
            elif isinstance(block, DrawingSheetNode):
                self._render_drawing_sheet(block, document.metadata)
            elif isinstance(block, FigureNode):
                caption = self._format_figure_caption(block)
                self._insert_image(
                    block.src,
                    caption,
                    explanatory_data=block.explanatory_data,
                )
            elif isinstance(block, PageBreakNode):
                self.doc.add_page_break()
            elif isinstance(block, FormulaNode):
                suppress_after = (
                    isinstance(next_block, FormulaNode)
                    and bool(next_block.consecutive_with)
                    and next_block.consecutive_with == block.id
                )
                self._render_formula(
                    block,
                    trailing_comma=suppress_after,
                    suppress_before=(
                        isinstance(previous_block, FormulaNode)
                        and bool(block.consecutive_with)
                        and block.consecutive_with == previous_block.id
                    ),
                    suppress_after=suppress_after,
                )
            elif isinstance(block, ListNode):
                self._render_list(block)
            elif isinstance(block, PosterNode):
                self._render_poster(block, document.metadata)
            elif isinstance(block, SlideDeckNode):
                self._render_slide_deck_placeholder(block)
            elif isinstance(block, AppendixNode):
                self._render_appendix(block, document.metadata)
            elif isinstance(block, TableOfContentsNode):
                self._render_table_of_contents(
                    block,
                    field=self._toc_field if self._toc_field is not None else None,
                )
            elif isinstance(block, BibliographyEntryNode):
                self._render_bibliography_entry(block)
            elif isinstance(block, SourceRecordNode):
                self._render_source_record(block)
            elif isinstance(block, FootnoteNode):
                continue
            elif isinstance(block, RawBlockNode):
                self._render_text_block(block.text)
            elif isinstance(block, ReferenceNode):
                self._render_text_block(f"[{block.target}]")
            elif isinstance(block, TitlePageNode):
                self._render_title_page(document.metadata, block.profile)
            elif isinstance(block, SectionSetupNode):
                self._render_section_setup(
                    block,
                    document.metadata,
                    default_project_designation=project_designation,
                )
            elif isinstance(block, ProjectDesignationNode):
                continue

    def _render_heading(self, block):
        if block.level is HeadingLevel.H1 and self._rendered_body_blocks:
            self.doc.add_page_break()

        if block.level is HeadingLevel.H2:
            self._add_empty_paragraph("empty_before_header")

        style_type = {
            HeadingLevel.H1: "h1",
            HeadingLevel.H2: "h2",
            HeadingLevel.H3: "h3",
            HeadingLevel.H4: "h4",
        }[block.level]
        p = self.doc.add_paragraph(self._heading_text(block))
        self._set_paragraph_format(p, style_type)
        self._add_empty_paragraph("empty_after_header")
        self._rendered_body_blocks = True

    def _heading_text(self, block):
        if block.number != "auto":
            return block.text

        number = self._numbering.next_section_number(block.level)
        title = block.text.rstrip().removesuffix(".")
        return f"{number} {title}"

    def _render_structural_section(self, block):
        if block.section_type is StructuralSectionType.CONTENTS:
            self._render_table_of_contents(
                TableOfContentsNode(title=block.title, source=block.source),
                field=self._toc_field if self._toc_field is not None else None,
            )
            return
        if block.section_type is StructuralSectionType.ABBREVIATIONS:
            self._render_abbreviations_section(block)
            return

        if self.config.STRUCTURAL_SECTION["page_break_before"]:
            self.doc.add_page_break()

        title = block.title.upper() if self.config.STRUCTURAL_SECTION["uppercase"] else block.title
        p = self.doc.add_paragraph()
        run = p.add_run(title)
        self._set_paragraph_format(p, "structural_section")
        run.underline = False
        self._add_empty_paragraph("empty_after_header")
        self._rendered_body_blocks = True

    def _render_abbreviations_section(self, block):
        if self.config.STRUCTURAL_SECTION["page_break_before"]:
            self.doc.add_page_break()

        title = block.title.upper() if self.config.STRUCTURAL_SECTION["uppercase"] else block.title
        p = self.doc.add_paragraph()
        run = p.add_run(title)
        self._set_paragraph_format(p, "structural_section")
        run.underline = False
        self._add_empty_paragraph("empty_after_header")

        if self._abbreviation_entries:
            self._create_abbreviations_table(self._abbreviation_entries)
            if not self._abbreviation_explicit:
                self._abbreviation_diagnostics.append(
                    Diagnostic(
                        code="ABBREVIATIONS_AUTO_DETECTED",
                        message="Abbreviations list generated from first-use introductions",
                        severity=Severity.INFO,
                        rule_id="common.abbreviations.two_column_layout",
                    )
                )
        self._rendered_body_blocks = True

    def _create_abbreviations_table(self, entries) -> None:
        rows = [[entry.short, f"{EM_DASH} {entry.long}"] for entry in entries]
        if not rows:
            return

        table = self.doc.add_table(rows=len(rows), cols=2)
        table.style = docx_styles.ABBREVIATIONS_TABLE
        table.autofit = False
        self._set_table_borders(table, borderless=True)

        widths = ABBREVIATIONS_COLUMN_WIDTHS
        for row_idx, row_data in enumerate(rows):
            for col_idx, text in enumerate(row_data):
                cell = table.rows[row_idx].cells[col_idx]
                cell.width = widths[col_idx]
                cell.text = text
                self._set_cell_margins(cell, top=0, bottom=0)
                for paragraph in cell.paragraphs:
                    paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    paragraph.paragraph_format.first_line_indent = NO_INDENT_CM
                    paragraph.paragraph_format.space_before = Pt(0)
                    paragraph.paragraph_format.space_after = Pt(0)
                    paragraph.paragraph_format.line_spacing = 1.0
                    for run in paragraph.runs:
                        self._set_run_style(run, bold=False)
                        run.font.size = Pt(14)

    def _render_title_page(self, metadata, profile_name=None, *, force: bool = False):
        if self._template_mode == "preserve-prefix":
            return
        if self._title_page_emitted and not force:
            return

        from sfu_converter.infrastructure.title_pages import (
            TitlePageLayout,
            select_title_page_form,
        )

        metadata = dict(metadata or {})
        profile = self._resolve_title_profile(profile_name)
        form = select_title_page_form(profile, metadata, override=profile_name)

        self._report_missing_title_metadata(metadata, form)
        layout = TitlePageLayout(self.doc, self.config)
        form.render(layout, metadata)

        self.doc.add_page_break()
        if not force:
            self._title_page_emitted = True

    def _resolve_title_profile(self, profile_name):
        if profile_name:
            try:
                return get_profile(profile_name)
            except KeyError:
                pass
        return self._profile

    def _report_missing_title_metadata(self, metadata, form):
        for field in form.required_metadata:
            if not metadata.get(field):
                self._title_page_diagnostics.append(
                    Diagnostic(
                        code=DiagnosticCodes.TXT_MISSING_METADATA,
                        message=f"Title page metadata field '{field}' is required for {form.form_id}",
                        severity=Severity.WARNING,
                        rule_id=form.form_id,
                        data={"field": field},
                    )
                )

    def _render_appendix(self, block, metadata=None):
        """Render appendix heading on a new page with the standard layout.

        STU 7.5-07-2021 requires: page break before, centered bold heading using
        Russian uppercase letters (no Ё, З, Й, О, Ч, Ъ, Ы, Ь), optional content
        type below, and an optional appendix title separated by one blank line.
        """

        if block.independent:
            self._render_title_page(_appendix_metadata(metadata, block), force=True)
        elif block.sheet_format is not SheetFormat.A4:
            self.doc.add_section(WD_SECTION.NEW_PAGE)
            section = self.doc.sections[-1]
            section_setup.configure(
                self.doc,
                section,
                SectionSetupNode(sheet_format=block.sheet_format),
            )
        else:
            self.doc.add_page_break()

        heading_text = block.title or APPENDIX_TITLE
        if block.letter and block.letter not in heading_text:
            heading_text = f"{APPENDIX_TITLE} {block.letter}"

        heading_para = self.doc.add_paragraph()
        heading_run = heading_para.add_run(heading_text.upper())
        self._set_paragraph_format(heading_para, "appendix_heading")
        heading_run.underline = False

        if block.appendix_type:
            type_para = self.doc.add_paragraph()
            type_run = type_para.add_run(f"({block.appendix_type})")
            self._set_paragraph_format(type_para, "caption_img")
            type_run.bold = False
            type_run.italic = False

        if block.subtitle:
            self._add_empty_paragraph("empty_after_header")
            subtitle_para = self.doc.add_paragraph()
            subtitle_run = subtitle_para.add_run(block.subtitle)
            self._set_paragraph_format(subtitle_para, "structural_section")
            subtitle_run.underline = False

        self._add_empty_paragraph("empty_after_header")
        if block.blocks:
            entered = False
            if block.letter:
                self._numbering.enter_appendix(block.letter)
                entered = True
            try:
                self._render_appendix_blocks(block, metadata)
            finally:
                if entered:
                    self._numbering.leave_appendix()
        self._rendered_body_blocks = True

    def _render_appendix_blocks(self, block: AppendixNode, metadata) -> None:
        chunks = _split_appendix_pages(block.blocks)
        if len(chunks) == 1:
            self._render_from_ast(Document(blocks=chunks[0], metadata=metadata or {}), allow_auto_toc=False)
            return

        for page_index, chunk in enumerate(chunks):
            if page_index > 0:
                self.doc.add_page_break()
                label = "Окончание" if page_index == len(chunks) - 1 else "Продолжение"
                self._render_appendix_continuation_label(f"{label} приложения {block.letter}")
            self._render_from_ast(Document(blocks=chunk, metadata=metadata or {}), allow_auto_toc=False)

    def _render_appendix_continuation_label(self, text: str) -> None:
        paragraph = self.doc.add_paragraph()
        run = paragraph.add_run(text)
        self._set_paragraph_format(paragraph, "appendix_heading")
        run.underline = False

    def _render_table_of_contents(self, block, *, field: TocField | None = None):
        """Insert a ``СОДЕРЖАНИЕ`` heading and a Word TOC field.

        ``python-docx`` does not support TOCs natively, so we emit raw OOXML
        ``w:fldChar`` / ``w:instrText`` elements that Word recognises and
        updates on open (or via Ctrl+A, F9).
        """

        self.doc.add_page_break()

        heading_para = self.doc.add_paragraph()
        heading_run = heading_para.add_run((block.title or "СОДЕРЖАНИЕ").upper())
        self._set_paragraph_format(heading_para, "toc_heading")
        heading_run.underline = False

        self._add_empty_paragraph("empty_after_header")

        toc_para = self.doc.add_paragraph()
        toc_para.paragraph_format.first_line_indent = NO_INDENT_CM
        begin_run = toc_para.add_run()
        self._set_run_style(begin_run, bold=False)
        fld_begin = OxmlElement("w:fldChar")
        fld_begin.set(qn("w:fldCharType"), "begin")
        begin_run._element.append(fld_begin)

        instr = OxmlElement("w:instrText")
        instr.set(qn("xml:space"), "preserve")
        instr.text = f' TOC \\o "1-{block.levels}" \\h \\z \\u '
        begin_run._element.append(instr)

        fld_separate = OxmlElement("w:fldChar")
        fld_separate.set(qn("w:fldCharType"), "separate")
        begin_run._element.append(fld_separate)

        placeholder_run = toc_para.add_run("Обновите оглавление (Ctrl+A, F9)")
        self._set_run_style(placeholder_run, bold=False)

        fld_end = OxmlElement("w:fldChar")
        fld_end.set(qn("w:fldCharType"), "end")
        placeholder_run._element.append(fld_end)

        for entry in (field.entries if field is not None else ()):
            self._render_toc_entry(entry)

        self._rendered_body_blocks = True

    def _render_toc_entry(self, entry: TocEntry) -> None:
        paragraph = self.doc.add_paragraph(entry.rendered_text)
        _apply_toc_style(self.doc, paragraph, entry.level)
        pf = paragraph.paragraph_format
        pf.first_line_indent = NO_INDENT_CM
        pf.left_indent = Cm(entry.left_indent_cm)
        pf.line_spacing = 1.0
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)
        pf.tab_stops.add_tab_stop(self.config.FORMULA["number_tab_pos"], WD_TAB_ALIGNMENT.RIGHT)
        for run in paragraph.runs:
            self._set_run_style(run, bold=False)

    def _apply_word_heading_style(self, paragraph, style_name):
        if self.doc is None:
            return
        docx_styles.apply_word_heading_style(self.doc, paragraph, style_name)

    def _render_paragraph(self, block):
        para = self.doc.add_paragraph()
        for text_run in block.runs:
            if isinstance(text_run, FootnoteAnchor):
                self._add_footnote_anchor(para, text_run)
            elif isinstance(text_run, CitationNode):
                para.add_run(format_citation_node(text_run))
            else:
                para.add_run(text_run.text)
        self._set_paragraph_format(para, "normal")
        for docx_run, text_run in zip(para.runs, block.runs, strict=False):
            if isinstance(text_run, FootnoteAnchor):
                continue
            if isinstance(text_run, CitationNode):
                self._set_run_style(docx_run, bold=False, italic=False)
                continue
            self._set_run_style(
                docx_run,
                bold=text_run.bold,
                italic=text_run.italic,
            )
        self._rendered_body_blocks = True

    def _render_text_block(self, text):
        p = self.doc.add_paragraph(text)
        self._set_paragraph_format(p, "normal")
        self._rendered_body_blocks = True

    def _render_formula(
        self,
        block,
        *,
        trailing_comma: bool = False,
        suppress_before: bool = False,
        suppress_after: bool = False,
    ):
        """Render a formula on its own line with right-aligned auto-number."""

        if block.number == "auto" or block.number is None:
            number_text = self._numbering.next_formula_number()
        else:
            number_text = str(block.number)
        if block.id:
            self._formula_numbers_by_id[block.id] = number_text

        if not suppress_before:
            self._add_empty_paragraph("empty_before_formula")

        para = self.doc.add_paragraph()
        self._set_paragraph_format(para, "formula")
        # Replace the placeholder run created by _set_paragraph_format with
        # explicit content + tab + number runs so styling stays consistent.
        for run in list(para.runs):
            run._element.getparent().remove(run._element)

        lines = split_formula_lines(block.content or "", block.continuation_lines)
        for line_index, line in enumerate(lines):
            if line_index:
                break_run = para.add_run()
                break_run.add_break()
                self._set_run_style(break_run, bold=False)
            content_run = para.add_run(line)
            self._set_run_style(content_run, bold=False)
        if trailing_comma:
            comma_run = para.add_run(",")
            self._set_run_style(comma_run, bold=False)

        number_run = para.add_run(f"\t({number_text})")
        self._set_run_style(number_run, bold=False)

        self._set_formula_number_tab(para)

        explanation_text = self._formula_explanation_text(block)
        if explanation_text:
            expl_para = self.doc.add_paragraph(explanation_text)
            self._set_paragraph_format(expl_para, "formula_explanation")
        for symbol in block.explanations:
            if not symbol.repeats and symbol.name:
                self._formula_symbol_numbers.setdefault(symbol.name, number_text)

        if not suppress_after:
            self._add_empty_paragraph("empty_after_formula")
        self._rendered_body_blocks = True

    def _add_footnote_anchor(self, para, anchor: FootnoteAnchor) -> None:
        run = para.add_run()
        self._set_run_style(run, bold=False)
        if docx_styles.FOOTNOTE_ANCHOR in [s.name for s in self.doc.styles]:
            run.style = self.doc.styles[docx_styles.FOOTNOTE_ANCHOR]
        footnote_id = self._footnote_ids.get(anchor.marker, _footnote_id(anchor.marker, len(self._footnote_ids) + 1))
        add_footnote_reference(run, footnote_id)
        if anchor.marker in self._footnote_bodies:
            self._rendered_footnotes[footnote_id] = self._footnote_bodies[anchor.marker]

    def _formula_explanation_text(self, block) -> str | None:
        if block.explanations:
            lines = explanation_lines(block.explanations, self._formula_symbol_numbers)
            return "\n".join(lines)
        return block.explanation

    def _set_formula_number_tab(self, para):
        """Add a right-aligned tab stop so ``\\t(N)`` lands at the right margin."""

        tab_pos = self.config.FORMULA["number_tab_pos"]
        ppr = para._p.get_or_add_pPr()
        existing = ppr.find(qn("w:tabs"))
        if existing is not None:
            ppr.remove(existing)

        tab_stops = OxmlElement("w:tabs")
        tab = OxmlElement("w:tab")
        tab.set(qn("w:val"), "right")
        tab.set(qn("w:pos"), str(int(tab_pos.emu / 635)))
        tab_stops.append(tab)
        ppr.append(tab_stops)

    def _render_bibliography_entry(self, block):
        """Render a single source list entry with hanging-paragraph layout."""

        text = f"{block.number} {block.text}"
        para = self.doc.add_paragraph(text)
        self._set_paragraph_format(para, "bibliography_entry")
        self._rendered_body_blocks = True

    def _render_source_record(self, block: SourceRecordNode):
        """Render a structured source record using the GOST formatter."""

        text = f"{block.number} {format_record(block)}"
        para = self.doc.add_paragraph(text)
        self._set_paragraph_format(para, "bibliography_entry")
        self._rendered_body_blocks = True

    def _render_section_setup(
        self,
        block: SectionSetupNode,
        metadata,
        *,
        default_project_designation: ProjectDesignationNode | None = None,
    ) -> None:
        project_designation = _first_project_designation(block.blocks) or default_project_designation
        if self._rendered_body_blocks:
            self.doc.add_section(WD_SECTION.NEW_PAGE)
        section = self.doc.sections[-1]
        section_setup.configure(self.doc, section, block)
        if block.frame is not FrameType.NONE:
            frames.draw(self.doc, section)

        self._render_from_ast(
            Document(blocks=block.blocks, metadata=metadata),
            allow_auto_toc=False,
            default_project_designation=project_designation,
        )

        if block.title_block_form is not None:
            main_inscription.render(
                self.doc,
                block.title_block_form.value,
                fields=_title_block_fields(metadata, designation=project_designation),
                page_number_in_graph_7=block.frame in {FrameType.TEXT_FIRST, FrameType.TEXT_FOLLOWING},
            )
        self._rendered_body_blocks = True

    def _render_drawing_sheet(self, block: DrawingSheetNode, metadata) -> None:
        if self._rendered_body_blocks:
            self.doc.add_section(WD_SECTION.NEW_PAGE)
        section = self.doc.sections[-1]
        setup = SectionSetupNode(
            orientation=SectionOrientation.LANDSCAPE,
            sheet_format=block.sheet_format,
            frame=block.frame,
            title_block_form=block.title_block_form,
        )
        section_setup.configure(self.doc, section, setup)
        if block.frame is not FrameType.NONE:
            frames.draw(self.doc, section)

        if block.src:
            self._insert_image(block.src, None)
        else:
            paragraph = self.doc.add_paragraph("Графический материал")
            self._set_paragraph_format(paragraph, "normal")

        form = block.title_block_form or TitleBlockForm.FORM_5
        fields = _title_block_fields(metadata)
        if block.designation:
            fields["2"] = block.designation
        main_inscription.render(self.doc, form.value, fields=fields, page_number_in_graph_7=False)
        self._rendered_body_blocks = True

    def _render_poster(self, block: PosterNode, metadata) -> None:
        if self._rendered_body_blocks:
            self.doc.add_section(WD_SECTION.NEW_PAGE)
        section = self.doc.sections[-1]
        setup = SectionSetupNode(
            orientation=SectionOrientation.LANDSCAPE,
            sheet_format=block.sheet_format,
            frame=FrameType.NONE,
        )
        section_setup.configure(self.doc, section, setup)

        if block.title:
            heading = self.doc.add_paragraph(block.title)
            self._set_paragraph_format(heading, "h1")
        if block.blocks:
            self._render_from_ast(
                Document(blocks=block.blocks, metadata=metadata),
                allow_auto_toc=False,
            )

        self.doc.add_section(WD_SECTION.NEW_PAGE)
        reverse = self.doc.sections[-1]
        section_setup.configure(self.doc, reverse, setup)
        reverse_label = self.doc.add_paragraph("Оборотная сторона")
        self._set_paragraph_format(reverse_label, "normal")
        main_inscription.render(
            self.doc,
            TitleBlockForm.FORM_5.value,
            fields=_title_block_fields(metadata),
            page_number_in_graph_7=False,
        )
        self._rendered_body_blocks = True

    def _render_slide_deck_placeholder(self, block: SlideDeckNode) -> None:
        for index, slide in enumerate(block.slides, start=1):
            title = slide.fields.get("title", f"Слайд {index}")
            paragraph = self.doc.add_paragraph(f"Слайд {index}: {title}")
            self._set_paragraph_format(paragraph, "h2" if index == 1 else "normal")
            for line in slide.body:
                body = self.doc.add_paragraph(line)
                self._set_paragraph_format(body, "normal")
        if block.slides:
            self._rendered_body_blocks = True

    def _render_list(self, block, *, level: int = 0):
        for index, item in enumerate(block.items):
            if isinstance(item, ListNode):
                self._render_list(item, level=level + 1)
                continue
            text = self._list_item_text(
                list_type=block.list_type,
                index=index,
                item_text=item.text,
                is_last=index == len(block.items) - 1,
            )
            p = self.doc.add_paragraph(text)
            self._set_paragraph_format(p, "list_item")
            apply_list_item_layout(p, block.list_type, level=level)
            for child in item.children:
                self._render_list(child, level=level + 1)
        if block.items:
            self._rendered_body_blocks = True

    def _list_item_text(self, list_type, index, item_text, is_last):
        text = _normalize_list_item_punctuation(
            item_text,
            list_type=list_type,
            is_last=is_last,
        )
        if list_type is ListType.BULLET:
            return f"- {text}"
        if list_type is ListType.LETTERED:
            return f"{list_marker(list_type, index)} {text}"
        if list_type is ListType.NUMBERED:
            return f"{list_marker(list_type, index)} {text}"
        raise ValueError(f"Unsupported list type: {list_type}")

    def _resolve_image_path(self, image_path):
        path = Path(image_path)
        if path.is_absolute():
            return path
        return self.base_dir / PathConfig.IMAGES_DIR / path

    def _resolve_template_path(self, template_path):
        path = Path(template_path)
        if path.is_absolute():
            return path

        for candidate in (
            self.base_dir / PathConfig.TEMPLATES_DIR / path,
            self.base_dir / path,
        ):
            if candidate.exists():
                return candidate
        return self.base_dir / PathConfig.TEMPLATES_DIR / path


def _normalize_list_item_punctuation(
    text: str,
    *,
    list_type: ListType,
    is_last: bool,
) -> str:
    stripped = text.strip()
    if not stripped or stripped.endswith((".", ";")):
        return stripped

    ending = "." if list_type is ListType.NUMBERED or is_last else ";"
    return f"{stripped}{ending}"


def _document_total_pages(document: Document) -> int | None:
    for key in ("total_pages", "page_count", "pages"):
        value = document.metadata.get(key)
        if value is None:
            continue
        try:
            total = int(str(value).strip())
        except ValueError:
            continue
        if total > 0:
            return total
    return None


def _core_metadata_keywords(metadata: dict[str, str]) -> str:
    priority = (
        "group",
        "supervisor",
        "teacher",
        "reviewer",
        "direction_code",
        "direction_name",
        "master_program_code",
        "master_program_name",
        "specialty_code",
        "consultants",
        "norm_controller",
    )
    ordered_keys = [key for key in priority if metadata.get(key)]
    ordered_keys.extend(key for key in sorted(metadata) if key not in ordered_keys and metadata.get(key))

    parts: list[str] = []
    length = 0
    for key in ordered_keys:
        part = f"{key}={metadata[key]}"
        next_length = length + len(part) + (2 if parts else 0)
        if next_length > 255:
            continue
        parts.append(part)
        length = next_length
    return "; ".join(parts)


def _title_block_fields(
    metadata,
    *,
    designation: ProjectDesignationNode | None = None,
) -> dict[str, str]:
    fields: dict[str, str] = {}
    for graph_number in range(1, 18):
        value = metadata.get(f"title_block_{graph_number}") if metadata else None
        if value is not None:
            fields[str(graph_number)] = str(value)
    if designation is not None:
        fields["2"] = format_designation(designation)
    return fields


def _split_appendix_pages(blocks: tuple) -> tuple[tuple, ...]:
    pages: list[list] = [[]]
    for block in blocks:
        if isinstance(block, PageBreakNode):
            pages.append([])
            continue
        pages[-1].append(block)
    return tuple(tuple(page) for page in pages)


def _appendix_metadata(metadata, block: AppendixNode) -> dict[str, str]:
    result = dict(metadata or {})
    for child in block.blocks:
        if isinstance(child, MetadataNode):
            result[child.key] = child.value
    result.setdefault("title", block.subtitle or block.title)
    return result


def _first_project_designation(blocks) -> ProjectDesignationNode | None:
    for block in blocks:
        if isinstance(block, ProjectDesignationNode):
            return block
    return None


def _is_toc_insertion_point(block) -> bool:
    return not isinstance(block, (MetadataNode, TitlePageNode, FootnoteNode))


def _apply_toc_style(document, paragraph, level: int) -> None:
    style_name = f"TOC {max(1, min(level, 4))}"
    try:
        paragraph.style = document.styles[style_name]
        return
    except KeyError:
        pass

    style = document.styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
    style.base_style = document.styles["Normal"]
    paragraph.style = style


def _footnote_id(marker: str, fallback: int) -> int:
    try:
        value = int(marker)
    except ValueError:
        return fallback
    return value if value > 0 else fallback


def _normalize_caption_dashes(text: str) -> str:
    """Replace ASCII hyphens used as dash separators with em-dash (U+2014)."""

    return re.sub(rf"\s[-{EN_DASH}]\s", DASH_SEPARATOR, text)

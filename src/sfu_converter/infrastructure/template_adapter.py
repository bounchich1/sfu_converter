"""DOCX template composition adapter.

Implements :class:`TemplatePort` so the converter can preserve a user-supplied
template (title page, front matter) and append generated content after a
configurable insertion point. The adapter never reformats template-owned pages;
it only splices the generated body in at the requested boundary.
"""

from __future__ import annotations

import logging
from copy import deepcopy
from io import BytesIO
from pathlib import Path

from docx import Document as DocxDocument
from docx.oxml.ns import qn

from sfu_converter.config import PathConfig
from sfu_converter.domain.diagnostics import Diagnostic, Severity
from sfu_converter.infrastructure import docx_styles
from sfu_converter.ports.template import (
    InsertionPoint,
    TemplateDocument,
    TemplatePort,
)
from sfu_converter.runtime_resources import packaged_template_bytes, packaged_template_label

_TEMPLATE_BOOKMARK_NOT_FOUND = "TEMPLATE_BOOKMARK_NOT_FOUND"
_TEMPLATE_PAGE_NOT_FOUND = "TEMPLATE_PAGE_NOT_FOUND"
_TEMPLATE_LOAD_FAILED = "TEMPLATE_LOAD_FAILED"


class DocxTemplateAdapter(TemplatePort):
    """python-docx implementation of :class:`TemplatePort`."""

    def __init__(self, base_dir: Path | str | None = None, logger=None):
        self.base_dir = Path(base_dir) if base_dir else Path.cwd()
        self.logger = logger or logging.getLogger(__name__)

    def load_template(self, path: str) -> TemplateDocument:
        resolved = self._resolve(path)
        if resolved.exists():
            doc = DocxDocument(str(resolved))
            return TemplateDocument(doc=doc, path=str(resolved))

        template_bytes = packaged_template_bytes(path)
        if template_bytes is not None:
            doc = DocxDocument(BytesIO(template_bytes))
            return TemplateDocument(doc=doc, path=packaged_template_label(path))

        doc = DocxDocument(str(resolved))
        return TemplateDocument(doc=doc, path=str(resolved))

    def find_insertion_point(
        self,
        template: TemplateDocument,
        mode: str,
        page: int | None = None,
        bookmark: str | None = None,
    ) -> InsertionPoint:
        if bookmark:
            return self._find_bookmark(template, bookmark)

        if mode == "preserve-prefix":
            if page is None:
                return InsertionPoint(
                    found=False,
                    diagnostic=Diagnostic(
                        code=_TEMPLATE_PAGE_NOT_FOUND,
                        message=(
                            "Template mode 'preserve-prefix' requires --insert-after-page or --insert-at-bookmark"
                        ),
                        severity=Severity.ERROR,
                    ),
                )
            return self._find_after_page(template, page)

        if mode == "replace-body":
            return InsertionPoint(found=True, element=None, truncate=True)

        return InsertionPoint(found=True, element=None, truncate=False)

    def compose(
        self,
        template: TemplateDocument,
        insertion_point: InsertionPoint,
        generated_content: bytes,
    ) -> bytes:
        if not insertion_point.found:
            raise ValueError("Cannot compose with an unresolved insertion point")

        generated_doc = DocxDocument(BytesIO(generated_content))
        template_doc = template.doc
        body = template_doc.element.body

        # The generated body references SFU paragraph styles by name. Copying
        # paragraphs across documents does not copy the style definitions, so
        # register them on the template or every styled paragraph silently
        # falls back to Normal and the validator loses its role signal.
        docx_styles.register_styles(template_doc)

        sect_pr = self._extract_sect_pr(body)
        anchor = insertion_point.element

        if insertion_point.truncate:
            if anchor is None:
                self._truncate_body(body, sect_pr)
            else:
                self._truncate_after(body, anchor, sect_pr)

        page_break = self._make_page_break_paragraph(template_doc)
        if page_break is not None:
            self._append_before_sect_pr(body, page_break, sect_pr)

        for element in self._iter_body_blocks(generated_doc.element.body):
            self._append_before_sect_pr(body, deepcopy(element), sect_pr)

        buffer = BytesIO()
        template_doc.save(buffer)
        return buffer.getvalue()

    # ----- helpers -----

    def _resolve(self, path: str) -> Path:
        candidate = Path(path)
        if candidate.is_absolute():
            return candidate
        for base in (
            self.base_dir / PathConfig.TEMPLATES_DIR / candidate,
            self.base_dir / candidate,
        ):
            if base.exists():
                return base
        return self.base_dir / PathConfig.TEMPLATES_DIR / candidate

    def _find_bookmark(
        self,
        template: TemplateDocument,
        bookmark: str,
    ) -> InsertionPoint:
        body = template.doc.element.body
        bookmarks = body.findall(".//" + qn("w:bookmarkStart"))
        for mark in bookmarks:
            if mark.get(qn("w:name")) == bookmark:
                paragraph = mark.getparent()
                while paragraph is not None and not paragraph.tag.endswith("}p"):
                    paragraph = paragraph.getparent()
                return InsertionPoint(found=True, element=paragraph)
        return InsertionPoint(
            found=False,
            diagnostic=Diagnostic(
                code=_TEMPLATE_BOOKMARK_NOT_FOUND,
                message=f'Bookmark "{bookmark}" not found in template',
                severity=Severity.ERROR,
            ),
        )

    def _find_after_page(
        self,
        template: TemplateDocument,
        page: int,
    ) -> InsertionPoint:
        if page < 1:
            return InsertionPoint(
                found=False,
                diagnostic=Diagnostic(
                    code=_TEMPLATE_PAGE_NOT_FOUND,
                    message=f"Page index must be >= 1, got {page}",
                    severity=Severity.ERROR,
                ),
            )
        body = template.doc.element.body
        breaks_seen = 0
        for paragraph in body.findall(qn("w:p")):
            for br in paragraph.findall(".//" + qn("w:br")):
                if br.get(qn("w:type")) == "page":
                    breaks_seen += 1
                    if breaks_seen == page:
                        return InsertionPoint(found=True, element=paragraph)
        if breaks_seen == 0 and page == 1:
            return InsertionPoint(found=True, element=None)
        return InsertionPoint(
            found=False,
            diagnostic=Diagnostic(
                code=_TEMPLATE_PAGE_NOT_FOUND,
                message=(f"Template has only {breaks_seen} explicit page break(s); cannot insert after page {page}"),
                severity=Severity.ERROR,
            ),
        )

    def _extract_sect_pr(self, body):
        sect_pr = body.find(qn("w:sectPr"))
        return sect_pr

    def _iter_body_blocks(self, body):
        sect_pr_tag = qn("w:sectPr")
        for child in list(body):
            if child.tag == sect_pr_tag:
                continue
            yield child

    def _truncate_body(self, body, sect_pr):
        for child in list(body):
            if child is sect_pr:
                continue
            body.remove(child)

    def _truncate_after(self, body, anchor, sect_pr):
        seen_anchor = False
        for child in list(body):
            if child is sect_pr:
                continue
            if seen_anchor:
                body.remove(child)
            elif child is anchor:
                seen_anchor = True

    def _append_before_sect_pr(self, body, element, sect_pr):
        if sect_pr is not None:
            sect_pr.addprevious(element)
        else:
            body.append(element)

    def _make_page_break_paragraph(self, doc):
        from docx.oxml import OxmlElement

        paragraph = OxmlElement("w:p")
        run = OxmlElement("w:r")
        br = OxmlElement("w:br")
        br.set(qn("w:type"), "page")
        run.append(br)
        paragraph.append(run)
        return paragraph

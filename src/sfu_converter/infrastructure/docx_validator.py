from __future__ import annotations

import re
import xml.etree.ElementTree as ET
from pathlib import Path
from zipfile import ZipFile

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm
from docx.oxml.ns import qn

from sfu_converter.application import metadata_check
from sfu_converter.config import SIBFUConfig
from sfu_converter.domain.ast_nodes import Document as AstDocument, SourceSpan
from sfu_converter.domain.diagnostics import Diagnostic, DiagnosticCodes, Severity
from sfu_converter.domain.formatting import FormattingProfile, FormattingRule, unsupported_rule_diagnostics
from sfu_converter.infrastructure.appendix import APPENDIX_LETTERS
from sfu_converter.infrastructure.frames import has_frame
from sfu_converter.infrastructure.formula_layout import OPERATOR_SIGNS
from sfu_converter.infrastructure.list_layout import RUSSIAN_LIST_LETTER_INDEX
from sfu_converter.infrastructure.paragraph_roles import ParagraphRole, classify
from sfu_converter.infrastructure.project_designation import validate_designation_text
from sfu_converter.infrastructure import docx_styles
from sfu_converter.registry import get_profile, get_rule

_LENGTH_TOLERANCE_EMU = 1000
_POINT_TOLERANCE = 0.5
_SPACING_TOLERANCE = 0.1
_TABLE_CELL_PADDING_TWIPS = 120
_ITALIC_LETTER_RE = re.compile(r"^[A-Za-zα-ωΑ-Ω]$")

_ALIGNMENT_BY_NAME = {
    "left": WD_ALIGN_PARAGRAPH.LEFT,
    "center": WD_ALIGN_PARAGRAPH.CENTER,
    "right": WD_ALIGN_PARAGRAPH.RIGHT,
    "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
}

_HEADING_ROLES = {
    ParagraphRole.HEADING_H1,
    ParagraphRole.HEADING_H2,
    ParagraphRole.HEADING_H3,
    ParagraphRole.HEADING_H4,
}

_HEADING_RULE_BY_ROLE = {
    ParagraphRole.HEADING_H1: "common.heading.h1",
    ParagraphRole.HEADING_H2: "common.heading.h2",
    ParagraphRole.HEADING_H3: "common.heading.h3",
    ParagraphRole.HEADING_H4: "common.heading.h4",
}


class DocxValidator:
    """Validate generated DOCX files against registry-backed formatting rules."""

    def __init__(
        self,
        profile: FormattingProfile | None = None,
        *,
        config_class=SIBFUConfig,
    ):
        self.profile = profile or get_profile("common")
        self.config = config_class
        self.diagnostics: list[Diagnostic] = []

    def validate_file(self, file_path: str) -> list[Diagnostic]:
        self.diagnostics = []
        doc_path = Path(file_path)
        if not doc_path.exists():
            self._add(
                code="DOCX_FILE_NOT_FOUND",
                message=f"Файл не найден: {file_path}",
                severity=Severity.ERROR,
            )
            return self.diagnostics

        try:
            doc = Document(str(doc_path))
        except Exception as exc:
            self._add(
                code="DOCX_OPEN_FAILED",
                message=f"Не удалось открыть файл: {exc}",
                severity=Severity.ERROR,
            )
            return self.diagnostics

        self.diagnostics.extend(unsupported_rule_diagnostics(self.profile, component="validator"))
        self.diagnostics.extend(
            metadata_check.run(
                AstDocument(blocks=(), metadata=_metadata_from_docx(doc)),
                self.profile,
                severity=Severity.WARNING,
            )
        )
        self._validate_margins(doc)
        self._validate_page_numbering(doc)
        self._seen_formula_symbols: set[str] = set()
        self._previous_non_empty_role: ParagraphRole | None = None
        self._previous_non_empty_paragraph = None
        self._list_expected_letter_by_indent: dict[float, int] = {}
        self._last_letter_indent_cm: float | None = None

        paragraphs = list(doc.paragraphs)
        for index, paragraph in enumerate(paragraphs, start=1):
            prev = paragraphs[index - 2] if index >= 2 else None
            nxt = paragraphs[index] if index < len(paragraphs) else None
            role = classify(paragraph, prev=prev, next=nxt, profile=self.profile)
            self._validate_paragraph(paragraph, index, role)
            if role in _HEADING_ROLES:
                self._validate_heading_blank_lines(paragraph, index, prev, nxt)
            if paragraph.text.strip():
                self._previous_non_empty_role = role
                self._previous_non_empty_paragraph = paragraph

        self._validate_heading_subpoint_structure(paragraphs)
        self._validate_toc(paragraphs)
        self._validate_appendix_compliance(paragraphs)
        self._validate_frame_requirements(doc)

        for table_index, table in enumerate(doc.tables, start=1):
            self._validate_table(table, table_index)

        self._validate_footnotes_part(doc_path)

        return self.diagnostics

    def _validate_page_numbering(self, doc) -> None:
        rule_id = "common.page.numbering"
        try:
            rule = self._rule(rule_id)
        except KeyError:
            return

        from docx.oxml.ns import qn as _qn

        for section_index, section in enumerate(doc.sections, start=1):
            sect_pr = section._sectPr

            if section_index == 1:
                title_pg = sect_pr.find(_qn("w:titlePg"))
                title_pg_active = title_pg is not None and title_pg.get(
                    _qn("w:val")
                ) not in ("0", "false")
                if not title_pg_active:
                    self._add(
                        code=DiagnosticCodes.FORMAT_PAGE_NUMBERING,
                        message=(
                            f"Section {section_index}: titlePg flag is missing — "
                            f"page number would appear on the title page"
                        ),
                        rule_id=rule_id,
                    )

            footer = section.footer
            footer_xml = footer._element.xml
            if " PAGE " not in footer_xml:
                continue

            for paragraph in footer.paragraphs:
                if " PAGE " not in paragraph._p.xml:
                    continue

                alignment = paragraph.paragraph_format.alignment
                if alignment != WD_ALIGN_PARAGRAPH.CENTER:
                    self._add(
                        code=DiagnosticCodes.FORMAT_PAGE_NUMBERING,
                        message=(
                            f"Section {section_index}: page number alignment "
                            "must be center"
                        ),
                        rule_id=rule_id,
                    )

                first_line_indent = paragraph.paragraph_format.first_line_indent
                if first_line_indent is not None and abs(first_line_indent - Cm(0)) > _LENGTH_TOLERANCE_EMU:
                    self._add(
                        code=DiagnosticCodes.FORMAT_PAGE_NUMBERING,
                        message=(
                            f"Section {section_index}: page number first-line "
                            f"indent {first_line_indent.cm:.2f} cm, expected 0 cm"
                        ),
                        rule_id=rule_id,
                    )

                for run in paragraph.runs:
                    font_name = run.font.name
                    if font_name and font_name != rule.parameters["font_name"]:
                        self._add(
                            code=DiagnosticCodes.FORMAT_PAGE_NUMBERING,
                            message=(
                                f"Section {section_index}: page number font "
                                f"'{font_name}', expected '{rule.parameters['font_name']}'"
                            ),
                            rule_id=rule_id,
                        )

                    font_size = _pt_value(run.font.size)
                    expected_size = rule.parameters["font_size_pt"]
                    if font_size and abs(font_size - expected_size) > _POINT_TOLERANCE:
                        self._add(
                            code=DiagnosticCodes.FORMAT_PAGE_NUMBERING,
                            message=(
                                f"Section {section_index}: page number size "
                                f"{font_size:.1f}pt, expected {expected_size}pt"
                            ),
                            rule_id=rule_id,
                        )

    def _validate_margins(self, doc) -> None:
        checks = (
            ("top_margin", "top_mm", DiagnosticCodes.FORMAT_MARGIN_TOP, "Top"),
            ("bottom_margin", "bottom_mm", DiagnosticCodes.FORMAT_MARGIN_BOTTOM, "Bottom"),
            ("left_margin", "left_mm", DiagnosticCodes.FORMAT_MARGIN_LEFT, "Left"),
            ("right_margin", "right_mm", DiagnosticCodes.FORMAT_MARGIN_RIGHT, "Right"),
        )

        for section_index, section in enumerate(doc.sections, start=1):
            rule_id = _margin_rule_id(section)
            params = self._rule(rule_id).parameters
            for attr_name, param_name, code, label in checks:
                actual = getattr(section, attr_name)
                expected = Cm(params[param_name] / 10)
                if abs(actual - expected) > _LENGTH_TOLERANCE_EMU:
                    self._add(
                        code=code,
                        message=(
                            f"{label} margin in section {section_index} is "
                            f"{actual.cm:.2f} cm, expected {expected.cm:.2f} cm"
                        ),
                        rule_id=rule_id,
                    )

    def _validate_paragraph(self, paragraph, index: int, role: ParagraphRole | None = None) -> None:
        if not paragraph.text.strip():
            return

        if role is None:
            role = classify(paragraph, profile=self.profile)

        if role is ParagraphRole.FIGURE_PLACEHOLDER:
            self._add(
                code=DiagnosticCodes.FIGURE_MISSING_IMAGE,
                message=f"Paragraph {index}: figure image is missing",
                severity=Severity.INFO,
                rule_id="common.figure.image",
                source=SourceSpan(index, index),
            )
            return
        if role is ParagraphRole.TABLE_UNIT:
            self._validate_table_unit(paragraph, index)
            return
        if role is ParagraphRole.FIGURE_EXPLANATORY:
            for run_index, run in enumerate(paragraph.runs, start=1):
                self._validate_run_font(run, index, run_index, validate_size=False)
            self._validate_styled_paragraph(paragraph, index, "common.figure.explanatory_data")
            self._validate_figure_explanatory(paragraph, index)
            return
        if role is ParagraphRole.FOOTNOTE_TEXT:
            self._validate_footnote_text_paragraph(paragraph, index)
            return

        for run_index, run in enumerate(paragraph.runs, start=1):
            self._validate_run_font(run, index, run_index)

        if role in _HEADING_ROLES:
            self._validate_heading(paragraph, index, role)
            return
        if role is ParagraphRole.STRUCTURAL_HEADING or role is ParagraphRole.TOC_HEADING:
            self._validate_styled_paragraph(paragraph, index, "common.structural.heading")
            return
        if role is ParagraphRole.APPENDIX_HEADING:
            rule_id = "common.appendix.heading"
            try:
                self._validate_styled_paragraph(paragraph, index, rule_id)
            except KeyError:
                self._validate_styled_paragraph(paragraph, index, "common.structural.heading")
            return
        if role is ParagraphRole.TOC_ENTRY:
            return
        if role is ParagraphRole.TABLE_CAPTION:
            self._validate_styled_paragraph(paragraph, index, "common.table.caption")
            return
        if role is ParagraphRole.FIGURE_CAPTION:
            self._validate_styled_paragraph(paragraph, index, "common.figure.caption")
            self._validate_figure_caption_semantics(paragraph, index)
            return
        if role is ParagraphRole.LIST_ITEM:
            self._validate_list_paragraph(paragraph, index)
            return
        if role is ParagraphRole.FORMULA_BODY:
            self._validate_styled_paragraph(paragraph, index, "common.formula.body")
            self._validate_formula_body_semantics(paragraph, index)
            return
        if role is ParagraphRole.FORMULA_EXPLANATION:
            self._validate_styled_paragraph(paragraph, index, "common.formula.explanation")
            self._validate_formula_explanation_semantics(paragraph, index)
            return
        if role is ParagraphRole.BIBLIOGRAPHY_ENTRY:
            self._validate_styled_paragraph(paragraph, index, "common.bibliography.entry")
            return
        if role is ParagraphRole.UNKNOWN:
            self._add(
                code=DiagnosticCodes.FORMAT_ROLE_UNRECOGNIZED,
                message=f"Paragraph {index}: paragraph role could not be classified",
                severity=Severity.WARNING,
                source=SourceSpan(index, index),
            )
            return

        self._validate_body_paragraph(paragraph, index)

    def _validate_styled_paragraph(self, paragraph, index: int, rule_id: str) -> None:
        """Validate ``paragraph`` against a non-body rule's paragraph format."""

        params = self._rule(rule_id).parameters
        if "indent_cm" in params:
            self._validate_first_line_indent(
                paragraph, index, rule_id=rule_id, expected_cm=params["indent_cm"]
            )
        if "line_spacing" in params:
            self._validate_line_spacing(
                paragraph, index, rule_id=rule_id, expected=params["line_spacing"]
            )
        if "alignment" in params:
            self._validate_alignment(
                paragraph, index, rule_id=rule_id, expected=params["alignment"]
            )
        self._validate_paragraph_spacing(paragraph, index, rule_id)

    def _validate_body_paragraph(self, paragraph, index: int) -> None:
        self._validate_first_line_indent(
            paragraph,
            index,
            rule_id="common.text.indent.first_line",
            expected_cm=self._rule("common.text.indent.first_line").parameters["indent_cm"],
        )
        self._validate_line_spacing(
            paragraph,
            index,
            rule_id="common.text.line_spacing",
            expected=self._rule("common.text.line_spacing").parameters["spacing"],
        )
        self._validate_alignment(
            paragraph,
            index,
            rule_id="common.text.alignment",
            expected=self._rule("common.text.alignment").parameters["alignment"],
        )
        self._validate_paragraph_spacing(paragraph, index, "common.text.line_spacing")

    def _validate_list_paragraph(self, paragraph, index: int) -> None:
        rule_id = "common.list.item"
        params = self._rule(rule_id).parameters
        if "line_spacing" in params:
            self._validate_line_spacing(paragraph, index, rule_id=rule_id, expected=params["line_spacing"])
        if "alignment" in params:
            self._validate_alignment(paragraph, index, rule_id=rule_id, expected=params["alignment"])
        self._validate_paragraph_spacing(paragraph, index, rule_id)
        self._validate_list_indent(paragraph, index)
        self._validate_list_marker_semantics(paragraph, index)

    def _validate_list_indent(self, paragraph, index: int) -> None:
        left_indent = paragraph.paragraph_format.left_indent
        first_line_indent = paragraph.paragraph_format.first_line_indent
        if left_indent is None:
            return
        marker = _list_marker(paragraph.text)
        expected_left = Cm(1.75) if marker and marker[:-1].isdigit() else Cm(1.25)
        if abs(left_indent - expected_left) > _LENGTH_TOLERANCE_EMU:
            self._add(
                code=DiagnosticCodes.FORMAT_INDENT,
                message=(
                    f"Paragraph {index}: list left indent {left_indent.cm:.2f} cm, "
                    f"expected {expected_left.cm:.2f} cm"
                ),
                rule_id="common.list.item",
                source=SourceSpan(index, index),
            )
        if first_line_indent is not None and abs(first_line_indent - Cm(-0.5)) > _LENGTH_TOLERANCE_EMU:
            self._add(
                code=DiagnosticCodes.FORMAT_INDENT,
                message=(
                    f"Paragraph {index}: list hanging indent {first_line_indent.cm:.2f} cm, "
                    "expected -0.50 cm"
                ),
                rule_id="common.list.item",
                source=SourceSpan(index, index),
            )

    def _validate_list_marker_semantics(self, paragraph, index: int) -> None:
        marker = _list_marker(paragraph.text)
        if marker is None:
            return

        if marker == "-":
            self._last_letter_indent_cm = None
            return

        marker_value = marker[:-1].casefold()
        left_indent = paragraph.paragraph_format.left_indent
        indent_cm = round(left_indent.cm if left_indent is not None else 0, 2)

        if marker_value.isdigit():
            self._validate_nested_numeric_indent(paragraph, index, indent_cm)
            return

        if marker_value not in RUSSIAN_LIST_LETTER_INDEX:
            self._add(
                code="LIST_MARKER_DISALLOWED_LETTER",
                message=f"Paragraph {index}: lettered list marker '{marker}' is not allowed",
                rule_id="common.list.lettered",
                source=SourceSpan(index, index),
            )
            return

        expected_index = self._list_expected_letter_by_indent.get(indent_cm, 0)
        actual_index = RUSSIAN_LIST_LETTER_INDEX[marker_value]
        if actual_index != expected_index:
            self._add(
                code="LIST_MARKER_OUT_OF_ORDER",
                message=f"Paragraph {index}: lettered list marker '{marker}' is out of order",
                rule_id="common.list.marker_alphabetical",
                source=SourceSpan(index, index),
            )
        self._list_expected_letter_by_indent[indent_cm] = actual_index + 1
        self._last_letter_indent_cm = indent_cm

    def _validate_nested_numeric_indent(self, paragraph, index: int, indent_cm: float) -> None:
        if self._last_letter_indent_cm is None:
            return
        expected = self._last_letter_indent_cm + 0.5
        if abs(indent_cm - expected) <= 0.05:
            return
        self._add(
            code="LIST_NESTED_NUMERIC_INDENT",
            message=(
                f"Paragraph {index}: nested numeric list indent {indent_cm:.2f} cm, "
                f"expected {expected:.2f} cm"
            ),
            rule_id="common.list.nested_numeric",
            source=SourceSpan(index, index),
        )

    def _validate_heading(self, paragraph, index: int, role: ParagraphRole | None = None) -> None:
        if role is None or role not in _HEADING_RULE_BY_ROLE:
            rule_id = self._heading_rule_id(paragraph)
        else:
            rule_id = _HEADING_RULE_BY_ROLE[role]
        params = self._rule(rule_id).parameters

        self._validate_first_line_indent(
            paragraph,
            index,
            rule_id=rule_id,
            expected_cm=params.get("indent_cm", 0),
        )
        self._validate_line_spacing(
            paragraph,
            index,
            rule_id=rule_id,
            expected=params.get("line_spacing", 1.0),
        )
        self._validate_alignment(
            paragraph,
            index,
            rule_id=rule_id,
            expected=params.get("alignment", "left"),
        )
        self._validate_paragraph_spacing(paragraph, index, rule_id)

        expected_bold = bool(params.get("bold", False))
        if expected_bold and paragraph.runs and not any(run.bold for run in paragraph.runs):
            self._add(
                code=DiagnosticCodes.FORMAT_HEADING_BOLD,
                message=f"Paragraph {index}: heading must be bold",
                rule_id=rule_id,
                source=SourceSpan(index, index),
            )

        if paragraph.text.strip().endswith("."):
            self._add(
                code=DiagnosticCodes.FORMAT_HEADING_NO_PERIOD,
                message=f"Paragraph {index}: heading must not end with a period",
                rule_id="common.heading.no_period",
                source=SourceSpan(index, index),
            )
        heading_text = paragraph.text.strip()
        if "\u00ad" in heading_text or "-\n" in heading_text or "-\r\n" in heading_text:
            self._add(
                code=DiagnosticCodes.HEADING_HYPHENATION,
                message=f"Paragraph {index}: heading must not contain word-break hyphenation",
                rule_id="common.heading.no_hyphenation",
                source=SourceSpan(index, index),
            )
        if _violates_two_sentence_heading(heading_text):
            self._add(
                code=DiagnosticCodes.HEADING_TWO_SENTENCE,
                message=f"Paragraph {index}: two-sentence heading separator is invalid",
                rule_id="common.heading.two_sentence_separator",
                source=SourceSpan(index, index),
            )

    def _validate_heading_blank_lines(self, paragraph, index: int, prev, nxt) -> None:
        if _requires_blank_line_before(prev, self.profile):
            self._add(
                code=DiagnosticCodes.HEADING_SPACING_BEFORE,
                message=f"Paragraph {index}: heading must be preceded by one blank line",
                rule_id="common.heading.spacing_before",
                source=SourceSpan(index, index),
            )
        if _requires_blank_line_after(nxt, self.profile):
            self._add(
                code=DiagnosticCodes.HEADING_SPACING_AFTER,
                message=f"Paragraph {index}: heading must be followed by one blank line",
                rule_id="common.heading.spacing_after",
                source=SourceSpan(index, index),
            )

    def _validate_heading_subpoint_structure(self, paragraphs) -> None:
        entries = [
            (index, classify(paragraph, profile=self.profile))
            for index, paragraph in enumerate(paragraphs, start=1)
            if paragraph.text.strip()
        ]
        min_subpoints = int(self._rule("common.heading.point_requires_subpoints").parameters.get("min_subpoints", 2))
        for group in _docx_h3_groups(entries):
            h3_count = sum(1 for _, role in group if role is ParagraphRole.HEADING_H3)
            has_h4 = any(role is ParagraphRole.HEADING_H4 for _, role in group)
            if h3_count <= 1 or not has_h4:
                continue
            for position, (paragraph_index, role) in enumerate(group):
                if role is not ParagraphRole.HEADING_H3:
                    continue
                direct_h4_count = 0
                for _, child_role in group[position + 1 :]:
                    if child_role is ParagraphRole.HEADING_H3:
                        break
                    if child_role is ParagraphRole.HEADING_H4:
                        direct_h4_count += 1
                if direct_h4_count == 0 or direct_h4_count >= min_subpoints:
                    continue
                self._add(
                    code=DiagnosticCodes.HEADING_POINT_REQUIRES_SUBPOINTS,
                    message=(
                        f"Paragraph {paragraph_index}: point heading has "
                        f"{direct_h4_count} subpoint, expected at least {min_subpoints}"
                    ),
                    rule_id="common.heading.point_requires_subpoints",
                    source=SourceSpan(paragraph_index, paragraph_index),
                    data={"subpoint_count": direct_h4_count},
                )

    def _validate_toc(self, paragraphs) -> None:
        entries: list[tuple[int, object]] = []
        heading_texts: set[str] = set()
        appendix_letters: list[str] = []

        for index, paragraph in enumerate(paragraphs, start=1):
            role = classify(paragraph, profile=self.profile)
            if role is ParagraphRole.TOC_ENTRY:
                entries.append((index, paragraph))
                continue
            if role in _HEADING_ROLES or role is ParagraphRole.STRUCTURAL_HEADING:
                heading_texts.add(_normalize_toc_text(paragraph.text))
                continue
            if role is ParagraphRole.APPENDIX_HEADING:
                heading_texts.add(_normalize_toc_text(paragraph.text))
                letter = _appendix_letter_from_heading(paragraph.text)
                if letter:
                    appendix_letters.append(letter)

        if not entries:
            return

        self._validate_toc_indents(entries)
        self._validate_toc_matches_headings(entries, heading_texts)
        self._validate_toc_appendix_grouping(entries, appendix_letters)

    def _validate_toc_indents(self, entries: list[tuple[int, object]]) -> None:
        for index, paragraph in entries:
            level = _toc_style_level(paragraph)
            if level is None:
                continue
            expected_cm = max(level - 1, 0) * 0.5
            left_indent = paragraph.paragraph_format.left_indent
            actual_cm = left_indent.cm if left_indent is not None else 0
            if abs(actual_cm - expected_cm) > 0.05:
                self._add(
                    code=DiagnosticCodes.TOC_INDENT_LEVEL,
                    message=(
                        f"Paragraph {index}: TOC level {level} indent "
                        f"{actual_cm:.2f} cm, expected {expected_cm:.2f} cm"
                    ),
                    rule_id="common.toc.indent_levels",
                    source=SourceSpan(index, index),
                )

    def _validate_toc_matches_headings(
        self,
        entries: list[tuple[int, object]],
        heading_texts: set[str],
    ) -> None:
        for index, paragraph in entries:
            entry_text = _toc_entry_heading_text(paragraph.text)
            normalized = _normalize_toc_text(entry_text)
            if not normalized or normalized.startswith("приложения "):
                continue
            if normalized.startswith("приложение "):
                continue
            if normalized not in heading_texts:
                self._add(
                    code=DiagnosticCodes.TOC_ENTRY_MISMATCH,
                    message=f"Paragraph {index}: TOC entry '{entry_text}' does not match a document heading",
                    rule_id="common.toc.matches_headings",
                    source=SourceSpan(index, index),
                )

    def _validate_toc_appendix_grouping(
        self,
        entries: list[tuple[int, object]],
        appendix_letters: list[str],
    ) -> None:
        if len(appendix_letters) < 3 or not _letters_are_contiguous(appendix_letters):
            return
        expected = f"Приложения {appendix_letters[0]}–{appendix_letters[-1]}"
        if any(_toc_entry_heading_text(paragraph.text).startswith(expected) for _, paragraph in entries):
            return
        self._add(
            code=DiagnosticCodes.TOC_APPENDIX_GROUPING,
            message=f"TOC must include grouped appendix entry '{expected}'",
            rule_id="common.toc.appendix_grouping",
        )

    def _validate_appendix_compliance(self, paragraphs) -> None:
        appendix_letters: list[str] = []
        current_letter: str | None = None
        label_letters: set[str] = set()

        for index, paragraph in enumerate(paragraphs, start=1):
            role = classify(paragraph, profile=self.profile)
            text = paragraph.text.strip()
            if role is ParagraphRole.APPENDIX_HEADING:
                letter = _appendix_letter_from_heading(text)
                if letter:
                    appendix_letters.append(letter)
                    current_letter = letter
                continue

            label_match = _APPENDIX_CONTINUATION_RE.match(text)
            if label_match is not None:
                label_letters.add(label_match.group("letter").upper())
                continue

            if current_letter and role in _HEADING_ROLES:
                expected_prefix = _expected_appendix_heading_prefix(current_letter, role)
                if expected_prefix and not text.startswith(expected_prefix):
                    self._add(
                        code=DiagnosticCodes.APPENDIX_SECTION_NUMBERING,
                        message=(
                            f"Paragraph {index}: appendix heading must start "
                            f"with '{expected_prefix}'"
                        ),
                        rule_id="common.appendix.section_numbering",
                        source=SourceSpan(index, index),
                        data={"letter": current_letter, "expected": expected_prefix},
                    )

        if appendix_letters and not _letters_are_contiguous(appendix_letters):
            self._add(
                code=DiagnosticCodes.APPENDIX_LETTER_SEQUENCE,
                message="Appendix letters must be contiguous Russian capitals with excluded letters skipped",
                rule_id="common.appendix.auto_letter",
                data={"letters": appendix_letters},
            )

        for label_letter in sorted(label_letters):
            if label_letter not in appendix_letters:
                self._add(
                    code=DiagnosticCodes.APPENDIX_CONTINUATION_LABEL,
                    message=f"Continuation label references unknown appendix {label_letter}",
                    rule_id="common.appendix.continuation_label",
                    data={"letter": label_letter},
                )

    def _validate_frame_requirements(self, doc) -> None:
        rule_ids = {rule.id for rule in self.profile.rules}
        framed = has_frame(doc)
        forms = _main_inscription_forms(doc)

        if "coursework.frame.course_project_explanatory_note" in rule_ids:
            if not framed or not forms.intersection({"form_1", "form_2", "form_3", "form_4"}):
                self._add(
                    code="FRAME_MISSING",
                    message="Coursework explanatory note must include a frame and form 1/3 title block",
                    rule_id="coursework.frame.course_project_explanatory_note",
                )

        if "project_designations.explanatory_note.frame" in rule_ids and not framed:
            self._add(
                code="FRAME_MISSING",
                message="Project explanatory note must include a framed sheet",
                rule_id="project_designations.explanatory_note.frame",
            )

        if (
            "graphic_and_demonstration_materials.sheet.frame" in rule_ids
            and self.profile.name == "graphic_and_demonstration_materials"
        ):
            if not framed or not forms.intersection({"form_5", "form_6"}):
                self._add(
                    code="FRAME_MISSING",
                    message="Graphic sheet must include a frame and form 5/6 title block",
                    rule_id="graphic_and_demonstration_materials.sheet.frame",
                )

        if "project_designations.title_block.letter_numeric_designation" in rule_ids:
            self._validate_project_designation_graphs(doc)

    def _validate_project_designation_graphs(self, doc) -> None:
        title_blocks = _main_inscription_tables(doc)
        if not title_blocks:
            self.diagnostics.extend(validate_designation_text(""))
            return
        for table in title_blocks:
            graph_2 = _main_inscription_graph_text(table, 2)
            self.diagnostics.extend(validate_designation_text(graph_2))

    def _validate_table(self, table, table_index: int) -> None:
        if _style_name(table) == docx_styles.ABBREVIATIONS_TABLE:
            self._validate_abbreviations_table(table, table_index)
            return

        rule = self._rule("common.table.font.size")
        min_size = rule.parameters["min_size_pt"]
        max_size = rule.parameters["max_size_pt"]
        self._validate_table_borders(table, table_index)
        self._validate_table_forbidden_headers(table, table_index)
        self._validate_table_header_periods(table, table_index)
        self._validate_table_diagonal_split(table, table_index)

        for row_index, row in enumerate(table.rows, start=1):
            for cell_index, cell in enumerate(row.cells, start=1):
                self._validate_table_cell_padding(table_index, row_index, cell_index, cell)
                for paragraph in cell.paragraphs:
                    stripped = paragraph.text.strip()
                    if _ITALIC_LETTER_RE.fullmatch(stripped):
                        self._validate_table_italic_letter(
                            table_index,
                            row_index,
                            cell_index,
                            paragraph,
                        )
                    for run_index, run in enumerate(paragraph.runs, start=1):
                        self._validate_run_font(
                            run,
                            f"table {table_index} row {row_index} cell {cell_index}",
                            run_index,
                            validate_size=False,
                        )
                        size = _pt_value(run.font.size)
                        if size and not (min_size <= size <= max_size):
                            self._add(
                                code=DiagnosticCodes.FORMAT_TABLE_FONT_SIZE,
                                message=(
                                    "Table "
                                    f"{table_index} row {row_index} cell {cell_index} "
                                    f"run {run_index}: size {size:.1f}pt, "
                                    f"expected {min_size}-{max_size}pt"
                                ),
                                rule_id=rule.id,
                            )

    def _validate_table_unit(self, paragraph, index: int) -> None:
        rule_id = "common.table.unit_label"
        if paragraph.paragraph_format.alignment != WD_ALIGN_PARAGRAPH.RIGHT:
            self._add(
                code=DiagnosticCodes.FORMAT_ALIGNMENT,
                message=f"Paragraph {index}: table unit label must be right-aligned",
                rule_id=rule_id,
                source=SourceSpan(index, index),
            )
        for run_index, run in enumerate(paragraph.runs, start=1):
            self._validate_run_font(run, index, run_index, validate_size=False)
            size = _pt_value(run.font.size)
            expected = self._rule(rule_id).parameters["font_size_pt"]
            if size and abs(size - expected) > _POINT_TOLERANCE:
                self._add(
                    code=DiagnosticCodes.FORMAT_FONT_SIZE,
                    message=(
                        f"Paragraph {index} run {run_index}: size {size:.1f}pt, "
                        f"expected {expected}pt"
                    ),
                    rule_id=rule_id,
                    source=SourceSpan(index, index),
                )

    def _validate_figure_explanatory(self, paragraph, index: int) -> None:
        rule_id = "common.figure.explanatory_data"
        if _style_name(paragraph) != docx_styles.FIGURE_EXPLANATORY:
            self._add(
                code=DiagnosticCodes.FORMAT_ROLE_UNRECOGNIZED,
                message=f"Paragraph {index}: figure explanatory data must use {docx_styles.FIGURE_EXPLANATORY}",
                rule_id=rule_id,
                source=SourceSpan(index, index),
            )
        expected_size = self._rule(rule_id).parameters["font_size_pt"]
        for run_index, run in enumerate(paragraph.runs, start=1):
            size = _pt_value(run.font.size)
            if size and abs(size - expected_size) > _POINT_TOLERANCE:
                self._add(
                    code=DiagnosticCodes.FORMAT_FONT_SIZE,
                    message=(
                        f"Paragraph {index} run {run_index}: size {size:.1f}pt, "
                        f"expected {expected_size}pt"
                    ),
                    rule_id=rule_id,
                    source=SourceSpan(index, index),
                )

    def _validate_figure_caption_semantics(self, paragraph, index: int) -> None:
        text = paragraph.text.strip()
        if re.search(r"\bлист\s+\d+\b", text, flags=re.IGNORECASE) and not re.match(
            r"^\s*Рисунок\s+\S+,\s*лист\s+\d+\s*$",
            text,
            flags=re.IGNORECASE,
        ):
            self._add(
                code=DiagnosticCodes.FIGURE_MULTI_SHEET_LABEL,
                message=f"Paragraph {index}: multi-sheet figure caption must use ', лист K'",
                rule_id="common.figure.multi_sheet_label",
                source=SourceSpan(index, index),
            )

    def _validate_formula_body_semantics(self, paragraph, index: int) -> None:
        self._validate_first_line_indent(
            paragraph,
            index,
            rule_id="common.formula.body_indent",
            expected_cm=self._rule("common.formula.body_indent").parameters["indent_cm"],
        )

        body_text = _formula_body_text(paragraph.text)
        if (
            self._previous_non_empty_role is ParagraphRole.FORMULA_BODY
            and self._previous_non_empty_paragraph is not None
            and not _formula_body_text(self._previous_non_empty_paragraph.text).rstrip().endswith(",")
        ):
            self._add(
                code=DiagnosticCodes.FORMULA_CONSECUTIVE_COMMA,
                message=f"Paragraph {index}: consecutive formulas must be separated by comma",
                rule_id="common.formula.consecutive_comma",
                source=SourceSpan(index, index),
            )

        if "\n" not in body_text and len(body_text) > 80 and any(sign in body_text for sign in OPERATOR_SIGNS):
            self._add(
                code=DiagnosticCodes.FORMULA_LINE_CONTINUATION,
                message=f"Paragraph {index}: long formula should break on an operator sign",
                severity=Severity.INFO,
                rule_id="common.formula.line_continuation",
                source=SourceSpan(index, index),
            )

    def _validate_formula_explanation_semantics(self, paragraph, index: int) -> None:
        lines = [line.strip() for line in paragraph.text.splitlines() if line.strip()]
        if not lines:
            return
        marker = lines[0].casefold()
        if marker.startswith("где:"):
            self._add(
                code=DiagnosticCodes.FORMULA_EXPLANATION_MARKER,
                message=f"Paragraph {index}: formula explanation marker must be 'где' without colon",
                rule_id="common.formula.explanation_marker",
                source=SourceSpan(index, index),
            )

        for line in lines[1:]:
            name = _symbol_name(line)
            if not name:
                continue
            if "то же, что и в формуле" in line.casefold():
                if name not in self._seen_formula_symbols:
                    self._add(
                        code=DiagnosticCodes.FORMULA_REPEATED_SYMBOL,
                        message=(
                            f"Paragraph {index}: repeated symbol '{name}' must "
                            "reference a symbol introduced in an earlier formula"
                        ),
                        rule_id="common.formula.repeated_symbol",
                        source=SourceSpan(index, index),
                    )
            else:
                self._seen_formula_symbols.add(name)

    def _validate_abbreviations_table(self, table, table_index: int) -> None:
        rule_id = "common.abbreviations.two_column_layout"
        for row_index, row in enumerate(table.rows, start=1):
            if len(row.cells) != 2:
                self._add(
                    code=DiagnosticCodes.FORMAT_ABBREVIATIONS_TABLE,
                    message=f"Abbreviations table {table_index} row {row_index}: expected 2 columns",
                    rule_id=rule_id,
                )
            for cell_index, cell in enumerate(row.cells, start=1):
                for paragraph in cell.paragraphs:
                    for run_index, run in enumerate(paragraph.runs, start=1):
                        self._validate_run_font(
                            run,
                            f"abbreviations table {table_index} row {row_index} cell {cell_index}",
                            run_index,
                            validate_size=False,
                        )
                        size = _pt_value(run.font.size)
                        if size and abs(size - 14) > _POINT_TOLERANCE:
                            self._add(
                                code=DiagnosticCodes.FORMAT_ABBREVIATIONS_TABLE,
                                message=(
                                    f"Abbreviations table {table_index} row {row_index} "
                                    f"cell {cell_index}: size {size:.1f}pt, expected 14pt"
                                ),
                                rule_id=rule_id,
                            )

    def _validate_footnote_text_paragraph(self, paragraph, index: int) -> None:
        self._validate_line_spacing(
            paragraph,
            index,
            rule_id="common.reference.footnote",
            expected=1.0,
        )
        for run_index, run in enumerate(paragraph.runs, start=1):
            self._validate_run_font(run, index, run_index, validate_size=False)
            size = _pt_value(run.font.size)
            if size and size > 12 + _POINT_TOLERANCE:
                self._add(
                    code=DiagnosticCodes.FOOTNOTE_FORMAT,
                    message=(
                        f"Paragraph {index} run {run_index}: footnote size "
                        f"{size:.1f}pt, expected 12pt or smaller"
                    ),
                    rule_id="common.reference.footnote",
                    source=SourceSpan(index, index),
                )

    def _validate_footnotes_part(self, doc_path: Path) -> None:
        try:
            with ZipFile(doc_path) as package:
                if "word/footnotes.xml" not in package.namelist():
                    return
                footnotes_xml = package.read("word/footnotes.xml")
        except Exception:
            return

        namespace = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
        root = ET.fromstring(footnotes_xml)
        if not any(
            footnote.get(qn("w:type")) == "separator"
            for footnote in root.findall("w:footnote", namespace)
        ):
            self._add(
                code=DiagnosticCodes.FOOTNOTE_FORMAT,
                message="Footnote separator line is missing",
                severity=Severity.WARNING,
                rule_id="common.reference.footnote",
            )

        for footnote in root.findall("w:footnote", namespace):
            note_type = footnote.get(qn("w:type"))
            if note_type in {"separator", "continuationSeparator"}:
                continue
            self._validate_footnote_xml(footnote, namespace)

    def _validate_footnote_xml(self, footnote, namespace: dict[str, str]) -> None:
        note_id = footnote.get(qn("w:id"), "?")
        spacing = footnote.find(".//w:spacing", namespace)
        if spacing is not None and spacing.get(qn("w:line")) not in {None, "240"}:
            self._add(
                code=DiagnosticCodes.FOOTNOTE_FORMAT,
                message=f"Footnote {note_id}: line spacing must be single",
                rule_id="common.reference.footnote",
            )
        for size in footnote.findall(".//w:sz", namespace):
            value = size.get(qn("w:val"))
            try:
                half_points = int(value) if value is not None else 0
            except ValueError:
                continue
            if half_points > 24:
                self._add(
                    code=DiagnosticCodes.FOOTNOTE_FORMAT,
                    message=f"Footnote {note_id}: text size must be 12pt or smaller",
                    rule_id="common.reference.footnote",
                )

    def _validate_table_borders(self, table, table_index: int) -> None:
        rule_id = "common.table.borders"
        xml = table._tbl.xml
        if 'w:val="nil"' not in xml or 'w:val="double"' not in xml:
            self._add(
                code=DiagnosticCodes.FORMAT_TABLE_BORDERS,
                message=(
                    f"Table {table_index}: expected no top border and double line "
                    "below the header"
                ),
                severity=Severity.WARNING,
                rule_id=rule_id,
            )

    def _validate_table_forbidden_headers(self, table, table_index: int) -> None:
        if not table.rows:
            return
        rule_id = "common.table.forbid_serial_column"
        forbidden = {
            " ".join(value.casefold().split())
            for value in self._rule(rule_id).parameters["forbidden_headers"]
        }
        for cell_index, cell in enumerate(table.rows[0].cells, start=1):
            text = " ".join(cell.text.casefold().split())
            if text in forbidden:
                self._add(
                    code=DiagnosticCodes.FORMAT_TABLE_FORBIDDEN_SERIAL_COLUMN,
                    message=f"Table {table_index} column {cell_index}: serial-number column is forbidden",
                    rule_id=rule_id,
                )

    def _validate_table_header_periods(self, table, table_index: int) -> None:
        if not table.rows:
            return
        rule_id = "common.table.no_period_in_header"
        for cell_index, cell in enumerate(table.rows[0].cells, start=1):
            if cell.text.strip().endswith("."):
                self._add(
                    code=DiagnosticCodes.FORMAT_TABLE_HEADER_PERIOD,
                    message=f"Table {table_index} header cell {cell_index}: header must not end with a period",
                    rule_id=rule_id,
                )

    def _validate_table_diagonal_split(self, table, table_index: int) -> None:
        xml = table._tbl.xml
        if any(marker in xml for marker in ("w:tl2br", "w:tr2bl", "<m:d", "<a:graphic")):
            self._add(
                code=DiagnosticCodes.FORMAT_TABLE_DIAGONAL_SPLIT,
                message=f"Table {table_index}: diagonal-split cells are forbidden",
                rule_id="common.table.no_diagonal_split",
            )

    def _validate_table_cell_padding(
        self,
        table_index: int,
        row_index: int,
        cell_index: int,
        cell,
    ) -> None:
        top = _cell_margin_twips(cell, "top")
        bottom = _cell_margin_twips(cell, "bottom")
        if top == _TABLE_CELL_PADDING_TWIPS and bottom == _TABLE_CELL_PADDING_TWIPS:
            return
        self._add(
            code=DiagnosticCodes.FORMAT_TABLE_CELL_PADDING,
            message=(
                f"Table {table_index} row {row_index} cell {cell_index}: "
                "vertical padding must be 6pt"
            ),
            severity=Severity.WARNING,
            rule_id="common.table.cell_padding",
        )

    def _validate_table_italic_letter(
        self,
        table_index: int,
        row_index: int,
        cell_index: int,
        paragraph,
    ) -> None:
        if any(run.italic or run.font.italic for run in paragraph.runs):
            return
        self._add(
            code=DiagnosticCodes.FORMAT_TABLE_ITALIC_LETTER,
            message=(
                f"Table {table_index} row {row_index} cell {cell_index}: "
                "single-letter designation must be italic"
            ),
            rule_id="common.table.italic_letters",
        )

    def _validate_run_font(
        self,
        run,
        paragraph_index,
        run_index: int,
        *,
        validate_size: bool = True,
    ) -> None:
        name_rule = self._rule("common.text.font.name")
        expected_name = name_rule.parameters["font_name"]
        if run.font.name and run.font.name != expected_name:
            self._add(
                code=DiagnosticCodes.FORMAT_FONT_NAME,
                message=(
                    f"Paragraph {paragraph_index} run {run_index}: font '{run.font.name}', expected '{expected_name}'"
                ),
                rule_id=name_rule.id,
                source=_source_for_index(paragraph_index),
            )

        if validate_size:
            size_rule = self._rule("common.text.font.size")
            expected_size = size_rule.parameters["font_size_pt"]
            current_size = _pt_value(run.font.size)
            if current_size and abs(current_size - expected_size) > _POINT_TOLERANCE:
                self._add(
                    code=DiagnosticCodes.FORMAT_FONT_SIZE,
                    message=(
                        f"Paragraph {paragraph_index} run {run_index}: "
                        f"size {current_size:.1f}pt, expected {expected_size}pt"
                    ),
                    rule_id=size_rule.id,
                    source=_source_for_index(paragraph_index),
                )

        color_rule = self._rule("common.text.font.color")
        expected_color = tuple(color_rule.parameters["color_rgb"])
        current_color = run.font.color.rgb
        if current_color is not None and tuple(current_color) != expected_color:
            self._add(
                code=DiagnosticCodes.FORMAT_FONT_COLOR,
                message=(
                    f"Paragraph {paragraph_index} run {run_index}: "
                    f"color {tuple(current_color)}, expected {expected_color}"
                ),
                rule_id=color_rule.id,
                source=_source_for_index(paragraph_index),
            )

    def _validate_first_line_indent(
        self,
        paragraph,
        index: int,
        *,
        rule_id: str,
        expected_cm: float,
    ) -> None:
        current = paragraph.paragraph_format.first_line_indent
        if current is None:
            return
        expected = Cm(expected_cm)
        if abs(current - expected) > _LENGTH_TOLERANCE_EMU:
            self._add(
                code=DiagnosticCodes.FORMAT_INDENT,
                message=(f"Paragraph {index}: first-line indent {current.pt:.1f}pt, expected {expected.pt:.1f}pt"),
                rule_id=rule_id,
                source=SourceSpan(index, index),
            )

    def _validate_line_spacing(
        self,
        paragraph,
        index: int,
        *,
        rule_id: str,
        expected: float,
    ) -> None:
        current = paragraph.paragraph_format.line_spacing
        if current is None:
            return
        spacing = _spacing_value(current)
        if abs(spacing - expected) > _SPACING_TOLERANCE:
            self._add(
                code=DiagnosticCodes.FORMAT_LINE_SPACING,
                message=f"Paragraph {index}: line spacing {spacing}, expected {expected}",
                rule_id=rule_id,
                source=SourceSpan(index, index),
            )

    def _validate_alignment(
        self,
        paragraph,
        index: int,
        *,
        rule_id: str,
        expected: str,
    ) -> None:
        current = paragraph.paragraph_format.alignment
        if current is None:
            return
        expected_alignment = _ALIGNMENT_BY_NAME[expected]
        if current != expected_alignment:
            self._add(
                code=DiagnosticCodes.FORMAT_ALIGNMENT,
                message=f"Paragraph {index}: alignment {current}, expected {expected}",
                rule_id=rule_id,
                source=SourceSpan(index, index),
            )

    def _validate_paragraph_spacing(self, paragraph, index: int, rule_id: str) -> None:
        pf = paragraph.paragraph_format
        for label, value in (("before", pf.space_before), ("after", pf.space_after)):
            current = _pt_value(value)
            if current > 1:
                self._add(
                    code=DiagnosticCodes.FORMAT_PARAGRAPH_SPACING,
                    message=(f"Paragraph {index}: spacing {label} {current:.1f}pt, expected 0pt"),
                    rule_id=rule_id,
                    source=SourceSpan(index, index),
                )

    def _is_heading_paragraph(self, paragraph) -> bool:
        style_name = paragraph.style.name if paragraph.style else ""
        return style_name.startswith("Heading") or (
            paragraph.paragraph_format.alignment == WD_ALIGN_PARAGRAPH.CENTER
            and any(run.bold for run in paragraph.runs)
        )

    def _heading_rule_id(self, paragraph) -> str:
        style_name = paragraph.style.name if paragraph.style else ""
        if style_name.startswith("Heading 2"):
            return "common.heading.h2"
        if style_name.startswith("Heading 3"):
            return "common.heading.h3"
        return "common.heading.h1"

    def _rule(self, rule_id: str) -> FormattingRule:
        for rule in self.profile.rules:
            if rule.id == rule_id:
                return rule
        return get_rule(rule_id)

    def _add(
        self,
        *,
        code: str,
        message: str,
        severity: Severity = Severity.ERROR,
        rule_id: str | None = None,
        source: SourceSpan | None = None,
        data: dict | None = None,
    ) -> None:
        self.diagnostics.append(
            Diagnostic(
                code=code,
                message=message,
                severity=severity,
                rule_id=rule_id,
                source=source,
                data=data,
            )
        )


def diagnostic_to_json(diagnostic: Diagnostic) -> dict[str, object]:
    source = _diagnostic_source(diagnostic)
    payload: dict[str, object] = {
        "code": diagnostic.code,
        "severity": diagnostic.severity.value,
        "message": diagnostic.message,
        "ruleId": diagnostic.rule_id,
        "source": source,
        "data": dict(diagnostic.data or {}),
    }
    if diagnostic.rule_id:
        try:
            rule = get_rule(diagnostic.rule_id)
        except KeyError:
            rule = None
        if rule is not None:
            payload["source"] = {
                **source,
                "document": rule.source_doc,
                "section": rule.source_section,
            }
    if diagnostic.target is not None:
        payload["target"] = diagnostic.target
    if diagnostic.suggestion:
        payload["suggestion"] = diagnostic.suggestion
    return payload


def _diagnostic_source(diagnostic: Diagnostic) -> dict[str, object]:
    span = diagnostic.source
    return {
        "document": span.filename if span is not None else None,
        "section": None,
        "lineStart": span.line_start if span is not None else None,
        "lineEnd": span.line_end if span is not None else None,
    }


def _metadata_from_docx(doc) -> dict[str, str]:
    metadata: dict[str, str] = {}
    core = doc.core_properties
    if core.title:
        metadata["title"] = core.title
    if core.author:
        metadata["student"] = core.author
    if core.subject:
        metadata["subject"] = core.subject
    metadata.update(_parse_metadata_payload(core.keywords or ""))

    for paragraph in doc.paragraphs:
        if _style_name(paragraph) != "SFUMetadata":
            continue
        metadata.update(_parse_metadata_payload(paragraph.text))

    return metadata


def _parse_metadata_payload(payload: str) -> dict[str, str]:
    metadata: dict[str, str] = {}
    for item in re.split(r"[;\n\r]+", payload or ""):
        key, separator, value = item.partition("=")
        key = key.strip()
        if not separator or not key:
            continue
        metadata[key] = value.strip()
    return metadata


def _source_for_index(index) -> SourceSpan | None:
    return SourceSpan(index, index) if isinstance(index, int) else None


def _margin_rule_id(section) -> str:
    if section.orientation == WD_ORIENT.LANDSCAPE or section.page_width > section.page_height:
        return "common.page.margins.landscape"
    return "common.page.margins.portrait"


_LIST_MARKER_RE = re.compile(r"^\s*(?P<marker>(?:[а-яА-Я]|\d+)\)|[-—–])\s+")


def _list_marker(text: str) -> str | None:
    match = _LIST_MARKER_RE.match(text or "")
    if match is None:
        return None
    marker = match.group("marker").replace("—", "-").replace("–", "-")
    if marker == "-":
        return marker
    return marker.casefold()


def _main_inscription_forms(doc) -> set[str]:
    forms: set[str] = set()
    for table in _main_inscription_tables(doc):
        text = table.rows[0].cells[0].text.casefold()
        match = re.search(r"форма\s+(form_[1-6])", text)
        if match is not None:
            forms.add(match.group(1))
    return forms


def _main_inscription_tables(doc) -> list[object]:
    tables = []
    for table in doc.tables:
        if not table.rows or not table.rows[0].cells:
            continue
        text = table.rows[0].cells[0].text.casefold()
        if re.search(r"форма\s+form_[1-6]", text):
            tables.append(table)
    return tables


def _main_inscription_graph_text(table, graph_number: int) -> str:
    index = graph_number - 1
    if len(table.rows) <= index or len(table.rows[index].cells) < 2:
        return ""
    return table.rows[index].cells[1].text.strip()


def _requires_blank_line_before(paragraph, profile) -> bool:
    return _requires_heading_blank_line(paragraph, profile)


def _requires_blank_line_after(paragraph, profile) -> bool:
    return _requires_heading_blank_line(paragraph, profile)


def _requires_heading_blank_line(paragraph, profile) -> bool:
    if paragraph is None or not paragraph.text.strip():
        return False
    role = classify(paragraph, profile=profile)
    return role not in _HEADING_ROLES and role not in {
        ParagraphRole.STRUCTURAL_HEADING,
        ParagraphRole.APPENDIX_HEADING,
        ParagraphRole.TOC_HEADING,
    }


def _violates_two_sentence_heading(text: str) -> bool:
    normalized = " ".join((text or "").strip().splitlines())
    first_period = normalized.find(".")
    if first_period < 0 or first_period == len(normalized) - 1:
        return False
    follows_single_space = (
        first_period + 2 < len(normalized)
        and normalized[first_period + 1] == " "
        and normalized[first_period + 2] != " "
    )
    return not follows_single_space or normalized.endswith(".")


def _docx_h3_groups(entries: list[tuple[int, ParagraphRole]]) -> list[list[tuple[int, ParagraphRole]]]:
    groups: list[list[tuple[int, ParagraphRole]]] = []
    current: list[tuple[int, ParagraphRole]] = []
    for entry in entries:
        _, role = entry
        if role in {ParagraphRole.HEADING_H1, ParagraphRole.HEADING_H2}:
            if current:
                groups.append(current)
                current = []
            continue
        if role in {ParagraphRole.HEADING_H3, ParagraphRole.HEADING_H4}:
            current.append(entry)
    if current:
        groups.append(current)
    return groups


def _formula_body_text(text: str) -> str:
    return (text or "").split("\t", 1)[0]


def _symbol_name(line: str) -> str | None:
    head, separator, _ = line.partition("—")
    if not separator:
        head, separator, _ = line.partition("-")
    name = head.strip()
    return name or None


def _style_name(element) -> str:
    style = getattr(element, "style", None)
    if style is None:
        return ""
    return getattr(style, "name", "") or ""


_APPENDIX_LETTERS = APPENDIX_LETTERS
_APPENDIX_LETTER_BY_VALUE = {letter: index for index, letter in enumerate(_APPENDIX_LETTERS)}
_APPENDIX_HEADING_LETTER_RE = re.compile(r"\bПРИЛОЖЕНИЕ\s+([А-Я])\b", re.IGNORECASE)
_APPENDIX_CONTINUATION_RE = re.compile(
    r"^(?:Продолжение|Окончание)\s+приложения\s+(?P<letter>[А-Я])$",
    re.IGNORECASE,
)
_APPENDIX_HEADING_PREFIX_BY_ROLE = {
    ParagraphRole.HEADING_H1: "{letter}.",
    ParagraphRole.HEADING_H2: "{letter}.",
    ParagraphRole.HEADING_H3: "{letter}.",
    ParagraphRole.HEADING_H4: "{letter}.",
}


def _toc_style_level(paragraph) -> int | None:
    match = re.match(r"TOC\s+(\d+)$", _style_name(paragraph))
    if match is None:
        return None
    return int(match.group(1))


def _toc_entry_heading_text(text: str) -> str:
    heading = (text or "").split("\t", 1)[0]
    heading = re.sub(r"\s*[.\u2024\u2025\u2026·…]{2,}\s*\d+(?:[-–—]\d+)?\s*$", "", heading)
    return " ".join(heading.strip().split())


def _normalize_toc_text(text: str) -> str:
    return _toc_entry_heading_text(text).casefold()


def _appendix_letter_from_heading(text: str) -> str | None:
    match = _APPENDIX_HEADING_LETTER_RE.search(text or "")
    if match is None:
        return None
    return match.group(1).upper()


def _expected_appendix_heading_prefix(letter: str, role: ParagraphRole) -> str | None:
    pattern = _APPENDIX_HEADING_PREFIX_BY_ROLE.get(role)
    if pattern is None:
        return None
    return pattern.format(letter=letter)


def _letters_are_contiguous(letters: list[str]) -> bool:
    if any(letter not in _APPENDIX_LETTER_BY_VALUE for letter in letters):
        return False
    positions = [_APPENDIX_LETTER_BY_VALUE[letter] for letter in letters]
    return positions == list(range(positions[0], positions[0] + len(positions)))


def _cell_margin_twips(cell, edge: str) -> int | None:
    tc_pr = cell._tc.tcPr
    if tc_pr is None:
        return None
    margins = tc_pr.find(qn("w:tcMar"))
    if margins is None:
        return None
    element = margins.find(qn(f"w:{edge}"))
    if element is None:
        return None
    value = element.get(qn("w:w"))
    if value is None:
        return None
    try:
        return int(value)
    except ValueError:
        return None


def _pt_value(value) -> float:
    if value is None:
        return 0
    if hasattr(value, "pt"):
        return float(value.pt)
    return float(value)


def _spacing_value(value) -> float:
    if hasattr(value, "pt"):
        return float(value.pt)
    return float(value)


def _slug(value: str) -> str:
    return value.strip().lower().replace(" ", "-")

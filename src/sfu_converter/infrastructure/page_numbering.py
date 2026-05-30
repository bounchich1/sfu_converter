"""Page numbering configuration for DOCX sections (STU 7.5-07-2021 §7.2).

Each :class:`PageNumberingSection` describes how one DOCX section participates
in document-wide numbering: where the number is rendered (footer or framed
title block field 7), whether the first page suppresses the visible number
(``<w:titlePg/>``), and whether the entire section omits the field while still
contributing to the global page counter.

The :func:`configure` function applies a list of section descriptors to a
``python-docx`` ``Document`` in order, so the caller is responsible for
emitting the appropriate section breaks before invoking it.
"""
from __future__ import annotations

from dataclasses import dataclass
from enum import Enum

from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


class Location(Enum):
    """Where to render the page number for a section."""

    FOOTER_CENTER = "footer_center"
    FRAME_FIELD_7 = "frame_field_7"


@dataclass(frozen=True)
class PageNumberingSection:
    """Per-section page numbering descriptor."""

    start_at: int | None
    hide_first_page: bool
    suppress_in_section: bool
    location: Location


def configure(doc, sections: list[PageNumberingSection], config) -> None:
    """Apply ``sections`` to ``doc.sections`` in order.

    For each DOCX section, install ``<w:pgNumType>`` if ``start_at`` is set,
    toggle ``<w:titlePg/>`` per ``hide_first_page``, and place a centered
    PAGE field in the footer unless ``suppress_in_section`` is true.
    """
    docx_sections = list(doc.sections)
    if len(sections) > len(docx_sections):
        raise ValueError(
            f"configure() received {len(sections)} descriptors but document "
            f"has only {len(docx_sections)} sections"
        )

    for descriptor, docx_section in zip(sections, docx_sections):
        _apply_section(descriptor, docx_section, config)


def _apply_section(descriptor: PageNumberingSection, section, config) -> None:
    section.different_first_page_header_footer = descriptor.hide_first_page
    if descriptor.start_at is not None:
        _set_page_number_start(section, descriptor.start_at)

    if descriptor.location is Location.FRAME_FIELD_7:
        # Task 24 wires field 7. Until it lands, fall back to footer.
        location = Location.FOOTER_CENTER
    else:
        location = descriptor.location

    if location is Location.FOOTER_CENTER:
        _populate_footer(section, descriptor, config)


def _set_page_number_start(section, start_at: int) -> None:
    sect_pr = section._sectPr
    existing = sect_pr.find(qn("w:pgNumType"))
    if existing is not None:
        sect_pr.remove(existing)
    pg_num_type = OxmlElement("w:pgNumType")
    pg_num_type.set(qn("w:start"), str(start_at))
    sect_pr.append(pg_num_type)


def _populate_footer(section, descriptor: PageNumberingSection, config) -> None:
    page_cfg = config.PAGE_NUMBERING

    footer = section.footer
    footer.is_linked_to_previous = False
    paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    paragraph.clear()
    pf = paragraph.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf.first_line_indent = Cm(0)
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)

    if descriptor.suppress_in_section:
        if descriptor.hide_first_page:
            _clear_first_page_footer(section)
        return

    run = paragraph.add_run()
    run.font.name = page_cfg["font_name"]
    run.font.size = page_cfg["font_size"]
    run.font.color.rgb = RGBColor(*config.FONT_COLOR_RGB)
    run._element.get_or_add_rPr().rFonts.set(
        qn("w:eastAsia"),
        page_cfg["font_name"],
    )

    field_begin = OxmlElement("w:fldChar")
    field_begin.set(qn("w:fldCharType"), "begin")
    run._element.append(field_begin)

    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    run._element.append(instruction)

    field_end = OxmlElement("w:fldChar")
    field_end.set(qn("w:fldCharType"), "end")
    run._element.append(field_end)

    if descriptor.hide_first_page:
        _clear_first_page_footer(section)


def _clear_first_page_footer(section) -> None:
    first_footer = section.first_page_footer
    first_footer.is_linked_to_previous = False
    first_paragraph = (
        first_footer.paragraphs[0] if first_footer.paragraphs else first_footer.add_paragraph()
    )
    first_paragraph.clear()

from __future__ import annotations

from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

from sfu_converter.config import SIBFUConfig
from sfu_converter.infrastructure import docx_styles


def render(
    document,
    form: str,
    *,
    fields: dict[str, str] | None = None,
    page_number_in_graph_7: bool = False,
):
    docx_styles.register_styles(document)
    table = document.add_table(rows=17, cols=2)
    table.style = docx_styles.FRAME_MAIN_INSCRIPTION
    table.autofit = False
    values = fields or {}

    for graph_number in range(1, 18):
        row = table.rows[graph_number - 1]
        row.cells[0].text = f"Форма {form}" if graph_number == 1 else str(graph_number)
        value_cell = row.cells[1]
        value_cell.text = ""
        paragraph = value_cell.paragraphs[0]
        if page_number_in_graph_7 and graph_number == 7:
            _add_page_field(paragraph, SIBFUConfig)
        else:
            paragraph.add_run(values.get(str(graph_number), ""))
        _format_cell(row.cells[0])
        _format_cell(value_cell)
    return table


def graph_text(table, graph_number: int) -> str:
    return table.rows[graph_number - 1].cells[1].text


def _add_page_field(paragraph, config) -> None:
    run = paragraph.add_run()
    run.font.name = config.PAGE_NUMBERING["font_name"]
    run.font.size = config.PAGE_NUMBERING["font_size"]
    run.font.color.rgb = RGBColor(*config.FONT_COLOR_RGB)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), config.PAGE_NUMBERING["font_name"])

    field_begin = OxmlElement("w:fldChar")
    field_begin.set(qn("w:fldCharType"), "begin")
    run._element.append(field_begin)

    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    run._element.append(instruction)

    field_end = OxmlElement("w:fldChar")
    field_end.set(qn("w:fldCharType"), "end")
    run._element.append(field_end)


def _format_cell(cell) -> None:
    for paragraph in cell.paragraphs:
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        for run in paragraph.runs:
            run.font.name = "Times New Roman"
            run.font.size = Pt(12)
            run.font.color.rgb = RGBColor(0, 0, 0)
            run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Times New Roman")

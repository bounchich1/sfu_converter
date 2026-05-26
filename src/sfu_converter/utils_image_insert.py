"""
Модуль для корректной вставки изображений в документ DOCX.
Поддерживает: конвертацию форматов, пропорциональное масштабирование, max_width.
"""

import io
import logging
from pathlib import Path
from typing import Any

from docx.document import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Emu, Inches, Pt
from PIL import Image

from sfu_converter.config import MeasurementConfig

logger = logging.getLogger(__name__)
DEFAULT_DPI = MeasurementConfig.DEFAULT_DPI


# ============================================================================
# 1. КОНВЕРТАЦИЯ ИЗОБРАЖЕНИЯ В RGB
# ============================================================================


def convert_image_to_rgb(img: Image.Image) -> Image.Image:
    """
    Конвертирует изображение в режим RGB с сохранением визуального вида.

    - RGBA/LA/P: добавляет белый фон под прозрачность
    - CMYK: конвертирует в RGB
    - Остальные режимы: принудительная конвертация в RGB

    Args:
        img: PIL.Image объект

    Returns:
        Image.Image в режиме RGB
    """
    if img.mode == "RGB":
        return img

    if img.mode in ("RGBA", "LA"):
        background = Image.new("RGB", img.size, (255, 255, 255))
        if img.mode == "LA":
            img = img.convert("RGBA")
        background.paste(img, mask=img.split()[-1])
        return background

    if img.mode == "P":
        img = img.convert("RGBA")
        background = Image.new("RGB", img.size, (255, 255, 255))
        background.paste(img, mask=img.split()[-1])
        return background

    if img.mode == "CMYK":
        return img.convert("RGB")

    # Все остальные режимы (1, L, etc.)
    return img.convert("RGB")


# ============================================================================
# 2. РАСЧЁТ РАЗМЕРОВ С УЧЁТОМ ПРОПОРЦИЙ И ОГРАНИЧЕНИЙ
# ============================================================================


def _to_emu(
    value: Cm | Inches | int | float | None,
    dpi: int = DEFAULT_DPI,
) -> int | None:
    """Приводит значение к EMU (English Metric Units) для внутренних расчётов."""
    if value is None:
        return None
    if isinstance(value, Cm):
        return int(value.emu)
    if isinstance(value, Inches):
        return int(value.emu)
    if isinstance(value, Emu):
        return int(value)
    # Если передано число — считаем, что это пиксели при заданном DPI
    return int(value * MeasurementConfig.EMU_PER_INCH / dpi)


def _from_emu_to_cm(emu: int) -> Cm:
    """Конвертирует EMU обратно в Cm."""
    return Cm(emu / MeasurementConfig.EMU_PER_CM)


def calculate_image_dimensions(
    original_size: tuple[int, int],
    width: Cm | Inches | int | None = None,
    height: Cm | Inches | int | None = None,
    max_width: Cm | Inches | int | None = None,
    dpi: int = DEFAULT_DPI,
) -> tuple[Cm | Inches | None, Cm | Inches | None]:
    """
    Рассчитывает финальные размеры изображения с учётом:
    - явных width/height из конфига
    - сохранения пропорций, если задано только одно измерение
    - ограничения max_width

    Args:
        original_size: (width_px, height_px) исходного изображения
        width: желаемая ширина (Cm/Inches/int) или None
        height: желаемая высота (Cm/Inches/int) или None
        max_width: максимальная ширина (Cm/Inches/int) или None
        dpi: DPI для расчётов (по умолчанию 96)

    Returns:
        (width, height) в тех же единицах, что и входные параметры, или None
    """
    orig_w, orig_h = original_size
    aspect_ratio = orig_h / orig_w if orig_w > 0 else 1

    # 1. Применяем max_width, если задан и ширина не фиксирована явно
    if max_width is not None and width is None:
        max_w_emu = _to_emu(max_width, dpi)
        current_w_emu = _to_emu(orig_w, dpi)

        if current_w_emu and max_w_emu and current_w_emu > max_w_emu:
            width = _from_emu_to_cm(max_w_emu)
            height = _from_emu_to_cm(int(max_w_emu * aspect_ratio))
            return width, height

    # 2. Если задана только ширина — вычисляем высоту
    if width is not None and height is None:
        if isinstance(width, Cm):
            height = Cm(width.emu * aspect_ratio / MeasurementConfig.EMU_PER_CM)
        elif isinstance(width, Inches):
            height = Inches(width.inches * aspect_ratio)
        else:
            height = width * aspect_ratio  # для int/float (пиксели)

    # 3. Если задана только высота — вычисляем ширину
    elif height is not None and width is None:
        inv_aspect = orig_w / orig_h if orig_h > 0 else 1
        if isinstance(height, Cm):
            width = Cm(height.emu * inv_aspect / MeasurementConfig.EMU_PER_CM)
        elif isinstance(height, Inches):
            width = Inches(height.inches * inv_aspect)
        else:
            width = height * inv_aspect

    return width, height


# ============================================================================
# 3. СОХРАНЕНИЕ ИЗОБРАЖЕНИЯ В БУФЕР
# ============================================================================


def save_image_to_buffer(
    img: Image.Image, format: str = "PNG", dpi: tuple[int, int] = (DEFAULT_DPI, DEFAULT_DPI), quality: int = 95
) -> io.BytesIO:
    """
    Сохраняет изображение в BytesIO-буфер в нужном формате.

    Args:
        img: PIL.Image в режиме RGB
        format: формат сохранения ('PNG', 'JPEG', 'TIFF')
        dpi: кортеж (x_dpi, y_dpi) — применяется только к JPEG/TIFF
        quality: качество для JPEG (1-100)

    Returns:
        io.BytesIO с данными изображения, готовый к чтению
    """
    buffer = io.BytesIO()
    save_format = format.upper()

    # JPEG не поддерживает прозрачность — гарантируем RGB
    if save_format == "JPEG" and img.mode != "RGB":
        img = img.convert("RGB")

    # DPI поддерживается только JPEG и TIFF
    if save_format in ("JPEG", "TIFF"):
        img.save(buffer, format=save_format, dpi=dpi, quality=quality)
    else:
        # PNG, BMP, etc. — DPI игнорируется, quality не применяется
        img.save(buffer, format=save_format)

    buffer.seek(0)
    return buffer


# ============================================================================
# 4. ВСТАВКА ИЗОБРАЖЕНИЯ В АБЗАЦЕЦ
# ============================================================================


def insert_image_into_paragraph(
    doc: Document,
    image_buffer: io.BytesIO,
    width: Cm | Inches | None = None,
    height: Cm | Inches | None = None,
    line_spacing: float = 1.5,
    alignment: WD_ALIGN_PARAGRAPH = WD_ALIGN_PARAGRAPH.CENTER,
    space_before=None,
    space_after=None,
) -> None:
    """
    Создаёт абзац с настройками и вставляет в него изображение.

    Args:
        doc: объект python-docx Document
        image_buffer: BytesIO с данными изображения
        width/height: размеры для add_picture (или None для авто)
        line_spacing: межстрочный интервал абзаца
        alignment: выравнивание абзаца
        space_before/after: отступы до/после абзаца
    """
    if space_before is None:
        space_before = Pt(0)
    if space_after is None:
        space_after = Pt(0)

    para = doc.add_paragraph()
    pf = para.paragraph_format
    pf.line_spacing = line_spacing
    pf.space_before = space_before
    pf.space_after = space_after
    pf.alignment = alignment
    pf.first_line_indent = Cm(0)

    run = para.add_run()
    run.add_picture(image_buffer, width=width, height=height)


# ============================================================================
# 5. ПУБЛИЧНЫЙ ИНТЕРФЕЙС: ОДНА ФУНКЦИЯ ДЛЯ ВСЕГО
# ============================================================================


def insert_image(
    doc: Document, image_path: str | Path, config: dict[str, Any], logger: logging.Logger | None = None
) -> bool:
    """
    Главная функция: вставляет изображение в документ с учётом конфига.

    Args:
        doc: python-docx Document
        image_path: путь к файлу изображения
        config: словарь с настройками (см. пример ниже)
        logger: опциональный логгер для отладки

    Returns:
        bool: True при успешной вставке, False при ошибке

    Пример конфига:
        {
            'line_spacing': 1.5,
            'alignment': WD_ALIGN_PARAGRAPH.CENTER,
            'width': None,          # или Cm(10)
            'height': None,         # или Cm(8)
            'dpi': (96, 96),
            'format': 'PNG',        # или 'JPEG'
            'max_width': Cm(15),    # ограничение по ширине
            'space_before': Pt(0),
            'space_after': Pt(0),
        }
    """
    log = logger or logging.getLogger(__name__)

    try:
        full_path = str(Path(image_path).resolve())

        # 1. Открываем и конвертируем изображение
        with Image.open(full_path) as img:
            img_rgb = convert_image_to_rgb(img)

            # 2. Рассчитываем размеры
            width, height = calculate_image_dimensions(
                original_size=img.size,
                width=config.get("width"),
                height=config.get("height"),
                max_width=config.get("max_width"),
                dpi=config.get("dpi", (DEFAULT_DPI, DEFAULT_DPI))[0],
            )

            # 3. Сохраняем в буфер
            buffer = save_image_to_buffer(
                img_rgb,
                format=config.get("format", "PNG"),
                dpi=config.get("dpi", (DEFAULT_DPI, DEFAULT_DPI)),
                quality=config.get("quality", 95),
            )

            # 4. Вставляем в документ
            insert_image_into_paragraph(
                doc=doc,
                image_buffer=buffer,
                width=width,
                height=height,
                line_spacing=config.get("line_spacing", 1.5),
                alignment=config.get("alignment", WD_ALIGN_PARAGRAPH.CENTER),
                space_before=config.get("space_before", Pt(0)),
                space_after=config.get("space_after", Pt(0)),
            )

        log.info(f"✓ Изображение вставлено: {image_path}")
        return True

    except FileNotFoundError:
        log.error(f"✗ Файл не найден: {image_path}")
        return False
    except Exception as e:
        log.error(f"✗ Ошибка при вставке {image_path}: {type(e).__name__}: {e}")
        return False

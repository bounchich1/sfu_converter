import pytest
from docx import Document as DocxDocument
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.shared import Cm, Pt, RGBColor
from PIL import Image

from sfu_converter.config import SIBFUConfig
from sfu_converter.infrastructure import docx_styles
from sfu_converter.domain.ast_nodes import (
    AppendixNode,
    BibliographyEntryNode,
    Document,
    FormulaNode,
    FigureNode,
    HeadingLevel,
    HeadingNode,
    ListItemNode,
    ListNode,
    ListType,
    MetadataNode,
    PageBreakNode,
    ParagraphNode,
    RawBlockNode,
    ReferenceNode,
    StructuralSectionNode,
    StructuralSectionType,
    TableCaptionNode,
    TableCell,
    TableNode,
    TableOfContentsNode,
    TableRow,
    TextRun,
    TitlePageNode,
)
from sfu_converter.domain.diagnostics import DiagnosticCodes
from sfu_converter.domain.formatting import FormattingProfile, FormattingRule, RuleSeverity, RuleStatus
import sfu_converter.infrastructure.docx_renderer as renderer_module
from sfu_converter.infrastructure.docx_renderer import DocxRenderer, SectionNumberer
from sfu_converter.infrastructure.docx_renderer import _normalize_list_item_punctuation
from sfu_converter.ports.renderer import RendererPort


def assert_close(actual, expected) -> None:
    assert abs(actual - expected) < 1000


def test_docx_renderer_renders_ast_to_file(tmp_path):
    renderer = DocxRenderer(config_class=SIBFUConfig, base_dir=tmp_path)
    profile = FormattingProfile(
        name="common",
        display_name="Common",
        source_docs=("standard",),
    )
    ast = Document(
        blocks=(
            HeadingNode(level=HeadingLevel.H1, text="Title"),
            ParagraphNode(runs=(TextRun("Body"),)),
            TableNode(
                caption="Table 1",
                rows=(
                    TableRow(cells=(TableCell("A"), TableCell("B"))),
                    TableRow(cells=(TableCell("1"), TableCell("2"))),
                ),
            ),
        )
    )
    output_path = tmp_path / "rendered.docx"

    diagnostics = renderer.render_to_file(ast, profile, str(output_path))

    assert diagnostics == []
    assert output_path.exists()

    doc = DocxDocument(str(output_path))
    assert [para.text for para in doc.paragraphs if para.text] == [
        "Title",
        "Body",
        "Таблица 1 — Table 1",
    ]
    assert doc.paragraphs[0].paragraph_format.alignment == WD_ALIGN_PARAGRAPH.CENTER
    assert len(doc.tables) == 1
    assert doc.tables[0].rows[1].cells[1].text == "2"


def test_docx_renderer_reports_unsupported_renderer_rules(tmp_path):
    unsupported_rule = FormattingRule(
        id="synthetic.renderer.rule",
        source_doc="docs/formatting requirements/common.md",
        source_section="Synthetic",
        severity=RuleSeverity.REQUIRED,
        renderer_status=RuleStatus.NOT_SUPPORTED,
        validator_status=RuleStatus.IMPLEMENTED,
    )
    profile = FormattingProfile(
        name="synthetic",
        display_name="Synthetic",
        source_docs=("standard",),
        rules=(unsupported_rule,),
    )
    renderer = DocxRenderer(config_class=SIBFUConfig, base_dir=tmp_path)
    output_path = tmp_path / "rendered.docx"

    diagnostics = renderer.render_to_file(
        Document(blocks=(ParagraphNode(runs=(TextRun("Body"),)),)),
        profile,
        str(output_path),
    )

    assert output_path.exists()
    assert [diagnostic.code for diagnostic in diagnostics] == [
        DiagnosticCodes.FORMAT_RULE_NOT_SUPPORTED
    ]
    assert diagnostics[0].rule_id == "synthetic.renderer.rule"


def test_section_numberer_tracks_hierarchical_numbers():
    numberer = SectionNumberer()

    assert numberer.next_number(3) == "1.1.1"
    numberer.reset()
    assert numberer.next_number(1) == "1"
    assert numberer.next_number(2) == "1.1"
    assert numberer.next_number(2) == "1.2"
    assert numberer.next_number(3) == "1.2.1"
    assert numberer.next_number(1) == "2"
    assert numberer.next_number(2) == "2.1"

    numberer.reset()

    assert numberer.next_number(1) == "1"


def test_section_numberer_tracks_subpoint_level_four():
    numberer = SectionNumberer()

    assert numberer.next_number(1) == "1"
    assert numberer.next_number(2) == "1.1"
    assert numberer.next_number(3) == "1.1.1"
    assert numberer.next_number(4) == "1.1.1.1"
    assert numberer.next_number(1) == "2"
    assert numberer.next_number(2) == "2.1"
    assert numberer.next_number(3) == "2.1.1"
    assert numberer.next_number(4) == "2.1.1.1"


def test_section_numberer_rejects_unsupported_levels():
    with pytest.raises(ValueError, match="Unsupported heading level"):
        SectionNumberer().next_number(5)


def test_docx_renderer_numbers_h4_subpoints_and_resets(tmp_path):
    renderer = DocxRenderer(config_class=SIBFUConfig, base_dir=tmp_path)
    profile = FormattingProfile(
        name="common",
        display_name="Common",
        source_docs=("standard",),
    )
    ast = Document(
        blocks=(
            HeadingNode(level=HeadingLevel.H1, text="One", number="auto"),
            HeadingNode(level=HeadingLevel.H2, text="One.One", number="auto"),
            HeadingNode(level=HeadingLevel.H3, text="One.One.One", number="auto"),
            HeadingNode(level=HeadingLevel.H4, text="One.One.One.One", number="auto"),
            HeadingNode(level=HeadingLevel.H1, text="Two", number="auto"),
            HeadingNode(level=HeadingLevel.H2, text="Two.One", number="auto"),
            HeadingNode(level=HeadingLevel.H3, text="Two.One.One", number="auto"),
            HeadingNode(level=HeadingLevel.H4, text="Two.One.One.One", number="auto"),
        )
    )
    output_path = tmp_path / "subpoints.docx"

    renderer.render_to_file(ast, profile, str(output_path))

    doc = DocxDocument(str(output_path))
    assert [paragraph.text for paragraph in doc.paragraphs if paragraph.text] == [
        "1 One",
        "1.1 One.One",
        "1.1.1 One.One.One",
        "1.1.1.1 One.One.One.One",
        "2 Two",
        "2.1 Two.One",
        "2.1.1 Two.One.One",
        "2.1.1.1 Two.One.One.One",
    ]


def test_docx_renderer_output_includes_heading_4_style(tmp_path):
    renderer = DocxRenderer(config_class=SIBFUConfig, base_dir=tmp_path)
    profile = FormattingProfile(
        name="common",
        display_name="Common",
        source_docs=("standard",),
    )
    ast = Document(
        blocks=(
            HeadingNode(level=HeadingLevel.H4, text="Subpoint", number="auto"),
        )
    )
    output_path = tmp_path / "h4_style.docx"

    renderer.render_to_file(ast, profile, str(output_path))

    doc = DocxDocument(str(output_path))
    assert "Heading 4" in [style.name for style in doc.styles]
    h4_para = next(p for p in doc.paragraphs if p.text)
    assert h4_para.style.name == "Heading 4"


def test_docx_renderer_render_returns_docx_bytes(tmp_path):
    renderer = DocxRenderer(config_class=SIBFUConfig, base_dir=tmp_path)
    payload = renderer.render(
        Document(blocks=(ParagraphNode(runs=(TextRun("Body"),)),)),
        _common_profile(),
    )

    assert payload.startswith(b"PK")


def test_docx_renderer_renders_auto_numbered_headings_without_trailing_period(tmp_path):
    renderer = DocxRenderer(config_class=SIBFUConfig, base_dir=tmp_path)
    profile = FormattingProfile(
        name="common",
        display_name="Common",
        source_docs=("standard",),
    )
    ast = Document(
        blocks=(
            HeadingNode(level=HeadingLevel.H1, text="Первый раздел", number="auto"),
            HeadingNode(level=HeadingLevel.H2, text="Первый подраздел", number="auto"),
            HeadingNode(level=HeadingLevel.H3, text="Первый пункт", number="auto"),
            HeadingNode(level=HeadingLevel.H1, text="Второй раздел", number="auto"),
            HeadingNode(level=HeadingLevel.H2, text="Новый подраздел", number="auto"),
        )
    )
    output_path = tmp_path / "numbered.docx"

    renderer.render_to_file(ast, profile, str(output_path))

    doc = DocxDocument(str(output_path))
    assert [paragraph.text for paragraph in doc.paragraphs if paragraph.text] == [
        "1 Первый раздел",
        "1.1 Первый подраздел",
        "1.1.1 Первый пункт",
        "2 Второй раздел",
        "2.1 Новый подраздел",
    ]


def test_docx_renderer_renders_list_items_with_sfu_markers_and_formatting(tmp_path):
    renderer = DocxRenderer(config_class=SIBFUConfig, base_dir=tmp_path)
    profile = FormattingProfile(
        name="common",
        display_name="Common",
        source_docs=("standard",),
    )
    ast = Document(
        blocks=(
            ListNode(
                list_type=ListType.BULLET,
                items=(
                    ListItemNode("первый элемент"),
                    ListItemNode("второй элемент"),
                ),
            ),
            ListNode(
                list_type=ListType.LETTERED,
                items=(
                    ListItemNode("первый вариант"),
                    ListItemNode("второй вариант"),
                ),
            ),
            ListNode(
                list_type=ListType.NUMBERED,
                items=(
                    ListItemNode("подпункт один"),
                    ListItemNode("подпункт два"),
                ),
            ),
        )
    )
    output_path = tmp_path / "lists.docx"

    renderer.render_to_file(ast, profile, str(output_path))

    doc = DocxDocument(str(output_path))
    paragraphs = [paragraph for paragraph in doc.paragraphs if paragraph.text]
    assert [paragraph.text for paragraph in paragraphs] == [
        "- первый элемент;",
        "- второй элемент.",
        "а) первый вариант;",
        "б) второй вариант.",
        "1) подпункт один.",
        "2) подпункт два.",
    ]
    for paragraph in paragraphs:
        assert paragraph.paragraph_format.alignment == WD_ALIGN_PARAGRAPH.JUSTIFY
        assert_close(paragraph.paragraph_format.first_line_indent, Cm(1.25))
        assert paragraph.paragraph_format.line_spacing == 1.5


def test_docx_renderer_adds_page_numbering_to_default_footer(tmp_path):
    renderer = DocxRenderer(config_class=SIBFUConfig, base_dir=tmp_path)

    renderer._initialize_document()

    section = renderer.doc.sections[0]
    assert section.different_first_page_header_footer is True

    footer = section.footer
    assert footer.is_linked_to_previous is False
    assert " PAGE " in footer._element.xml

    paragraph = footer.paragraphs[0]
    assert paragraph.paragraph_format.alignment == WD_ALIGN_PARAGRAPH.CENTER
    assert_close(paragraph.paragraph_format.first_line_indent, Cm(0))

    run = paragraph.runs[0]
    assert run.font.name == SIBFUConfig.FONT_NAME
    assert run.font.size == Pt(14)
    assert run.font.color.rgb == RGBColor(0, 0, 0)

    first_page_footer = section.first_page_footer
    assert first_page_footer.paragraphs[0].text == ""
    assert " PAGE " not in first_page_footer._element.xml


def test_docx_renderer_renders_structural_sections_with_special_formatting(tmp_path):
    renderer = DocxRenderer(config_class=SIBFUConfig, base_dir=tmp_path)
    profile = FormattingProfile(
        name="common",
        display_name="Common",
        source_docs=("standard",),
    )
    ast = Document(
        blocks=(
            StructuralSectionNode(
                section_type=StructuralSectionType.INTRODUCTION,
                title="введение",
            ),
            ParagraphNode(runs=(TextRun("Body"),)),
        )
    )
    output_path = tmp_path / "structural.docx"

    renderer.render_to_file(ast, profile, str(output_path))

    doc = DocxDocument(str(output_path))
    texts = [paragraph.text for paragraph in doc.paragraphs]
    heading_index = texts.index("ВВЕДЕНИЕ")
    heading = doc.paragraphs[heading_index]

    assert heading_index > 0
    assert 'w:type="page"' in doc.paragraphs[heading_index - 1]._p.xml
    assert doc.paragraphs[heading_index + 1].text == ""
    assert "Body" in texts

    assert heading.paragraph_format.alignment == WD_ALIGN_PARAGRAPH.CENTER
    assert_close(heading.paragraph_format.first_line_indent, Cm(0))
    assert heading.paragraph_format.line_spacing == 1.0
    assert heading.runs[0].bold is True
    assert heading.runs[0].underline is False


def test_docx_renderer_implements_renderer_port():
    assert issubclass(DocxRenderer, RendererPort)


def test_docx_renderer_renders_paragraph_with_bold_and_italic_runs(tmp_path):
    renderer = DocxRenderer(config_class=SIBFUConfig, base_dir=tmp_path)
    profile = FormattingProfile(
        name="common",
        display_name="Common",
        source_docs=("standard",),
    )
    ast = Document(
        blocks=(
            ParagraphNode(
                runs=(
                    TextRun(text="Plain "),
                    TextRun(text="bold", bold=True),
                    TextRun(text=" and "),
                    TextRun(text="italic", italic=True),
                    TextRun(text=" and "),
                    TextRun(text="both", bold=True, italic=True),
                    TextRun(text="."),
                ),
            ),
        )
    )
    output_path = tmp_path / "inline.docx"

    renderer.render_to_file(ast, profile, str(output_path))

    doc = DocxDocument(str(output_path))
    paragraph = next(p for p in doc.paragraphs if p.text)
    assert paragraph.text == "Plain bold and italic and both."
    assert [
        (run.text, bool(run.bold), bool(run.font.italic))
        for run in paragraph.runs
    ] == [
        ("Plain ", False, False),
        ("bold", True, False),
        (" and ", False, False),
        ("italic", False, True),
        (" and ", False, False),
        ("both", True, True),
        (".", False, False),
    ]


def test_docx_renderer_paragraph_runs_share_font_settings(tmp_path):
    renderer = DocxRenderer(config_class=SIBFUConfig, base_dir=tmp_path)
    profile = FormattingProfile(
        name="common",
        display_name="Common",
        source_docs=("standard",),
    )
    ast = Document(
        blocks=(
            ParagraphNode(
                runs=(
                    TextRun(text="Жирный фрагмент", bold=True),
                    TextRun(text=" и обычный текст."),
                ),
            ),
        )
    )
    output_path = tmp_path / "inline_font.docx"

    renderer.render_to_file(ast, profile, str(output_path))

    doc = DocxDocument(str(output_path))
    paragraph = next(p for p in doc.paragraphs if p.text)
    for run in paragraph.runs:
        assert run.font.name == SIBFUConfig.FONT_NAME
        assert run.font.size == SIBFUConfig.FONT_SIZE
        assert run.font.color.rgb == RGBColor(*SIBFUConfig.FONT_COLOR_RGB)


def test_docx_renderer_table_uses_configured_font_size_and_header_bold(tmp_path):
    renderer = DocxRenderer(config_class=SIBFUConfig, base_dir=tmp_path)
    profile = FormattingProfile(
        name="common",
        display_name="Common",
        source_docs=("standard",),
    )
    ast = Document(
        blocks=(
            TableNode(
                rows=(
                    TableRow(cells=(TableCell("H1"), TableCell("H2"))),
                    TableRow(cells=(TableCell("A"), TableCell("B"))),
                ),
            ),
        )
    )
    output_path = tmp_path / "table_font.docx"

    renderer.render_to_file(ast, profile, str(output_path))

    doc = DocxDocument(str(output_path))
    table = doc.tables[0]
    header_run = table.rows[0].cells[0].paragraphs[0].runs[0]
    body_run = table.rows[1].cells[0].paragraphs[0].runs[0]

    assert header_run.font.size == SIBFUConfig.TABLE["font_size"]
    assert body_run.font.size == SIBFUConfig.TABLE["font_size"]
    assert header_run.bold is True
    assert body_run.bold is False
    assert header_run.font.italic is not True
    body_para = table.rows[1].cells[0].paragraphs[0]
    assert body_para.paragraph_format.line_spacing == SIBFUConfig.TABLE["line_spacing"]


def test_docx_renderer_table_caption_normalizes_to_em_dash(tmp_path):
    renderer = DocxRenderer(config_class=SIBFUConfig, base_dir=tmp_path)
    profile = FormattingProfile(
        name="common",
        display_name="Common",
        source_docs=("standard",),
    )
    ast = Document(
        blocks=(
            TableNode(
                caption="Данные",
                rows=(TableRow(cells=(TableCell("A"),)),),
            ),
            TableNode(
                caption="Таблица 7 - Существующая",
                rows=(TableRow(cells=(TableCell("B"),)),),
            ),
        )
    )
    output_path = tmp_path / "captions.docx"

    renderer.render_to_file(ast, profile, str(output_path))

    doc = DocxDocument(str(output_path))
    captions = [p.text for p in doc.paragraphs if p.text.startswith("Таблица")]
    assert captions == [
        "Таблица 1 — Данные",
        "Таблица 7 — Существующая",
    ]


def test_docx_renderer_sets_repeat_header_row_property_on_tables(tmp_path):
    renderer = DocxRenderer(config_class=SIBFUConfig, base_dir=tmp_path)
    profile = FormattingProfile(
        name="common",
        display_name="Common",
        source_docs=("standard",),
    )
    ast = Document(
        blocks=(
            TableNode(
                rows=(
                    TableRow(cells=(TableCell("H1"), TableCell("H2"))),
                    TableRow(cells=(TableCell("A"), TableCell("B"))),
                ),
            ),
        )
    )
    output_path = tmp_path / "repeat_header.docx"

    renderer.render_to_file(ast, profile, str(output_path))

    doc = DocxDocument(str(output_path))
    header_row = doc.tables[0].rows[0]
    assert "tblHeader" in header_row._tr.xml
    body_row = doc.tables[0].rows[1]
    assert "tblHeader" not in body_row._tr.xml


def _common_profile():
    return FormattingProfile(
        name="common",
        display_name="Common",
        source_docs=("standard",),
    )


def test_register_styles_exposes_all_sfu_styles_on_rendered_document(tmp_path):
    renderer = DocxRenderer(config_class=SIBFUConfig, base_dir=tmp_path)
    output_path = tmp_path / "styles.docx"

    renderer.render_to_file(
        Document(blocks=(ParagraphNode(runs=(TextRun("Body"),)),)),
        _common_profile(),
        str(output_path),
    )

    doc = DocxDocument(str(output_path))
    names = [style.name for style in doc.styles]
    for name in docx_styles.ALL_SFU_STYLES:
        assert name in names, f"{name} missing from rendered document"
        assert doc.styles[name] is not None


def test_caption_styles_are_based_on_caption_builtin():
    doc = DocxDocument()
    docx_styles.register_styles(doc)

    for name in (docx_styles.TABLE_CAPTION, docx_styles.FIGURE_CAPTION):
        base = doc.styles[name].base_style
        assert base is not None, f"{name} has no base style"
        assert base.name == "Caption"


def test_toc_heading_is_based_on_structural_heading():
    doc = DocxDocument()
    docx_styles.register_styles(doc)

    base = doc.styles[docx_styles.TOC_HEADING].base_style
    assert base is not None
    assert base.name == docx_styles.STRUCTURAL_HEADING


def test_formula_body_style_carries_right_aligned_tab_stop():
    doc = DocxDocument()
    docx_styles.register_styles(doc)

    tab_stops = list(doc.styles[docx_styles.FORMULA_BODY].paragraph_format.tab_stops)
    assert tab_stops, "SFUFormulaBody has no tab stops"
    assert any(stop.alignment == WD_TAB_ALIGNMENT.RIGHT for stop in tab_stops)


def test_register_styles_is_idempotent():
    doc = DocxDocument()
    docx_styles.register_styles(doc)
    before = len([s.name for s in doc.styles])
    docx_styles.register_styles(doc)
    after = len([s.name for s in doc.styles])
    assert before == after


def test_docx_renderer_centers_formula_with_auto_number(tmp_path):
    renderer = DocxRenderer(config_class=SIBFUConfig, base_dir=tmp_path)
    ast = Document(blocks=(FormulaNode(content="E = mc^2"),))
    output_path = tmp_path / "formula.docx"

    renderer.render_to_file(ast, _common_profile(), str(output_path))

    doc = DocxDocument(str(output_path))
    formula_para = next(p for p in doc.paragraphs if p.text)
    assert formula_para.text == "E = mc^2\t(1)"
    assert formula_para.paragraph_format.alignment == WD_ALIGN_PARAGRAPH.CENTER
    assert formula_para.paragraph_format.line_spacing == 1.5
    assert "<w:tabs" in formula_para._p.xml
    assert 'w:val="right"' in formula_para._p.xml


def test_docx_renderer_assigns_sequential_formula_numbers(tmp_path):
    renderer = DocxRenderer(config_class=SIBFUConfig, base_dir=tmp_path)
    ast = Document(
        blocks=(
            FormulaNode(content="x + y = z"),
            ParagraphNode(runs=(TextRun("Между формулами"),)),
            FormulaNode(content="a^2 + b^2 = c^2"),
        )
    )
    output_path = tmp_path / "formulas.docx"

    renderer.render_to_file(ast, _common_profile(), str(output_path))

    doc = DocxDocument(str(output_path))
    texts = [p.text for p in doc.paragraphs if p.text]
    assert "x + y = z\t(1)" in texts
    assert "a^2 + b^2 = c^2\t(2)" in texts


def test_docx_renderer_renders_formula_explanation_with_left_alignment(tmp_path):
    renderer = DocxRenderer(config_class=SIBFUConfig, base_dir=tmp_path)
    explanation = "где E — энергия, Дж;\n    m — масса, кг;\n    c — скорость света, м/с."
    ast = Document(
        blocks=(FormulaNode(content="E = mc^2", explanation=explanation),)
    )
    output_path = tmp_path / "formula_explanation.docx"

    renderer.render_to_file(ast, _common_profile(), str(output_path))

    doc = DocxDocument(str(output_path))
    paragraphs = [p for p in doc.paragraphs if p.text]
    formula_para = next(p for p in paragraphs if p.text.endswith("(1)"))
    explanation_para = next(p for p in paragraphs if p.text.startswith("где"))

    assert formula_para is not paragraphs[-1] or len(paragraphs) >= 2
    assert explanation_para.paragraph_format.alignment == WD_ALIGN_PARAGRAPH.LEFT
    assert_close(explanation_para.paragraph_format.first_line_indent, Cm(0))
    assert explanation_para.paragraph_format.line_spacing == 1.5
    assert "масса, кг" in explanation_para.text


def test_docx_renderer_resets_formula_counter_per_document(tmp_path):
    renderer = DocxRenderer(config_class=SIBFUConfig, base_dir=tmp_path)

    first_path = tmp_path / "first.docx"
    second_path = tmp_path / "second.docx"
    profile = _common_profile()
    renderer.render_to_file(
        Document(blocks=(FormulaNode(content="A"),)), profile, str(first_path)
    )
    renderer.render_to_file(
        Document(blocks=(FormulaNode(content="B"),)), profile, str(second_path)
    )

    second = DocxDocument(str(second_path))
    formula_para = next(p for p in second.paragraphs if p.text)
    assert formula_para.text == "B\t(1)"


def test_docx_renderer_renders_bibliography_entries_with_paragraph_indent(tmp_path):
    renderer = DocxRenderer(config_class=SIBFUConfig, base_dir=tmp_path)
    ast = Document(
        blocks=(
            StructuralSectionNode(
                section_type=StructuralSectionType.SOURCES,
                title="СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ",
            ),
            BibliographyEntryNode(
                number=1,
                text="Иванов И.И. Основы программирования. — М.: Наука, 2023.",
            ),
            BibliographyEntryNode(
                number=2,
                text="Петров П.П. Алгоритмы и структуры данных. — СПб.: БХВ, 2022.",
            ),
        )
    )
    output_path = tmp_path / "sources.docx"

    renderer.render_to_file(ast, _common_profile(), str(output_path))

    doc = DocxDocument(str(output_path))
    entries = [p for p in doc.paragraphs if p.text.startswith(("1 ", "2 "))]
    assert [p.text for p in entries] == [
        "1 Иванов И.И. Основы программирования. — М.: Наука, 2023.",
        "2 Петров П.П. Алгоритмы и структуры данных. — СПб.: БХВ, 2022.",
    ]
    for entry in entries:
        assert entry.paragraph_format.alignment == WD_ALIGN_PARAGRAPH.JUSTIFY
        assert_close(entry.paragraph_format.first_line_indent, Cm(1.25))
        assert entry.paragraph_format.line_spacing == 1.5


def test_docx_renderer_renders_appendix_on_new_page_with_centered_heading(tmp_path):
    renderer = DocxRenderer(config_class=SIBFUConfig, base_dir=tmp_path)
    ast = Document(
        blocks=(
            ParagraphNode(runs=(TextRun("Перед приложением."),)),
            AppendixNode(
                title="ПРИЛОЖЕНИЕ А",
                letter="А",
                appendix_type="справочное",
                subtitle="Исходные данные",
            ),
        )
    )
    output_path = tmp_path / "appendix.docx"

    renderer.render_to_file(ast, _common_profile(), str(output_path))

    doc = DocxDocument(str(output_path))
    texts = [p.text for p in doc.paragraphs]
    heading_index = texts.index("ПРИЛОЖЕНИЕ А")
    heading = doc.paragraphs[heading_index]
    type_para = doc.paragraphs[heading_index + 1]
    subtitle_index = texts.index("Исходные данные")
    subtitle = doc.paragraphs[subtitle_index]

    assert 'w:type="page"' in doc.paragraphs[heading_index - 1]._p.xml
    assert heading.paragraph_format.alignment == WD_ALIGN_PARAGRAPH.CENTER
    assert_close(heading.paragraph_format.first_line_indent, Cm(0))
    assert heading.runs[0].bold is True
    assert "Heading 1" in heading.style.name
    assert type_para.text == "(справочное)"
    assert subtitle.paragraph_format.alignment == WD_ALIGN_PARAGRAPH.CENTER


def test_docx_renderer_renders_appendix_without_optional_fields(tmp_path):
    renderer = DocxRenderer(config_class=SIBFUConfig, base_dir=tmp_path)
    ast = Document(
        blocks=(
            AppendixNode(title="ПРИЛОЖЕНИЕ Б", letter="Б"),
        )
    )
    output_path = tmp_path / "appendix_min.docx"

    renderer.render_to_file(ast, _common_profile(), str(output_path))

    doc = DocxDocument(str(output_path))
    texts = [p.text for p in doc.paragraphs]
    assert "ПРИЛОЖЕНИЕ Б" in texts
    assert not any(text.startswith("(") for text in texts if text)


def test_docx_renderer_renders_toc_field_with_placeholder_and_heading_styles(tmp_path):
    renderer = DocxRenderer(config_class=SIBFUConfig, base_dir=tmp_path)
    ast = Document(
        blocks=(
            TableOfContentsNode(),
            HeadingNode(level=HeadingLevel.H1, text="Раздел", number="auto"),
        )
    )
    output_path = tmp_path / "toc.docx"

    renderer.render_to_file(ast, _common_profile(), str(output_path))

    doc = DocxDocument(str(output_path))
    body_xml = doc.element.body.xml
    assert " TOC " in body_xml
    assert 'w:fldCharType="begin"' in body_xml
    assert 'w:fldCharType="end"' in body_xml
    texts = [p.text for p in doc.paragraphs]
    assert "СОДЕРЖАНИЕ" in texts
    assert any("Обновите оглавление" in text for text in texts)


def test_docx_renderer_treats_contents_structural_section_as_toc(tmp_path):
    renderer = DocxRenderer(config_class=SIBFUConfig, base_dir=tmp_path)
    ast = Document(
        blocks=(
            StructuralSectionNode(
                section_type=StructuralSectionType.CONTENTS,
                title="Содержание",
            ),
        )
    )
    output_path = tmp_path / "contents.docx"

    renderer.render_to_file(ast, _common_profile(), str(output_path))

    doc = DocxDocument(str(output_path))
    body_xml = doc.element.body.xml
    assert " TOC " in body_xml
    assert 'w:fldCharType="begin"' in body_xml


def test_docx_renderer_attaches_word_heading_style_to_h1_for_toc(tmp_path):
    renderer = DocxRenderer(config_class=SIBFUConfig, base_dir=tmp_path)
    ast = Document(
        blocks=(
            HeadingNode(level=HeadingLevel.H1, text="Раздел", number="auto"),
        )
    )
    output_path = tmp_path / "h1_style.docx"

    renderer.render_to_file(ast, _common_profile(), str(output_path))

    doc = DocxDocument(str(output_path))
    heading = next(p for p in doc.paragraphs if p.text.startswith("1 "))
    assert "Heading 1" in heading.style.name


def test_docx_renderer_covers_image_and_table_edge_paths(tmp_path, monkeypatch, caplog):
    renderer = DocxRenderer(config_class=SIBFUConfig, base_dir=tmp_path)
    renderer._initialize_document()

    renderer._set_paragraph_format(renderer.doc.add_paragraph("x"), "unknown")
    assert "Unknown style_type" in caplog.text

    renderer._style_map["minimal"] = {}
    renderer._set_paragraph_format(renderer.doc.add_paragraph(), "minimal")

    renderer.doc = None
    renderer._set_paragraph_format(DocxDocument().add_paragraph("heading"), "h1")
    renderer._initialize_document()

    renderer._create_table([], caption="ignored")
    assert renderer._format_table_caption("", 1) is None
    assert renderer._format_table_caption("   ", 1) is None

    renderer._insert_image(None, None)
    renderer._insert_image(None, "Caption only")
    renderer._insert_image("missing.png", "Missing caption")

    monkeypatch.setattr(renderer_module, "insert_image", lambda **_kwargs: True)
    image_path = tmp_path / "images" / "image.png"
    image_path.parent.mkdir()
    Image.new("RGB", (2, 2)).save(str(image_path))
    renderer._insert_image("image.png", None)

    monkeypatch.setattr(renderer_module, "insert_image", lambda **_kwargs: False)
    renderer._insert_image("image.png", None)

    def exploding_insert_image(**_kwargs):
        raise RuntimeError("boom")

    monkeypatch.setattr(renderer_module, "insert_image", exploding_insert_image)
    renderer._insert_image("image.png", None)

    texts = [paragraph.text for paragraph in renderer.doc.paragraphs]
    assert "Caption only" in texts
    assert "[Изображение не найдено: missing.png]" in texts
    assert "[Ошибка: image.png]" in texts


def test_docx_renderer_table_and_config_false_branches(tmp_path):
    class NoRepeatConfig(SIBFUConfig):
        TABLE = {**SIBFUConfig.TABLE, "header_repeat_on_pages": False}
        STRUCTURAL_SECTION = {
            **SIBFUConfig.STRUCTURAL_SECTION,
            "page_break_before": False,
            "uppercase": False,
        }

    renderer = DocxRenderer(config_class=NoRepeatConfig, base_dir=tmp_path)
    renderer._initialize_document()
    renderer._create_table([["H"], ["A", "ignored extra"]])
    renderer._create_table([[""]])
    assert "tblHeader" not in renderer.doc.tables[0].rows[0]._tr.xml

    renderer._render_structural_section(
        StructuralSectionNode(
            section_type=StructuralSectionType.INTRODUCTION,
            title="Mixed Case",
        )
    )
    assert any(paragraph.text == "Mixed Case" for paragraph in renderer.doc.paragraphs)


def test_docx_renderer_covers_remaining_ast_blocks_and_title_options(tmp_path):
    renderer = DocxRenderer(config_class=SIBFUConfig, base_dir=tmp_path)
    ast = Document(
        metadata={
            "ministry": "Ministry",
            "university": "University",
            "institute": "Institute",
            "department": "Department",
            "title": "Title",
            "subject": "Subject",
            "student": "Student",
            "group": "Group",
            "supervisor": "Supervisor",
            "supervisor_title": "Professor",
            "city": "City",
            "year": "2026",
        },
        blocks=(
            TitlePageNode(profile="common"),
            MetadataNode(key="ignored", value="ignored"),
            TableCaptionNode(text="Standalone table caption"),
            FigureNode(src=None, caption="Standalone figure caption"),
            PageBreakNode(),
            RawBlockNode(text="Raw text"),
            ReferenceNode(target="fig:overview"),
            FormulaNode(content="x = y", number="A.1"),
            AppendixNode(
                title="Appendix",
                letter="A",
                blocks=(ParagraphNode(runs=(TextRun("Nested appendix body"),)),),
            ),
        ),
    )
    output_path = tmp_path / "remaining.docx"

    renderer.render_to_file(ast, _common_profile(), str(output_path))

    doc = DocxDocument(str(output_path))
    texts = [paragraph.text for paragraph in doc.paragraphs if paragraph.text]
    assert "UNIVERSITY" in texts
    assert "Institute" in texts
    assert "Department" in texts
    assert "SUBJECT" in texts
    assert "Title" in texts
    assert "Руководитель, Professor ____________ Supervisor" in texts
    assert "Студент Group ____________ Student" in texts
    assert "City 2026" in texts
    assert "Standalone table caption" in texts
    assert "Standalone figure caption" in texts
    assert "Raw text" in texts
    assert "[fig:overview]" in texts
    assert "x = y\t(A.1)" in texts
    assert "Nested appendix body" in texts


def test_docx_renderer_title_page_skip_and_helper_edges(tmp_path):
    renderer = DocxRenderer(config_class=SIBFUConfig, base_dir=tmp_path)
    renderer._initialize_document(template_mode="preserve-prefix")
    renderer._render_title_page({"title": "Skipped"})
    assert "Skipped" not in [paragraph.text for paragraph in renderer.doc.paragraphs]

    renderer._initialize_document()
    renderer._render_title_page({"title": "Once"})
    renderer._render_title_page({"title": "Twice"})
    assert [paragraph.text for paragraph in renderer.doc.paragraphs].count("Twice") == 0

    renderer._initialize_document()
    renderer._render_title_page(
        {
            "institute": "",
            "department": "",
            "title": "",
            "subject": "",
            "student": "Student",
            "group": "",
            "supervisor": "Supervisor",
            "supervisor_title": "",
            "city": "",
            "year": "",
        }
    )
    texts = [paragraph.text for paragraph in renderer.doc.paragraphs]
    assert "Руководитель ____________ Supervisor" in texts
    assert "Студент ____________ Student" in texts

    paragraph = renderer.doc.add_paragraph()
    renderer._apply_word_heading_style(paragraph, "Missing Word Style")
    renderer.doc = None
    renderer._apply_word_heading_style(paragraph, "Heading 1")

    renderer = DocxRenderer(config_class=SIBFUConfig, base_dir=tmp_path)
    renderer._initialize_document()
    formula_para = renderer.doc.add_paragraph()
    renderer._set_formula_number_tab(formula_para)
    renderer._set_formula_number_tab(formula_para)
    assert formula_para._p.xml.count("w:tabs") == 2

    assert renderer._resolve_image_path(tmp_path / "image.png") == tmp_path / "image.png"
    assert renderer._resolve_template_path(tmp_path / "template.docx") == tmp_path / "template.docx"
    assert renderer._resolve_template_path("missing.docx") == tmp_path / "templates" / "missing.docx"

    renderer._load_template("missing.docx")
    assert renderer.doc is not None

    renderer._initialize_document()
    renderer._render_appendix(AppendixNode(title="Data", letter="А"))
    assert "ПРИЛОЖЕНИЕ А" in [paragraph.text for paragraph in renderer.doc.paragraphs]

    renderer.render_to_file(
        Document(blocks=(ListNode(list_type=ListType.BULLET, items=()),)),
        _common_profile(),
        str(tmp_path / "empty-list.docx"),
    )

    renderer.render_to_file(
        Document(blocks=(MetadataNode(key="only", value="metadata"),)),
        _common_profile(),
        str(tmp_path / "metadata-only.docx"),
    )
    renderer._initialize_document()
    renderer._render_from_ast(Document(blocks=(MetadataNode(key="direct", value="metadata"),)))

    template = tmp_path / "template.docx"
    DocxDocument().save(str(template))
    assert renderer._resolve_template_path("template.docx") == template

    assert renderer._list_item_text(ListType.LETTERED, 50, "item", True) == "51) item."
    with pytest.raises(ValueError, match="Unsupported list type"):
        renderer._list_item_text(object(), 0, "item", True)

    assert _normalize_list_item_punctuation("", list_type=ListType.BULLET, is_last=False) == ""
    assert (
        _normalize_list_item_punctuation("done.", list_type=ListType.BULLET, is_last=False)
        == "done."
    )

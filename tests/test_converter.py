import pytest
import shutil
import uuid
from types import SimpleNamespace
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

from sfu_converter.domain.ast_nodes import Document as AstDocument, ParagraphNode, SourceSpan, TextRun
from sfu_converter.domain.diagnostics import Diagnostic, Severity
from sfu_converter.converter import TextToDocxConverter
import sfu_converter.converter as converter_module


@pytest.fixture
def tmp_path():
    """Создает временную директорию внутри репозитория, обходя проблемы прав доступа pytest."""
    root = Path(__file__).resolve().parent.parent / '.tmp' / 'test_converter'
    root.mkdir(parents=True, exist_ok=True)
    path = root / uuid.uuid4().hex
    path.mkdir()
    try:
        yield path
    finally:
        shutil.rmtree(path, ignore_errors=True)


class TestTextToDocxConverter:
    """Тесты для конвертера TXT в DOCX"""
    
    @pytest.fixture
    def converter(self, tmp_path):
        """Создает конвертер с временной директорией"""
        return TextToDocxConverter(base_dir=tmp_path)
    
    @pytest.fixture
    def sample_txt_file(self, tmp_path):
        """Создает тестовый TXT файл с произвольным путем"""
        input_dir = tmp_path / 'inputs'
        input_dir.mkdir()
        txt_file = input_dir / 'test.txt'
        txt_file.write_text("[H1] Тестовый заголовок\n\nОбычный текст", encoding='utf-8')
        return txt_file
    
    def test_converter_initialization(self, tmp_path):
        """Тест: Инициализация конвертера с базовой директорией"""
        converter = TextToDocxConverter(base_dir=tmp_path)
        assert converter.base_dir == tmp_path
        assert converter.doc is None
    
    def test_set_run_style_applies_font(self, converter):
        """Тест: Применение шрифта Times New Roman к тексту"""
        converter.doc = Document()
        para = converter.doc.add_paragraph("Тест")
        run = para.runs[0]
        
        converter._set_run_style(run, bold=True)
        
        assert run.font.name == 'Times New Roman'
        assert run.font.size == Pt(14)
        assert run.bold is True
    
    def test_set_paragraph_format_normal_style(self, converter):
        """Тест: Форматирование обычного абзаца с отступом 1.25 см"""
        converter.doc = Document()
        para = converter.doc.add_paragraph("Текст")
        
        converter._set_paragraph_format(para, 'normal')
        pf = para.paragraph_format
        
        assert pf.line_spacing == 1.5
        assert pf.alignment == WD_ALIGN_PARAGRAPH.JUSTIFY
        # Используем допустимую погрешность для EMU (1000 единиц = ~0.03 мм)
        assert abs(pf.first_line_indent - Cm(1.25)) < 1000
    
    def test_set_paragraph_format_h1_style(self, converter):
        """Тест: Форматирование заголовка H1 (по центру, жирный)"""
        converter.doc = Document()
        para = converter.doc.add_paragraph("Заголовок")
        
        converter._set_paragraph_format(para, 'h1')
        pf = para.paragraph_format
        
        assert pf.alignment == WD_ALIGN_PARAGRAPH.CENTER
        assert pf.line_spacing == 1.0
        assert abs(pf.first_line_indent - Cm(0)) < 1000
    
    def test_add_empty_paragraph(self, converter):
        """Тест: Добавление пустого абзаца с правильными интервалами"""
        converter.doc = Document()
        
        converter._add_empty_paragraph('empty_after_header')
        
        assert len(converter.doc.paragraphs) == 1
        pf = converter.doc.paragraphs[0].paragraph_format
        assert pf.space_before == Pt(0)
        assert pf.space_after == Pt(0)
    
    def test_parse_table_line_valid(self, converter):
        """Тест: Разбор строки таблицы с корректным форматом"""
        line = "| Ячейка 1 | Ячейка 2 | Ячейка 3 |"
        result = converter._parse_table_line(line)
        
        assert result == ['Ячейка 1', 'Ячейка 2', 'Ячейка 3']
    
    def test_parse_table_line_invalid(self, converter):
        """Тест: Разбор некорректной строки таблицы"""
        line = "Неправильный формат"
        result = converter._parse_table_line(line)
        
        assert result is None
    
    def test_create_table(self, converter, tmp_path):
        """Тест: Создание таблицы 3x3 с подписью"""
        converter.doc = Document()
        
        rows = [
            ['Заголовок 1', 'Заголовок 2', 'Заголовок 3'],
            ['Ячейка 1.1', 'Ячейка 1.2', 'Ячейка 1.3'],
            ['Ячейка 2.1', 'Ячейка 2.2', 'Ячейка 2.3']
        ]
        
        converter._create_table(rows, caption="Таблица 1 - Тест")
        
        assert len(converter.doc.tables) == 1
        table = converter.doc.tables[0]
        assert len(table.rows) == 3
        assert len(table.columns) == 3
    
    def test_convert_file_reads_and_writes_explicit_paths(self, converter, sample_txt_file, tmp_path):
        """Тест: Конвертер принимает явные пути входа и выхода"""
        output_file = tmp_path / 'artifacts' / 'output.docx'

        output_path = converter.convert_file(sample_txt_file, output_file)

        assert output_path == str(output_file)
        assert output_file.exists()

        doc = Document(output_path)
        assert [para.text for para in doc.paragraphs if para.text] == [
            "1 Тестовый заголовок",
            "Обычный текст",
        ]

    def test_convert_file_creates_output_directory(self, converter, sample_txt_file, tmp_path):
        """Тест: Конвертер создает директорию для выходного файла"""
        output_file = tmp_path / 'nested' / 'results' / 'heading.docx'

        converter.convert_file(sample_txt_file, output_file)

        assert output_file.parent.exists()
        assert output_file.exists()

    def test_convert_preserves_menu_compatibility(self, converter, tmp_path):
        """Тест: Старый API продолжает работать с examples/ и results/"""
        examples_dir = tmp_path / 'examples'
        examples_dir.mkdir()
        txt_file = examples_dir / 'heading.txt'
        txt_file.write_text("[H1] Важный заголовок", encoding='utf-8')

        output_path = converter.convert('heading.txt', 'heading.docx')
        doc = Document(output_path)

        found_h1 = False
        for para in doc.paragraphs:
            if para.text == "1 Важный заголовок":
                assert para.paragraph_format.alignment == WD_ALIGN_PARAGRAPH.CENTER
                found_h1 = True

        assert found_h1, "Заголовок H1 не найден"

    def test_convert_file_with_table(self, converter, tmp_path):
        """Тест: Конвертация таблицы через явные пути"""
        input_dir = tmp_path / 'fixtures'
        input_dir.mkdir()
        txt_file = input_dir / 'table.txt'
        txt_content = """[TABLE_START]
[TABLE_CAPTION] Таблица 1
| Кол1 | Кол2 |
| А | Б |
[TABLE_END]"""
        txt_file.write_text(txt_content, encoding='utf-8')

        output_file = tmp_path / 'exports' / 'table.docx'

        converter.convert_file(txt_file, output_file)
        doc = Document(str(output_file))

        assert len(doc.tables) == 1
        table = doc.tables[0]
        assert table.rows[0].cells[0].text == 'Кол1'
        assert table.rows[1].cells[1].text == 'Б'
    
    def test_setup_document_margins(self, converter):
        """Тест: Настройка полей страницы с допустимой погрешностью"""
        converter.doc = Document()
        
        converter._setup_document_margins()
        
        section = converter.doc.sections[0]
        # Используем допустимую погрешность для EMU
        assert abs(section.top_margin - Cm(2)) < 1000
        assert abs(section.bottom_margin - Cm(2)) < 1000
        assert abs(section.left_margin - Cm(3)) < 1000
        assert abs(section.right_margin - Cm(1)) < 1000

    def test_converter_wrapper_methods_and_line_sources(self, converter, tmp_path, caplog):
        converter._initialize_document()
        converter._insert_image(None, "Caption")
        converter._load_template("missing-template.docx")

        assert converter._lines_to_source("already text") == "already text"
        assert converter._lines_to_source(["a\n", "b\n"]) == "a\nb\n"
        assert converter._lines_to_source(["a", "b"]) == "a\nb"

        ast = AstDocument(blocks=(ParagraphNode(runs=(TextRun("From AST"),)),))
        converter._render_lines(ast)
        converter._render_lines(["Plain text"])
        assert any(paragraph.text == "From AST" for paragraph in converter.doc.paragraphs)
        assert any(paragraph.text == "Plain text" for paragraph in converter.doc.paragraphs)

        converter._log_parser_diagnostics(
            [
                Diagnostic(
                    code="ERR",
                    message="error",
                    severity=Severity.ERROR,
                    source=SourceSpan(1, 1),
                ),
                Diagnostic(code="WARN", message="warning", severity=Severity.WARNING),
            ]
        )
        assert "ERR at line 1" in caplog.text
        assert "WARN at line ?" in caplog.text

        assert converter.convert("input.txt") is None

    def test_converter_template_composition_missing_diagnostic_branch(self, converter, tmp_path, monkeypatch):
        input_file = tmp_path / "input.txt"
        input_file.write_text("Body", encoding="utf-8")
        output_file = tmp_path / "output.docx"

        class UseCase:
            def __init__(self, *_args, **_kwargs):
                pass

            def execute(self, *, output_path, **_kwargs):
                doc = Document()
                doc.add_paragraph("Generated")
                doc.save(output_path)
                return []

        class Adapter:
            def __init__(self, *_args, **_kwargs):
                pass

            def load_template(self, _template):
                return object()

            def find_insertion_point(self, *_args, **_kwargs):
                return SimpleNamespace(found=False, diagnostic=None)

            def compose(self, *_args, **_kwargs):
                raise AssertionError("compose should not be called")

        monkeypatch.setattr(converter_module, "ConvertTextToDocx", UseCase)
        monkeypatch.setattr(
            "sfu_converter.infrastructure.template_adapter.DocxTemplateAdapter",
            Adapter,
        )

        assert converter.convert_file(
            input_file,
            output_file,
            template="template.docx",
            template_mode="preserve-prefix",
            insert_after_page=1,
        ) == str(output_file)

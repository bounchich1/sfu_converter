from docx import Document as DocxDocument

from sfu_converter.domain.ast_nodes import (
    AppendixNode,
    Document,
    HeadingLevel,
    HeadingNode,
    PageBreakNode,
    ParagraphNode,
    SheetFormat,
    TextRun,
)
from sfu_converter.domain.diagnostics import DiagnosticCodes, Severity
from sfu_converter.infrastructure.appendix import assign_appendix_letters
from sfu_converter.infrastructure.docx_renderer import DocxRenderer
from sfu_converter.parser import V2Parser
from sfu_converter.registry import get_profile


def _texts(path):
    return [paragraph.text for paragraph in DocxDocument(str(path)).paragraphs if paragraph.text]


def test_v2_parser_preserves_appendix_sheet_and_independent_flags():
    result = V2Parser().parse(
        '[APPENDIX id=app:a letter=А type=reference subtitle="Data" sheet=A3x4 independent=true]'
    )

    appendix = result.document.blocks[0]

    assert isinstance(appendix, AppendixNode)
    assert appendix.letter == "А"
    assert appendix.appendix_type == "reference"
    assert appendix.subtitle == "Data"
    assert appendix.sheet_format is SheetFormat.A3X4
    assert appendix.independent is True


def test_appendix_auto_letter_assigns_next_letter_after_explicit_letter():
    document = Document(
        blocks=(
            AppendixNode(title="ПРИЛОЖЕНИЕ Б", letter="Б"),
            AppendixNode(title="ПРИЛОЖЕНИЕ"),
        )
    )

    assigned, diagnostics = assign_appendix_letters(document)

    assert assigned.blocks[1].letter == "В"
    assert assigned.blocks[1].id == "app:в"
    assert diagnostics[0].code == DiagnosticCodes.APPENDIX_AUTOLETTER_ASSIGNED
    assert diagnostics[0].severity is Severity.INFO
    assert diagnostics[0].data == {"letter": "В", "position": 2}


def test_appendix_auto_letter_skips_excluded_letters():
    document = Document(
        blocks=(
            AppendixNode(title="ПРИЛОЖЕНИЕ Е", letter="Е"),
            AppendixNode(title="ПРИЛОЖЕНИЕ"),
        )
    )

    assigned, _ = assign_appendix_letters(document)

    assert assigned.blocks[1].letter == "Ж"


def test_docx_renderer_numbers_appendix_headings_with_letter_prefix(tmp_path):
    document = Document(
        blocks=(
            AppendixNode(
                title="ПРИЛОЖЕНИЕ А",
                letter="А",
                blocks=(
                    HeadingNode(level=HeadingLevel.H1, text="Раздел", number="auto"),
                    HeadingNode(level=HeadingLevel.H2, text="Подраздел", number="auto"),
                    HeadingNode(level=HeadingLevel.H3, text="Пункт", number="auto"),
                    HeadingNode(level=HeadingLevel.H4, text="Подпункт", number="auto"),
                ),
            ),
        )
    )
    output = tmp_path / "appendix_numbering.docx"

    DocxRenderer(base_dir=tmp_path).render_to_file(document, get_profile("common"), str(output))

    texts = _texts(output)
    assert "А.1 Раздел" in texts
    assert "А.1.1.1.1 Подпункт" in texts


def test_docx_renderer_emits_continuation_and_final_labels_for_page_broken_appendix(tmp_path):
    document = Document(
        blocks=(
            AppendixNode(
                title="ПРИЛОЖЕНИЕ А",
                letter="А",
                blocks=(
                    ParagraphNode(runs=(TextRun("Первая страница"),)),
                    PageBreakNode(),
                    ParagraphNode(runs=(TextRun("Вторая страница"),)),
                    PageBreakNode(),
                    ParagraphNode(runs=(TextRun("Третья страница"),)),
                ),
            ),
        )
    )
    output = tmp_path / "appendix_continuation.docx"

    DocxRenderer(base_dir=tmp_path).render_to_file(document, get_profile("common"), str(output))

    texts = _texts(output)
    assert "Продолжение приложения А" in texts
    assert "Окончание приложения А" in texts


def test_independent_appendix_renders_own_title_page(tmp_path):
    document = Document(
        metadata={
            "title": "Основной документ",
            "student": "Иванов И.И.",
            "group": "КИ20-01",
            "supervisor": "Петров П.П.",
        },
        blocks=(
            AppendixNode(
                title="ПРИЛОЖЕНИЕ А",
                letter="А",
                independent=True,
                blocks=(ParagraphNode(runs=(TextRun("Текст приложения"),)),),
            ),
        ),
    )
    output = tmp_path / "independent_appendix.docx"

    DocxRenderer(base_dir=tmp_path).render_to_file(document, get_profile("coursework"), str(output))

    texts = _texts(output)
    assert any("Основной документ" in text for text in texts)
    assert "ПРИЛОЖЕНИЕ А" in texts

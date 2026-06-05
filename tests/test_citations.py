from io import BytesIO

from docx import Document as DocxDocument

from sfu_converter.domain.ast_nodes import Citation, CitationNode, Document, ParagraphNode, TextRun
from sfu_converter.domain.diagnostics import DiagnosticCodes
from sfu_converter.infrastructure.docx_renderer import DocxRenderer
from sfu_converter.parser.citations import format_citation_node, parse_citation_text
from sfu_converter.parser.v1_parser import V1Parser
from sfu_converter.parser.v2_parser import V2Parser
from sfu_converter.registry import get_profile


def test_parse_supported_source_citation_forms():
    assert parse_citation_text("[20]") == CitationNode((Citation(number=20),))
    assert parse_citation_text("[20, с. 29]") == CitationNode((Citation(number=20, pages=29),))
    assert parse_citation_text("[18, т. 1, с. 75]") == CitationNode((Citation(number=18, volume=1, pages=75),))
    assert parse_citation_text("[59; 67, с. 40-46; 82]") == CitationNode(
        (
            Citation(number=59),
            Citation(number=67, pages=(40, 46)),
            Citation(number=82),
        )
    )


def test_v1_parser_replaces_citation_text_with_structured_node():
    result = V1Parser().parse("См. источник [20, с. 29].", filename="report.txt")

    assert result.diagnostics == []
    (paragraph,) = result.document.blocks
    assert isinstance(paragraph, ParagraphNode)
    assert paragraph.runs == (
        TextRun("См. источник ", source=paragraph.source),
        CitationNode((Citation(number=20, pages=29),), source=paragraph.source),
        TextRun(".", source=paragraph.source),
    )


def test_v2_parser_replaces_citation_text_with_structured_node():
    result = V2Parser().parse("[P] См. источники [59; 67, с. 40-46; 82].")

    assert result.diagnostics == []
    (paragraph,) = result.document.blocks
    citation = paragraph.runs[1]
    assert isinstance(citation, CitationNode)
    assert citation.citations == (
        Citation(number=59),
        Citation(number=67, pages=(40, 46)),
        Citation(number=82),
    )


def test_malformed_citation_reports_diagnostic_and_keeps_literal_text():
    result = V1Parser().parse("Bad citation [20, p. 29].")

    assert result.document.blocks[0].runs[0].text == "Bad citation [20, p. 29]."
    assert result.diagnostics[0].code == DiagnosticCodes.CITATION_MALFORMED


def test_reversed_page_range_and_duplicate_number_report_diagnostics():
    reversed_range = V1Parser().parse("[10, с. 50-30]")
    duplicated = V1Parser().parse("[7; 7, с. 9]")

    assert reversed_range.diagnostics[0].code == DiagnosticCodes.CITATION_PAGE_RANGE_REVERSED
    assert duplicated.diagnostics[0].code == DiagnosticCodes.CITATION_NUMBER_DUPLICATED


def test_citation_round_trip_normalizes_page_range_dash(tmp_path):
    citation = parse_citation_text("[59; 67, с. 40-46; 82]")

    assert format_citation_node(citation) == "[59; 67, с. 40–46; 82]"

    renderer = DocxRenderer(base_dir=tmp_path)
    paragraph = ParagraphNode((TextRun("Источник "), citation))
    docx_bytes = renderer.render(
        document=Document((paragraph,)),
        profile=get_profile("common"),
    )
    doc = DocxDocument(BytesIO(docx_bytes))
    assert doc.paragraphs[0].text == "Источник [59; 67, с. 40–46; 82]"

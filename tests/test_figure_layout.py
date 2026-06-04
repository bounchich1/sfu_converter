from docx import Document as DocxDocument
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

from sfu_converter.domain.ast_nodes import (
    AppendixNode,
    Document,
    FigureNode,
    ParagraphNode,
    ReferenceNode,
    TextRun,
)
from sfu_converter.domain.diagnostics import DiagnosticCodes, Severity
from sfu_converter.infrastructure import docx_styles
from sfu_converter.infrastructure.docx_renderer import DocxRenderer
from sfu_converter.registry import get_profile


def _render(tmp_path, document: Document):
    output_path = tmp_path / "figures.docx"
    diagnostics = DocxRenderer(base_dir=tmp_path).render_to_file(
        document,
        get_profile("common"),
        str(output_path),
    )
    return DocxDocument(str(output_path)), diagnostics


def test_figure_inside_appendix_uses_appendix_number(tmp_path):
    doc, _ = _render(
        tmp_path,
        Document(
            blocks=(
                AppendixNode(
                    title="ПРИЛОЖЕНИЕ",
                    letter="А",
                    blocks=(FigureNode(src=None, caption="Схема"),),
                ),
            )
        ),
    )

    assert "Рисунок А.1 — Схема" in [paragraph.text for paragraph in doc.paragraphs]


def test_multisheet_figure_uses_sheet_caption(tmp_path):
    doc, _ = _render(
        tmp_path,
        Document(
            blocks=(
                FigureNode(src=None, caption="Предыдущий 1"),
                FigureNode(src=None, caption="Предыдущий 2"),
                FigureNode(src=None, caption="Предыдущий 3"),
                FigureNode(src=None, caption="Предыдущий 4"),
                FigureNode(src=None, caption="Схема", sheet=2, total_sheets=3),
            )
        ),
    )

    assert "Рисунок 5, лист 2" in [paragraph.text for paragraph in doc.paragraphs]


def test_figure_explanatory_data_renders_centered_12pt_above_caption(tmp_path):
    doc, _ = _render(
        tmp_path,
        Document(
            blocks=(
                FigureNode(
                    src=None,
                    caption="Схема",
                    explanatory_data=("1 — модуль ввода",),
                ),
            )
        ),
    )

    texts = [p.text for p in doc.paragraphs]
    explanatory_index = texts.index("1 — модуль ввода")
    caption_index = next(index for index, text in enumerate(texts) if text.startswith("Рисунок"))
    explanatory = doc.paragraphs[explanatory_index]
    assert explanatory_index < caption_index
    assert explanatory.style.name == docx_styles.FIGURE_EXPLANATORY
    assert explanatory.paragraph_format.alignment == WD_ALIGN_PARAGRAPH.CENTER
    assert explanatory.runs[0].font.size == Pt(12)


def test_unreferenced_figure_emits_warning_diagnostic(tmp_path):
    _, diagnostics = _render(
        tmp_path,
        Document(blocks=(FigureNode(src=None, caption="Схема", id="fig:missing"),)),
    )

    diagnostic = next(
        item for item in diagnostics if item.code == DiagnosticCodes.FIGURE_NEVER_REFERENCED
    )
    assert diagnostic.severity is Severity.WARNING
    assert diagnostic.rule_id == "common.reference.figure_table_formula"


def test_figure_before_late_first_reference_emits_placement_info(tmp_path):
    _, diagnostics = _render(
        tmp_path,
        Document(
            blocks=(
                FigureNode(src=None, caption="Схема", id="fig:late"),
                ParagraphNode(runs=(TextRun("1"),)),
                ParagraphNode(runs=(TextRun("2"),)),
                ParagraphNode(runs=(TextRun("3"),)),
                ParagraphNode(runs=(TextRun("4"),)),
                ReferenceNode(target="fig:late"),
            )
        ),
    )

    diagnostic = next(
        item
        for item in diagnostics
        if item.code == DiagnosticCodes.FIGURE_PLACEMENT_NEXT_PAGE
    )
    assert diagnostic.severity is Severity.INFO
    assert diagnostic.rule_id == "common.figure.placement_after_reference"

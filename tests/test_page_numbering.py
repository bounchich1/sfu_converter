"""Tests for page numbering (Task 12).

Covers:
- Title page counted but number NOT printed (titlePg flag).
- Body pages display Arabic numerals starting at 2, bottom center, TNR 14pt.
- Validator detects missing titlePg flag.
- Validator detects wrong footer font.
- page_numbering.configure() wires sections correctly.
"""
from docx import Document as DocxDocument
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

from sfu_converter.config import SIBFUConfig
from sfu_converter.domain.ast_nodes import (
    Document,
    HeadingLevel,
    HeadingNode,
    ParagraphNode,
    TextRun,
    TitlePageNode,
)
from sfu_converter.domain.diagnostics import DiagnosticCodes
from sfu_converter.domain.formatting import FormattingProfile
from sfu_converter.infrastructure.docx_renderer import DocxRenderer
from sfu_converter.infrastructure.docx_validator import DocxValidator
from sfu_converter.infrastructure.page_numbering import (
    Location,
    PageNumberingSection,
    configure,
)
from sfu_converter.registry import get_profile


def _common_profile():
    return FormattingProfile(
        name="common",
        display_name="Common",
        source_docs=("standard",),
    )


# ---------------------------------------------------------------------------
# page_numbering module unit tests
# ---------------------------------------------------------------------------

def test_configure_sets_title_pg_on_first_section():
    """First section (title page) must have different_first_page_header_footer."""
    doc = DocxDocument()
    sections = [
        PageNumberingSection(
            start_at=1,
            hide_first_page=True,
            suppress_in_section=False,
            location=Location.FOOTER_CENTER,
        ),
    ]
    configure(doc, sections, SIBFUConfig)

    section = doc.sections[0]
    assert section.different_first_page_header_footer is True


def test_configure_adds_page_field_to_footer():
    """Footer must contain a PAGE field instruction."""
    doc = DocxDocument()
    sections = [
        PageNumberingSection(
            start_at=1,
            hide_first_page=True,
            suppress_in_section=False,
            location=Location.FOOTER_CENTER,
        ),
    ]
    configure(doc, sections, SIBFUConfig)

    footer = doc.sections[0].footer
    assert " PAGE " in footer._element.xml


def test_configure_footer_uses_tnr_14pt_centered():
    """Footer page number paragraph: TNR, 14pt, centered, no indent."""
    doc = DocxDocument()
    sections = [
        PageNumberingSection(
            start_at=1,
            hide_first_page=False,
            suppress_in_section=False,
            location=Location.FOOTER_CENTER,
        ),
    ]
    configure(doc, sections, SIBFUConfig)

    paragraph = doc.sections[0].footer.paragraphs[0]
    assert paragraph.paragraph_format.alignment == WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.runs[0]
    assert run.font.name == "Times New Roman"
    assert run.font.size == Pt(14)


def test_configure_first_page_footer_empty_when_hidden():
    """When hide_first_page=True, first-page footer has no PAGE field."""
    doc = DocxDocument()
    sections = [
        PageNumberingSection(
            start_at=1,
            hide_first_page=True,
            suppress_in_section=False,
            location=Location.FOOTER_CENTER,
        ),
    ]
    configure(doc, sections, SIBFUConfig)

    first_footer = doc.sections[0].first_page_footer
    assert " PAGE " not in first_footer._element.xml


def test_configure_suppressed_section_has_no_page_field():
    """suppress_in_section=True means no PAGE field in footer at all."""
    doc = DocxDocument()
    # Need a second section in the DOCX
    body = doc.element.body
    sectPr = body.makeelement(qn("w:sectPr"), {})
    body.append(sectPr)

    sections = [
        PageNumberingSection(
            start_at=1,
            hide_first_page=True,
            suppress_in_section=False,
            location=Location.FOOTER_CENTER,
        ),
        PageNumberingSection(
            start_at=None,
            hide_first_page=False,
            suppress_in_section=True,
            location=Location.FOOTER_CENTER,
        ),
    ]
    configure(doc, sections, SIBFUConfig)

    assert " PAGE " not in doc.sections[1].footer._element.xml


# ---------------------------------------------------------------------------
# Renderer integration — page numbering after rendering
# ---------------------------------------------------------------------------

def test_rendered_document_hides_page_number_on_title_page(tmp_path):
    """Title page is counted (start_at=1) but page number not printed."""
    renderer = DocxRenderer(config_class=SIBFUConfig, base_dir=tmp_path)
    ast = Document(
        metadata={"title": "Test", "student": "Student", "group": "G1",
                  "supervisor": "Super", "city": "City", "year": "2026"},
        blocks=(
            TitlePageNode(profile="common"),
            HeadingNode(level=HeadingLevel.H1, text="Intro", number="auto"),
            ParagraphNode(runs=(TextRun("Body text"),)),
        ),
    )
    output_path = tmp_path / "page_numbered.docx"
    renderer.render_to_file(ast, _common_profile(), str(output_path))

    doc = DocxDocument(str(output_path))
    first_section = doc.sections[0]
    assert first_section.different_first_page_header_footer is True
    first_footer_xml = first_section.first_page_footer._element.xml
    assert " PAGE " not in first_footer_xml


def test_rendered_document_body_pages_have_page_field(tmp_path):
    """Body section footer must contain a PAGE field."""
    renderer = DocxRenderer(config_class=SIBFUConfig, base_dir=tmp_path)
    ast = Document(
        blocks=(
            ParagraphNode(runs=(TextRun("Body"),)),
        ),
    )
    output_path = tmp_path / "page_body.docx"
    renderer.render_to_file(ast, _common_profile(), str(output_path))

    doc = DocxDocument(str(output_path))
    section = doc.sections[0]
    assert " PAGE " in section.footer._element.xml


# ---------------------------------------------------------------------------
# Validator — common.page.numbering checks
# ---------------------------------------------------------------------------

def test_validator_passes_when_page_numbering_correct(tmp_path):
    """No common.page.numbering diagnostic on well-formed doc."""
    renderer = DocxRenderer(config_class=SIBFUConfig, base_dir=tmp_path)
    ast = Document(
        blocks=(ParagraphNode(runs=(TextRun("Body"),)),),
    )
    output_path = tmp_path / "valid_numbering.docx"
    renderer.render_to_file(ast, _common_profile(), str(output_path))

    diagnostics = DocxValidator(get_profile("common")).validate_file(str(output_path))
    page_diags = [
        d for d in diagnostics
        if d.rule_id == "common.page.numbering"
    ]
    assert page_diags == []


def test_validator_reports_missing_title_pg_flag(tmp_path):
    """Validator catches when titlePg flag is missing on first section."""
    doc = DocxDocument()
    section = doc.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(3)
    section.right_margin = Cm(1)
    section.different_first_page_header_footer = False

    footer = section.footer
    footer.is_linked_to_previous = False
    paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    paragraph.clear()
    paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.first_line_indent = Cm(0)
    run = paragraph.add_run()
    run.font.name = "Times New Roman"
    run.font.size = Pt(14)
    from docx.oxml import OxmlElement
    field_begin = OxmlElement("w:fldChar")
    field_begin.set(qn("w:fldCharType"), "begin")
    run._element.append(field_begin)
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    run._element.append(instruction)
    field_end = OxmlElement("w:fldChar")
    field_end.set(qn("w:fldCharType"), "end")
    run._element.append(field_end)

    path = tmp_path / "no_title_pg.docx"
    doc.save(str(path))

    diagnostics = DocxValidator(get_profile("common")).validate_file(str(path))
    assert any(
        d.rule_id == "common.page.numbering"
        and "titlepg" in d.message.lower()
        for d in diagnostics
    )


def test_validator_reports_wrong_footer_font(tmp_path):
    """Validator catches non-TNR font in page number footer."""
    doc = DocxDocument()
    section = doc.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(3)
    section.right_margin = Cm(1)
    section.different_first_page_header_footer = True

    footer = section.footer
    footer.is_linked_to_previous = False
    paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    paragraph.clear()
    paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.first_line_indent = Cm(0)
    run = paragraph.add_run()
    run.font.name = "Arial"
    run.font.size = Pt(14)
    from docx.oxml import OxmlElement
    field_begin = OxmlElement("w:fldChar")
    field_begin.set(qn("w:fldCharType"), "begin")
    run._element.append(field_begin)
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    run._element.append(instruction)
    field_end = OxmlElement("w:fldChar")
    field_end.set(qn("w:fldCharType"), "end")
    run._element.append(field_end)

    first_footer = section.first_page_footer
    first_footer.is_linked_to_previous = False
    fp = first_footer.paragraphs[0] if first_footer.paragraphs else first_footer.add_paragraph()
    fp.clear()

    path = tmp_path / "wrong_font.docx"
    doc.save(str(path))

    diagnostics = DocxValidator(get_profile("common")).validate_file(str(path))
    assert any(
        d.rule_id == "common.page.numbering"
        and "font" in d.message.lower()
        for d in diagnostics
    )


def test_validator_reports_wrong_footer_alignment(tmp_path):
    """Validator catches a PAGE field that is not bottom-centred."""
    doc = DocxDocument()
    section = doc.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(3)
    section.right_margin = Cm(1)
    section.different_first_page_header_footer = True

    footer = section.footer
    footer.is_linked_to_previous = False
    paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    paragraph.clear()
    paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    paragraph.paragraph_format.first_line_indent = Cm(0)
    run = paragraph.add_run()
    run.font.name = "Times New Roman"
    run.font.size = Pt(14)
    from docx.oxml import OxmlElement
    field_begin = OxmlElement("w:fldChar")
    field_begin.set(qn("w:fldCharType"), "begin")
    run._element.append(field_begin)
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    run._element.append(instruction)
    field_end = OxmlElement("w:fldChar")
    field_end.set(qn("w:fldCharType"), "end")
    run._element.append(field_end)

    first_footer = section.first_page_footer
    first_footer.is_linked_to_previous = False
    fp = first_footer.paragraphs[0] if first_footer.paragraphs else first_footer.add_paragraph()
    fp.clear()

    path = tmp_path / "wrong_alignment.docx"
    doc.save(str(path))

    diagnostics = DocxValidator(get_profile("common")).validate_file(str(path))
    assert any(
        d.rule_id == "common.page.numbering"
        and "alignment" in d.message.lower()
        for d in diagnostics
    )


def test_validator_reports_footer_first_line_indent(tmp_path):
    """Validator catches a PAGE field paragraph with first-line indent."""
    doc = DocxDocument()
    section = doc.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(3)
    section.right_margin = Cm(1)
    section.different_first_page_header_footer = True

    footer = section.footer
    footer.is_linked_to_previous = False
    paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    paragraph.clear()
    paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.first_line_indent = Cm(1)
    run = paragraph.add_run()
    run.font.name = "Times New Roman"
    run.font.size = Pt(14)
    from docx.oxml import OxmlElement
    field_begin = OxmlElement("w:fldChar")
    field_begin.set(qn("w:fldCharType"), "begin")
    run._element.append(field_begin)
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    run._element.append(instruction)
    field_end = OxmlElement("w:fldChar")
    field_end.set(qn("w:fldCharType"), "end")
    run._element.append(field_end)

    first_footer = section.first_page_footer
    first_footer.is_linked_to_previous = False
    fp = first_footer.paragraphs[0] if first_footer.paragraphs else first_footer.add_paragraph()
    fp.clear()

    path = tmp_path / "wrong_indent.docx"
    doc.save(str(path))

    diagnostics = DocxValidator(get_profile("common")).validate_file(str(path))
    assert any(
        d.rule_id == "common.page.numbering"
        and "indent" in d.message.lower()
        for d in diagnostics
    )


def test_validator_page_numbering_no_longer_unsupported(tmp_path):
    """common.page.numbering should NOT appear in unsupported list."""
    doc = DocxDocument()
    section = doc.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(3)
    section.right_margin = Cm(1)
    section.different_first_page_header_footer = True

    footer = section.footer
    footer.is_linked_to_previous = False
    paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    paragraph.clear()
    paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.first_line_indent = Cm(0)
    run = paragraph.add_run()
    run.font.name = "Times New Roman"
    run.font.size = Pt(14)
    from docx.oxml import OxmlElement
    field_begin = OxmlElement("w:fldChar")
    field_begin.set(qn("w:fldCharType"), "begin")
    run._element.append(field_begin)
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    run._element.append(instruction)
    field_end = OxmlElement("w:fldChar")
    field_end.set(qn("w:fldCharType"), "end")
    run._element.append(field_end)

    first_footer = section.first_page_footer
    first_footer.is_linked_to_previous = False
    fp = first_footer.paragraphs[0] if first_footer.paragraphs else first_footer.add_paragraph()
    fp.clear()

    path = tmp_path / "correct.docx"
    doc.save(str(path))

    diagnostics = DocxValidator(get_profile("common")).validate_file(str(path))
    unsupported_ids = {
        d.rule_id for d in diagnostics
        if d.code == DiagnosticCodes.FORMAT_RULE_NOT_SUPPORTED
    }
    assert "common.page.numbering" not in unsupported_ids

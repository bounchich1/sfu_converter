from __future__ import annotations

import importlib

from docx import Document as DocxDocument
import pytest

from sfu_converter.domain import ast_nodes
from sfu_converter.domain.ast_nodes import Document
from sfu_converter.domain.diagnostics import DiagnosticCodes
from sfu_converter.infrastructure.docx_renderer import DocxRenderer
from sfu_converter.parser import V2Parser
from sfu_converter.registry import get_profile


def _bibliography_module():
    try:
        return importlib.import_module("sfu_converter.infrastructure.bibliography")
    except ImportError as exc:
        pytest.fail(f"bibliography module missing: {exc}")


def _source_record_type():
    record_type = getattr(ast_nodes, "SourceRecordType", None)
    assert record_type is not None, "SourceRecordType is missing"
    return record_type


def _source_record_node():
    record_node = getattr(ast_nodes, "SourceRecordNode", None)
    assert record_node is not None, "SourceRecordNode is missing"
    return record_node


def _book(number: int, authors: str, title: str, *, year: int = 2023, language: str = "ru"):
    return _source_record_node()(
        number=number,
        record_type=_source_record_type().BOOK_TWO_AUTHORS,
        fields={
            "authors": authors,
            "title": title,
            "city": "Красноярск",
            "publisher": "СФУ",
            "year": str(year),
            "pages": "320",
        },
        language=language,
    )


def test_book_two_authors_record_formats_as_gost_string():
    bibliography = _bibliography_module()
    record = _book(1, "Иванов И. И., Петров П. П.", "Анализ данных")

    assert (
        bibliography.format_record(record)
        == "Иванов, И. И. Анализ данных / И. И. Иванов, П. П. Петров. — Красноярск: СФУ, 2023. — 320 с."
    )


def test_v2_parser_parses_typed_source_record_fields():
    result = V2Parser().parse(
        "\n".join(
            [
                "[SOURCE number=1 type=book_two_authors lang=ru]",
                'authors="Иванов И. И., Петров П. П."',
                'title="Анализ данных"',
                'city="Красноярск" publisher="СФУ" year=2023 pages=320',
                "[/SOURCE]",
            ]
        )
    )

    assert result.diagnostics == []
    (record,) = result.document.blocks
    assert isinstance(record, _source_record_node())
    assert record.number == 1
    assert record.record_type is _source_record_type().BOOK_TWO_AUTHORS
    assert record.language == "ru"
    assert record.fields["authors"] == "Иванов И. И., Петров П. П."
    assert record.fields["pages"] == "320"


def test_patent_record_missing_required_field_produces_diagnostic():
    bibliography = _bibliography_module()
    record = _source_record_node()(
        number=2,
        record_type=_source_record_type().PATENT,
        fields={
            "country": "RU",
            "number": "123456",
            "title": "Способ обработки данных",
            "mki": "G06F 17/00",
            "applied_date": "01.02.2023",
            "published_date": "10.08.2023",
        },
        language="ru",
    )

    diagnostics = bibliography.validate_records((record,))

    assert any(
        diagnostic.code == DiagnosticCodes.BIBLIOGRAPHY_MISSING_FIELD
        and diagnostic.rule_id == "common.bibliography.gost_record"
        and "applicant" in diagnostic.message
        for diagnostic in diagnostics
    )


def test_mixed_bibliography_grouping_emits_single_warning_with_methods():
    bibliography = _bibliography_module()
    records = (
        _book(1, "Иванов И. И., Петров П. П.", "Анализ данных", year=2023),
        _book(2, "Петров П. П., Иванов И. И.", "Базы данных", year=2021),
    )

    diagnostics = bibliography.validate_records(
        records,
        grouping_method=bibliography.BibliographyGroupingMethod.ALPHABETICAL,
    )

    grouping = [
        diagnostic
        for diagnostic in diagnostics
        if diagnostic.rule_id == "common.bibliography.grouping_method"
    ]
    assert len(grouping) == 1
    assert grouping[0].code == DiagnosticCodes.BIBLIOGRAPHY_GROUPING_METHOD
    assert "alphabetical" in grouping[0].message
    assert "chronological" in grouping[0].message


def test_renderer_emits_russian_first_warning_for_foreign_entry_before_russian(tmp_path):
    renderer = DocxRenderer(base_dir=tmp_path)
    output_path = tmp_path / "sources.docx"
    ast = Document(
        blocks=(
            _book(1, "Smith J., Brown A.", "Data Analysis", language="en"),
            _book(2, "Иванов И. И., Петров П. П.", "Анализ данных", language="ru"),
        )
    )

    diagnostics = renderer.render_to_file(ast, get_profile("common"), str(output_path))

    assert any(
        diagnostic.code == DiagnosticCodes.BIBLIOGRAPHY_RUSSIAN_FIRST
        and diagnostic.rule_id == "common.bibliography.russian_first"
        for diagnostic in diagnostics
    )
    doc = DocxDocument(str(output_path))
    entries = [paragraph.text for paragraph in doc.paragraphs if paragraph.text.startswith(("1 ", "2 "))]
    assert entries == [
        "1 Smith, J. Data Analysis / J. Smith, A. Brown. — Красноярск: СФУ, 2023. — 320 с.",
        "2 Иванов, И. И. Анализ данных / И. И. Иванов, П. П. Петров. — Красноярск: СФУ, 2023. — 320 с.",
    ]

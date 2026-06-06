from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

from sfu_converter.domain.diagnostics import Diagnostic, DiagnosticCodes, Severity
from sfu_converter.domain.formatting import FormattingProfile
from sfu_converter.domain.ast_nodes import (
    BibliographyEntryNode,
    Document as AstDocument,
    FigureNode,
    FormulaNode,
    ListItemNode,
    ListNode,
    ListType,
    TableCell,
    TableNode,
    TableRow,
)
from sfu_converter.infrastructure.docx_renderer import DocxRenderer
from sfu_converter.infrastructure import docx_styles
from sfu_converter.infrastructure.docx_validator import (
    DocxValidator,
    _pt_value,
    _slug,
    _spacing_value,
    diagnostic_to_json,
)
from sfu_converter.infrastructure.main_inscription import render as render_main_inscription
from sfu_converter.parser import V2Parser
from sfu_converter.registry import get_profile


def _save_doc(tmp_path, doc, name="document.docx"):
    path = tmp_path / name
    doc.save(str(path))
    return path


def _body_paragraph(doc, text="Body"):
    paragraph = doc.add_paragraph(text)
    paragraph.paragraph_format.first_line_indent = Cm(1.25)
    paragraph.paragraph_format.line_spacing = 1.5
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.runs[0]
    run.font.name = "Times New Roman"
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0, 0, 0)
    return paragraph


def test_docx_validator_returns_structured_margin_diagnostic_with_rule_id(tmp_path):
    doc = Document()
    section = doc.sections[0]
    section.left_margin = Cm(2)
    section.right_margin = Cm(1)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    _body_paragraph(doc)
    path = _save_doc(tmp_path, doc)

    diagnostics = DocxValidator(get_profile("common")).validate_file(str(path))

    diagnostic = next(
        diag for diag in diagnostics if diag.code == DiagnosticCodes.FORMAT_MARGIN_LEFT
    )
    assert diagnostic.rule_id == "common.page.margins.portrait"

    payload = diagnostic_to_json(diagnostic)
    assert payload["code"] == "FORMAT_MARGIN_LEFT"
    assert payload["severity"] == "error"
    assert payload["ruleId"] == "common.page.margins.portrait"
    assert payload["source"]["document"] == "docs/formatting requirements/common.md"
    assert payload["source"]["section"] == "Page and Paper Setup"


def test_docx_validator_reports_common_unsupported_validator_rules(tmp_path):
    doc = Document()
    _body_paragraph(doc)
    path = _save_doc(tmp_path, doc)

    diagnostics = DocxValidator(get_profile("common")).validate_file(str(path))
    unsupported = [
        diagnostic
        for diagnostic in diagnostics
        if diagnostic.code == DiagnosticCodes.FORMAT_RULE_NOT_SUPPORTED
    ]

    assert len(unsupported) == 9
    assert all(diagnostic.severity is Severity.WARNING for diagnostic in unsupported)
    assert all("not supported by the validator" in diagnostic.message for diagnostic in unsupported)
    unsupported_ids = {diagnostic.rule_id for diagnostic in unsupported}
    assert "common.figure.caption" not in unsupported_ids
    assert "common.table.caption" not in unsupported_ids
    assert "common.formula.body" not in unsupported_ids
    assert "common.formula.explanation" not in unsupported_ids
    assert "common.figure.image" not in unsupported_ids
    assert "common.figure.explanatory_data" not in unsupported_ids
    assert "common.figure.multi_sheet_label" not in unsupported_ids
    assert "common.formula.body_indent" not in unsupported_ids
    assert "common.formula.line_continuation" not in unsupported_ids
    assert "common.formula.explanation_marker" not in unsupported_ids
    assert "common.formula.repeated_symbol" not in unsupported_ids
    assert "common.formula.consecutive_comma" not in unsupported_ids
    assert "common.heading.h4" not in unsupported_ids
    assert "common.heading.spacing_before" not in unsupported_ids
    assert "common.heading.spacing_after" not in unsupported_ids
    assert "common.heading.no_hyphenation" not in unsupported_ids
    assert "common.heading.two_sentence_separator" not in unsupported_ids
    assert "common.heading.point_requires_subpoints" not in unsupported_ids
    assert "common.list.item" not in unsupported_ids
    assert "common.list.lettered" not in unsupported_ids
    assert "common.list.nested_numeric" not in unsupported_ids
    assert "common.list.marker_alphabetical" not in unsupported_ids
    assert "common.page.margins.landscape" not in unsupported_ids
    assert "common.bibliography.entry" not in unsupported_ids
    assert "common.bibliography.gost_record" not in unsupported_ids
    assert "common.bibliography.gost_abbreviations" not in unsupported_ids
    assert "common.bibliography.grouping_method" not in unsupported_ids
    assert "common.bibliography.russian_first" not in unsupported_ids
    assert "common.reference.footnote" not in unsupported_ids
    assert "common.reference.in_text_simple" not in unsupported_ids
    assert "common.reference.in_text_pages" not in unsupported_ids
    assert "common.reference.in_text_volume" not in unsupported_ids
    assert "common.reference.in_text_group" not in unsupported_ids
    assert "common.reference.cross_check" not in unsupported_ids
    assert "common.reference.figure_table_formula" not in unsupported_ids
    assert "common.appendix.in_text_reference" not in unsupported_ids
    assert "common.style.abbreviation_introduction" not in unsupported_ids
    assert "common.style.no_abbreviations_in_headings" not in unsupported_ids
    assert "common.style.unit_consistency" not in unsupported_ids

    payload = diagnostic_to_json(unsupported[0])
    assert payload["ruleId"] == unsupported[0].rule_id
    assert payload["source"]["document"] == "docs/formatting requirements/common.md"
    assert payload["source"]["section"]


def test_docx_validator_checks_line_spacing_regression(tmp_path):
    doc = Document()
    paragraph = _body_paragraph(doc)
    paragraph.paragraph_format.line_spacing = 1.0
    path = _save_doc(tmp_path, doc)

    diagnostics = DocxValidator(get_profile("common")).validate_file(str(path))

    assert any(
        diagnostic.code == DiagnosticCodes.FORMAT_LINE_SPACING
        and diagnostic.rule_id == "common.text.line_spacing"
        for diagnostic in diagnostics
    )


def test_docx_validator_checks_every_run_in_paragraph(tmp_path):
    doc = Document()
    paragraph = _body_paragraph(doc, "Good")
    bad_run = paragraph.add_run(" bad")
    bad_run.font.name = "Arial"
    bad_run.font.size = Pt(14)
    path = _save_doc(tmp_path, doc)

    diagnostics = DocxValidator(get_profile("common")).validate_file(str(path))

    assert any(
        diagnostic.code == DiagnosticCodes.FORMAT_FONT_NAME
        and "run 2" in diagnostic.message
        for diagnostic in diagnostics
    )


def test_docx_validator_heading_detection_uses_style_or_bold_centering():
    validator = DocxValidator(get_profile("common"))
    doc = Document()

    centered_caption = doc.add_paragraph("Рисунок 1 - Caption")
    centered_caption.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    centered_caption.runs[0].bold = False

    heading_style = doc.add_paragraph("Heading")
    heading_style.style = doc.styles["Heading 1"]

    centered_bold = doc.add_paragraph("Centered bold")
    centered_bold.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    centered_bold.runs[0].bold = True

    assert validator._is_heading_paragraph(centered_caption) is False
    assert validator._is_heading_paragraph(heading_style) is True
    assert validator._is_heading_paragraph(centered_bold) is True


def test_docx_validator_checks_table_font_size_range(tmp_path):
    doc = Document()
    table = doc.add_table(rows=1, cols=1)
    run = table.cell(0, 0).paragraphs[0].add_run("oversized")
    run.font.name = "Times New Roman"
    run.font.size = Pt(14)
    path = _save_doc(tmp_path, doc)

    diagnostics = DocxValidator(get_profile("common")).validate_file(str(path))

    assert any(
        diagnostic.code == DiagnosticCodes.FORMAT_TABLE_FONT_SIZE
        and diagnostic.rule_id == "common.table.font.size"
        for diagnostic in diagnostics
    )


def test_docx_validator_rejects_heading_period(tmp_path):
    doc = Document()
    heading = doc.add_paragraph("Heading.")
    heading.style = doc.styles["Heading 1"]
    heading.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    heading.paragraph_format.first_line_indent = Cm(0)
    heading.paragraph_format.line_spacing = 1.0
    heading.runs[0].font.name = "Times New Roman"
    heading.runs[0].font.size = Pt(14)
    heading.runs[0].bold = True
    path = _save_doc(tmp_path, doc)

    diagnostics = DocxValidator(get_profile("common")).validate_file(str(path))

    assert any(
        diagnostic.code == DiagnosticCodes.FORMAT_HEADING_NO_PERIOD
        and diagnostic.rule_id == "common.heading.no_period"
        for diagnostic in diagnostics
    )


def test_docx_validator_reports_heading_blank_line_mismatches(tmp_path):
    doc = Document()
    _body_paragraph(doc, "Before")
    heading = doc.add_paragraph("Heading")
    heading.style = doc.styles["Heading 2"]
    heading.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    heading.paragraph_format.first_line_indent = Cm(0)
    heading.paragraph_format.line_spacing = 1.0
    heading.paragraph_format.space_before = Pt(0)
    heading.paragraph_format.space_after = Pt(0)
    heading.runs[0].font.name = "Times New Roman"
    heading.runs[0].font.size = Pt(14)
    heading.runs[0].font.color.rgb = RGBColor(0, 0, 0)
    heading.runs[0].bold = True
    _body_paragraph(doc, "After")
    path = _save_doc(tmp_path, doc, "heading_spacing.docx")

    diagnostics = DocxValidator(get_profile("common")).validate_file(str(path))

    assert any(
        diagnostic.code == DiagnosticCodes.HEADING_SPACING_BEFORE
        and diagnostic.rule_id == "common.heading.spacing_before"
        and diagnostic.source.line_start == 2
        for diagnostic in diagnostics
    )
    assert any(
        diagnostic.code == DiagnosticCodes.HEADING_SPACING_AFTER
        and diagnostic.rule_id == "common.heading.spacing_after"
        and diagnostic.source.line_start == 2
        for diagnostic in diagnostics
    )


def test_docx_validator_reports_missing_and_unopenable_files(tmp_path):
    validator = DocxValidator(get_profile("common"))

    missing = validator.validate_file(str(tmp_path / "missing.docx"))
    assert missing[0].code == "DOCX_FILE_NOT_FOUND"

    invalid = tmp_path / "invalid.docx"
    invalid.write_text("not a docx", encoding="utf-8")
    failed = validator.validate_file(str(invalid))
    assert failed[0].code == "DOCX_OPEN_FAILED"


def test_docx_validator_covers_heading_variants_and_format_errors(tmp_path):
    doc = Document()
    h2 = doc.add_paragraph("Heading 2")
    h2.style = doc.styles["Heading 2"]
    h2.paragraph_format.first_line_indent = Cm(1)
    h2.paragraph_format.line_spacing = 2.0
    h2.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    h2.runs[0].font.name = "Arial"
    h2.runs[0].font.size = Pt(16)
    h2.runs[0].font.color.rgb = RGBColor(255, 0, 0)
    h2.runs[0].bold = False

    h3 = doc.add_paragraph("Heading 3")
    h3.style = doc.styles["Heading 3"]
    h3.paragraph_format.first_line_indent = Cm(0)
    h3.paragraph_format.line_spacing = 1.0
    h3.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    h3.paragraph_format.space_before = Pt(2)
    h3.paragraph_format.space_after = Pt(2)
    h3.runs[0].font.name = "Times New Roman"
    h3.runs[0].font.size = Pt(14)

    path = _save_doc(tmp_path, doc)
    diagnostics = DocxValidator(get_profile("common")).validate_file(str(path))
    codes = {diagnostic.code for diagnostic in diagnostics}

    assert DiagnosticCodes.FORMAT_FONT_NAME in codes
    assert DiagnosticCodes.FORMAT_FONT_SIZE in codes
    assert DiagnosticCodes.FORMAT_FONT_COLOR in codes
    assert DiagnosticCodes.FORMAT_INDENT in codes
    assert DiagnosticCodes.FORMAT_ALIGNMENT in codes
    assert DiagnosticCodes.FORMAT_HEADING_BOLD in codes
    assert DiagnosticCodes.FORMAT_PARAGRAPH_SPACING in codes


def test_docx_validator_helpers_cover_optional_json_fields():
    diagnostic = Diagnostic(
        code="CODE",
        message="Message",
        severity=Severity.WARNING,
        rule_id="common.text.font.name",
        suggestion="Suggestion",
    )

    payload = diagnostic_to_json(diagnostic)
    assert payload["suggestion"] == "Suggestion"
    assert payload["source"]["document"] == "docs/formatting requirements/common.md"
    assert payload["source"]["section"] == "Page and Paper Setup"
    assert _pt_value(None) == 0
    assert _pt_value(2.5) == 2.5
    assert _spacing_value(Pt(12)) == 12
    assert _spacing_value(1.5) == 1.5
    assert _slug(" Section Name ") == "section-name"

    without_rule = diagnostic_to_json(
        Diagnostic(code="NO_RULE", message="Message", severity=Severity.ERROR)
    )
    assert without_rule["ruleId"] is None
    assert without_rule["source"] == {
        "document": None,
        "section": None,
        "lineStart": None,
        "lineEnd": None,
    }

    assert (
        DocxValidator(FormattingProfile("empty", "Empty", ()))._rule(
            "common.text.font.name"
        ).id
        == "common.text.font.name"
    )


def test_docx_validator_routes_generated_special_blocks_without_body_indent_errors(tmp_path):
    ast = AstDocument(
        blocks=(
            TableNode(
                caption="Данные",
                rows=(
                    TableRow(cells=(TableCell("Параметр"), TableCell("Значение"))),
                    TableRow(cells=(TableCell("A"), TableCell("1"))),
                ),
            ),
            FigureNode(src="missing.png", caption="Рисунок 1 — Схема установки"),
            FormulaNode(content="E = mc^2", explanation="где E — энергия, Дж;"),
            ListNode(
                list_type=ListType.BULLET,
                items=(ListItemNode("первый элемент"), ListItemNode("второй элемент")),
            ),
            BibliographyEntryNode(number=1, text="Иванов И.И. Источник. — М.: Наука, 2024."),
        )
    )
    output = tmp_path / "special_blocks.docx"
    DocxRenderer(base_dir=tmp_path).render_to_file(ast, get_profile("common"), str(output))

    diagnostics = DocxValidator(get_profile("common")).validate_file(str(output))

    assert not any(
        diagnostic.code == DiagnosticCodes.FORMAT_INDENT
        and diagnostic.rule_id == "common.text.indent.first_line"
        for diagnostic in diagnostics
    )


def test_docx_validator_reports_figure_caption_alignment_against_figure_rule(tmp_path):
    doc = Document()
    paragraph = _special_paragraph(doc, "Рисунок 1 — Схема", WD_ALIGN_PARAGRAPH.LEFT, indent_cm=0)
    path = _save_doc(tmp_path, doc, "figure_caption.docx")

    diagnostics = DocxValidator(get_profile("common")).validate_file(str(path))

    assert paragraph.text
    assert any(
        diagnostic.code == DiagnosticCodes.FORMAT_ALIGNMENT
        and diagnostic.rule_id == "common.figure.caption"
        for diagnostic in diagnostics
    )


def test_docx_validator_reports_table_caption_alignment_against_table_rule(tmp_path):
    doc = Document()
    _special_paragraph(doc, "Таблица 1 — Данные", WD_ALIGN_PARAGRAPH.CENTER, indent_cm=0)
    path = _save_doc(tmp_path, doc, "table_caption.docx")

    diagnostics = DocxValidator(get_profile("common")).validate_file(str(path))

    assert any(
        diagnostic.code == DiagnosticCodes.FORMAT_ALIGNMENT
        and diagnostic.rule_id == "common.table.caption"
        for diagnostic in diagnostics
    )


def test_docx_validator_reports_formula_explanation_indent_against_formula_rule(tmp_path):
    doc = Document()
    _special_paragraph(doc, "где E — энергия, Дж;", WD_ALIGN_PARAGRAPH.LEFT, indent_cm=1.25)
    path = _save_doc(tmp_path, doc, "formula_explanation.docx")

    diagnostics = DocxValidator(get_profile("common")).validate_file(str(path))

    assert any(
        diagnostic.code == DiagnosticCodes.FORMAT_INDENT
        and diagnostic.rule_id == "common.formula.explanation"
        for diagnostic in diagnostics
    )


def test_docx_validator_reports_formula_explanation_marker_colon(tmp_path):
    doc = Document()
    docx_styles.register_styles(doc)
    paragraph = _special_paragraph(doc, "где:", WD_ALIGN_PARAGRAPH.LEFT, indent_cm=0)
    paragraph.style = doc.styles[docx_styles.FORMULA_EXPLANATION]
    path = _save_doc(tmp_path, doc, "formula_marker.docx")

    diagnostics = DocxValidator(get_profile("common")).validate_file(str(path))

    assert any(
        diagnostic.code == DiagnosticCodes.FORMULA_EXPLANATION_MARKER
        and diagnostic.rule_id == "common.formula.explanation_marker"
        for diagnostic in diagnostics
    )


def test_docx_validator_reports_repeated_symbol_without_prior_definition(tmp_path):
    doc = Document()
    docx_styles.register_styles(doc)
    paragraph = _special_paragraph(
        doc,
        "где\nc — то же, что и в формуле (1)",
        WD_ALIGN_PARAGRAPH.LEFT,
        indent_cm=0,
    )
    paragraph.style = doc.styles[docx_styles.FORMULA_EXPLANATION]
    path = _save_doc(tmp_path, doc, "formula_repeated_symbol.docx")

    diagnostics = DocxValidator(get_profile("common")).validate_file(str(path))

    assert any(
        diagnostic.code == DiagnosticCodes.FORMULA_REPEATED_SYMBOL
        and diagnostic.rule_id == "common.formula.repeated_symbol"
        for diagnostic in diagnostics
    )


def test_docx_validator_reports_consecutive_formula_without_comma(tmp_path):
    doc = Document()
    docx_styles.register_styles(doc)
    first = _special_paragraph(doc, "a = b\t(1)", WD_ALIGN_PARAGRAPH.CENTER, indent_cm=1.25)
    second = _special_paragraph(doc, "c = d\t(2)", WD_ALIGN_PARAGRAPH.CENTER, indent_cm=1.25)
    first.style = doc.styles[docx_styles.FORMULA_BODY]
    second.style = doc.styles[docx_styles.FORMULA_BODY]
    path = _save_doc(tmp_path, doc, "formula_consecutive.docx")

    diagnostics = DocxValidator(get_profile("common")).validate_file(str(path))

    assert any(
        diagnostic.code == DiagnosticCodes.FORMULA_CONSECUTIVE_COMMA
        and diagnostic.rule_id == "common.formula.consecutive_comma"
        for diagnostic in diagnostics
    )


def test_docx_validator_reports_long_formula_without_operator_break(tmp_path):
    doc = Document()
    docx_styles.register_styles(doc)
    paragraph = _special_paragraph(
        doc,
        (
            "result = alpha + beta + gamma + delta + epsilon + zeta + eta + "
            "theta + iota + kappa + lambda\t(1)"
        ),
        WD_ALIGN_PARAGRAPH.CENTER,
        indent_cm=1.25,
    )
    paragraph.style = doc.styles[docx_styles.FORMULA_BODY]
    path = _save_doc(tmp_path, doc, "formula_long.docx")

    diagnostics = DocxValidator(get_profile("common")).validate_file(str(path))

    diagnostic = next(
        item
        for item in diagnostics
        if item.code == DiagnosticCodes.FORMULA_LINE_CONTINUATION
    )
    assert diagnostic.severity is Severity.INFO
    assert diagnostic.rule_id == "common.formula.line_continuation"


def test_docx_validator_reports_bad_multisheet_figure_caption(tmp_path):
    doc = Document()
    docx_styles.register_styles(doc)
    paragraph = _special_paragraph(
        doc,
        "Рисунок 5 — Схема, лист 2",
        WD_ALIGN_PARAGRAPH.CENTER,
        indent_cm=0,
    )
    paragraph.style = doc.styles[docx_styles.FIGURE_CAPTION]
    path = _save_doc(tmp_path, doc, "bad_multisheet_caption.docx")

    diagnostics = DocxValidator(get_profile("common")).validate_file(str(path))

    assert any(
        diagnostic.code == DiagnosticCodes.FIGURE_MULTI_SHEET_LABEL
        and diagnostic.rule_id == "common.figure.multi_sheet_label"
        for diagnostic in diagnostics
    )


def test_docx_validator_reports_table_header_period_and_serial_column(tmp_path):
    doc = Document()
    table = doc.add_table(rows=2, cols=2)
    table.rows[0].cells[0].text = "№ п/п"
    table.rows[0].cells[1].text = "Параметр."
    table.rows[1].cells[0].text = "1"
    table.rows[1].cells[1].text = "A"
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = "Times New Roman"
                    run.font.size = Pt(12)
    path = _save_doc(tmp_path, doc, "bad_table_headers.docx")

    diagnostics = DocxValidator(get_profile("common")).validate_file(str(path))
    rule_ids = {diagnostic.rule_id for diagnostic in diagnostics}

    assert "common.table.forbid_serial_column" in rule_ids
    assert "common.table.no_period_in_header" in rule_ids


def test_docx_validator_reports_diagonal_split_and_nonitalic_letter_cells(tmp_path):
    doc = Document()
    table = doc.add_table(rows=2, cols=1)
    table.rows[0].cells[0].text = "Параметр"
    table.rows[1].cells[0].text = "A"
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = "Times New Roman"
                    run.font.size = Pt(12)

    tc_pr = table.rows[0].cells[0]._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    diagonal = OxmlElement("w:tl2br")
    diagonal.set(qn("w:val"), "single")
    borders.append(diagonal)
    tc_pr.append(borders)
    path = _save_doc(tmp_path, doc, "bad_table_diagonal.docx")

    diagnostics = DocxValidator(get_profile("common")).validate_file(str(path))
    rule_ids = {diagnostic.rule_id for diagnostic in diagnostics}

    assert "common.table.no_diagonal_split" in rule_ids
    assert "common.table.italic_letters" in rule_ids


def test_docx_validator_accepts_full_grid_table_borders(tmp_path):
    doc = Document()
    table = doc.add_table(rows=2, cols=2)
    table.rows[0].cells[0].text = "Параметр"
    table.rows[0].cells[1].text = "Значение"
    table.rows[1].cells[0].text = "A"
    table.rows[1].cells[1].text = "10"

    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = OxmlElement(f"w:{edge}")
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "000000")
        borders.append(element)
    table._tbl.tblPr.append(borders)

    for cell in table.rows[0].cells:
        tc_pr = cell._tc.get_or_add_tcPr()
        cell_borders = OxmlElement("w:tcBorders")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "double")
        bottom.set(qn("w:sz"), "6")
        bottom.set(qn("w:color"), "000000")
        cell_borders.append(bottom)
        tc_pr.append(cell_borders)

    path = _save_doc(tmp_path, doc, "full_grid_table.docx")

    diagnostics = DocxValidator(get_profile("common")).validate_file(str(path))

    assert not any(diagnostic.rule_id == "common.table.borders" for diagnostic in diagnostics)


def test_docx_validator_accepts_generated_table_unit_label_at_12pt(tmp_path):
    ast = V2Parser().parse(
        "\n".join(
            [
                '[TABLE caption="Параметры" unit="МПа"]',
                "| Параметр | Значение |",
                "| Давление | 10 |",
                "[TABLE_END]",
            ]
        )
    ).document
    output = tmp_path / "unit_label.docx"
    DocxRenderer(base_dir=tmp_path).render_to_file(ast, get_profile("common"), str(output))

    doc = Document(str(output))
    assert any(paragraph.text == ", МПа" for paragraph in doc.paragraphs)

    diagnostics = DocxValidator(get_profile("common")).validate_file(str(output))

    assert not any(
        diagnostic.code == DiagnosticCodes.FORMAT_FONT_SIZE
        and diagnostic.rule_id == "common.text.font.size"
        and "12.0pt" in diagnostic.message
        for diagnostic in diagnostics
    )
    assert not any(
        diagnostic.rule_id == "common.table.unit_label"
        for diagnostic in diagnostics
    )


def test_docx_validator_reports_list_marker_errors(tmp_path):
    doc = Document()
    docx_styles.register_styles(doc)
    good = _list_paragraph(doc, "а) первый", left_indent_cm=1.25)
    skipped = _list_paragraph(doc, "в) третий", left_indent_cm=1.25)
    disallowed = _list_paragraph(doc, "й) запрещенная буква", left_indent_cm=1.25)
    for paragraph in (good, skipped, disallowed):
        paragraph.style = doc.styles[docx_styles.LIST_ITEM]
    path = _save_doc(tmp_path, doc, "bad_list_markers.docx")

    diagnostics = DocxValidator(get_profile("common")).validate_file(str(path))

    assert any(
        diagnostic.code == "LIST_MARKER_OUT_OF_ORDER"
        and diagnostic.rule_id == "common.list.marker_alphabetical"
        for diagnostic in diagnostics
    )
    assert any(
        diagnostic.code == "LIST_MARKER_DISALLOWED_LETTER"
        and diagnostic.rule_id == "common.list.lettered"
        for diagnostic in diagnostics
    )


def test_docx_validator_reports_bad_nested_numeric_indent(tmp_path):
    doc = Document()
    docx_styles.register_styles(doc)
    parent = _list_paragraph(doc, "а) основной пункт", left_indent_cm=1.25)
    nested = _list_paragraph(doc, "1) вложенный пункт", left_indent_cm=1.40)
    parent.style = doc.styles[docx_styles.LIST_ITEM]
    nested.style = doc.styles[docx_styles.LIST_ITEM]
    path = _save_doc(tmp_path, doc, "bad_nested_indent.docx")

    diagnostics = DocxValidator(get_profile("common")).validate_file(str(path))

    assert any(
        diagnostic.code == "LIST_NESTED_NUMERIC_INDENT"
        and diagnostic.rule_id == "common.list.nested_numeric"
        for diagnostic in diagnostics
    )


def test_docx_validator_uses_landscape_margin_rule_for_landscape_sections(tmp_path):
    doc = Document()
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Cm(29.7)
    section.page_height = Cm(21)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(1)
    _body_paragraph(doc)
    path = _save_doc(tmp_path, doc, "bad_landscape.docx")

    diagnostics = DocxValidator(get_profile("common")).validate_file(str(path))

    assert any(
        diagnostic.code == DiagnosticCodes.FORMAT_MARGIN_TOP
        and diagnostic.rule_id == "common.page.margins.landscape"
        for diagnostic in diagnostics
    )


def test_docx_validator_reports_missing_coursework_frame(tmp_path):
    doc = Document()
    _body_paragraph(doc)
    path = _save_doc(tmp_path, doc, "unframed_coursework.docx")

    diagnostics = DocxValidator(get_profile("coursework")).validate_file(str(path))

    assert any(
        diagnostic.code == "FRAME_MISSING"
        and diagnostic.rule_id == "coursework.frame.course_project_explanatory_note"
        for diagnostic in diagnostics
    )


def test_docx_validator_reports_missing_recovered_metadata_as_warning(tmp_path):
    doc = Document()
    _body_paragraph(doc)
    path = _save_doc(tmp_path, doc, "missing_metadata.docx")

    diagnostics = DocxValidator(get_profile("coursework")).validate_file(str(path))

    assert any(
        diagnostic.code == DiagnosticCodes.TXT_MISSING_METADATA
        and diagnostic.rule_id == "coursework.metadata.required"
        and diagnostic.severity is Severity.WARNING
        and "supervisor" in diagnostic.data["missing"]
        for diagnostic in diagnostics
    )


def test_docx_validator_warns_when_project_title_block_graph_two_is_empty(tmp_path):
    doc = Document()
    render_main_inscription(doc, "form_1", fields={"1": "Пояснительная записка"})
    path = _save_doc(tmp_path, doc, "empty_designation.docx")

    diagnostics = DocxValidator(get_profile("coursework")).validate_file(str(path))

    assert any(
        diagnostic.code == DiagnosticCodes.PROJECT_DESIGNATION_MISSING
        and diagnostic.severity is Severity.WARNING
        and diagnostic.rule_id == "project_designations.title_block.letter_numeric_designation"
        for diagnostic in diagnostics
    )


def _special_paragraph(doc, text, alignment, *, indent_cm: float):
    paragraph = doc.add_paragraph(text)
    paragraph.paragraph_format.alignment = alignment
    paragraph.paragraph_format.first_line_indent = Cm(indent_cm)
    paragraph.paragraph_format.line_spacing = 1.5
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.runs[0]
    run.font.name = "Times New Roman"
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0, 0, 0)
    return paragraph


def _list_paragraph(doc, text, *, left_indent_cm: float):
    paragraph = _special_paragraph(doc, text, WD_ALIGN_PARAGRAPH.JUSTIFY, indent_cm=-0.5)
    paragraph.paragraph_format.left_indent = Cm(left_indent_cm)
    return paragraph

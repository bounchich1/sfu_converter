from docx import Document as DocxDocument
from docx.shared import Cm

from sfu_converter.domain.ast_nodes import Document, FormulaNode, FormulaSymbol
from sfu_converter.infrastructure.docx_renderer import DocxRenderer
from sfu_converter.infrastructure.formula_layout import split_formula_lines
from sfu_converter.registry import get_profile


def _render(tmp_path, document: Document) -> DocxDocument:
    output_path = tmp_path / "formulas.docx"
    DocxRenderer(base_dir=tmp_path).render_to_file(
        document,
        get_profile("common"),
        str(output_path),
    )
    return DocxDocument(str(output_path))


def test_split_formula_lines_repeats_operator_on_continuation():
    lines = split_formula_lines("a + b + c + d", max_chars=8)

    assert lines[0].endswith("+")
    assert lines[1].startswith("+")


def test_formula_body_uses_standard_indent_and_right_tab(tmp_path):
    doc = _render(tmp_path, Document(blocks=(FormulaNode(content="E = mc^2"),)))

    formula = next(p for p in doc.paragraphs if p.text.endswith("(1)"))
    assert abs(formula.paragraph_format.first_line_indent - Cm(1.25)) < 1000
    assert "<w:tabs" in formula._p.xml
    assert 'w:val="right"' in formula._p.xml


def test_formula_symbols_render_repeated_shorthand(tmp_path):
    doc = _render(
        tmp_path,
        Document(
            blocks=(
                FormulaNode(
                    content="E = m * c**2",
                    explanations=(FormulaSymbol(name="c", description="скорость света"),),
                ),
                FormulaNode(
                    content="p = m * c",
                    explanations=(FormulaSymbol(name="c", description="", repeats=True),),
                ),
            )
        ),
    )

    assert any(
        "c — то же, что и в формуле (1)" in paragraph.text
        for paragraph in doc.paragraphs
    )


def test_consecutive_formulas_render_without_blank_line_and_with_comma(tmp_path):
    doc = _render(
        tmp_path,
        Document(
            blocks=(
                FormulaNode(id="f1", content="a = b"),
                FormulaNode(id="f2", content="c = d", consecutive_with="f1"),
            )
        ),
    )

    formula_indexes = [
        index
        for index, paragraph in enumerate(doc.paragraphs)
        if paragraph.text.endswith("(1)") or paragraph.text.endswith("(2)")
    ]
    first = doc.paragraphs[formula_indexes[0]]
    assert first.text == "a = b,\t(1)"
    assert formula_indexes[1] == formula_indexes[0] + 1


def test_long_formula_body_breaks_on_operator(tmp_path):
    doc = _render(
        tmp_path,
        Document(
            blocks=(
                FormulaNode(
                    content=(
                        "result = alpha + beta + gamma + delta + epsilon + zeta + "
                        "eta + theta + iota + kappa + lambda"
                    ),
                ),
            )
        ),
    )

    formula = next(p for p in doc.paragraphs if p.text.endswith("(1)"))
    assert "\n+" in formula.text

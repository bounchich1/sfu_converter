from pathlib import Path

from docx import Document as DocxDocument

from sfu_converter.converter import TextToDocxConverter
from sfu_converter.domain.ast_nodes import (
    AppendixNode,
    BibliographyEntryNode,
    FigureNode,
    FormulaNode,
    FormulaSymbol,
    HeadingLevel,
    HeadingNode,
    ListNode,
    ListType,
    ParagraphNode,
    StructuralSectionNode,
    StructuralSectionType,
    TableNode,
    TableOfContentsNode,
)
from sfu_converter.domain.diagnostics import DiagnosticCodes, Severity
from sfu_converter.parser.v1_parser import APPENDIX_LETTERS, V1Parser


def test_parse_v1_headings_plain_text_image_and_table():
    source = "\n".join(
        [
            "[H1] Введение",
            "Обычный текст",
            "[H2] Глава 1",
            "[H3] Подраздел",
            "[IMAGE=diagram.png]",
            "Рисунок 1 - Схема",
            "[TABLE_START]",
            "[TABLE_CAPTION] Таблица 1 - Данные",
            "| A | B |",
            "| C | D |",
            "[TABLE_END]",
        ]
    )

    result = V1Parser().parse(source, filename="report.txt")

    assert result.diagnostics == []
    assert result.document.syntax_version == 1
    assert result.document.source_file == "report.txt"
    assert [type(block) for block in result.document.blocks] == [
        StructuralSectionNode,
        ParagraphNode,
        HeadingNode,
        HeadingNode,
        FigureNode,
        TableNode,
    ]
    assert result.document.blocks[0].section_type is StructuralSectionType.INTRODUCTION
    assert result.document.blocks[1].runs[0].text == "Обычный текст"
    assert result.document.blocks[4].src == "diagram.png"
    assert result.document.blocks[4].caption == "Рисунок 1 - Схема"
    assert result.document.blocks[5].caption == "Таблица 1 - Данные"
    assert result.document.blocks[5].rows[1].cells[1].text == "D"


def test_parse_complete_fixture_file_to_expected_logical_blocks():
    source = Path("tests/test_input.txt").read_text(encoding="utf-8")

    result = V1Parser().parse(source, filename="tests/test_input.txt")

    assert result.diagnostics == []
    assert len(result.document.blocks) == 12
    assert [type(block).__name__ for block in result.document.blocks] == [
        "StructuralSectionNode",
        "ParagraphNode",
        "HeadingNode",
        "ParagraphNode",
        "ParagraphNode",
        "FigureNode",
        "ParagraphNode",
        "TableNode",
        "ParagraphNode",
        "ParagraphNode",
        "HeadingNode",
        "ParagraphNode",
    ]


def test_parser_converts_known_h1_to_structural_section():
    result = V1Parser().parse("[H1] введение\n[H1] Обычный раздел")

    assert result.diagnostics == []
    structural, regular = result.document.blocks
    assert isinstance(structural, StructuralSectionNode)
    assert structural.section_type is StructuralSectionType.INTRODUCTION
    assert structural.title == "введение"
    assert isinstance(regular, HeadingNode)
    assert regular.level is HeadingLevel.H1
    assert regular.text == "Обычный раздел"
    assert regular.number == "auto"


def test_parser_marks_regular_headings_for_auto_numbering():
    result = V1Parser().parse(
        "\n".join(
            [
                "[H1] Раздел",
                "[H2] Подраздел",
                "[H3] Пункт",
            ]
        )
    )

    assert result.diagnostics == []
    assert [
        (block.level, block.text, block.number)
        for block in result.document.blocks
    ] == [
        (HeadingLevel.H1, "Раздел", "auto"),
        (HeadingLevel.H2, "Подраздел", "auto"),
        (HeadingLevel.H3, "Пункт", "auto"),
    ]


def test_parser_accepts_h4_subpoint_marker():
    result = V1Parser().parse(
        "\n".join(
            [
                "[H1] One",
                "[H2] One.One",
                "[H3] One.One.One",
                "[H4] One.One.One.One",
            ]
        )
    )

    assert result.diagnostics == []
    assert [
        (block.level, block.text, block.number)
        for block in result.document.blocks
    ] == [
        (HeadingLevel.H1, "One", "auto"),
        (HeadingLevel.H2, "One.One", "auto"),
        (HeadingLevel.H3, "One.One.One", "auto"),
        (HeadingLevel.H4, "One.One.One.One", "auto"),
    ]


def test_parser_rejects_heading_level_above_four():
    result = V1Parser().parse("[H5] Too deep")

    assert [block for block in result.document.blocks] == []
    codes = [diagnostic.code for diagnostic in result.diagnostics]
    assert DiagnosticCodes.INVALID_HEADING_LEVEL in codes


def test_parser_warns_on_skipped_heading_level():
    result = V1Parser().parse(
        "\n".join(
            [
                "[H1] One",
                "[H4] Skipped to four",
            ]
        )
    )

    skipped = [
        diagnostic
        for diagnostic in result.diagnostics
        if diagnostic.code == DiagnosticCodes.HEADING_LEVEL_SKIPPED
    ]
    assert len(skipped) == 1
    assert skipped[0].severity is Severity.WARNING
    assert [block.level for block in result.document.blocks] == [
        HeadingLevel.H1,
        HeadingLevel.H4,
    ]


def test_parser_groups_consecutive_dash_lines_into_bullet_list():
    result = V1Parser().parse(
        "\n".join(
            [
                "Перед списком",
                "- первый элемент",
                "- второй элемент;",
                "- третий элемент.",
                "После списка",
            ]
        ),
        filename="list.txt",
    )

    assert result.diagnostics == []
    before, list_block, after = result.document.blocks
    assert before.runs[0].text == "Перед списком"
    assert after.runs[0].text == "После списка"
    assert isinstance(list_block, ListNode)
    assert list_block.list_type is ListType.BULLET
    assert [item.text for item in list_block.items] == [
        "первый элемент",
        "второй элемент;",
        "третий элемент.",
    ]
    assert list_block.source.line_start == 2
    assert list_block.source.line_end == 4


def test_parser_preserves_indented_nested_numeric_children_in_dash_list():
    result = V1Parser().parse(
        "\n".join(
            [
                "- первый пункт",
                "- второй пункт",
                "  1) подпункт один",
                "  2) подпункт два",
            ]
        ),
        filename="nested-list.txt",
    )

    assert result.diagnostics == []
    (list_block,) = result.document.blocks
    assert isinstance(list_block, ListNode)
    second_item = list_block.items[1]
    assert second_item.text == "второй пункт"
    assert len(second_item.children) == 1
    nested = second_item.children[0]
    assert nested.list_type is ListType.NUMBERED
    assert [item.text for item in nested.items] == [
        "подпункт один",
        "подпункт два",
    ]
    assert nested.source.line_start == 3
    assert nested.source.line_end == 4


def test_parser_parses_explicit_lettered_and_numbered_lists():
    result = V1Parser().parse(
        "\n".join(
            [
                "[LIST type=letter]",
                "[а)] первый элемент",
                "[б)] второй элемент",
                "[LIST_END]",
                "[LIST type=number]",
                "[1)] первый подпункт",
                "[2)] второй подпункт",
                "[LIST_END]",
            ]
        )
    )

    assert result.diagnostics == []
    lettered, numbered = result.document.blocks
    assert lettered.list_type is ListType.LETTERED
    assert [item.text for item in lettered.items] == [
        "первый элемент",
        "второй элемент",
    ]
    assert numbered.list_type is ListType.NUMBERED
    assert [item.text for item in numbered.items] == [
        "первый подпункт",
        "второй подпункт",
    ]


def test_parser_accepts_empty_explicit_list():
    result = V1Parser().parse("[LIST type=bullet]\n[LIST_END]")

    assert result.diagnostics == []
    (list_block,) = result.document.blocks
    assert isinstance(list_block, ListNode)
    assert list_block.list_type is ListType.BULLET
    assert list_block.items == ()


def test_parser_recognizes_explicit_structural_markers():
    result = V1Parser().parse(
        "\n".join(
            [
                "[SECTION type=conclusion]",
                '[STRUCTURAL title="СПИСОК СОКРАЩЕНИЙ"]',
            ]
        )
    )

    assert result.diagnostics == []
    first, second = result.document.blocks
    assert first == StructuralSectionNode(
        section_type=StructuralSectionType.CONCLUSION,
        title="ЗАКЛЮЧЕНИЕ",
        source=first.source,
    )
    assert second == StructuralSectionNode(
        section_type=StructuralSectionType.ABBREVIATIONS,
        title="СПИСОК СОКРАЩЕНИЙ",
        source=second.source,
    )


def test_parser_reports_cyrillic_marker_lookalikes_with_line_numbers():
    result = V1Parser().parse("[Н1] Wrong marker", filename="bad.txt")

    assert any(
        diagnostic.code == DiagnosticCodes.TXT_CYRILLIC_IN_MARKER
        and diagnostic.severity is Severity.ERROR
        and diagnostic.source.line_start == 1
        and diagnostic.source.filename == "bad.txt"
        for diagnostic in result.diagnostics
    )
    assert result.has_errors is True


def test_parser_reports_unknown_marker_warning():
    result = V1Parser().parse("[UNKNOWN] value")

    assert len(result.diagnostics) == 1
    diagnostic = result.diagnostics[0]
    assert diagnostic.code == DiagnosticCodes.TXT_UNKNOWN_MARKER
    assert diagnostic.severity is Severity.WARNING
    assert diagnostic.source.line_start == 1
    assert result.has_errors is False


def test_parser_reports_missing_table_end_and_preserves_rows():
    result = V1Parser().parse(
        "\n".join(
            [
                "Before",
                "[TABLE_START]",
                "| A | B |",
                "| C | D |",
            ]
        )
    )

    assert result.document.blocks[1].rows[0].cells[0].text == "A"
    assert result.diagnostics[0].code == DiagnosticCodes.TXT_MISSING_BLOCK_END
    assert result.diagnostics[0].severity is Severity.ERROR
    assert result.diagnostics[0].source.line_start == 2
    assert result.has_errors is True


def test_parser_reports_invalid_table_shape():
    result = V1Parser().parse(
        "\n".join(
            [
                "[TABLE_START]",
                "| A | B |",
                "| C |",
                "[TABLE_END]",
            ]
        )
    )

    assert result.diagnostics[0].code == DiagnosticCodes.TXT_INVALID_TABLE_SHAPE
    assert result.diagnostics[0].source.line_start == 3
    assert result.has_errors is True


def test_image_without_caption_and_empty_image_marker():
    result = V1Parser().parse("[IMAGE=chart.png]\n\n[IMAGE]")

    first, second = result.document.blocks
    assert first.src == "chart.png"
    assert first.caption is None
    assert first.source.line_start == 1
    assert second.src is None
    assert second.caption is None
    assert second.source.line_start == 3


def test_parser_reports_malformed_image_tag():
    result = V1Parser().parse("[IMAGE=chart.png")

    assert result.document.blocks == ()
    assert result.diagnostics[0].code == DiagnosticCodes.TXT_MALFORMED_ATTRIBUTE
    assert result.diagnostics[0].severity is Severity.ERROR
    assert result.diagnostics[0].source.line_start == 1


def test_parser_splits_paragraph_into_bold_italic_runs():
    result = V1Parser().parse(
        "Hello **bold** and *italic* and ***bold italic*** done."
    )

    assert result.diagnostics == []
    (paragraph,) = result.document.blocks
    assert isinstance(paragraph, ParagraphNode)
    assert [
        (run.text, run.bold, run.italic) for run in paragraph.runs
    ] == [
        ("Hello ", False, False),
        ("bold", True, False),
        (" and ", False, False),
        ("italic", False, True),
        (" and ", False, False),
        ("bold italic", True, True),
        (" done.", False, False),
    ]


def test_parser_returns_single_run_when_no_inline_formatting():
    result = V1Parser().parse("Просто обычный текст без форматирования.")

    (paragraph,) = result.document.blocks
    assert len(paragraph.runs) == 1
    assert paragraph.runs[0].text == "Просто обычный текст без форматирования."
    assert paragraph.runs[0].bold is False
    assert paragraph.runs[0].italic is False


def test_parser_treats_unmatched_asterisks_as_literal_text():
    result = V1Parser().parse("Простой * одиночный знак и **незакрытый текст")

    (paragraph,) = result.document.blocks
    rendered = "".join(run.text for run in paragraph.runs)
    assert rendered == "Простой * одиночный знак и **незакрытый текст"
    assert all(run.bold is False and run.italic is False for run in paragraph.runs)


def test_empty_input_returns_empty_document_without_diagnostics():
    result = V1Parser().parse("")

    assert result.document.blocks == ()
    assert result.document.syntax_version == 1
    assert result.diagnostics == []
    assert result.has_errors is False


def test_parser_module_is_independent_from_python_docx():
    parser_source = Path("src/sfu_converter/parser/v1_parser.py").read_text(encoding="utf-8")

    assert "docx" not in parser_source.lower()


def test_converter_renders_preparsed_ast(tmp_path):
    parser_result = V1Parser().parse(
        "\n".join(
            [
                "[H1] Заголовок",
                "Абзац",
                "[TABLE_START]",
                "| A | B |",
                "| C | D |",
                "[TABLE_END]",
            ]
        )
    )
    output_file = tmp_path / "ast.docx"
    converter = TextToDocxConverter(base_dir=tmp_path)
    converter._initialize_document()

    converter._render_lines(parser_result.document)
    converter.doc.save(str(output_file))

    doc = DocxDocument(str(output_file))
    assert [paragraph.text for paragraph in doc.paragraphs if paragraph.text] == [
        "1 Заголовок",
        "Абзац",
    ]
    assert doc.tables[0].rows[1].cells[1].text == "D"


def test_parser_parses_formula_block_with_explanation():
    result = V1Parser().parse(
        "\n".join(
            [
                "[FORMULA]",
                "E = mc^2",
                "[FORMULA_END]",
                "[FORMULA_EXPLANATION]",
                "где E — энергия, Дж;",
                "    m — масса, кг;",
                "[FORMULA_EXPLANATION_END]",
            ]
        )
    )

    assert result.diagnostics == []
    (formula,) = result.document.blocks
    assert isinstance(formula, FormulaNode)
    assert formula.content == "E = mc^2"
    assert formula.explanation == "где E — энергия, Дж;\n    m — масса, кг;"
    assert formula.source.line_start == 1
    assert formula.source.line_end == 7


def test_parser_parses_formula_without_explanation():
    result = V1Parser().parse(
        "\n".join(
            [
                "Перед формулой",
                "[FORMULA]",
                "a + b = c",
                "[FORMULA_END]",
                "После формулы",
            ]
        )
    )

    assert result.diagnostics == []
    before, formula, after = result.document.blocks
    assert before.runs[0].text == "Перед формулой"
    assert isinstance(formula, FormulaNode)
    assert formula.content == "a + b = c"
    assert formula.explanation is None
    assert after.runs[0].text == "После формулы"


def test_parser_supports_formula_attributes():
    result = V1Parser().parse(
        "\n".join(
            [
                "[FORMULA id=eq:energy number=auto]",
                "E = mc^2",
                "[FORMULA_END]",
            ]
        )
    )

    assert result.diagnostics == []
    (formula,) = result.document.blocks
    assert formula.id == "eq:energy"
    assert formula.number == "auto"


def test_parser_parses_formula_symbol_lines():
    result = V1Parser().parse(
        "\n".join(
            [
                "[FORMULA id=eq:momentum]",
                "p = m * c",
                '[FORMULA_SYMBOL name=m text="масса, кг"]',
                '[FORMULA_SYMBOL name=c text="" repeats=true]',
                "[FORMULA_END]",
            ]
        )
    )

    assert result.diagnostics == []
    (formula,) = result.document.blocks
    assert formula.content == "p = m * c"
    assert formula.explanations == (
        FormulaSymbol(name="m", description="масса, кг"),
        FormulaSymbol(name="c", description="", repeats=True),
    )


def test_parser_reports_unterminated_formula_block():
    result = V1Parser().parse("[FORMULA]\nE = mc^2\n")

    assert result.has_errors is True
    assert result.diagnostics[0].code == DiagnosticCodes.TXT_MISSING_BLOCK_END
    assert result.diagnostics[0].source.line_start == 1
    (formula,) = result.document.blocks
    assert formula.content == "E = mc^2"


def test_parser_reports_orphan_formula_explanation_marker():
    result = V1Parser().parse("[FORMULA_EXPLANATION]\nне ожидается")

    assert result.has_errors is True
    assert result.diagnostics[0].code == DiagnosticCodes.TXT_UNKNOWN_MARKER
    assert "Unexpected marker" in result.diagnostics[0].message


def test_parser_parses_numbered_lines_after_sources_heading_as_bibliography_entries():
    result = V1Parser().parse(
        "\n".join(
            [
                "[H1] СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ",
                "1 Иванов И.И. Основы программирования. — М.: Наука, 2023.",
                "2 Петров П.П. Алгоритмы и структуры данных. — СПб.: БХВ, 2022.",
            ]
        )
    )

    assert result.diagnostics == []
    sources_heading, first, second = result.document.blocks
    assert isinstance(sources_heading, StructuralSectionNode)
    assert sources_heading.section_type is StructuralSectionType.SOURCES
    assert isinstance(first, BibliographyEntryNode)
    assert first.number == 1
    assert first.text == "Иванов И.И. Основы программирования. — М.: Наука, 2023."
    assert isinstance(second, BibliographyEntryNode)
    assert second.number == 2


def test_parser_does_not_treat_numbered_lines_outside_sources_as_bibliography():
    result = V1Parser().parse(
        "\n".join(
            [
                "1 это просто абзац который начинается с цифры",
            ]
        )
    )

    assert result.diagnostics == []
    (paragraph,) = result.document.blocks
    assert isinstance(paragraph, ParagraphNode)


def test_parser_resets_bibliography_mode_on_next_h1():
    result = V1Parser().parse(
        "\n".join(
            [
                "[H1] СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ",
                "1 Иванов И.И. Книга. — М., 2023.",
                "[H1] Дополнения",
                "1 этот текст не должен стать записью библиографии",
            ]
        )
    )

    assert result.diagnostics == []
    sources, entry, heading, paragraph = result.document.blocks
    assert isinstance(sources, StructuralSectionNode)
    assert isinstance(entry, BibliographyEntryNode)
    assert isinstance(heading, HeadingNode)
    assert heading.level is HeadingLevel.H1
    assert isinstance(paragraph, ParagraphNode)
    assert paragraph.runs[0].text.startswith("1 этот текст")


def test_parser_recognises_section_marker_for_bibliography_mode():
    result = V1Parser().parse(
        "\n".join(
            [
                "[SECTION type=sources]",
                "1 Иванов И.И. Книга. — М., 2023.",
            ]
        )
    )

    assert result.diagnostics == []
    section, entry = result.document.blocks
    assert section.section_type is StructuralSectionType.SOURCES
    assert isinstance(entry, BibliographyEntryNode)
    assert entry.number == 1


def test_appendix_letters_skip_excluded_cyrillic_letters():
    excluded = {"Ё", "З", "Й", "О", "Ч", "Ъ", "Ы", "Ь"}
    assert excluded.isdisjoint(APPENDIX_LETTERS)
    assert APPENDIX_LETTERS[0] == "А"
    assert "Б" in APPENDIX_LETTERS


def test_parser_recognises_appendix_h1_with_letter_type_and_subtitle():
    result = V1Parser().parse(
        "\n".join(
            [
                "[H1] ПРИЛОЖЕНИЕ А",
                "Справочное",
                "Исходные данные эксперимента",
                "Содержание приложения.",
            ]
        )
    )

    assert result.diagnostics == []
    appendix, paragraph = result.document.blocks
    assert isinstance(appendix, AppendixNode)
    assert appendix.letter == "А"
    assert appendix.appendix_type == "справочное"
    assert appendix.subtitle == "Исходные данные эксперимента"
    assert appendix.title == "ПРИЛОЖЕНИЕ А"
    assert appendix.id == "app:а"
    assert appendix.source.line_start == 1
    assert appendix.source.line_end == 3
    assert isinstance(paragraph, ParagraphNode)
    assert paragraph.runs[0].text == "Содержание приложения."


def test_parser_appendix_without_type_or_subtitle_consumes_only_heading():
    result = V1Parser().parse(
        "\n".join(
            [
                "[H1] ПРИЛОЖЕНИЕ Б",
                "Тело текста приложения, не подзаголовок.",
            ]
        )
    )

    assert result.diagnostics == []
    appendix, paragraph = result.document.blocks
    assert isinstance(appendix, AppendixNode)
    assert appendix.letter == "Б"
    assert appendix.appendix_type is None
    assert appendix.subtitle is None
    assert isinstance(paragraph, ParagraphNode)


def test_parser_excludes_letters_not_allowed_for_appendix_designation():
    """Letters Ё, З, Й, О, Ч, Ъ, Ы, Ь fall back to a regular H1 heading."""

    result = V1Parser().parse("[H1] ПРИЛОЖЕНИЕ З")

    assert result.diagnostics == []
    (heading,) = result.document.blocks
    assert isinstance(heading, HeadingNode)
    assert heading.text == "ПРИЛОЖЕНИЕ З"


def test_parser_recognises_toc_marker():
    result = V1Parser().parse('[TOC levels=2 title="ОГЛАВЛЕНИЕ"]')

    assert result.diagnostics == []
    (toc,) = result.document.blocks
    assert isinstance(toc, TableOfContentsNode)
    assert toc.levels == 2
    assert toc.title == "ОГЛАВЛЕНИЕ"


def test_parser_default_toc_marker_uses_three_levels_and_default_title():
    result = V1Parser().parse("[TOC]")

    assert result.diagnostics == []
    (toc,) = result.document.blocks
    assert toc.levels == 3
    assert toc.title == "СОДЕРЖАНИЕ"


def test_parser_recognises_metadata_marker_and_collects_metadata():
    source = "\n".join(
        [
            '[META key=title value="Отчёт"]',
            '[META key=year value="2026"]',
            "Текст",
        ]
    )

    result = V1Parser().parse(source)

    assert result.diagnostics == []
    assert result.document.metadata["title"] == "Отчёт"
    assert result.document.metadata["year"] == "2026"


def test_parser_meta_marker_without_key_emits_diagnostic():
    result = V1Parser().parse('[META value="Без ключа"]')

    assert any(
        diag.code == "TXT_MALFORMED_ATTRIBUTE" for diag in result.diagnostics
    )


def test_parser_recognises_title_page_marker():
    from sfu_converter.domain.ast_nodes import TitlePageNode

    result = V1Parser().parse("[TITLE_PAGE]")

    assert result.diagnostics == []
    (title_page,) = result.document.blocks
    assert isinstance(title_page, TitlePageNode)
    assert title_page.profile is None


def test_parser_title_page_marker_accepts_profile_attribute():
    from sfu_converter.domain.ast_nodes import TitlePageNode

    result = V1Parser().parse("[TITLE_PAGE profile=lab_practical_project_reports]")

    assert result.diagnostics == []
    (title_page,) = result.document.blocks
    assert isinstance(title_page, TitlePageNode)
    assert title_page.profile == "lab_practical_project_reports"

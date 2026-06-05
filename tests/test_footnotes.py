from __future__ import annotations

from zipfile import ZipFile

from docx import Document as DocxDocument
from docx.shared import Pt

from sfu_converter.domain import ast_nodes
from sfu_converter.domain.ast_nodes import Document, ParagraphNode, TextRun
from sfu_converter.domain.diagnostics import DiagnosticCodes
from sfu_converter.infrastructure import docx_styles
from sfu_converter.infrastructure.docx_renderer import DocxRenderer
from sfu_converter.infrastructure.docx_validator import DocxValidator
from sfu_converter.parser import V2Parser
from sfu_converter.registry import get_profile


def _footnote_node():
    footnote_node = getattr(ast_nodes, "FootnoteNode", None)
    assert footnote_node is not None, "FootnoteNode is missing"
    return footnote_node


def test_v2_parser_parses_inline_footnote_anchor_and_body():
    result = V2Parser().parse('[P] Текст [FN id=1 text="Источник"] после.')

    assert result.diagnostics == []
    paragraph, footnote = result.document.blocks
    assert isinstance(paragraph, ParagraphNode)
    assert [type(run).__name__ for run in paragraph.runs] == ["TextRun", "FootnoteAnchor", "TextRun"]
    assert "".join(getattr(run, "text", "") for run in paragraph.runs) == "Текст  после."
    assert paragraph.runs[1].marker == "1"
    assert isinstance(footnote, _footnote_node())
    assert footnote.marker == "1"
    assert footnote.text == "Источник"


def test_docx_renderer_writes_footnote_reference_and_footnotes_part(tmp_path):
    renderer = DocxRenderer(base_dir=tmp_path)
    output_path = tmp_path / "footnote.docx"
    ast = Document(
        blocks=(
            ParagraphNode(runs=(TextRun("Текст "), TextRun(" после."))),
            _footnote_node()(marker="1", text="Источник"),
        )
    )
    anchor = V2Parser().parse('[P] Текст [FN_ANCHOR id=1] после.').document.blocks[0]
    ast = Document(blocks=(anchor, ast.blocks[1]))

    renderer.render_to_file(ast, get_profile("common"), str(output_path))

    with ZipFile(output_path) as package:
        names = set(package.namelist())
        document_xml = package.read("word/document.xml").decode("utf-8")
        footnotes_xml = package.read("word/footnotes.xml").decode("utf-8")

    assert "word/footnotes.xml" in names
    assert "<w:footnoteReference" in document_xml
    assert ">1<" in footnotes_xml
    assert "Источник" in footnotes_xml


def test_duplicate_footnote_anchors_produce_diagnostic(tmp_path):
    renderer = DocxRenderer(base_dir=tmp_path)
    output_path = tmp_path / "duplicate.docx"
    first = V2Parser().parse('[P] Один [FN_ANCHOR id=1]').document.blocks[0]
    second = V2Parser().parse('[P] Два [FN_ANCHOR id=1]').document.blocks[0]
    ast = Document(blocks=(first, second, _footnote_node()(marker="1", text="Источник")))

    diagnostics = renderer.render_to_file(ast, get_profile("common"), str(output_path))

    assert any(
        diagnostic.code == DiagnosticCodes.FOOTNOTE_DUPLICATE
        and diagnostic.rule_id == "common.reference.footnote"
        for diagnostic in diagnostics
    )


def test_docx_validator_reports_oversized_footnote_text(tmp_path):
    doc = DocxDocument()
    docx_styles.register_styles(doc)
    paragraph = doc.add_paragraph("1 Источник")
    paragraph.style = doc.styles[docx_styles.FOOTNOTE_TEXT]
    paragraph.paragraph_format.line_spacing = 1.0
    paragraph.runs[0].font.name = "Times New Roman"
    paragraph.runs[0].font.size = Pt(16)
    path = tmp_path / "bad_footnote.docx"
    doc.save(str(path))

    diagnostics = DocxValidator(get_profile("common")).validate_file(str(path))

    assert any(
        diagnostic.code == DiagnosticCodes.FOOTNOTE_FORMAT
        and diagnostic.rule_id == "common.reference.footnote"
        for diagnostic in diagnostics
    )

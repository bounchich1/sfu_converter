from docx import Document as DocxDocument
from docx.enum.text import WD_ALIGN_PARAGRAPH

import pytest

from sfu_converter.config import SIBFUConfig
from sfu_converter.domain.ast_nodes import (
    Document,
    HeadingLevel,
    HeadingNode,
    TitlePageNode,
)
from sfu_converter.domain.diagnostics import DiagnosticCodes
from sfu_converter.domain.formatting import FormattingProfile
from sfu_converter.infrastructure.docx_renderer import DocxRenderer
from sfu_converter.infrastructure.title_pages import (
    FORM_REGISTRY,
    GENERIC_FORM,
    select_title_page_form,
)
from sfu_converter.registry import get_profile


def _profile():
    return FormattingProfile(
        name="common",
        display_name="Common",
        source_docs=("standard",),
    )


def test_title_page_renders_full_metadata_with_centered_layout(tmp_path):
    metadata = {
        "university": "Сибирский федеральный университет",
        "institute": "Институт космических и информационных технологий",
        "department": "Кафедра вычислительной техники",
        "title": "Отчёт по лабораторной работе №1",
        "subject": "Программирование",
        "student": "Иванов И.И.",
        "group": "КИ22-01Б",
        "supervisor": "Петров П.П.",
        "supervisor_title": "доцент, канд. техн. наук",
        "city": "Красноярск",
        "year": "2026",
    }
    ast = Document(
        blocks=(TitlePageNode(),),
        metadata=metadata,
    )
    output = tmp_path / "title.docx"

    renderer = DocxRenderer(config_class=SIBFUConfig, base_dir=tmp_path)
    renderer.render_to_file(ast, _profile(), str(output))

    doc = DocxDocument(str(output))
    texts = [p.text for p in doc.paragraphs]
    body_xml = doc.element.body.xml

    assert any("МИНИСТЕРСТВО" in t for t in texts)
    assert any("СИБИРСКИЙ ФЕДЕРАЛЬНЫЙ УНИВЕРСИТЕТ" in t for t in texts)
    assert "Институт космических и информационных технологий" in texts
    assert "Кафедра вычислительной техники" in texts
    assert "Отчёт по лабораторной работе №1" in texts
    assert any("Иванов И.И." in t and "Студент" in t for t in texts)
    assert any("Петров П.П." in t and "Руководитель" in t for t in texts)
    assert any("Красноярск" in t and "2026" in t for t in texts)
    assert 'w:type="page"' in body_xml


def test_title_page_renders_with_minimal_metadata(tmp_path):
    metadata = {"title": "Курсовая работа"}
    ast = Document(blocks=(TitlePageNode(),), metadata=metadata)
    output = tmp_path / "title_min.docx"

    renderer = DocxRenderer(config_class=SIBFUConfig, base_dir=tmp_path)
    renderer.render_to_file(ast, _profile(), str(output))

    doc = DocxDocument(str(output))
    texts = [p.text for p in doc.paragraphs]
    assert "Курсовая работа" in texts
    # Default city should always appear
    assert any("Красноярск" in t for t in texts)


def test_title_page_paragraphs_are_centered_and_use_sfu_font(tmp_path):
    ast = Document(
        blocks=(TitlePageNode(),),
        metadata={"title": "Тест", "year": "2026"},
    )
    output = tmp_path / "title_font.docx"

    renderer = DocxRenderer(config_class=SIBFUConfig, base_dir=tmp_path)
    renderer.render_to_file(ast, _profile(), str(output))

    doc = DocxDocument(str(output))
    title_para = next(p for p in doc.paragraphs if p.text == "Тест")
    assert title_para.paragraph_format.alignment == WD_ALIGN_PARAGRAPH.CENTER
    assert title_para.runs[0].font.name == SIBFUConfig.FONT_NAME
    assert title_para.runs[0].bold is True


def test_title_page_skipped_when_template_mode_is_preserve_prefix(tmp_path):
    ast = Document(
        blocks=(
            TitlePageNode(),
            HeadingNode(level=HeadingLevel.H1, text="Раздел", number="auto"),
        ),
        metadata={"title": "Не должно появиться"},
    )
    output = tmp_path / "title_skipped.docx"

    renderer = DocxRenderer(config_class=SIBFUConfig, base_dir=tmp_path)
    renderer.render_to_file(
        ast,
        _profile(),
        str(output),
        template_mode="preserve-prefix",
    )

    doc = DocxDocument(str(output))
    texts = [p.text for p in doc.paragraphs]
    assert "Не должно появиться" not in texts
    # The body heading still renders
    assert any("Раздел" in t for t in texts)


def test_title_page_first_page_footer_remains_blank(tmp_path):
    ast = Document(
        blocks=(TitlePageNode(),),
        metadata={"title": "Документ"},
    )
    output = tmp_path / "title_footer.docx"

    renderer = DocxRenderer(config_class=SIBFUConfig, base_dir=tmp_path)
    renderer.render_to_file(ast, _profile(), str(output))

    doc = DocxDocument(str(output))
    section = doc.sections[0]
    assert section.different_first_page_header_footer is True
    first_page_footer = section.first_page_footer
    assert first_page_footer.paragraphs[0].text == ""
    assert " PAGE " not in first_page_footer._element.xml


def test_lab_report_title_page_uses_profile_form_m(tmp_path):
    metadata = {
        "institute": "Институт инженерной физики",
        "department": "Кафедра физики",
        "document_type": "ОТЧЕТ ПО ЛАБОРАТОРНОЙ РАБОТЕ",
        "title": "Измерение сопротивления",
        "teacher": "Петров П.П.",
        "student": "Иванов И.И.",
        "group": "ФИ22-01Б",
        "record_book": "123456",
        "city": "Красноярск",
        "year": "2026",
    }
    ast = Document(blocks=(TitlePageNode(),), metadata=metadata)
    output = tmp_path / "lab_title.docx"

    renderer = DocxRenderer(config_class=SIBFUConfig, base_dir=tmp_path)
    diagnostics = renderer.render_to_file(ast, get_profile("lab_practical_project_reports"), str(output))

    assert not any(diagnostic.code == DiagnosticCodes.TXT_MISSING_METADATA for diagnostic in diagnostics)
    doc = DocxDocument(str(output))
    texts = [p.text for p in doc.paragraphs]
    assert "ОТЧЕТ ПО ЛАБОРАТОРНОЙ РАБОТЕ" in texts
    assert "на тему: Измерение сопротивления" in texts
    assert any("Преподаватель" in text and "Петров П.П." in text for text in texts)
    assert any("Студент" in text and "ФИ22-01Б" in text and "123456" in text for text in texts)


def test_profile_title_pages_render_distinct_document_labels(tmp_path):
    expectations = {
        "practice_reports": "ОТЧЕТ О ПРАКТИКЕ",
        "research_reports": "ОТЧЕТ О НАУЧНО-ИССЛЕДОВАТЕЛЬСКОЙ РАБОТЕ",
        "small_written_works": "РЕФЕРАТ",
        "coursework": "КУРСОВАЯ РАБОТА",
    }

    for profile_name, expected_label in expectations.items():
        output = tmp_path / f"{profile_name}.docx"
        ast = Document(
            blocks=(TitlePageNode(),),
            metadata={
                "title": "Тема",
                "student": "Иванов И.И.",
                "group": "КИ22-01Б",
                "supervisor": "Петров П.П.",
                "teacher": "Сидоров С.С.",
                "city": "Красноярск",
                "year": "2026",
            },
        )

        renderer = DocxRenderer(config_class=SIBFUConfig, base_dir=tmp_path)
        renderer.render_to_file(ast, get_profile(profile_name), str(output))

        doc = DocxDocument(str(output))
        texts = [p.text for p in doc.paragraphs]
        assert expected_label in texts


def test_missing_title_page_metadata_reports_diagnostic(tmp_path):
    ast = Document(blocks=(TitlePageNode(),), metadata={"title": "Без студента"})
    output = tmp_path / "missing_metadata.docx"

    renderer = DocxRenderer(config_class=SIBFUConfig, base_dir=tmp_path)
    diagnostics = renderer.render_to_file(ast, get_profile("coursework"), str(output))

    missing = [
        diagnostic
        for diagnostic in diagnostics
        if diagnostic.code == DiagnosticCodes.TXT_MISSING_METADATA
    ]
    assert missing
    assert missing[0].rule_id == "form_i"
    assert "student" in missing[0].message
    assert output.exists()


def test_title_page_node_profile_overrides_existing_document_profile(tmp_path):
    ast = Document(
        blocks=(TitlePageNode(profile="coursework"),),
        metadata={
            "title": "Профиль из узла",
            "student": "Иванов И.И.",
            "group": "КИ22-01Б",
            "supervisor": "Петров П.П.",
            "year": "2026",
        },
    )
    output = tmp_path / "override.docx"

    renderer = DocxRenderer(config_class=SIBFUConfig, base_dir=tmp_path)
    renderer.render_to_file(ast, get_profile("common"), str(output))

    doc = DocxDocument(str(output))
    assert "КУРСОВАЯ РАБОТА" in [p.text for p in doc.paragraphs]


def test_unknown_title_page_node_profile_falls_back_to_document_profile(tmp_path):
    ast = Document(
        blocks=(TitlePageNode(profile="missing_profile"),),
        metadata={
            "title": "Общий титульный лист",
            "student": "Иванов И.И.",
            "year": "2026",
        },
    )
    output = tmp_path / "missing_override.docx"

    renderer = DocxRenderer(config_class=SIBFUConfig, base_dir=tmp_path)
    renderer.render_to_file(ast, get_profile("common"), str(output))

    doc = DocxDocument(str(output))
    texts = [p.text for p in doc.paragraphs]
    assert "Общий титульный лист" in texts
    assert "КУРСОВАЯ РАБОТА" not in texts


def _vkr_metadata(extra: dict | None = None) -> dict:
    base = {
        "institute": "Институт космических и информационных технологий",
        "department": "Кафедра ВТ",
        "title": "Тема выпускной работы",
        "student": "Иванов И.И.",
        "supervisor": "Петров П.П.",
        "supervisor_title": "доцент",
        "department_head": "Сидоров С.С.",
        "approval_date": "«25» мая 2026 г.",
        "city": "Красноярск",
        "year": "2026",
    }
    if extra:
        base.update(extra)
    return base


def test_form_b_master_dissertation_renders_required_strings(tmp_path):
    metadata = _vkr_metadata({
        "direction_code": "09.04.01",
        "direction_name": "Информатика и вычислительная техника",
        "master_program_code": "09.04.01-01",
        "master_program_name": "Высокопроизводительные вычисления",
        "reviewer": "Кузнецов К.К.",
    })
    ast = Document(
        blocks=(TitlePageNode(profile="form_b"),),
        metadata=metadata,
    )
    output = tmp_path / "form_b.docx"

    renderer = DocxRenderer(config_class=SIBFUConfig, base_dir=tmp_path)
    diagnostics = renderer.render_to_file(ast, get_profile("graduation_qualification_work"), str(output))

    assert not [d for d in diagnostics if d.code == DiagnosticCodes.TXT_MISSING_METADATA]
    texts = [p.text for p in DocxDocument(str(output)).paragraphs]
    assert "МАГИСТЕРСКАЯ ДИССЕРТАЦИЯ" in texts
    assert any("по направлению 09.04.01" in t for t in texts)
    assert any("по магистерской программе 09.04.01-01" in t for t in texts)
    assert any("Научный руководитель" in t and "Петров П.П." in t for t in texts)
    assert any("Рецензент" in t and "Кузнецов К.К." in t for t in texts)
    assert any("Заведующий кафедрой" in t for t in texts)


def test_form_v_diploma_project_includes_explanatory_subtitle(tmp_path):
    metadata = _vkr_metadata({
        "specialty_code": "10.05.03",
        "specialty_name": "Информационная безопасность",
        "consultants": "Иванова А.А.; Сидорова Б.Б.",
        "norm_controller": "Контролёров Н.Н.",
    })
    ast = Document(blocks=(TitlePageNode(profile="form_v"),), metadata=metadata)
    output = tmp_path / "form_v.docx"

    renderer = DocxRenderer(config_class=SIBFUConfig, base_dir=tmp_path)
    renderer.render_to_file(ast, get_profile("graduation_qualification_work"), str(output))

    texts = [p.text for p in DocxDocument(str(output)).paragraphs]
    assert "ДИПЛОМНЫЙ ПРОЕКТ" in texts
    assert "Пояснительная записка" in texts
    assert any("по специальности 10.05.03" in t for t in texts)
    assert any("Консультант" in t and "Иванова А.А." in t for t in texts)
    assert any("Консультант" in t and "Сидорова Б.Б." in t for t in texts)
    assert any("Нормоконтролёр" in t for t in texts)
    assert any("Дипломник" in t and "Иванов И.И." in t for t in texts)


def test_form_g_diploma_work_has_reviewer(tmp_path):
    metadata = _vkr_metadata({"reviewer": "Рецензентов Р.Р."})
    ast = Document(blocks=(TitlePageNode(profile="form_g"),), metadata=metadata)
    output = tmp_path / "form_g.docx"

    renderer = DocxRenderer(config_class=SIBFUConfig, base_dir=tmp_path)
    renderer.render_to_file(ast, get_profile("graduation_qualification_work"), str(output))

    texts = [p.text for p in DocxDocument(str(output)).paragraphs]
    assert "ДИПЛОМНАЯ РАБОТА" in texts
    assert any("Рецензент" in t and "Рецензентов Р.Р." in t for t in texts)
    assert any("Выпускник" in t and "Иванов И.И." in t for t in texts)


def test_form_d_bachelor_work_has_no_reviewer_by_default(tmp_path):
    metadata = _vkr_metadata()
    ast = Document(blocks=(TitlePageNode(profile="form_d"),), metadata=metadata)
    output = tmp_path / "form_d.docx"

    renderer = DocxRenderer(config_class=SIBFUConfig, base_dir=tmp_path)
    renderer.render_to_file(ast, get_profile("graduation_qualification_work"), str(output))

    texts = [p.text for p in DocxDocument(str(output)).paragraphs]
    assert "БАКАЛАВРСКАЯ РАБОТА" in texts
    assert not any("Рецензент" in t for t in texts)


def test_form_e_continuation_lists_consultants_and_norm_controller(tmp_path):
    metadata = {
        "title": "Продолжение титульного листа",
        "consultants": "Иванова А.А.; Сидорова Б.Б.",
        "norm_controller": "Контролёров Н.Н.",
    }
    ast = Document(blocks=(TitlePageNode(profile="form_e"),), metadata=metadata)
    output = tmp_path / "form_e.docx"

    renderer = DocxRenderer(config_class=SIBFUConfig, base_dir=tmp_path)
    renderer.render_to_file(ast, get_profile("graduation_qualification_work"), str(output))

    texts = [p.text for p in DocxDocument(str(output)).paragraphs]
    assert "Консультанты по разделам:" in texts
    assert any("Иванова А.А." in t for t in texts)
    assert any("Нормоконтролёр" in t and "Контролёров Н.Н." in t for t in texts)


def test_form_zh_combined_single_page(tmp_path):
    metadata = _vkr_metadata({
        "reviewer": "Рецензентов Р.Р.",
        "consultants": "Иванова А.А.",
        "norm_controller": "Контролёров Н.Н.",
    })
    ast = Document(blocks=(TitlePageNode(profile="form_zh"),), metadata=metadata)
    output = tmp_path / "form_zh.docx"

    renderer = DocxRenderer(config_class=SIBFUConfig, base_dir=tmp_path)
    renderer.render_to_file(ast, get_profile("graduation_qualification_work"), str(output))

    texts = [p.text for p in DocxDocument(str(output)).paragraphs]
    assert any("Рецензент" in t for t in texts)
    assert any("Консультант" in t for t in texts)
    assert any("Нормоконтролёр" in t for t in texts)
    assert any("Выпускник" in t for t in texts)


def test_form_k_practice_report_uses_practice_kind(tmp_path):
    metadata = {
        "document_type": "ОТЧЕТ О ПРОИЗВОДСТВЕННОЙ ПРАКТИКЕ",
        "title": "Практика на ПАО",
        "practice_place": "ПАО «Завод»",
        "university_supervisor": "Петров П.П.",
        "enterprise_supervisor": "Заводов З.З.",
        "student": "Иванов И.И.",
        "group": "КИ22-01Б",
        "record_book": "654321",
        "city": "Красноярск",
        "year": "2026",
    }
    ast = Document(blocks=(TitlePageNode(),), metadata=metadata)
    output = tmp_path / "form_k.docx"

    renderer = DocxRenderer(config_class=SIBFUConfig, base_dir=tmp_path)
    renderer.render_to_file(ast, get_profile("practice_reports"), str(output))

    texts = [p.text for p in DocxDocument(str(output)).paragraphs]
    assert "ОТЧЕТ О ПРОИЗВОДСТВЕННОЙ ПРАКТИКЕ" in texts
    assert any("Место прохождения практики" in t for t in texts)
    assert any("Руководитель от университета" in t and "Петров П.П." in t for t in texts)
    assert any("Руководитель от предприятия" in t and "Заводов З.З." in t for t in texts)


def test_form_l_master_research_report_has_program_head_block(tmp_path):
    metadata = {
        "title": "Исследование",
        "master_program": "09.04.01-01 Высокопроизводительные вычисления",
        "program_head": "Главов Г.Г.",
        "approval_date": "«25» мая 2026 г.",
        "supervisor": "Петров П.П.",
        "student": "Иванов И.И.",
        "group": "КИ22-01М",
        "city": "Красноярск",
        "year": "2026",
    }
    ast = Document(blocks=(TitlePageNode(),), metadata=metadata)
    output = tmp_path / "form_l.docx"

    renderer = DocxRenderer(config_class=SIBFUConfig, base_dir=tmp_path)
    renderer.render_to_file(ast, get_profile("research_reports"), str(output))

    texts = [p.text for p in DocxDocument(str(output)).paragraphs]
    assert "ОТЧЕТ О НАУЧНО-ИССЛЕДОВАТЕЛЬСКОЙ РАБОТЕ" in texts
    assert "УТВЕРЖДАЮ" in texts
    assert any("Главов Г.Г." in t for t in texts)


def test_form_n_referat_uses_discipline_subtitle(tmp_path):
    metadata = {
        "document_type": "РЕФЕРАТ",
        "discipline": "Архитектура ЭВМ",
        "title": "Кэширование",
        "teacher": "Сидоров С.С.",
        "student": "Иванов И.И.",
        "group": "КИ22-01Б",
        "city": "Красноярск",
        "year": "2026",
    }
    ast = Document(blocks=(TitlePageNode(),), metadata=metadata)
    output = tmp_path / "form_n.docx"

    renderer = DocxRenderer(config_class=SIBFUConfig, base_dir=tmp_path)
    renderer.render_to_file(ast, get_profile("small_written_works"), str(output))

    texts = [p.text for p in DocxDocument(str(output)).paragraphs]
    assert "РЕФЕРАТ" in texts
    assert any("по дисциплине: Архитектура ЭВМ" in t for t in texts)
    assert any("на тему: Кэширование" in t for t in texts)


@pytest.mark.parametrize(
    "form_id,required_field",
    [
        ("form_b", "title"),
        ("form_v", "supervisor"),
        ("form_g", "student"),
        ("form_d", "title"),
        ("form_i", "group"),
        ("form_k", "title"),
        ("form_m", "teacher"),
        ("form_n", "student"),
    ],
)
def test_form_reports_missing_required_field(tmp_path, form_id, required_field):
    metadata = {
        "title": "X",
        "student": "S",
        "supervisor": "Sup",
        "group": "G",
        "teacher": "T",
    }
    metadata[required_field] = ""
    ast = Document(blocks=(TitlePageNode(profile=form_id),), metadata=metadata)
    output = tmp_path / f"missing_{form_id}.docx"

    renderer = DocxRenderer(config_class=SIBFUConfig, base_dir=tmp_path)
    diagnostics = renderer.render_to_file(ast, get_profile("common"), str(output))

    missing = [d for d in diagnostics if d.code == DiagnosticCodes.TXT_MISSING_METADATA and d.rule_id == form_id]
    assert any(d.data and d.data.get("field") == required_field for d in missing)
    assert output.exists()


def test_select_title_page_form_returns_generic_for_common_profile():
    form = select_title_page_form(get_profile("common"), {})
    assert form is GENERIC_FORM


def test_select_title_page_form_resolves_profile_form():
    form = select_title_page_form(get_profile("graduation_qualification_work"), {})
    assert form is FORM_REGISTRY["form_b"]


def test_select_title_page_form_override_uses_named_form():
    form = select_title_page_form(get_profile("common"), {}, override="form_n")
    assert form.form_id == "form_n"


def test_form_i_node_override_inside_lab_profile(tmp_path):
    metadata = {
        "title": "Курсовая внутри лабы",
        "student": "Иванов И.И.",
        "group": "КИ22-01Б",
        "supervisor": "Петров П.П.",
        "teacher": "Сидоров С.С.",
    }
    ast = Document(blocks=(TitlePageNode(profile="form_i"),), metadata=metadata)
    output = tmp_path / "override_form_i.docx"

    renderer = DocxRenderer(config_class=SIBFUConfig, base_dir=tmp_path)
    renderer.render_to_file(ast, get_profile("lab_practical_project_reports"), str(output))

    texts = [p.text for p in DocxDocument(str(output)).paragraphs]
    assert "КУРСОВАЯ РАБОТА" in texts
    assert "ОТЧЕТ ПО ЛАБОРАТОРНОЙ РАБОТЕ" not in texts


def test_first_page_excluded_from_page_numbering(tmp_path):
    ast = Document(blocks=(TitlePageNode(),), metadata={"title": "Документ"})
    output = tmp_path / "first_page.docx"

    renderer = DocxRenderer(config_class=SIBFUConfig, base_dir=tmp_path)
    renderer.render_to_file(ast, get_profile("graduation_qualification_work"), str(output))

    doc = DocxDocument(str(output))
    section = doc.sections[0]
    assert section.different_first_page_header_footer is True
    first_page_footer = section.first_page_footer
    assert " PAGE " not in first_page_footer._element.xml

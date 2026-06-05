import json

from docx import Document

from sfu_converter import cli
from sfu_converter.registry import PROFILES


def test_create_parser_parses_convert_command(tmp_path):
    parser = cli.create_parser()

    args = parser.parse_args(
        [
            "--format",
            "json",
            "--workdir",
            str(tmp_path),
            "convert",
            "--input",
            "input.txt",
            "--output",
            "output.docx",
            "--strict",
        ]
    )

    assert args.command == "convert"
    assert args.format == "json"
    assert args.input.name == "input.txt"
    assert args.output.name == "output.docx"
    assert args.strict is True


def test_convert_command_writes_docx_and_json_result(tmp_path, capsys):
    input_file = tmp_path / "report.txt"
    input_file.write_text("[H1] Report title\n\nBody text", encoding="utf-8")

    exit_code = cli.main(
        [
            "--format",
            "json",
            "--workdir",
            str(tmp_path),
            "convert",
            "--input",
            "report.txt",
            "--output",
            "out/report.docx",
        ]
    )

    captured = capsys.readouterr()
    payload = json.loads(captured.out)
    output_file = tmp_path / "out" / "report.docx"

    assert exit_code == cli.ExitCodes.SUCCESS
    assert output_file.exists()
    assert payload["ok"] is True
    assert payload["command"] == "convert"
    assert payload["diagnostics"] == []
    assert payload["outputs"]["docx"] == str(output_file)

    doc = Document(str(output_file))
    assert [para.text for para in doc.paragraphs if para.text] == [
        "1 Report title",
        "Body text",
    ]


def test_convert_command_passes_selected_profile_to_converter(tmp_path, capsys, monkeypatch):
    input_file = tmp_path / "report.txt"
    input_file.write_text("Body text", encoding="utf-8")
    calls = []

    def fake_convert_file(self, input_path, output_path, **kwargs):
        calls.append(kwargs["profile"].name)
        self.diagnostics = []
        output_path.parent.mkdir(parents=True, exist_ok=True)
        output_path.write_bytes(b"fake")
        return str(output_path)

    monkeypatch.setattr(
        "sfu_converter.converter.TextToDocxConverter.convert_file",
        fake_convert_file,
    )

    exit_code = cli.main(
        [
            "--format",
            "json",
            "--workdir",
            str(tmp_path),
            "convert",
            "--input",
            "report.txt",
            "--output",
            "out/report.docx",
            "--profile",
            "research_reports",
        ]
    )

    captured = capsys.readouterr()
    payload = json.loads(captured.out)

    assert exit_code == cli.ExitCodes.SUCCESS
    assert calls == ["research_reports"]
    assert payload["profile"] == "research_reports"


def test_convert_command_uses_v2_parser_when_requested(tmp_path, capsys):
    input_file = tmp_path / "report-v2.txt"
    input_file.write_text(
        "\n".join(
            [
                "[DOC syntax=2]",
                '[H level=1 title="Report title" number=auto]',
                "[P] Body text",
            ]
        ),
        encoding="utf-8",
    )

    exit_code = cli.main(
        [
            "--format",
            "json",
            "--workdir",
            str(tmp_path),
            "convert",
            "--input",
            "report-v2.txt",
            "--output",
            "out/report-v2.docx",
            "--syntax-version",
            "2",
        ]
    )

    captured = capsys.readouterr()
    payload = json.loads(captured.out)
    output_file = tmp_path / "out" / "report-v2.docx"

    assert exit_code == cli.ExitCodes.SUCCESS
    assert payload["syntaxVersion"] == 2

    doc = Document(str(output_file))
    assert [para.text for para in doc.paragraphs if para.text] == [
        "1 Report title",
        "Body text",
    ]


def test_convert_accepts_common_options_after_subcommand(tmp_path, capsys):
    input_file = tmp_path / "report.txt"
    input_file.write_text("Body text", encoding="utf-8")

    exit_code = cli.main(
        [
            "convert",
            "--input",
            str(input_file),
            "--output",
            str(tmp_path / "report.docx"),
            "--format",
            "json",
        ]
    )

    captured = capsys.readouterr()
    payload = json.loads(captured.out)

    assert exit_code == cli.ExitCodes.SUCCESS
    assert payload["ok"] is True
    assert payload["command"] == "convert"


def test_convert_command_missing_input_returns_missing_resource(tmp_path, capsys):
    exit_code = cli.main(
        [
            "--format",
            "json",
            "--workdir",
            str(tmp_path),
            "convert",
            "--input",
            "missing.txt",
            "--output",
            "out.docx",
        ]
    )

    captured = capsys.readouterr()
    payload = json.loads(captured.out)

    assert exit_code == cli.ExitCodes.MISSING_RESOURCE
    assert payload["ok"] is False
    assert payload["command"] == "convert"
    assert payload["diagnostics"][0]["code"] == "MISSING_INPUT"


def test_validate_docx_command_reports_valid_document(tmp_path, capsys):
    input_file = tmp_path / "report.txt"
    output_file = tmp_path / "report.docx"
    input_file.write_text("[H1] Report title\n\nBody text", encoding="utf-8")
    assert (
        cli.main(
            [
                "--workdir",
                str(tmp_path),
                "convert",
                "--input",
                input_file.name,
                "--output",
                output_file.name,
            ]
        )
        == cli.ExitCodes.SUCCESS
    )

    exit_code = cli.main(
        [
            "--format",
            "json",
            "--workdir",
            str(tmp_path),
            "validate-docx",
            "--input",
            output_file.name,
        ]
    )

    captured = capsys.readouterr()
    payload = json.loads(captured.out.splitlines()[-1])

    assert exit_code == cli.ExitCodes.SUCCESS
    assert payload["ok"] is True
    assert payload["command"] == "validate-docx"
    assert all(diagnostic["severity"] == "warning" for diagnostic in payload["diagnostics"])


def test_validate_docx_command_uses_selected_profile(tmp_path, capsys, monkeypatch):
    input_file = tmp_path / "report.docx"
    input_file.write_bytes(b"placeholder")
    calls = []

    class FakeStyleValidator:
        def __init__(self, config_class, *, profile=None, profile_name=None):
            calls.append(profile.name if profile is not None else profile_name)

        def validate_file(self, file_path):
            return True

        def get_report(self):
            return {"diagnostics": []}

    monkeypatch.setattr("sfu_converter.validator.StyleValidator", FakeStyleValidator)

    exit_code = cli.main(
        [
            "--format",
            "json",
            "--workdir",
            str(tmp_path),
            "validate-docx",
            "--input",
            "report.docx",
            "--profile",
            "coursework",
        ]
    )

    captured = capsys.readouterr()
    payload = json.loads(captured.out)

    assert exit_code == cli.ExitCodes.SUCCESS
    assert calls == ["coursework"]
    assert payload["inputs"]["profile"] == "coursework"


def test_unknown_profile_returns_missing_profile_diagnostic(tmp_path, capsys):
    input_file = tmp_path / "report.txt"
    input_file.write_text("Body text", encoding="utf-8")

    exit_code = cli.main(
        [
            "--format",
            "json",
            "--workdir",
            str(tmp_path),
            "convert",
            "--input",
            "report.txt",
            "--output",
            "out/report.docx",
            "--profile",
            "missing_profile",
        ]
    )

    captured = capsys.readouterr()
    payload = json.loads(captured.out)

    assert exit_code == cli.ExitCodes.MISSING_RESOURCE
    assert payload["ok"] is False
    diagnostic = payload["diagnostics"][0]
    assert diagnostic["code"] == "MISSING_PROFILE"
    assert diagnostic["ruleId"] is None
    assert diagnostic["source"] == {
        "document": None,
        "section": None,
        "lineStart": None,
        "lineEnd": None,
    }


def test_validate_docx_command_missing_input_returns_missing_resource(tmp_path, capsys):
    exit_code = cli.main(
        [
            "--format",
            "json",
            "--workdir",
            str(tmp_path),
            "validate-docx",
            "--input",
            "missing.docx",
        ]
    )

    captured = capsys.readouterr()
    payload = json.loads(captured.out)

    assert exit_code == cli.ExitCodes.MISSING_RESOURCE
    assert payload["diagnostics"][0]["code"] == "MISSING_INPUT"


def test_interactive_command_launches_legacy_menu(monkeypatch):
    calls = []

    def fake_legacy_main():
        calls.append("called")

    monkeypatch.setattr("sfu_converter.main.main", fake_legacy_main)

    exit_code = cli.main(["interactive"])

    assert exit_code == cli.ExitCodes.SUCCESS
    assert calls == ["called"]


def test_parse_command_outputs_json_ast(tmp_path, capsys):
    input_file = tmp_path / "report.txt"
    input_file.write_text("[H1] Title\n\nBody text", encoding="utf-8")

    exit_code = cli.main(
        [
            "--format",
            "json",
            "--workdir",
            str(tmp_path),
            "parse",
            "--input",
            "report.txt",
        ]
    )

    captured = capsys.readouterr()
    payload = json.loads(captured.out)

    assert exit_code == cli.ExitCodes.SUCCESS
    assert payload["ok"] is True
    assert payload["command"] == "parse"
    assert payload["syntaxVersion"] == 1
    assert payload["profile"] == "common"
    assert payload["ast"]["type"] == "DOCUMENT"
    assert payload["ast"]["blocks"][0]["type"] == "HEADING"
    assert payload["diagnostics"] == []


def test_parse_unknown_syntax_version_returns_structured_diagnostic(tmp_path, capsys):
    input_file = tmp_path / "report.txt"
    input_file.write_text("Body text", encoding="utf-8")

    exit_code = cli.main(
        [
            "--format",
            "json",
            "--workdir",
            str(tmp_path),
            "parse",
            "--input",
            "report.txt",
            "--syntax-version",
            "99",
        ]
    )

    captured = capsys.readouterr()
    payload = json.loads(captured.out)

    assert exit_code == cli.ExitCodes.MISSING_RESOURCE
    assert payload["ok"] is False
    assert payload["diagnostics"][0]["code"] == "UNKNOWN_SYNTAX_VERSION"


def test_lint_reports_malformed_markers_without_writing_docx(tmp_path, capsys):
    input_file = tmp_path / "broken.txt"
    input_file.write_text("[TABLE_START]\n| A |\n", encoding="utf-8")

    exit_code = cli.main(
        [
            "--format",
            "json",
            "--workdir",
            str(tmp_path),
            "lint",
            "--input",
            "broken.txt",
        ]
    )

    captured = capsys.readouterr()
    payload = json.loads(captured.out)

    assert exit_code == cli.ExitCodes.VALIDATION_ERROR
    assert payload["ok"] is False
    assert payload["summary"]["errors"] >= 1
    assert any(diagnostic["code"] == "TXT_MISSING_BLOCK_END" for diagnostic in payload["diagnostics"])
    assert not list(tmp_path.glob("*.docx"))


def test_lint_strict_returns_warning_exit_for_unsupported_rules(tmp_path, capsys):
    input_file = tmp_path / "report.txt"
    input_file.write_text("Body text", encoding="utf-8")

    exit_code = cli.main(
        [
            "--format",
            "json",
            "--workdir",
            str(tmp_path),
            "lint",
            "--input",
            "report.txt",
            "--strict",
        ]
    )

    captured = capsys.readouterr()
    payload = json.loads(captured.out)

    assert exit_code == cli.ExitCodes.WARNINGS_STRICT
    assert payload["ok"] is False
    assert any(
        diagnostic["code"] == "FORMAT_RULE_NOT_SUPPORTED"
        for diagnostic in payload["diagnostics"]
    )


def test_list_profiles_outputs_registered_profile_keys(capsys):
    exit_code = cli.main(["--format", "json", "list-profiles"])

    captured = capsys.readouterr()
    payload = json.loads(captured.out)

    assert exit_code == cli.ExitCodes.SUCCESS
    assert {profile["name"] for profile in payload["profiles"]} == set(PROFILES)
    assert payload["total"] == len(PROFILES)
    common = next(profile for profile in payload["profiles"] if profile["name"] == "common")
    assert common["ruleCount"] > 0
    assert "rendererSupport" in common
    assert "validatorSupport" in common
    assert "support" in common
    assert "parserSupport" in common["support"]
    assert "ruleSupport" in common
    assert "unsupportedRendererRuleIds" in common
    assert "unsupportedValidatorRuleIds" in common

    coursework = next(profile for profile in payload["profiles"] if profile["name"] == "coursework")
    assert "coursework.frame.course_project_explanatory_note" not in coursework["unsupportedRendererRuleIds"]
    assert "coursework.title_block.field_2_designation" in coursework["unsupportedRendererRuleIds"]
    assert "student" in coursework["requiredMetadata"]
    assert coursework["titlePageForm"] == "appendix_i"


def test_export_coverage_outputs_markdown_and_json(capsys):
    md_exit = cli.main(["export-coverage", "--format", "md"])
    md_output = capsys.readouterr().out

    json_exit = cli.main(["export-coverage", "--format", "json"])
    json_output = capsys.readouterr().out
    payload = json.loads(json_output)

    assert md_exit == cli.ExitCodes.SUCCESS
    assert md_output.startswith("# SFU Standard Coverage Matrix\n")
    assert "| Rule ID | Profiles | Source | Severity | Parser | Renderer | Validator | Tests |" in md_output
    assert json_exit == cli.ExitCodes.SUCCESS
    assert payload["rows"]
    assert payload["rows"][0]["ruleId"]


def test_export_schema_diagnostics_contains_required_fields(capsys):
    exit_code = cli.main(
        [
            "--format",
            "json",
            "export-schema",
            "--schema",
            "diagnostics",
        ]
    )

    captured = capsys.readouterr()
    schema = json.loads(captured.out)

    assert exit_code == cli.ExitCodes.SUCCESS
    required = set(schema["required"])
    assert {"code", "severity", "message", "ruleId", "source"} <= required
    assert schema["$schema"] == "https://json-schema.org/draft/2020-12/schema"

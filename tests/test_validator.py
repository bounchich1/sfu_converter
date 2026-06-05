import pytest
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

from sfu_converter.domain.diagnostics import DiagnosticCodes
from sfu_converter.registry import get_profile
from sfu_converter.validator import StyleValidator


def _set_standard_margins(doc):
    section = doc.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(3)
    section.right_margin = Cm(1)
    section.different_first_page_header_footer = True


class TestStyleValidator:
    """Тесты для валидатора стилей"""
    
    @pytest.fixture
    def validator(self):
        """Создает валидатор"""
        return StyleValidator()
    
    @pytest.fixture
    def valid_doc(self, tmp_path):
        """Создает валидный документ с правильными стилями"""
        doc = Document()
        _set_standard_margins(doc)
        para = doc.add_paragraph("Тестовый текст")
        run = para.runs[0]
        run.font.name = 'Times New Roman'
        run.font.size = Pt(14)
        
        pf = para.paragraph_format
        pf.first_line_indent = Cm(1.25)
        pf.line_spacing = 1.5
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)
        
        doc_path = tmp_path / 'valid.docx'
        doc.save(str(doc_path))
        return doc_path
    
    @pytest.fixture
    def invalid_doc_wrong_font(self, tmp_path):
        """Создает документ с неправильным шрифтом"""
        doc = Document()
        _set_standard_margins(doc)
        para = doc.add_paragraph("Текст")
        run = para.runs[0]
        run.font.name = 'Arial'
        run.font.size = Pt(14)
        
        doc_path = tmp_path / 'invalid_font.docx'
        doc.save(str(doc_path))
        return doc_path
    
    @pytest.fixture
    def invalid_doc_spacing(self, tmp_path):
        """Создает документ с неправильными интервалами"""
        doc = Document()
        _set_standard_margins(doc)
        para = doc.add_paragraph("Текст")
        run = para.runs[0]
        run.font.name = 'Times New Roman'
        run.font.size = Pt(14)
        
        pf = para.paragraph_format
        pf.space_before = Pt(12)
        pf.space_after = Pt(12)
        
        doc_path = tmp_path / 'invalid_spacing.docx'
        doc.save(str(doc_path))
        return doc_path
    
    def test_validator_initialization(self, validator):
        """Тест: Инициализация валидатора"""
        assert validator.errors == []
        assert validator.warnings == []

    def test_validator_accepts_profile_name_and_profile_object(self):
        """Тест: Инициализация валидатора с профилем"""
        by_name = StyleValidator(profile_name="coursework")
        by_object = StyleValidator(profile=get_profile("research_reports"))

        assert by_name._validator.profile.name == "coursework"
        assert by_object._validator.profile.name == "research_reports"
    
    def test_get_pt_value_with_emu(self, validator):
        """Тест: Конвертация EMU в пункты"""
        value = Pt(14)
        result = validator._get_pt_value(value)
        assert result == 14
    
    def test_get_pt_value_with_none(self, validator):
        """Тест: Обработка None значения"""
        result = validator._get_pt_value(None)
        assert result == 0

    def test_get_pt_value_with_raw_number(self, validator):
        """Тест: Обработка числового значения"""
        assert validator._get_pt_value(2.5) == 2.5
    
    def test_validate_font_correct(self, validator):
        """Тест: Валидация правильного шрифта"""
        doc = Document()
        para = doc.add_paragraph("Тест")
        run = para.runs[0]
        run.font.name = 'Times New Roman'
        run.font.size = Pt(14)
        
        issues = validator.validate_font(run, 1)
        assert len(issues) == 0
    
    def test_validate_font_wrong_name(self, validator):
        """Тест: Валидация неправильного шрифта"""
        doc = Document()
        para = doc.add_paragraph("Тест")
        run = para.runs[0]
        run.font.name = 'Arial'
        
        issues = validator.validate_font(run, 1)
        assert len(issues) == 1
        assert 'Arial' in issues[0]

    def test_validate_font_wrong_size(self, validator):
        """Тест: Валидация неправильного размера шрифта"""
        doc = Document()
        para = doc.add_paragraph("Тест")
        run = para.runs[0]
        run.font.name = 'Times New Roman'
        run.font.size = Pt(16)

        issues = validator.validate_font(run, 1)

        assert len(issues) == 1
        assert "Размер" in issues[0]
    
    def test_validate_paragraph_spacing_correct(self, validator):
        """Тест: Валидация правильных интервалов (0)"""
        doc = Document()
        para = doc.add_paragraph("Тест")
        pf = para.paragraph_format
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)
        
        issues = validator.validate_paragraph_spacing(para, 1)
        assert len(issues) == 0
    
    def test_validate_paragraph_spacing_wrong(self, validator):
        """Тест: Валидация неправильных интервалов"""
        doc = Document()
        para = doc.add_paragraph("Тест")
        pf = para.paragraph_format
        pf.space_before = Pt(12)
        pf.space_after = Pt(12)
        
        issues = validator.validate_paragraph_spacing(para, 1)
        assert len(issues) == 2
        assert any('Интервал перед' in issue for issue in issues)
    
    def test_validate_file_valid_document(self, validator, valid_doc):
        """Тест: Валидация корректного документа"""
        result = validator.validate_file(str(valid_doc))
        assert result is True
        assert len(validator.errors) == 0
    
    def test_validate_file_wrong_font(self, validator, invalid_doc_wrong_font):
        """Тест: Валидация документа с неправильным шрифтом"""
        result = validator.validate_file(str(invalid_doc_wrong_font))
        assert result is False
        assert len(validator.errors) > 0
    
    def test_validate_file_wrong_spacing(self, validator, invalid_doc_spacing):
        """Тест: Валидация документа с неправильными интервалами"""
        result = validator.validate_file(str(invalid_doc_spacing))
        assert result is False
        assert any(
            diagnostic.code == DiagnosticCodes.FORMAT_PARAGRAPH_SPACING
            for diagnostic in validator.diagnostics
        )
    
    def test_validate_file_not_exists(self, validator):
        """Тест: Валидация несуществующего файла"""
        result = validator.validate_file('/nonexistent/path.docx')
        assert result is False
        assert len(validator.errors) == 1
    
    def test_get_report(self, validator, invalid_doc_wrong_font):
        """Тест: Получение отчета о валидации"""
        validator.validate_file(str(invalid_doc_wrong_font))
        report = validator.get_report()
        
        assert 'errors' in report
        assert 'warnings' in report
        assert report['errors'] > 0
    
    def test_validate_document_with_table(self, validator, tmp_path):
        """Тест: Валидация документа с таблицей"""
        doc = Document()
        _set_standard_margins(doc)
        table = doc.add_table(rows=2, cols=2)
        
        for row in table.rows:
            for cell in row.cells:
                para = cell.paragraphs[0]
                run = para.runs[0] if para.runs else para.add_run()
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)
        
        doc_path = tmp_path / 'with_table.docx'
        doc.save(str(doc_path))
        
        result = validator.validate_file(str(doc_path))
        assert result is True
    
    def test_validate_line_spacing(self, validator):
        """Тест: Валидация межстрочного интервала"""
        doc = Document()
        para = doc.add_paragraph("Тест")
        pf = para.paragraph_format
        pf.line_spacing = 1.5
        
        issues = validator.validate_line_spacing(para, 1, 1.5)
        assert len(issues) == 0

    def test_legacy_indent_and_line_spacing_helpers_report_errors(self, validator):
        doc = Document()
        heading = doc.add_paragraph("Heading")
        heading.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        heading.paragraph_format.first_line_indent = Cm(1)
        heading.runs[0].bold = True

        body = doc.add_paragraph("Body")
        body.paragraph_format.first_line_indent = Cm(0)
        body.paragraph_format.line_spacing = 1.0

        assert validator._is_header_paragraph(heading) is True
        assert "Заголовок имеет отступ" in validator.validate_first_line_indent(heading, 1)[0]
        assert "Отступ" in validator.validate_first_line_indent(body, 2)[0]
        assert "Интервал" in validator.validate_line_spacing(body, 2, 1.5)[0]

        heading.paragraph_format.first_line_indent = Cm(0)
        body.paragraph_format.first_line_indent = Cm(1.25)
        body.paragraph_format.line_spacing = None

        assert validator.validate_first_line_indent(heading, 1) == []
        assert validator.validate_first_line_indent(body, 2) == []
        assert validator.validate_line_spacing(body, 2, 1.5) == []
    
    def test_full_validation_workflow(self, tmp_path):
        """Тест: Полный рабочий процесс валидации с заголовком и обычным текстом"""
        doc = Document()
        _set_standard_margins(doc)
        
        # Заголовок (по центру, без отступа первой строки)
        para1 = doc.add_paragraph("Заголовок")
        para1.runs[0].font.name = 'Times New Roman'
        para1.runs[0].font.size = Pt(14)
        para1.runs[0].bold = True
        para1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para1.paragraph_format.first_line_indent = Cm(0)

        doc.add_paragraph()
        
        # Обычный текст (с отступом первой строки 1.25 см)
        para2 = doc.add_paragraph("Обычный текст")
        para2.runs[0].font.name = 'Times New Roman'
        para2.runs[0].font.size = Pt(14)
        para2.paragraph_format.first_line_indent = Cm(1.25)
        para2.paragraph_format.line_spacing = 1.5
        para2.paragraph_format.space_before = Pt(0)
        para2.paragraph_format.space_after = Pt(0)
        
        doc_path = tmp_path / 'full_test.docx'
        doc.save(str(doc_path))
        
        validator = StyleValidator()
        result = validator.validate_file(str(doc_path))
        
        report = validator.get_report()
        assert report['errors'] == 0

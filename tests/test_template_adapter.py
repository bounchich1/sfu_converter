from io import BytesIO
from pathlib import Path

import pytest
from docx import Document as DocxDocument
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from sfu_converter.domain.diagnostics import Severity
from sfu_converter.infrastructure.template_adapter import DocxTemplateAdapter
from sfu_converter.ports.template import TemplateDocument, TemplatePort


def _new_template(tmp_path, *, pages=1, bookmark_name=None):
    """Build a small DOCX with the requested number of explicit page breaks.

    ``pages=1`` writes only the first page (no page break). ``pages=2`` writes
    "Page 1" then a page break and "Page 2", etc.
    """

    doc = DocxDocument()
    for index in range(pages):
        para = doc.add_paragraph(f"Template page {index + 1}")
        if index < pages - 1:
            run = para.add_run()
            br = OxmlElement("w:br")
            br.set(qn("w:type"), "page")
            run._element.append(br)

    if bookmark_name is not None:
        para = doc.add_paragraph()
        bookmark_start = OxmlElement("w:bookmarkStart")
        bookmark_start.set(qn("w:id"), "0")
        bookmark_start.set(qn("w:name"), bookmark_name)
        bookmark_end = OxmlElement("w:bookmarkEnd")
        bookmark_end.set(qn("w:id"), "0")
        para._p.append(bookmark_start)
        para._p.append(bookmark_end)

    path = tmp_path / "template.docx"
    doc.save(str(path))
    return path


def _generated_doc_bytes(text="Generated body"):
    doc = DocxDocument()
    doc.add_paragraph(text)
    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def test_docx_template_adapter_implements_template_port():
    assert issubclass(DocxTemplateAdapter, TemplatePort)


def test_load_template_returns_template_document(tmp_path):
    template_path = _new_template(tmp_path, pages=1)
    adapter = DocxTemplateAdapter(base_dir=tmp_path)

    template = adapter.load_template(str(template_path))

    assert isinstance(template, TemplateDocument)
    assert template.path == str(template_path)


def test_find_insertion_point_after_page_one(tmp_path):
    template_path = _new_template(tmp_path, pages=2)
    adapter = DocxTemplateAdapter(base_dir=tmp_path)
    template = adapter.load_template(str(template_path))

    insertion = adapter.find_insertion_point(
        template, mode="preserve-prefix", page=1
    )

    assert insertion.found is True
    assert insertion.element is not None


def test_find_insertion_point_after_page_two(tmp_path):
    template_path = _new_template(tmp_path, pages=3)
    adapter = DocxTemplateAdapter(base_dir=tmp_path)
    template = adapter.load_template(str(template_path))

    insertion = adapter.find_insertion_point(
        template, mode="preserve-prefix", page=2
    )

    assert insertion.found is True
    assert insertion.element is not None


def test_find_insertion_point_with_bookmark(tmp_path):
    template_path = _new_template(tmp_path, pages=1, bookmark_name="CONTENT_START")
    adapter = DocxTemplateAdapter(base_dir=tmp_path)
    template = adapter.load_template(str(template_path))

    insertion = adapter.find_insertion_point(
        template, mode="append", bookmark="CONTENT_START"
    )

    assert insertion.found is True
    assert insertion.element is not None


def test_find_insertion_point_with_missing_bookmark_returns_diagnostic(tmp_path):
    template_path = _new_template(tmp_path, pages=1)
    adapter = DocxTemplateAdapter(base_dir=tmp_path)
    template = adapter.load_template(str(template_path))

    insertion = adapter.find_insertion_point(
        template, mode="append", bookmark="DOES_NOT_EXIST"
    )

    assert insertion.found is False
    assert insertion.diagnostic is not None
    assert insertion.diagnostic.code == "TEMPLATE_BOOKMARK_NOT_FOUND"
    assert insertion.diagnostic.severity is Severity.ERROR


def test_find_insertion_point_with_too_high_page_returns_diagnostic(tmp_path):
    template_path = _new_template(tmp_path, pages=1)
    adapter = DocxTemplateAdapter(base_dir=tmp_path)
    template = adapter.load_template(str(template_path))

    insertion = adapter.find_insertion_point(
        template, mode="preserve-prefix", page=5
    )

    assert insertion.found is False
    assert insertion.diagnostic is not None
    assert insertion.diagnostic.code == "TEMPLATE_PAGE_NOT_FOUND"


def test_find_insertion_point_preserve_prefix_requires_page_or_bookmark(tmp_path):
    template_path = _new_template(tmp_path, pages=1)
    adapter = DocxTemplateAdapter(base_dir=tmp_path)
    template = adapter.load_template(str(template_path))

    insertion = adapter.find_insertion_point(template, mode="preserve-prefix")

    assert insertion.found is False
    assert insertion.diagnostic is not None
    assert insertion.diagnostic.code == "TEMPLATE_PAGE_NOT_FOUND"


def test_find_insertion_point_default_appends_at_end(tmp_path):
    template_path = _new_template(tmp_path, pages=1)
    adapter = DocxTemplateAdapter(base_dir=tmp_path)
    template = adapter.load_template(str(template_path))

    insertion = adapter.find_insertion_point(template, mode="append")

    assert insertion.found is True
    assert insertion.element is None


def test_compose_appends_generated_content_after_template(tmp_path):
    template_path = _new_template(tmp_path, pages=1)
    adapter = DocxTemplateAdapter(base_dir=tmp_path)
    template = adapter.load_template(str(template_path))
    insertion = adapter.find_insertion_point(template, mode="append")

    composed_bytes = adapter.compose(
        template, insertion, _generated_doc_bytes("Generated content")
    )

    composed = DocxDocument(BytesIO(composed_bytes))
    texts = [p.text for p in composed.paragraphs]
    assert "Template page 1" in texts
    assert "Generated content" in texts
    template_index = texts.index("Template page 1")
    generated_index = texts.index("Generated content")
    assert template_index < generated_index


def test_compose_preserves_template_page_break_before_generated(tmp_path):
    template_path = _new_template(tmp_path, pages=1)
    adapter = DocxTemplateAdapter(base_dir=tmp_path)
    template = adapter.load_template(str(template_path))
    insertion = adapter.find_insertion_point(template, mode="append")

    composed_bytes = adapter.compose(
        template, insertion, _generated_doc_bytes("Generated content")
    )

    composed = DocxDocument(BytesIO(composed_bytes))
    body_xml = composed.element.body.xml
    assert 'w:type="page"' in body_xml


def test_compose_truncates_template_after_preserve_prefix_anchor(tmp_path):
    template_path = _new_template(tmp_path, pages=2)
    adapter = DocxTemplateAdapter(base_dir=tmp_path)
    template = adapter.load_template(str(template_path))

    insertion = adapter.find_insertion_point(
        template, mode="preserve-prefix", page=1
    )
    composed_bytes = adapter.compose(
        template, insertion, _generated_doc_bytes("Generated body")
    )

    composed = DocxDocument(BytesIO(composed_bytes))
    texts = [p.text for p in composed.paragraphs if p.text]
    # Template page 1 stays, page 2 is replaced by generated content
    assert "Template page 1" in texts
    assert "Template page 2" not in texts
    assert "Generated body" in texts


def test_compose_with_unresolved_insertion_point_raises(tmp_path):
    template_path = _new_template(tmp_path, pages=1)
    adapter = DocxTemplateAdapter(base_dir=tmp_path)
    template = adapter.load_template(str(template_path))
    insertion = adapter.find_insertion_point(
        template, mode="append", bookmark="MISSING"
    )

    with pytest.raises(ValueError):
        adapter.compose(template, insertion, _generated_doc_bytes())


def test_template_adapter_resolves_relative_path_under_templates_dir(tmp_path):
    templates_dir = tmp_path / "templates"
    templates_dir.mkdir()
    doc = DocxDocument()
    doc.add_paragraph("Resolved")
    doc.save(str(templates_dir / "rel.docx"))
    adapter = DocxTemplateAdapter(base_dir=tmp_path)

    template = adapter.load_template("rel.docx")

    assert Path(template.path).name == "rel.docx"
    assert template.doc.paragraphs[0].text == "Resolved"


def test_template_adapter_covers_replace_body_and_page_edge_paths(tmp_path):
    template_path = _new_template(tmp_path, pages=1)
    adapter = DocxTemplateAdapter(base_dir=tmp_path)
    template = adapter.load_template(str(template_path))

    invalid_page = adapter.find_insertion_point(template, mode="preserve-prefix", page=0)
    assert invalid_page.found is False
    assert invalid_page.diagnostic.code == "TEMPLATE_PAGE_NOT_FOUND"

    first_page = adapter.find_insertion_point(template, mode="preserve-prefix", page=1)
    assert first_page.found is True
    assert first_page.element is None

    replace_body = adapter.find_insertion_point(template, mode="replace-body")
    assert replace_body.found is True
    assert replace_body.truncate is True

    composed_bytes = adapter.compose(
        template,
        replace_body,
        _generated_doc_bytes("Replacement"),
    )
    composed = DocxDocument(BytesIO(composed_bytes))
    texts = [paragraph.text for paragraph in composed.paragraphs if paragraph.text]
    assert "Template page 1" not in texts
    assert "Replacement" in texts


def test_template_adapter_resolve_fallbacks_and_nested_bookmark(tmp_path):
    direct_template = tmp_path / "direct.docx"
    _new_template(tmp_path, pages=1).replace(direct_template)
    adapter = DocxTemplateAdapter(base_dir=tmp_path)

    assert adapter._resolve(str(direct_template)) == direct_template
    assert adapter._resolve("missing.docx") == tmp_path / "templates" / "missing.docx"

    doc = DocxDocument()
    paragraph = doc.add_paragraph()
    run = paragraph.add_run("anchor")
    bookmark_start = OxmlElement("w:bookmarkStart")
    bookmark_start.set(qn("w:id"), "1")
    bookmark_start.set(qn("w:name"), "NESTED")
    run._r.append(bookmark_start)
    template = TemplateDocument(doc=doc, path="memory")

    insertion = adapter.find_insertion_point(template, mode="append", bookmark="NESTED")
    assert insertion.found is True
    assert insertion.element.tag.endswith("}p")

    body_bookmark = OxmlElement("w:bookmarkStart")
    body_bookmark.set(qn("w:id"), "2")
    body_bookmark.set(qn("w:name"), "BODY")
    doc.element.body.append(body_bookmark)
    body_insertion = adapter.find_insertion_point(template, mode="append", bookmark="BODY")
    assert body_insertion.found is True
    assert body_insertion.element is None


def test_template_adapter_non_page_break_and_compose_without_page_break(tmp_path, monkeypatch):
    doc = DocxDocument()
    paragraph = doc.add_paragraph("Line break")
    run = paragraph.add_run()
    br = OxmlElement("w:br")
    run._element.append(br)
    template = TemplateDocument(doc=doc, path="memory")
    adapter = DocxTemplateAdapter(base_dir=tmp_path)

    insertion = adapter.find_insertion_point(template, mode="preserve-prefix", page=1)
    assert insertion.found is True
    assert insertion.element is None

    monkeypatch.setattr(adapter, "_make_page_break_paragraph", lambda _doc: None)
    composed_bytes = adapter.compose(
        template,
        adapter.find_insertion_point(template, mode="append"),
        _generated_doc_bytes("Generated"),
    )
    composed = DocxDocument(BytesIO(composed_bytes))
    assert "Generated" in [paragraph.text for paragraph in composed.paragraphs]


def test_template_adapter_append_without_section_properties(tmp_path):
    adapter = DocxTemplateAdapter(base_dir=tmp_path)
    doc = DocxDocument()
    body = doc.element.body
    sect_pr = body.find(qn("w:sectPr"))
    if sect_pr is not None:
        body.remove(sect_pr)

    paragraph = doc.add_paragraph("Appended")
    adapter._append_before_sect_pr(body, paragraph._p, None)

    assert body[-1] is paragraph._p

    external_anchor = OxmlElement("w:p")
    adapter._truncate_after(body, external_anchor, None)
    assert body[-1] is paragraph._p

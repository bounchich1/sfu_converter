"""Tests for ``sfu_converter.infrastructure.paragraph_roles`` (Task 05)."""

from __future__ import annotations

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt

from sfu_converter.config import SIBFUConfig
from sfu_converter.domain.ast_nodes import (
    BibliographyEntryNode,
    Document as AstDocument,
    FigureNode,
    FormulaNode,
    ListItemNode,
    ListNode,
    ListType,
    ParagraphNode,
    TableCell,
    TableNode,
    TableRow,
    TextRun,
)
from sfu_converter.domain.formatting import FormattingProfile
from sfu_converter.infrastructure import docx_styles
from sfu_converter.infrastructure.docx_renderer import DocxRenderer
from sfu_converter.infrastructure.paragraph_roles import ParagraphRole, classify
from sfu_converter.registry import get_profile


def _para(doc, text, *, style_name=None, alignment=None, bold=False, indent_cm=None):
    p = doc.add_paragraph(text)
    if style_name and style_name in [s.name for s in doc.styles]:
        p.style = doc.styles[style_name]
    if alignment is not None:
        p.paragraph_format.alignment = alignment
    if indent_cm is not None:
        p.paragraph_format.first_line_indent = Cm(indent_cm)
    if p.runs:
        run = p.runs[0]
        run.font.name = "Times New Roman"
        run.font.size = Pt(14)
        run.bold = bold
    return p


def test_classify_recognises_word_heading_styles():
    doc = Document()
    h1 = _para(doc, "Введение", style_name="Heading 1")
    h2 = _para(doc, "Подраздел", style_name="Heading 2")
    h3 = _para(doc, "Пункт", style_name="Heading 3")
    h4 = _para(doc, "Подпункт", style_name="Heading 4")

    assert classify(h1) is ParagraphRole.HEADING_H1
    assert classify(h2) is ParagraphRole.HEADING_H2
    assert classify(h3) is ParagraphRole.HEADING_H3
    assert classify(h4) is ParagraphRole.HEADING_H4


def test_classify_recognises_caption_word_style_for_figure():
    doc = Document()
    p = _para(doc, "Рисунок 1 — Foo", style_name="Caption", alignment=WD_ALIGN_PARAGRAPH.CENTER, indent_cm=0)
    assert classify(p) is ParagraphRole.FIGURE_CAPTION


def test_classify_recognises_sfu_custom_styles():
    doc = Document()
    docx_styles.register_styles(doc)

    pairs = (
        (docx_styles.TABLE_CAPTION, ParagraphRole.TABLE_CAPTION),
        (docx_styles.FIGURE_CAPTION, ParagraphRole.FIGURE_CAPTION),
        (docx_styles.FIGURE_EXPLANATORY, ParagraphRole.FIGURE_EXPLANATORY),
        (docx_styles.FIGURE_PLACEHOLDER, ParagraphRole.FIGURE_PLACEHOLDER),
        (docx_styles.FORMULA_BODY, ParagraphRole.FORMULA_BODY),
        (docx_styles.FORMULA_EXPLANATION, ParagraphRole.FORMULA_EXPLANATION),
        (docx_styles.BIBLIOGRAPHY_ENTRY, ParagraphRole.BIBLIOGRAPHY_ENTRY),
        (docx_styles.LIST_ITEM, ParagraphRole.LIST_ITEM),
        (docx_styles.STRUCTURAL_HEADING, ParagraphRole.STRUCTURAL_HEADING),
        (docx_styles.TOC_HEADING, ParagraphRole.TOC_HEADING),
        (docx_styles.APPENDIX_HEADING, ParagraphRole.APPENDIX_HEADING),
    )
    for style_name, role in pairs:
        p = _para(doc, "Body", style_name=style_name)
        assert classify(p) is role, f"{style_name} did not classify as {role}"


def test_classify_falls_back_to_text_regex_for_table_caption():
    doc = Document()
    p = _para(doc, "Таблица 1 — Данные")
    assert classify(p) is ParagraphRole.TABLE_CAPTION

    p2 = _para(doc, "Продолжение таблицы 1")
    assert classify(p2) is ParagraphRole.TABLE_CAPTION


def test_classify_falls_back_to_text_regex_for_figure_caption():
    doc = Document()
    p = _para(doc, "Рисунок 1 — Схема")
    assert classify(p) is ParagraphRole.FIGURE_CAPTION


def test_classify_falls_back_to_text_regex_for_placeholder():
    doc = Document()
    p = _para(doc, "[Изображение не найдено: missing.png]")
    assert classify(p) is ParagraphRole.FIGURE_PLACEHOLDER


def test_classify_recognises_formula_explanation_text():
    doc = Document()
    p = _para(doc, "где E — энергия, Дж;")
    assert classify(p) is ParagraphRole.FORMULA_EXPLANATION


def test_classify_recognises_list_items():
    doc = Document()
    bullet = _para(doc, "- первый элемент")
    lettered = _para(doc, "а) второй элемент")
    nested = _para(doc, "1) вложенный элемент")
    assert classify(bullet) is ParagraphRole.LIST_ITEM
    assert classify(lettered) is ParagraphRole.LIST_ITEM
    assert classify(nested) is ParagraphRole.LIST_ITEM


def test_classify_recognises_structural_headings_by_text():
    doc = Document()
    p = _para(doc, "СОДЕРЖАНИЕ")
    assert classify(p) is ParagraphRole.STRUCTURAL_HEADING


def test_classify_recognises_appendix_heading_by_text():
    doc = Document()
    p = _para(doc, "ПРИЛОЖЕНИЕ А")
    assert classify(p) is ParagraphRole.APPENDIX_HEADING


def test_classify_centered_bold_uppercase_starting_with_digit_is_structural_heading():
    doc = Document()
    p = _para(
        doc,
        "1 ВВЕДЕНИЕ В ТЕМУ",
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        bold=True,
    )
    assert classify(p) is ParagraphRole.STRUCTURAL_HEADING


def test_classify_returns_body_for_plain_paragraph():
    doc = Document()
    p = _para(doc, "Обычный абзац без особенностей.", indent_cm=1.25)
    assert classify(p) is ParagraphRole.BODY


def test_classify_returns_body_for_empty_paragraph():
    doc = Document()
    p = _para(doc, "")
    assert classify(p) is ParagraphRole.BODY


def test_classify_recognises_numbered_headings_by_text():
    doc = Document()
    h1 = _para(doc, "1 Введение")
    h2 = _para(doc, "1.1 Подраздел")
    h3 = _para(doc, "1.1.1 Пункт")
    h4 = _para(doc, "1.1.1.1 Подпункт")
    assert classify(h1) is ParagraphRole.HEADING_H1
    assert classify(h2) is ParagraphRole.HEADING_H2
    assert classify(h3) is ParagraphRole.HEADING_H3
    assert classify(h4) is ParagraphRole.HEADING_H4


# Map every SFU style the renderer applies to the role the validator expects.
# Structural/TOC/appendix headings are intentionally excluded: the renderer
# layers the built-in ``Heading 1`` style on top so Word's TOC field can find
# them, which makes them classify as HEADING_H1 by design.
_GOLDEN_STYLE_ROLES = {
    docx_styles.TABLE_CAPTION: ParagraphRole.TABLE_CAPTION,
    docx_styles.FIGURE_CAPTION: ParagraphRole.FIGURE_CAPTION,
    docx_styles.FORMULA_BODY: ParagraphRole.FORMULA_BODY,
    docx_styles.FORMULA_EXPLANATION: ParagraphRole.FORMULA_EXPLANATION,
    docx_styles.BIBLIOGRAPHY_ENTRY: ParagraphRole.BIBLIOGRAPHY_ENTRY,
    docx_styles.LIST_ITEM: ParagraphRole.LIST_ITEM,
}


def test_classify_matches_renderer_emitted_styles(tmp_path):
    """Golden check: every styled paragraph the renderer emits classifies back
    to the role the validator expects, purely from ``paragraph.style.name``."""

    renderer = DocxRenderer(config_class=SIBFUConfig, base_dir=tmp_path)
    profile = FormattingProfile(
        name="common", display_name="Common", source_docs=("standard",)
    )
    ast = AstDocument(
        blocks=(
            ParagraphNode(runs=(TextRun("Обычный абзац основного текста."),)),
            TableNode(
                caption="Данные",
                rows=(TableRow(cells=(TableCell("A"), TableCell("B"))),),
            ),
            FigureNode(src=None, caption="Рисунок 1 — Схема"),
            FormulaNode(content="E = mc^2", explanation="где E — энергия, Дж."),
            BibliographyEntryNode(number=1, text="Иванов И.И. Основы. — М., 2023."),
            ListNode(
                list_type=ListType.LETTERED,
                items=(ListItemNode("первый"), ListItemNode("второй")),
            ),
        )
    )
    output_path = tmp_path / "golden.docx"
    renderer.render_to_file(ast, profile, str(output_path))

    doc = Document(str(output_path))
    seen_roles: set[ParagraphRole] = set()
    for paragraph in doc.paragraphs:
        style_name = paragraph.style.name if paragraph.style else ""
        expected = _GOLDEN_STYLE_ROLES.get(style_name)
        if expected is None:
            continue
        assert classify(paragraph) is expected, (
            f"paragraph {paragraph.text!r} styled {style_name} "
            f"classified as {classify(paragraph)}, expected {expected}"
        )
        seen_roles.add(expected)

    assert seen_roles == set(_GOLDEN_STYLE_ROLES.values()), (
        f"renderer did not emit all expected styled paragraphs; saw {seen_roles}"
    )

from docx import Document as DocxDocument
from docx.shared import Pt

from sfu_converter.domain import ast_nodes
from sfu_converter.domain.ast_nodes import (
    Document,
    ParagraphNode,
    StructuralSectionNode,
    StructuralSectionType,
    TextRun,
)
from sfu_converter.infrastructure.docx_renderer import DocxRenderer
from sfu_converter.registry import get_profile


def test_abbreviations_section_renders_auto_detected_sorted_rows(tmp_path):
    ast = Document(
        blocks=(
            ParagraphNode(
                runs=(
                    TextRun(
                        "В работе используется информационно-аналитический комплекс (ИАК) "
                        "для мониторинга."
                    ),
                ),
            ),
            ParagraphNode(
                runs=(
                    TextRun(
                        "Центр обработки данных (ЦОД) интегрирован с ИАК и "
                        "информационно-аналитическим комплексом (ИАК)."
                    ),
                ),
            ),
            StructuralSectionNode(
                section_type=StructuralSectionType.ABBREVIATIONS,
                title="СПИСОК СОКРАЩЕНИЙ",
            ),
        )
    )
    output = tmp_path / "abbreviations.docx"

    diagnostics = DocxRenderer(base_dir=tmp_path).render_to_file(
        ast,
        get_profile("common"),
        str(output),
    )

    doc = DocxDocument(str(output))
    assert [(row.cells[0].text, row.cells[1].text) for row in doc.tables[0].rows] == [
        ("ИАК", "— информационно-аналитический комплекс"),
        ("ЦОД", "— Центр обработки данных"),
    ]
    assert doc.tables[0].style.name == "SFUAbbreviationsTable"
    assert all(
        run.font.size == Pt(14)
        for row in doc.tables[0].rows
        for cell in row.cells
        for paragraph in cell.paragraphs
        for run in paragraph.runs
    )
    assert any(
        diagnostic.rule_id == "common.abbreviations.two_column_layout"
        and diagnostic.severity.value == "info"
        for diagnostic in diagnostics
    )


def test_explicit_abbreviations_list_overrides_auto_detection(tmp_path):
    assert hasattr(ast_nodes, "AbbreviationEntryNode")
    assert hasattr(ast_nodes, "AbbreviationsListNode")

    ast = Document(
        blocks=(
            ast_nodes.AbbreviationsListNode(
                entries=(
                    ast_nodes.AbbreviationEntryNode("БД", "база данных"),
                ),
            ),
            ParagraphNode(runs=(TextRun("Центр обработки данных (ЦОД) не должен попасть в список."),)),
            StructuralSectionNode(
                section_type=StructuralSectionType.ABBREVIATIONS,
                title="СПИСОК СОКРАЩЕНИЙ",
            ),
        )
    )
    output = tmp_path / "explicit_abbreviations.docx"

    diagnostics = DocxRenderer(base_dir=tmp_path).render_to_file(
        ast,
        get_profile("common"),
        str(output),
    )

    doc = DocxDocument(str(output))
    assert [(row.cells[0].text, row.cells[1].text) for row in doc.tables[0].rows] == [
        ("БД", "— база данных"),
    ]
    assert not any(diagnostic.code == "ABBREVIATIONS_AUTO_DETECTED" for diagnostic in diagnostics)


def test_abbreviation_collector_ignores_excluded_standard_terms():
    assert hasattr(ast_nodes, "AbbreviationEntryNode")

    from sfu_converter.infrastructure.abbreviations import collect_abbreviations

    document = Document(
        blocks=(
            ParagraphNode(runs=(TextRun("ГОСТ (государственный стандарт) не включается."),)),
            ParagraphNode(runs=(TextRun("Методика использует центр обработки данных (ЦОД)."),)),
        )
    )

    assert collect_abbreviations(document) == (
        ast_nodes.AbbreviationEntryNode("ЦОД", "центр обработки данных"),
    )

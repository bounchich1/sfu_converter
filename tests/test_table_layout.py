from docx import Document as DocxDocument
from docx.shared import Pt

from sfu_converter.infrastructure.docx_renderer import DocxRenderer
from sfu_converter.parser import V2Parser
from sfu_converter.registry import get_profile


def _parse_v2(source: str):
    result = V2Parser().parse(source)
    assert result.diagnostics == []
    return result.document


def test_renderer_outputs_compliant_table_unit_header_borders_and_notes(tmp_path):
    document = _parse_v2(
        "\n".join(
            [
                '[TABLE caption="Параметры" unit="МПа" column_units="-,Гц,°С" header_rows=2]',
                "| Параметр | Частота | Температура |",
                "|----------|---------|-------------|",
                "| Обозначение | f | T |",
                "| Давление | 50 | 20 |",
                '[TABLE_NOTE marker="*" text="Значение приведено при нормальных условиях"]',
                "[TABLE_END]",
            ]
        )
    )
    output = tmp_path / "table_layout.docx"

    DocxRenderer(base_dir=tmp_path).render_to_file(document, get_profile("common"), str(output))

    doc = DocxDocument(str(output))
    unit = next(paragraph for paragraph in doc.paragraphs if paragraph.text == ", МПа")
    assert unit.style.name == "SFUTableUnit"
    assert unit.paragraph_format.alignment.name == "RIGHT"
    assert unit.runs[0].font.size == Pt(12)

    table = doc.tables[0]
    assert table.style.name == "SFUTable"
    assert "Table Grid" not in table.style.name
    assert 'w:val="nil"' in table._tbl.xml
    assert 'w:val="double"' in table.rows[1]._tr.xml
    assert "tblHeader" in table.rows[0]._tr.xml
    assert "tblHeader" in table.rows[1]._tr.xml
    assert table.rows[0].cells[1].text == "Частота, Гц"
    assert table.rows[0].cells[2].text == "Температура, °С"

    note_cell = table.rows[-1].cells[0]
    assert "Значение приведено при нормальных условиях" in note_cell.text
    assert note_cell.paragraphs[0].runs[0].font.superscript is True


def test_renderer_outputs_continuation_label_and_numbering_row(tmp_path):
    document = _parse_v2(
        "\n".join(
            [
                '[TABLE caption="Параметры" number=3 continuation=final]',
                "| Параметр | Значение |",
                "|----------|----------|",
                "| Давление | 10 |",
                "[TABLE_END]",
            ]
        )
    )
    output = tmp_path / "table_continuation.docx"

    DocxRenderer(base_dir=tmp_path).render_to_file(document, get_profile("common"), str(output))

    doc = DocxDocument(str(output))
    assert "Окончание таблицы 3" in [paragraph.text for paragraph in doc.paragraphs]
    assert [cell.text for cell in doc.tables[0].rows[0].cells] == ["1", "2"]
    assert not any(paragraph.text == "Таблица 3 — Параметры" for paragraph in doc.paragraphs)

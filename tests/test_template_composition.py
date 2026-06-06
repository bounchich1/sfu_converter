"""End-to-end converter tests for template composition (Task 23).

Builds a DOCX template on the fly, runs the converter with various
``--template-mode`` settings, and asserts that the produced file preserves
template content while appending generated body content.
"""

from io import BytesIO
from pathlib import Path

import pytest
from docx import Document as DocxDocument
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from sfu_converter.converter import TextToDocxConverter
from sfu_converter.registry import get_profile


@pytest.fixture
def workdir(tmp_path):
    (tmp_path / "templates").mkdir()
    (tmp_path / "examples").mkdir()
    (tmp_path / "results").mkdir()
    return tmp_path


def _add_page_break(paragraph):
    run = paragraph.add_run()
    br = OxmlElement("w:br")
    br.set(qn("w:type"), "page")
    run._element.append(br)


def _build_template(workdir: Path, *, pages: int = 1) -> Path:
    doc = DocxDocument()
    for index in range(pages):
        para = doc.add_paragraph(f"Шаблон страница {index + 1}")
        if index < pages - 1:
            _add_page_break(para)
    template_path = workdir / "templates" / "front.docx"
    doc.save(str(template_path))
    return template_path


def _input_file(
    workdir: Path,
    source: str = '[H level=1 title="Раздел" number=auto]\n[P] Тело документа.',
) -> Path:
    input_path = workdir / "examples" / "input.txt"
    input_path.write_text(source, encoding="utf-8")
    return input_path


def test_converter_appends_after_template_in_append_mode(workdir):
    template = _build_template(workdir, pages=1)
    input_path = _input_file(workdir)
    output_path = workdir / "results" / "appended.docx"

    converter = TextToDocxConverter(base_dir=workdir)
    converter.convert_file(
        input_path,
        output_path,
        template=str(template),
        template_mode="append",
        profile=get_profile("common"),
    )

    composed = DocxDocument(str(output_path))
    texts = [p.text for p in composed.paragraphs if p.text]
    assert "Шаблон страница 1" in texts
    assert any("Тело документа" in t for t in texts)


def test_converter_preserve_prefix_truncates_template_after_anchor(workdir):
    template = _build_template(workdir, pages=2)
    input_path = _input_file(workdir)
    output_path = workdir / "results" / "preserve.docx"

    converter = TextToDocxConverter(base_dir=workdir)
    converter.convert_file(
        input_path,
        output_path,
        template=str(template),
        template_mode="preserve-prefix",
        insert_after_page=1,
        profile=get_profile("common"),
    )

    composed = DocxDocument(str(output_path))
    texts = [p.text for p in composed.paragraphs if p.text]
    assert "Шаблон страница 1" in texts
    assert "Шаблон страница 2" not in texts
    assert any("Тело документа" in t for t in texts)


def test_converter_writes_pristine_output_when_bookmark_missing(workdir, caplog):
    template = _build_template(workdir, pages=1)
    input_path = _input_file(workdir)
    output_path = workdir / "results" / "missing_bookmark.docx"

    converter = TextToDocxConverter(base_dir=workdir)
    with caplog.at_level("ERROR"):
        converter.convert_file(
            input_path,
            output_path,
            template=str(template),
            insert_at_bookmark="DOES_NOT_EXIST",
            profile=get_profile("common"),
        )

    # Output exists (the pristine generated body) but template was not composed.
    assert output_path.exists()
    composed = DocxDocument(str(output_path))
    texts = [p.text for p in composed.paragraphs if p.text]
    assert any("Тело документа" in t for t in texts)
    assert not any("Шаблон" in t for t in texts)

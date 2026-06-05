from docx import Document as DocxDocument
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Cm

from sfu_converter.domain.ast_nodes import (
    AppendixNode,
    Document,
    HeadingLevel,
    HeadingNode,
    TableOfContentsNode,
)
from sfu_converter.domain.diagnostics import DiagnosticCodes, Severity
from sfu_converter.infrastructure import docx_styles
from sfu_converter.infrastructure.docx_renderer import DocxRenderer
from sfu_converter.infrastructure.docx_validator import DocxValidator
from sfu_converter.infrastructure.toc import build_toc_field
from sfu_converter.registry import get_profile


def test_short_document_without_explicit_toc_skips_and_reports_info():
    field = build_toc_field(
        Document(blocks=(HeadingNode(level=HeadingLevel.H1, text="Раздел", number="auto"),)),
        profile=get_profile("common"),
        total_pages=12,
    )

    assert field.should_insert is False
    assert field.diagnostics[0].code == DiagnosticCodes.TOC_NOT_REQUIRED_FOR_SHORT_DOCUMENT
    assert field.diagnostics[0].severity is Severity.INFO


def test_short_document_with_explicit_toc_still_inserts():
    field = build_toc_field(
        Document(
            blocks=(
                TableOfContentsNode(),
                HeadingNode(level=HeadingLevel.H1, text="Раздел", number="auto"),
            )
        ),
        profile=get_profile("common"),
        total_pages=12,
    )

    assert field.should_insert is True
    assert [entry.text for entry in field.entries] == ["1 Раздел"]


def test_renderer_inserts_toc_entries_with_required_indents(tmp_path):
    ast = Document(
        blocks=(
            HeadingNode(level=HeadingLevel.H1, text="Раздел", number="auto"),
            HeadingNode(level=HeadingLevel.H2, text="Подраздел", number="auto"),
        ),
        metadata={"total_pages": "30"},
    )
    output_path = tmp_path / "toc_entries.docx"

    diagnostics = DocxRenderer(base_dir=tmp_path).render_to_file(
        ast,
        get_profile("common"),
        str(output_path),
    )

    assert not any(diagnostic.code == DiagnosticCodes.TOC_NOT_REQUIRED_FOR_SHORT_DOCUMENT for diagnostic in diagnostics)
    doc = DocxDocument(str(output_path))
    body_xml = doc.element.body.xml
    assert " TOC " in body_xml
    toc_entries = [paragraph for paragraph in doc.paragraphs if paragraph.style.name.startswith("TOC ")]
    assert [paragraph.text.split("\t", 1)[0] for paragraph in toc_entries[:2]] == [
        "1 Раздел",
        "1.1 Подраздел",
    ]
    assert toc_entries[0].paragraph_format.first_line_indent == Cm(0)
    assert toc_entries[1].paragraph_format.first_line_indent == Cm(0)
    assert abs(toc_entries[0].paragraph_format.left_indent - Cm(0)) < 1000
    assert abs(toc_entries[1].paragraph_format.left_indent - Cm(0.5)) < 1000


def test_coursework_course_work_uses_form_t_layout():
    field = build_toc_field(
        Document(
            blocks=(HeadingNode(level=HeadingLevel.H1, text="Общие сведения", number="auto"),),
            metadata={"course_work": "true"},
        ),
        profile=get_profile("coursework"),
        total_pages=30,
    )

    assert field.layout == "form_t"
    assert field.title == "СОДЕРЖАНИЕ"


def test_appendix_grouping_collapses_three_contiguous_appendices():
    field = build_toc_field(
        Document(
            blocks=(
                AppendixNode(title="ПРИЛОЖЕНИЕ А", letter="А", subtitle="Данные"),
                AppendixNode(title="ПРИЛОЖЕНИЕ Б", letter="Б", subtitle="Схемы"),
                AppendixNode(title="ПРИЛОЖЕНИЕ В", letter="В", subtitle="Таблицы"),
            )
        ),
        profile=get_profile("common"),
        total_pages=30,
    )

    assert [entry.text for entry in field.entries] == ["Приложения А–В"]


def test_validator_reports_toc_entry_that_does_not_match_heading(tmp_path):
    doc = DocxDocument()
    docx_styles.register_styles(doc)
    doc.styles.add_style("TOC 1", WD_STYLE_TYPE.PARAGRAPH)
    toc_entry = doc.add_paragraph("Другой раздел\t1")
    toc_entry.style = doc.styles["TOC 1"]
    toc_entry.paragraph_format.first_line_indent = Cm(0)
    toc_entry.paragraph_format.left_indent = Cm(0)
    heading = doc.add_paragraph("1 Правильный раздел")
    heading.style = doc.styles["Heading 1"]
    output_path = tmp_path / "bad_toc.docx"
    doc.save(str(output_path))

    diagnostics = DocxValidator(get_profile("common")).validate_file(str(output_path))

    assert any(
        diagnostic.code == DiagnosticCodes.TOC_ENTRY_MISMATCH
        and diagnostic.rule_id == "common.toc.matches_headings"
        for diagnostic in diagnostics
    )

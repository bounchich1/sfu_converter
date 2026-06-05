from docx import Document as DocxDocument

from sfu_converter.domain.ast_nodes import (
    Document,
    ParagraphNode,
    ProjectDesignationNode,
    SectionSetupNode,
    TextRun,
    TitleBlockForm,
)
from sfu_converter.domain.diagnostics import DiagnosticCodes, Severity
from sfu_converter.infrastructure.docx_renderer import DocxRenderer
from sfu_converter.infrastructure.main_inscription import graph_text
from sfu_converter.infrastructure.project_designation import (
    format_designation,
    validate_designation,
    validate_document_designations,
)
from sfu_converter.parser import V2Parser
from sfu_converter.registry import get_profile


def test_v2_designation_formats_as_letter_numeric_code():
    result = V2Parser().parse(
        "[DESIGNATION prefix=ДП specialty=23.05.02 "
        "group=ABCDEF.001 document=СБ year=2021]"
    )

    assert result.diagnostics == []
    (designation,) = result.document.blocks
    assert isinstance(designation, ProjectDesignationNode)
    assert format_designation(designation) == "ДП-23.05.02-2021 ABCDEF.001 СБ"


def test_project_designation_reports_unknown_document_code_with_suggestions():
    node = ProjectDesignationNode(
        prefix="ДП",
        specialty_code="23.05.02",
        group_code="ABCDEF.001",
        document_code="ZZ",
    )

    diagnostics = validate_designation(node)

    diagnostic = next(
        item for item in diagnostics if item.code == DiagnosticCodes.PROJECT_DESIGNATION_CODE
    )
    assert diagnostic.rule_id == "project_designations.code.dictionary"
    assert "Closest matches" in diagnostic.message
    assert "ПЗ" in diagnostic.message


def test_project_designation_formats_gost_r_designation_without_group_code():
    node = ProjectDesignationNode(
        prefix="ДП",
        specialty_code="08.05.01",
        document_code="ПЗ",
    )

    assert format_designation(node) == "ДП-08.05.01 ПЗ"


def test_project_designation_rejects_out_of_range_year():
    node = ProjectDesignationNode(
        prefix="ДП",
        specialty_code="23.05.02",
        group_code="ABCDEF.001",
        document_code="СБ",
        year="1899",
    )

    diagnostics = validate_designation(node)

    assert any(
        diagnostic.code == DiagnosticCodes.PROJECT_DESIGNATION_YEAR
        and diagnostic.rule_id == "project_designations.code.format"
        for diagnostic in diagnostics
    )


def test_project_designation_rejects_unknown_schema_code():
    result = V2Parser().parse(
        "[DESIGNATION prefix=ДП specialty=23.05.02 "
        "group=ABCDEF.001 document=СБ schema=Х5]"
    )

    diagnostics = validate_designation(result.document.blocks[0])

    assert any(
        diagnostic.code == DiagnosticCodes.PROJECT_DESIGNATION_SCHEMA
        and diagnostic.rule_id == "project_designations.code.dictionary"
        for diagnostic in diagnostics
    )


def test_coursework_without_project_designation_warns():
    document = Document(blocks=(ParagraphNode(runs=(TextRun("Body"),)),))

    diagnostics = validate_document_designations(document, get_profile("coursework"))

    assert any(
        diagnostic.code == DiagnosticCodes.PROJECT_DESIGNATION_MISSING
        and diagnostic.severity is Severity.WARNING
        and diagnostic.rule_id == "project_designations.title_block.letter_numeric_designation"
        for diagnostic in diagnostics
    )


def test_renderer_places_designation_in_title_block_graph_two(tmp_path):
    designation = ProjectDesignationNode(
        prefix="ДП",
        specialty_code="23.05.02",
        group_code="ABCDEF.001",
        document_code="СБ",
        year="2021",
    )
    document = Document(
        blocks=(
            designation,
            SectionSetupNode(
                title_block_form=TitleBlockForm.FORM_1,
                blocks=(ParagraphNode(runs=(TextRun("Framed body"),)),),
            ),
        )
    )
    output = tmp_path / "designation.docx"

    DocxRenderer(base_dir=tmp_path).render_to_file(
        document, get_profile("coursework"), str(output)
    )

    rendered = DocxDocument(str(output))
    assert graph_text(rendered.tables[-1], 2) == "ДП-23.05.02-2021 ABCDEF.001 СБ"

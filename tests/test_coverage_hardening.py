import json
import logging
from pathlib import Path

import pytest
from docx import Document
from docx.shared import Cm, Emu, Inches
from PIL import Image

import sfu_converter.main as legacy_main
import sfu_converter.tools.check_cyrillic_markers as check_tool
import sfu_converter.tools.fix_cyrillic_markers as fix_tool
from sfu_converter import cli
from sfu_converter.infrastructure.filesystem import LocalFilesystem
from sfu_converter.menu import ConsoleMenu
from sfu_converter.parser.syntax_spec import get_syntax_spec
from sfu_converter.utils_image_insert import (
    _to_emu,
    calculate_image_dimensions,
    convert_image_to_rgb,
    insert_image,
    insert_image_into_paragraph,
    save_image_to_buffer,
)


def _queue_inputs(monkeypatch, *answers):
    values = iter(answers)
    monkeypatch.setattr("builtins.input", lambda _prompt="": next(values))


def _make_docx(path: Path, text: str = "Document") -> None:
    doc = Document()
    doc.add_paragraph(text)
    doc.save(str(path))


def test_syntax_spec_rejects_unknown_version():
    with pytest.raises(ValueError, match="Unsupported syntax version"):
        get_syntax_spec(99)


def test_local_filesystem_reads_text_and_writes_bytes(tmp_path):
    fs = LocalFilesystem()
    source = tmp_path / "in.txt"
    output = tmp_path / "nested" / "out.bin"
    source.write_text("text", encoding="utf-8")

    assert fs.read_text(str(source)) == "text"

    fs.write_bytes(str(output), b"bytes")
    assert output.read_bytes() == b"bytes"


def test_cli_parser_error_uses_invalid_usage_exit(capsys):
    with pytest.raises(SystemExit) as exc_info:
        cli.main([])

    captured = capsys.readouterr()
    assert exc_info.value.code == cli.ExitCodes.INVALID_USAGE
    assert "error:" in captured.err


def test_cli_convert_missing_template_reports_text_error(tmp_path, capsys):
    input_file = tmp_path / "report.txt"
    input_file.write_text("Body", encoding="utf-8")

    exit_code = cli.main(
        [
            "--workdir",
            str(tmp_path),
            "convert",
            "--input",
            input_file.name,
            "--output",
            "out.docx",
            "--template",
            "missing.docx",
        ]
    )

    captured = capsys.readouterr()
    assert exit_code == cli.ExitCodes.MISSING_RESOURCE
    assert "Template not found" in captured.err


def test_cli_convert_write_failure_reports_json(tmp_path, monkeypatch, capsys):
    input_file = tmp_path / "report.txt"
    input_file.write_text("Body", encoding="utf-8")

    class FailingConverter:
        def __init__(self, *_args, **_kwargs):
            pass

        def convert_file(self, *_args, **_kwargs):
            raise OSError("disk full")

    monkeypatch.setattr("sfu_converter.converter.TextToDocxConverter", FailingConverter)

    exit_code = cli.main(
        [
            "--format",
            "json",
            "--workdir",
            str(tmp_path),
            "convert",
            "--input",
            input_file.name,
            "--output",
            "out.docx",
        ]
    )

    payload = json.loads(capsys.readouterr().out)
    assert exit_code == cli.ExitCodes.WRITE_FAILURE
    assert payload["diagnostics"][0]["code"] == "WRITE_FAILURE"


def test_cli_convert_validation_failure_and_quiet_text_branches(tmp_path, monkeypatch, capsys):
    input_file = tmp_path / "report.txt"
    input_file.write_text("Body", encoding="utf-8")

    class InvalidValidator:
        def __init__(self, *_args, **_kwargs):
            pass

        def validate_file(self, _path):
            return False

        def get_report(self):
            return {"error_list": ["bad"], "warning_list": ["warn"]}

    monkeypatch.setattr("sfu_converter.validator.StyleValidator", InvalidValidator)

    exit_code = cli.main(
        [
            "--format",
            "json",
            "--workdir",
            str(tmp_path),
            "convert",
            "--input",
            input_file.name,
            "--output",
            "validated.docx",
            "--validate-output",
        ]
    )
    payload = json.loads(capsys.readouterr().out)
    assert exit_code == cli.ExitCodes.VALIDATION_ERROR
    assert payload["diagnostics"][0]["message"] == "bad"

    class ValidValidator:
        def __init__(self, *_args, **_kwargs):
            pass

        def validate_file(self, _path):
            return True

        def get_report(self):
            return {"error_list": [], "warning_list": []}

    monkeypatch.setattr("sfu_converter.validator.StyleValidator", ValidValidator)
    exit_code = cli.main(
        [
            "--format",
            "json",
            "--workdir",
            str(tmp_path),
            "convert",
            "--input",
            input_file.name,
            "--output",
            "valid-output.docx",
            "--validate-output",
        ]
    )
    assert exit_code == cli.ExitCodes.SUCCESS
    capsys.readouterr()

    exit_code = cli.main(
        [
            "--quiet",
            "--workdir",
            str(tmp_path),
            "convert",
            "--input",
            input_file.name,
            "--output",
            "quiet.docx",
        ]
    )
    assert exit_code == cli.ExitCodes.SUCCESS
    assert capsys.readouterr().out == ""


def test_cli_convert_write_failure_reports_text(tmp_path, monkeypatch, capsys):
    input_file = tmp_path / "report.txt"
    input_file.write_text("Body", encoding="utf-8")

    class FailingConverter:
        def __init__(self, *_args, **_kwargs):
            pass

        def convert_file(self, *_args, **_kwargs):
            raise OSError("disk full")

    monkeypatch.setattr("sfu_converter.converter.TextToDocxConverter", FailingConverter)

    exit_code = cli.main(
        [
            "--workdir",
            str(tmp_path),
            "convert",
            "--input",
            input_file.name,
            "--output",
            "out.docx",
        ]
    )

    assert exit_code == cli.ExitCodes.WRITE_FAILURE
    assert "disk full" in capsys.readouterr().err

    exit_code = cli.main(
        [
            "--quiet",
            "--workdir",
            str(tmp_path),
            "convert",
            "--input",
            input_file.name,
            "--output",
            "out.docx",
        ]
    )

    assert exit_code == cli.ExitCodes.WRITE_FAILURE
    assert capsys.readouterr().err == ""


def test_cli_validate_text_output_for_invalid_document(tmp_path, monkeypatch, capsys):
    docx_path = tmp_path / "report.docx"
    _make_docx(docx_path)

    class InvalidValidator:
        def __init__(self, *_args, **_kwargs):
            pass

        def validate_file(self, _path):
            return False

        def get_report(self):
            return {"error_list": ["bad"], "warning_list": []}

    monkeypatch.setattr("sfu_converter.validator.StyleValidator", InvalidValidator)

    exit_code = cli.main(["validate-docx", "--input", str(docx_path)])

    captured = capsys.readouterr()
    assert exit_code == cli.ExitCodes.VALIDATION_ERROR
    assert "Validation failed" in captured.out
    assert "ERROR: bad" in captured.out


def test_cli_validate_text_valid_and_quiet_branches(tmp_path, monkeypatch, capsys):
    docx_path = tmp_path / "report.docx"
    _make_docx(docx_path)

    class ValidValidator:
        def __init__(self, *_args, **_kwargs):
            pass

        def validate_file(self, _path):
            return True

        def get_report(self):
            return {"error_list": [], "warning_list": []}

    monkeypatch.setattr("sfu_converter.validator.StyleValidator", ValidValidator)

    assert cli.main(["validate-docx", "--input", str(docx_path)]) == cli.ExitCodes.SUCCESS
    assert "Validation passed" in capsys.readouterr().out

    assert (
        cli.main(["--quiet", "validate-docx", "--input", str(docx_path)])
        == cli.ExitCodes.SUCCESS
    )
    assert capsys.readouterr().out == ""


def test_cli_helpers_cover_absolute_template_and_legacy_diagnostics(tmp_path):
    template = tmp_path / "template.docx"
    _make_docx(template)

    assert cli._template_exists(tmp_path, template) is True
    assert cli._validation_diagnostics(
        {"error_list": ["error"], "warning_list": ["warning"]}
    ) == [
        {"code": "VALIDATION_ERROR", "message": "error", "severity": "error"},
        {"code": "VALIDATION_WARNING", "message": "warning", "severity": "warning"},
    ]


def test_cli_quiet_missing_resource_and_explain_syntax(monkeypatch, tmp_path, capsys):
    assert (
        cli.main(
            [
                "--quiet",
                "--workdir",
                str(tmp_path),
                "convert",
                "--input",
                "missing.txt",
                "--output",
                "out.docx",
            ]
        )
        == cli.ExitCodes.MISSING_RESOURCE
    )
    assert capsys.readouterr().err == ""

    assert cli.main(["--quiet", "explain-syntax"]) == cli.ExitCodes.SUCCESS
    assert capsys.readouterr().out == ""

    class Parser:
        def parse_args(self, _argv):
            return type("Args", (), {"command": "unknown"})()

        def print_help(self):
            print("help")

    monkeypatch.setattr(cli, "create_parser", lambda: Parser())
    assert cli.main([]) == cli.ExitCodes.INVALID_USAGE
    assert "help" in capsys.readouterr().out


def test_menu_lists_files_and_templates(tmp_path):
    menu = ConsoleMenu(tmp_path)
    assert menu.get_txt_files() == []
    assert menu.get_templates() == []

    examples = tmp_path / "examples"
    templates = tmp_path / "templates"
    examples.mkdir()
    templates.mkdir()
    (examples / "b.txt").write_text("", encoding="utf-8")
    (examples / "a.txt").write_text("", encoding="utf-8")
    _make_docx(templates / "template.docx")

    assert menu.get_txt_files() == ["a.txt", "b.txt"]
    assert menu.get_templates() == ["template.docx"]


def test_menu_select_template_variants(tmp_path, monkeypatch, capsys):
    menu = ConsoleMenu(tmp_path)
    _queue_inputs(monkeypatch, "")
    assert menu.select_template() is None

    templates = tmp_path / "templates"
    templates.mkdir()
    _make_docx(templates / "template.docx")

    _queue_inputs(monkeypatch, "1", "")
    assert menu.select_template() == "template.docx"

    _queue_inputs(monkeypatch, "0", "")
    assert menu.select_template() is None

    _queue_inputs(monkeypatch, "9", "")
    assert menu.select_template() is None
    assert "Неверный выбор" in capsys.readouterr().out

    _queue_inputs(monkeypatch, "abc", "")
    assert menu.select_template() is None


def test_menu_select_files_variants(tmp_path, monkeypatch):
    menu = ConsoleMenu(tmp_path)
    _queue_inputs(monkeypatch, "")
    assert menu.select_files() == []

    examples = tmp_path / "examples"
    examples.mkdir()
    (examples / "a.txt").write_text("", encoding="utf-8")
    (examples / "b.txt").write_text("", encoding="utf-8")

    _queue_inputs(monkeypatch, "a", "")
    assert menu.select_files() == ["a.txt", "b.txt"]

    _queue_inputs(monkeypatch, "0")
    assert menu.select_files() == []

    _queue_inputs(monkeypatch, "1,2", "")
    assert menu.select_files() == ["a.txt", "b.txt"]

    _queue_inputs(monkeypatch, "bad")
    assert menu.select_files() == []


def test_menu_main_menu_and_exit_path(tmp_path, monkeypatch, capsys):
    menu = ConsoleMenu(tmp_path)
    menu.selected_template = "template.docx"
    menu.selected_files = ["a.txt"]

    _queue_inputs(monkeypatch, "0")
    assert menu.show_main_menu() == "0"

    _queue_inputs(monkeypatch, "0")
    with pytest.raises(SystemExit):
        menu.run(object(), object())

    assert "До свидания" in capsys.readouterr().out


def test_menu_run_dispatches_template_and_file_selection(tmp_path, monkeypatch):
    examples = tmp_path / "examples"
    templates = tmp_path / "templates"
    examples.mkdir()
    templates.mkdir()
    (examples / "a.txt").write_text("", encoding="utf-8")
    _make_docx(templates / "template.docx")

    menu = ConsoleMenu(tmp_path)
    _queue_inputs(monkeypatch, "1", "1", "", "2", "1", "", "0")
    with pytest.raises(SystemExit):
        menu.run(object(), object())

    assert menu.selected_template == "template.docx"
    assert menu.selected_files == ["a.txt"]


def test_menu_run_conversion_success_and_error(tmp_path, monkeypatch, capsys):
    menu = ConsoleMenu(tmp_path)
    menu.selected_template = "template.docx"
    menu.selected_files = ["ok.txt", "bad.txt"]

    class Converter:
        def convert_file(self, input_file, output_file, template):
            assert template == "template.docx"
            if input_file.name == "bad.txt":
                raise RuntimeError("boom")
            output_file.parent.mkdir(parents=True, exist_ok=True)
            _make_docx(output_file)

    _queue_inputs(monkeypatch, "3", "", "0")

    with pytest.raises(SystemExit):
        menu.run(Converter(), object())

    output = capsys.readouterr().out
    assert "Успешно" in output
    assert "Ошибка: boom" in output


def test_menu_run_no_files_invalid_choice_and_validation_paths(tmp_path, monkeypatch, capsys):
    menu = ConsoleMenu(tmp_path)

    _queue_inputs(monkeypatch, "3", "", "bad", "", "4", "", "0")
    with pytest.raises(SystemExit):
        menu.run(object(), object())

    output = capsys.readouterr().out
    assert "Сначала выберите файлы" in output
    assert "Неверный выбор" in output
    assert "Нет сгенерированных файлов" in output

    results = tmp_path / "results"
    results.mkdir()
    _make_docx(results / "last.docx")
    menu = ConsoleMenu(tmp_path)

    class Validator:
        def validate_file(self, path):
            assert path.endswith("last.docx")
            return False

        def get_report(self):
            return {"errors": 2}

    _queue_inputs(monkeypatch, "4", "", "0")
    with pytest.raises(SystemExit):
        menu.run(object(), Validator())

    assert "Найдено ошибок: 2" in capsys.readouterr().out

    class ValidValidator:
        def validate_file(self, _path):
            return True

    _queue_inputs(monkeypatch, "4", "", "0")
    with pytest.raises(SystemExit):
        menu.run(object(), ValidValidator())

    assert "соответствует стандарту" in capsys.readouterr().out


def test_legacy_setup_logging_and_main_paths(tmp_path, monkeypatch, capsys):
    logger = legacy_main.setup_logging(tmp_path / "logs")
    assert isinstance(logger, logging.Logger)
    assert (tmp_path / "logs" / "converter.log").exists()

    class Menu:
        def __init__(self, _base_dir):
            pass

        def run(self, _converter, _validator):
            return None

    monkeypatch.chdir(tmp_path)
    monkeypatch.setattr(legacy_main, "ConsoleMenu", Menu)
    legacy_main.main()

    class InterruptingMenu(Menu):
        def run(self, _converter, _validator):
            raise KeyboardInterrupt

    monkeypatch.setattr(legacy_main, "ConsoleMenu", InterruptingMenu)
    legacy_main.main()
    assert "завершено" in capsys.readouterr().out

    class ExplodingMenu(Menu):
        def run(self, _converter, _validator):
            raise RuntimeError("boom")

    monkeypatch.setattr(legacy_main, "ConsoleMenu", ExplodingMenu)
    legacy_main.main()
    assert "Ошибка: boom" in capsys.readouterr().out


def test_check_cyrillic_main_paths(tmp_path, monkeypatch, capsys):
    monkeypatch.chdir(tmp_path)
    check_tool.main()
    assert "Директория не найдена" in capsys.readouterr().out

    examples = tmp_path / "examples"
    examples.mkdir()
    check_tool.main()
    assert "TXT файлы не найдены" in capsys.readouterr().out

    (examples / "clean.txt").write_text("[H1] Title\n", encoding="utf-8")
    (examples / "bad.txt").write_text("[Н1] Title\n", encoding="utf-8")
    check_tool.main()
    output = capsys.readouterr().out
    assert "ошибок нет" in output
    assert "найдено проблем" in output

    (examples / "bad.txt").write_text("[H1] Fixed\n", encoding="utf-8")
    check_tool.main()
    assert "Все файлы корректны" in capsys.readouterr().out


def test_fix_cyrillic_main_paths(tmp_path, monkeypatch, capsys):
    monkeypatch.chdir(tmp_path)
    fix_tool.main()
    assert "Директория не найдена" in capsys.readouterr().out

    examples = tmp_path / "examples"
    examples.mkdir()
    fix_tool.main()
    assert "TXT файлы не найдены" in capsys.readouterr().out

    (examples / "clean.txt").write_text("[H1] Title\n", encoding="utf-8")
    (examples / "bad.txt").write_text("[Н1] [ТABLE_START] [В] [С] [Е] [К]\n", encoding="utf-8")
    fix_tool.main()
    output = capsys.readouterr().out
    assert "Исправлено замен" in output
    assert "ошибок не найдено" in output
    assert "Все файлы исправлены" in output
    assert "и ещё 1 замен" in output

    clean_file = examples / "already_clean.txt"
    clean_file.write_text("No markers\n", encoding="utf-8")
    assert fix_tool.fix_file(clean_file, create_backup=False)["fixed"] == 0

    remaining_file = examples / "remaining.txt"
    remaining_file.write_text("[Н1]\n", encoding="utf-8")
    assert fix_tool.verify_file(remaining_file) == [
        {"line": 1, "marker": "[Н1]", "char": "Н"}
    ]


def test_fix_cyrillic_main_reports_remaining_errors(tmp_path, monkeypatch, capsys):
    examples = tmp_path / "examples"
    examples.mkdir()
    (examples / "bad.txt").write_text("[Н1]\n", encoding="utf-8")
    monkeypatch.chdir(tmp_path)

    monkeypatch.setattr(
        fix_tool,
        "verify_file",
        lambda _path: [
            {"line": 1, "marker": "[Н1]", "char": "Н"},
            {"line": 2, "marker": "[Т]", "char": "Т"},
            {"line": 3, "marker": "[В]", "char": "В"},
            {"line": 4, "marker": "[С]", "char": "С"},
        ],
    )

    fix_tool.main()
    output = capsys.readouterr().out
    assert "осталось проблем" in output
    assert "и ещё 1" in output
    assert "Некоторые ошибки" in output

    monkeypatch.setattr(
        fix_tool,
        "verify_file",
        lambda _path: [
            {"line": 1, "marker": "[Н1]", "char": "Н"},
            {"line": 2, "marker": "[Т]", "char": "Т"},
            {"line": 3, "marker": "[В]", "char": "В"},
        ],
    )
    fix_tool.main()
    assert "и ещё" not in capsys.readouterr().out


def test_image_conversion_and_dimension_edge_cases(tmp_path):
    assert convert_image_to_rgb(Image.new("LA", (1, 1))).mode == "RGB"
    assert convert_image_to_rgb(Image.new("P", (1, 1))).mode == "RGB"
    assert convert_image_to_rgb(Image.new("CMYK", (1, 1))).mode == "RGB"
    assert convert_image_to_rgb(Image.new("L", (1, 1))).mode == "RGB"

    assert _to_emu(None) is None
    assert _to_emu(Cm(1)) == Cm(1).emu
    assert _to_emu(Inches(1)) == Inches(1).emu
    assert _to_emu(Emu(42)) == 42

    assert calculate_image_dimensions((0, 100), width=Cm(1))[1] == Cm(1)
    assert calculate_image_dimensions((100, 50), width=Cm(2))[1] == Cm(1)
    assert calculate_image_dimensions((100, 50), width=Inches(2))[1] == Inches(1)
    assert calculate_image_dimensions((100, 50), width=200)[1] == 100
    assert calculate_image_dimensions((100, 0), height=Cm(2))[0] == Cm(2)
    assert calculate_image_dimensions((100, 50), height=Inches(1))[0] == Inches(2)
    assert calculate_image_dimensions((100, 50), height=50)[0] == 100
    assert calculate_image_dimensions((10, 5), max_width=Cm(15)) == (None, None)


def test_image_buffer_and_insert_error_paths(tmp_path):
    jpeg_buffer = save_image_to_buffer(Image.new("RGBA", (2, 2)), format="JPEG")
    tiff_buffer = save_image_to_buffer(Image.new("RGB", (2, 2)), format="TIFF")
    assert jpeg_buffer.getbuffer().nbytes > 0
    assert tiff_buffer.getbuffer().nbytes > 0

    doc = Document()
    insert_image_into_paragraph(doc, save_image_to_buffer(Image.new("RGB", (2, 2))))
    assert len(doc.inline_shapes) == 1

    assert insert_image(doc, tmp_path / "missing.png", {}) is False

    bad_image = tmp_path / "bad.png"
    bad_image.write_text("not an image", encoding="utf-8")
    assert insert_image(doc, bad_image, {}) is False
